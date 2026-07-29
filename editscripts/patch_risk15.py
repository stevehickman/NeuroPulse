"""
⚠ SUPERSEDED (2026-07-28): the 5-layer zone-keying system this script writes (mechanical
type-differentiating key + ZONE_ID resistor ladder + per-zone braille/tactile markers +
bone-conduction zone-name announcement) is retired by the hex-tile/socket architecture
(NP-HEX-ZM-001) — SMART-1 makes every socket I2C/TIA-capable, so there is no "wrong zone"
left to key against, and modules carry only a universal orientation-only key. Do not re-run
this script; it targets an already-patched document with insert-only logic (re-running
duplicates content) and, more importantly, would reintroduce retired content. The corrected
RISK-15 state is applied by editscripts/patch_hexzm_risk_corrections.py. Kept here as the
historical record of what was originally shipped and why (RISK-15 rationale, ISO reference
numbers for the braille/tactile approach) — still useful context, not current guidance.

Patch RISK-15:
Mitigate zone keying for colour-vision-deficient AND blind users.

Five-layer mitigation:
  Layer 1: Mechanical asymmetric key — physically impossible to mis-insert (blind-safe)
  Layer 2: Zone ID resistor on SPARE_1 (pin 18) — firmware-verified identity
  Layer 3: Tactile markers on module face (braille + raised numeral) — blind-readable
  Layer 4: Tactile position markers on shell near each slot — blind-readable
  Layer 5: Audio confirmation via existing bone conduction speaker on correct/wrong insert

Updates:
  - Risk register: RISK-15 → MITIGATED with full multi-layer description
  - FPC spec §5.2 pin table: SPARE_1 → ZONE_ID (unique resistor per zone)
  - FPC spec §11.4: add zone ID, braille, tactile shell, audio confirmation requirements
  - Coordination checklist NP-COORD-001 G2-02: already references §11.4; §11.4 is the
    definitive location; the coord checklist is updated by reference
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def ins_before(doc, target_para, text, bold=False, bullet=False, color=None, size=10):
    new_p = OxmlElement("w:p")
    target_para._p.addprevious(new_p)
    for p_obj in doc.paragraphs:
        if p_obj._p is new_p:
            if bullet:
                try:
                    p_obj.style = doc.styles["List Bullet"]
                except Exception:
                    pass
            run = p_obj.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p_obj
    return None

# ─────────────────────────────────────────────────────────────────────────────
# 1.  FPC Spec — update pin 18 in the pinout table
# ─────────────────────────────────────────────────────────────────────────────

SPEC_PATH = "docs/neurone_fpc_zone_module_spec_revA.docx"
spec = Document(SPEC_PATH)

pin18_updated = False
for ti, tbl in enumerate(spec.tables):
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 5 and cells[0].text.strip() == "18" and "SPARE" in cells[1].text:
            # Columns: Pin | Signal | Direction | Current | Notes
            set_cell_text(cells[1], "ZONE_ID", bold=True)
            set_cell_text(cells[2], "ZM → Hub")
            set_cell_text(cells[3], "Signal")
            set_cell_text(cells[4],
                "Zone identity resistor. Unique 1% 0402 resistor per zone module: "
                "ZM-01=10 kΩ, ZM-02=22 kΩ, ZM-03=47 kΩ, ZM-04=100 kΩ, ZM-05=220 kΩ. "
                "Hub MCU reads ADC voltage via 10 kΩ pull-up at session start; "
                "zone ID confirmed before ENABLE is asserted. "
                "Mismatch (module in wrong slot) → FAULT asserted, audio alert via "
                "bone conduction, session blocked until correct module inserted. "
                "BOM delta: one 0402 resistor per zone module (~$0.03 each).")
            pin18_updated = True
            print(f"  Pin 18 updated in Table {ti}")
            break
    if pin18_updated:
        break

if not pin18_updated:
    print("  WARNING: Pin 18 SPARE_1 row not found in spec tables")

# ─────────────────────────────────────────────────────────────────────────────
# 2.  FPC Spec §11.4 — insert full multi-layer keying requirements
#     Insert after the existing §11.4 paragraphs, before §12
# ─────────────────────────────────────────────────────────────────────────────

# Find the §12 paragraph (insertion point)
s12_para = None
for para in spec.paragraphs:
    if para.text.strip().startswith("12."):
        s12_para = para
        break

print(f"  §12 anchor found: {s12_para.text[:60]!r}" if s12_para else "  WARNING: §12 not found")

if s12_para:
    LAYER_CONTENT = [
        # (text, bold, bullet, color)

        # Layer 5 — audio (last inserted = first rendered before §12 because we prepend)
        ("Layer 5 — Audio/haptic confirmation (blind-accessible): "
         "On zone module insertion, the Hub MCU reads the ZONE_ID resistor (pin 18) and "
         "compares it to the expected zone for that slot. "
         "If correct: bone conduction speaker announces the zone name "
         "(e.g., 'Frontal Left connected') — spoken-word audio, not a tone. "
         "If incorrect: bone conduction announces 'Wrong module — expected [zone name]' "
         "and ENABLE is held de-asserted; session cannot start. "
         "If module absent: bone conduction announces the missing zone by name when the "
         "user initiates a session. "
         "Audio confirmation uses the existing bone conduction transducer (already in BOM "
         "for Neural Audio Entrainment modality) — zero additional BOM cost. "
         "Firmware coordinate item: NP-COORD-001 G2-02 and new G2-10 (audio zone ID).",
         False, False, None),

        # Layer 4 — tactile shell markers
        ("[DRAWING REQUIREMENT — SHELL TOOLING] "
         "Layer 4 — Tactile shell position markers (blind-accessible): "
         "Each zone slot opening in the CFRP shell shall have a unique tactile indicator "
         "moulded into the shell surface adjacent to the slot (within 10 mm of slot edge, "
         "in a consistent position across all 5 slots for easy finger-finding). "
         "Indicator pattern: N raised dots (dome, 1.5 mm diameter, 0.8 mm height) where "
         "N = zone number: ZM-01=1 dot, ZM-02=2 dots, ZM-03=3 dots, ZM-04=4 dots, "
         "ZM-05=5 dots. "
         "Dot spacing: 2.5 mm centre-to-centre (ISO 9241-920 tactile element guidance). "
         "Moulded into shell at zero incremental cost if specified before first-cut. "
         "Cannot be retrofitted after tooling — must be in shell tooling spec.",
         True, False, (0xC0, 0x00, 0x00)),

        # Layer 3 — tactile module markers
        ("[DRAWING REQUIREMENT — ZONE MODULE TOOLING] "
         "Layer 3 — Tactile zone module identification (blind-accessible): "
         "Each zone module body shall carry: "
         "(a) a raised numeral (1–5) on the face opposite the LED array, "
         "6 mm height, 1 mm relief, moulded into the body — no paint or label required; "
         "(b) Grade 1 braille abbreviation for the zone position moulded immediately below "
         "the numeral: ZM-01 = 'FL' (Frontal Left), ZM-02 = 'FR' (Frontal Right), "
         "ZM-03 = 'CL' (Central Left), ZM-04 = 'CR' (Central Right), "
         "ZM-05 = 'PO' (Parietal/Occipital). "
         "Braille cell dimensions per ISO 17049:2013 (dot diameter 1.5 mm, "
         "inter-dot spacing 2.5 mm, inter-cell spacing 6 mm). "
         "Moulded into zone module body at zero incremental cost if specified before tooling. "
         "In addition: large-print text identifier (e.g., 'ZM-01 FL') in contrasting colour "
         "for low-vision and sighted users.",
         True, False, (0xC0, 0x00, 0x00)),

        # Layer 2 — zone ID resistor
        ("[DRAWING REQUIREMENT] "
         "Layer 2 — Electronic zone identity (ZONE_ID resistor, pin 18): "
         "Each zone module PCB shall include a unique 1% tolerance 0402 SMD resistor "
         "between ZONE_ID (pin 18) and GND (AGND, pin 17), with values: "
         "ZM-01=10 kΩ, ZM-02=22 kΩ, ZM-03=47 kΩ, ZM-04=100 kΩ, ZM-05=220 kΩ. "
         "The Hub MCU reads ADC(ZONE_ID) through a 10 kΩ pull-up at session start "
         "before asserting ENABLE. ADC thresholds: <14 kΩ=ZM-01, 14–32 kΩ=ZM-02, "
         "32–68 kΩ=ZM-03, 68–150 kΩ=ZM-04, >150 kΩ=ZM-05. "
         "Unrecognised value (e.g., uncertified third-party module, open circuit): "
         "FAULT asserted, session blocked, user notified via audio and app. "
         "BOM delta: +$0.03 per zone module (one 0402 resistor).",
         True, False, (0xC0, 0x00, 0x00)),

        # Layer 1 — mechanical key (already in §11.4, this is the reinforcement)
        ("Layer 1 — Mechanical asymmetric key (primary, blind-safe): "
         "The mechanical key (unique asymmetric tab/notch per zone) is the primary "
         "keying mechanism and the only layer that is purely passive — it requires no "
         "firmware, no reading, and no user attention. A blind user cannot install the "
         "wrong module because the geometry physically prevents it. "
         "All other layers (2–5) provide additional confirmation and user feedback, "
         "but Layer 1 is the safety guarantee. "
         "Key geometry must be defined in shell tooling spec (NP-COORD-001 G2-02) "
         "and verified in FAI (NP-FAI-ZM-001 FAI-A01–A08).",
         False, False, None),

        # Section heading
        ("Multi-layer zone keying — accessibility for colour-vision-deficient and blind users:",
         True, False, (0x1F, 0x49, 0x7D)),
    ]

    for (text, bold, bullet, color) in LAYER_CONTENT:
        ins_before(spec, s12_para, text, bold=bold, bullet=bullet, color=color)

    # Add blank spacer
    ins_before(spec, s12_para, "")

    print("  §11.4 multi-layer keying content inserted")

spec.save(SPEC_PATH)
print(f"  FPC spec saved: {SPEC_PATH}")

# ─────────────────────────────────────────────────────────────────────────────
# 3.  Risk Register — RISK-15 → MITIGATED
# ─────────────────────────────────────────────────────────────────────────────

RISK_PATH = "docs/neurone_fpc_zone_module_risks_revA.docx"
risk_doc = Document(RISK_PATH)

tbl = risk_doc.tables[0]
for row in tbl.rows:
    if row.cells[0].text.strip() == "RISK-15":
        # Severity
        set_cell_text(row.cells[1], "LOW\n(MITIGATED)", bold=True,
                      color=(0x00, 0x70, 0x00), size=9)
        set_cell_bg(row.cells[1], "E2EFDA")
        # Required Mitigation
        set_cell_text(row.cells[5],
            "MITIGATED: Five-layer keying system — accessible to sighted, "
            "colour-vision-deficient, AND blind users. "
            "Layer 1 (blind-safe): Unique asymmetric mechanical key per zone — physically "
            "impossible to mis-insert regardless of vision. "
            "Layer 2: ZONE_ID resistor on pin 18 (SPARE_1) — Hub MCU reads zone identity "
            "at session start; ENABLE blocked if wrong module detected. "
            "Layer 3: Braille abbreviation + raised numeral moulded into module face — "
            "zero BOM cost if specified before tooling. "
            "Layer 4: Raised tactile dot pattern (N dots = zone N) moulded into shell "
            "adjacent to slot — zero BOM cost if specified before tooling. "
            "Layer 5: Bone conduction audio confirmation — spoken zone name on correct "
            "insert; spoken alert on wrong insert; uses existing BOM transducer. "
            "Total BOM delta: +$0.03/zone module (ZONE_ID resistor only). "
            "All other layers are zero-cost if specified before tooling cut. "
            "See FPC spec §11.4 for full requirements.", size=9)
        # Status
        set_cell_text(row.cells[8],
            "MITIGATED — 5-layer keying (mechanical + electronic + braille + "
            "tactile shell + audio); all layers work for blind users",
            bold=True, color=(0x00, 0x70, 0x00), size=9)
        set_cell_bg(row.cells[8], "92D050")
        print("  RISK-15 summary row updated")
        break

# ── Add RISK-15 detailed section before "OPEN ISSUES" ──
open_issues_idx = None
for i, para in enumerate(risk_doc.paragraphs):
    if "OPEN ISSUES" in para.text:
        open_issues_idx = i
        break

target = risk_doc.paragraphs[open_issues_idx]

ins_before(risk_doc, target, "")  # spacer

ins_before(risk_doc, target,
    "→ Action 5 (OPEN — firmware): Implement bone conduction audio zone confirmation "
    "(Layer 5). Firmware reads ZONE_ID resistor, compares to slot map, speaks zone name "
    "or error via bone conduction. Tracked as NP-COORD-001 G2-10.",
    bullet=True, color=(0xC0, 0x60, 0x00))

ins_before(risk_doc, target,
    "→ Action 4 (OPEN — shell tooling spec): Tactile dot patterns adjacent to all 5 "
    "zone slots (Layer 4). Must be in shell tooling spec before first-cut. "
    "Tracked as NP-COORD-001 G2-02.",
    bullet=True, color=(0xC0, 0x60, 0x00))

ins_before(risk_doc, target,
    "→ Action 3 (OPEN — zone module tooling spec): Braille + raised numeral on module "
    "face (Layer 3). Must be in zone module tooling spec before first-cut. "
    "Zero cost if specified now. Tracked as NP-COORD-001 G2-02.",
    bullet=True, color=(0xC0, 0x60, 0x00))

ins_before(risk_doc, target,
    "→ Action 2 (OPEN — PCB layout): Place 0402 ZONE_ID resistor on each zone module "
    "PCB, connected between pin 18 (ZONE_ID) and AGND (pin 17). "
    "Values: ZM-01=10k, ZM-02=22k, ZM-03=47k, ZM-04=100k, ZM-05=220k (all 1%). "
    "Hub MCU firmware reads ADC on ZONE_ID before every ENABLE assertion.",
    bullet=True, color=(0xC0, 0x60, 0x00))

ins_before(risk_doc, target,
    "→ Action 1 (OPEN — shell tooling spec): Unique asymmetric mechanical key geometry "
    "for all 5 zone slots (Layer 1). Must be in shell tooling spec before first-cut. "
    "This is the only layer that is entirely passive and requires no firmware or user "
    "action — it is the safety guarantee for all users including blind. "
    "Tracked as NP-COORD-001 G2-02.",
    bullet=True, color=(0xC0, 0x60, 0x00))

ins_before(risk_doc, target,
    "The original risk considered only colour-vision deficiency (~8% males). The mitigation "
    "is extended to cover completely blind users, since NeurOne's clinical positioning "
    "(neurological conditions including visual processing disorders) makes it plausible that "
    "a proportion of users will have significant visual impairment. The five-layer approach "
    "costs ~$0.15 total per headset (5 ZONE_ID resistors at $0.03 each) if the mould "
    "features (Layers 3 and 4) are specified before tooling. After tooling, Layers 3 and 4 "
    "are lost; only the ZONE_ID resistor and audio can be added, which provides electronic "
    "and audio protection but not tactile identification. The mechanical key (Layer 1) "
    "must be in tooling regardless — it is the blocking physical safety guarantee.")

ins_before(risk_doc, target,
    "RISK-15  |  LOW (MITIGATED)  |  Five-Layer Zone Keying — Colour-Blind and Blind-Accessible",
    bold=True, color=(0x1F, 0x49, 0x7D))

risk_doc.save(RISK_PATH)
print(f"  Risk register saved: {RISK_PATH}")

# ─────────────────────────────────────────────────────────────────────────────
# 4.  Engineering Coordination Checklist — add G2-10 (audio zone ID firmware)
#     The .docx is a generated document; we patch by finding the G2-02 row in
#     Table 0 (the G2-A coordination table) and appending G2-10 nearby.
#     Simpler: append a note in the RISK-15 entry in the risk register
#     and add G2-10 as a note to the coord checklist revision history.
#     Full coord checklist patch would require regeneration — instead we'll
#     note the addition in the revision history table.
# ─────────────────────────────────────────────────────────────────────────────

COORD_PATH = "docs/neurone_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

# Find G2-02 row and update its text to reference blind accessibility
for tbl in coord.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "G2-02":
            # Update description cell (col 2)
            old = row.cells[2].text
            set_cell_text(row.cells[2],
                "RISK-15 — Unique asymmetric mechanical key geometry for each of 5 zone slots "
                "(Layer 1 of 5-layer accessibility system). Key must physically prevent wrong-zone "
                "insertion for ALL users including blind. Also in this item: "
                "(a) tactile dot patterns adjacent to each slot in shell (N dots = zone N, "
                "Layer 4, ISO 9241-920); "
                "(b) braille + raised numeral on zone module face "
                "(Layer 3, ISO 17049:2013). "
                "All three features must be in shell/module tooling spec before first-cut. "
                "Zero BOM cost if specified now. See NP-HW-FPC-001 §11.4.", size=8)
            # Update deliverable cell (col 3)
            set_cell_text(row.cells[3],
                "Key geometry drawings all 5 zones. "
                "Shell drawing showing tactile dot positions (N dots per zone). "
                "Zone module drawing showing braille cell and raised numeral. "
                "Written confirmation no two key profiles are interchangeable.", size=8)
            print(f"  NP-COORD-001 G2-02 row updated")
            break

# Add G2-10 after G2-09 in the dashboard table (Section 5)
# Find the dashboard table by looking for the ID column with "G2-09"
g209_row_ref = None
dash_tbl = None
for tbl in coord.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "G2-09":
            g209_row_ref = row
            dash_tbl = tbl
            break
    if g209_row_ref:
        break

if dash_tbl and g209_row_ref:
    # Add a new row after G2-09
    # python-docx doesn't have a clean "insert row after" method, so we add to end
    # of table and accept ordering is slightly off, or we copy the row XML.
    new_row_xml = OxmlElement("w:tr")
    g209_row_ref._tr.addnext(new_row_xml)

    # Find the new row
    new_row = None
    for row in dash_tbl.rows:
        if row._tr is new_row_xml:
            new_row = row
            break

    if new_row is None:
        # The row was appended; find it as the last row
        new_row = dash_tbl.rows[-1]

    # We need to add cells to the new row
    STATUS_BG = "FFE699"
    col_data = [
        ("G2-10", "G2", "FW + HW EE", "Audio zone ID confirmation via bone conduction — spoken zone name on insert/wrong-module alert (RISK-15 Layer 5)", "OPEN"),
    ]
    widths = [0.6, 0.5, 1.35, 4.0, 0.85]

    # The row was just created empty; add cells via XML
    for c_idx, txt in enumerate(col_data[0]):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        w = OxmlElement("w:tcW")
        w.set(qn("w:w"), str(int(widths[c_idx] * 1440)))
        w.set(qn("w:type"), "dxa")
        tc_pr.append(w)
        tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")  # 8pt
        rpr.append(sz)
        r_el.append(rpr)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve")
        t_el.text = txt
        r_el.append(t_el)
        p_el.append(r_el)
        tc.append(p_el)
        new_row._tr.append(tc)

    # Add shading to status cell (last)
    if new_row.cells:
        set_cell_bg(new_row.cells[-1], STATUS_BG)

    print("  NP-COORD-001 G2-10 dashboard row added")
else:
    print("  WARNING: G2-09 dashboard row not found — G2-10 not added to dashboard")

# Update revision history
for tbl in coord.tables:
    for row in tbl.rows:
        for cell in row.cells:
            if "Initial release" in cell.text and "RISK-13" in cell.text:
                # Add Rev B note to next row or update this cell
                existing = cell.text
                set_cell_text(cell,
                    existing + " | Rev A.1 update: G2-02 extended with blind-user "
                    "accessibility requirements (RISK-15): tactile dots, braille, "
                    "raised numeral. G2-10 added: audio zone ID firmware.", size=8)
                print("  Revision history updated")
                break

coord.save(COORD_PATH)
print(f"  Coord checklist saved: {COORD_PATH}")

print("\nAll patches complete.")
