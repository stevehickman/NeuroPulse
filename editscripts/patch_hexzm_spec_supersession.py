"""
Add a SUPERSEDED banner to NP-HW-FPC-001 Rev 4 (docs/superseded/np_hw_fpc_001.docx).

docs/status/document-register.md already flags this .docx as superseded
(2026-07-28, alongside its markdown sibling docs/superseded/np_hw_fpc_001.md, see commit
b96f2d4). That pass explicitly deferred opening/editing the .docx itself
("out of scope for this pass"). This script closes that gap by inserting the
same style of in-document supersession banner used on the markdown siblings
(np_hw_fpc_001.md, np_tool_zm_sm_001.md, etc.) — distinguishing retired vs.
still-reusable content, not deleting anything.

Retired: 66x78mm zone module footprint (§4), 20-pin Hirose-per-zone pinout
(§2-3), ZONE_ID resistor-ladder detection (§3.2, retired outright by
np_module_map's UID-based auto-inventory), 5-slot I2C/TIA counts (§3.1, §6.4),
5-layer zone keying incl. the type-differentiating mechanical key (§11.4 —
also see the RISK-15 correction in patch_hexzm_risk_corrections.py).

Still-reusable, not yet re-validated against the hex-tile form factor:
on-module driver architecture, InGaAs PD choice, PDMS bonding process.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SPEC_PATH = "docs/superseded/np_hw_fpc_001.docx"

BANNER_MARKER = "SUPERSEDED (2026-07-28)"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def already_patched(doc):
    return any(BANNER_MARKER in p.text for p in doc.paragraphs)


def insert_banner(doc):
    # Anchor: the "1.  PURPOSE AND SCOPE" heading — insert the banner
    # immediately before it, after the cover-page metadata block.
    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "1.  PURPOSE AND SCOPE":
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("Could not find '1.  PURPOSE AND SCOPE' anchor paragraph")

    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "FCE4D6")
    para = cell.paragraphs[0]

    def run(text, bold=False, size=10, color=(0xC0, 0x00, 0x00)):
        r = para.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color)
        return r

    run(
        "⚠ SUPERSEDED (2026-07-28) — mechanical/addressing model retired, do not use for new design work.  ",
        bold=True,
    )
    para.add_run().add_break()
    run(
        "This document specifies the position-unique, large-format (66×78mm) zone-module "
        "architecture: one Hirose FH34S 20-pin connector per zone, five fixed zone-slot positions "
        "(ZM-01..05), and ZONE_ID resistor-ladder detection. NP-HEX-ZM-001 (2026-07-15) replaced this "
        "entirely with a universal 40mm hex-tile SKU tiling ~80 sockets, addressed by UID-based "
        "auto-inventory (np_module_map), not a resistor ladder — see docs/np_hex_zm_001.md §4a.",
        bold=False, color=(0x40, 0x40, 0x40),
    )
    para.add_run().add_break()
    run(
        "Retired outright: the 66×78mm FPC footprint and LED counts/pitch derived from it (§4); "
        "the 20-pin connector and full pinout (§2-3); the ZONE_ID resistor-ladder table (§3.2, pin 18); "
        "per-slot I2C bus/pull-up counts sized for 5 slots (§3.1, §6.4); the type-differentiating "
        "§11.4 five-layer keying system (RISK-15) — SMART-1 makes every socket I2C/TIA-capable, so "
        "there is no \"wrong zone\" left to key against; the module retains only a universal, "
        "orientation-only mechanical key (see docs/np_hex_zm_001.md §4a, "
        "\"No type-differentiating mechanical key\").",
        bold=False, color=(0x40, 0x40, 0x40),
    )
    para.add_run().add_break()
    run(
        "Likely still-reusable, NOT yet re-validated against the hex-tile form factor: on-module "
        "driver architecture (ATtiny402 + FETs), InGaAs PD choice for dose metering, PDMS bonding "
        "process/material spec.",
        bold=False, color=(0x40, 0x40, 0x40),
    )
    para.add_run().add_break()
    run(
        "Open gap this note does not fill: no document yet specifies the T1-A/B/C hex-tile FPC "
        "pinout or per-socket electrical layout at the detail level this document provided for the "
        "retired architecture (see MECH-1/MECH-2/SMART-1 in docs/np_hex_zm_001.md §7).",
        bold=False, color=(0x40, 0x40, 0x40),
    )

    anchor._p.addprevious(t._tbl)
    # spacer paragraph between banner and heading
    spacer = OxmlElement("w:p")
    anchor._p.addprevious(spacer)


def main():
    doc = Document(SPEC_PATH)
    if already_patched(doc):
        print(f"  Already patched — skipping: {SPEC_PATH}")
        return
    insert_banner(doc)
    doc.save(SPEC_PATH)
    print(f"  Supersession banner inserted: {SPEC_PATH}")


if __name__ == "__main__":
    main()
