"""
RISK-04 mitigation: PDMS thermal cycling qualification.
Changes:
  - NP-FAI-ZM-001: fix FAI-M02/M03 thresholds; add §3e thermal cycling
    qualification section (FAI-TC01 through FAI-TC06); add RISK-04 to §9
  - NP-PROC-SUP-001: fix SUP-C-04 peel threshold; add SUP-C-08 thermal
    cycling qualification requirement
  - NP-COORD-001: add G1-08 PDMS qualification gate item
  - Risk register: RISK-04 → MITIGATED
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ───────────────────────────────────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def append_xml_row(tbl, data, widths_in, bold_col0=True):
    """Append a row using raw XML (avoids python-docx row count issues)."""
    new_tr = OxmlElement("w:tr")
    for c_idx, txt in enumerate(data):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        w_el = OxmlElement("w:tcW")
        w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
        w_el.set(qn("w:type"), "dxa")
        tc_pr.append(w_el)
        tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr_el = OxmlElement("w:rPr")
        if bold_col0 and c_idx == 0:
            b_el = OxmlElement("w:b")
            rpr_el.append(b_el)
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), "16")
        rpr_el.append(sz_el)
        r_el.append(rpr_el)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve")
        t_el.text = txt
        r_el.append(t_el)
        p_el.append(r_el)
        tc.append(p_el)
        new_tr.append(tc)
    tbl._tbl.append(new_tr)
    return new_tr

def insert_xml_row_after(tbl, anchor_tr, data, widths_in, bold_col0=True):
    new_tr = OxmlElement("w:tr")
    for c_idx, txt in enumerate(data):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        w_el = OxmlElement("w:tcW")
        w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
        w_el.set(qn("w:type"), "dxa")
        tc_pr.append(w_el)
        tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr_el = OxmlElement("w:rPr")
        if bold_col0 and c_idx == 0:
            b_el = OxmlElement("w:b")
            rpr_el.append(b_el)
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), "16")
        rpr_el.append(sz_el)
        r_el.append(rpr_el)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve")
        t_el.text = txt
        r_el.append(t_el)
        p_el.append(r_el)
        tc.append(p_el)
        new_tr.append(tc)
    anchor_tr.addnext(new_tr)
    return new_tr

# ══════════════════════════════════════════════════════════════════════════════
# 1.  NP-FAI-ZM-001
# ══════════════════════════════════════════════════════════════════════════════
FAI_PATH = "docs/superseded/np_fai_zm_001.docx"
fai = Document(FAI_PATH)

# 1a. Fix FAI-M02 (peel force) and FAI-M03 (thermal cycles) in §3d table
for tbl in fai.tables:
    for row in tbl.rows:
        rid = row.cells[0].text.strip()
        if rid == "FAI-M02":
            set_text(row.cells[1],
                "Peel test (witness coupon from same bonding batch as production units): "
                "90° peel test on 25 mm × 25 mm PDMS–PI bonded coupon. "
                "Coupon must be processed through full SiO₂ sputter + O₂ plasma + cure "
                "sequence identically to production FPCs.",
                size=8)
            set_text(row.cells[3],
                "Peel force ≥ 150 N/m (per NP-HW-FPC-001 §9.3 and IPC-TM-650 2.4.9). "
                "Target ≥ 400 N/m. "
                "100–149 N/m → conditional accept, review bonding process, re-run 3 coupons. "
                "< 100 N/m → reject batch; halt production; SiO₂ interlayer process "
                "non-conformant — escalate to Manufacturing Eng.",
                size=8)
            print("  FAI-M02 peel threshold corrected to ≥150 N/m")
        if rid == "FAI-M03":
            set_text(row.cells[1],
                "Accelerated thermal cycling screen (production batch check — 5 cycles): "
                "5 cycles −20°C to +70°C on 1 witness coupon per production batch. "
                "This is a fast screen only — it does NOT replace the full 200-cycle "
                "IEC 60068-2-14 process qualification (see §3e FAI-TC items, which must "
                "pass before any production begins).",
                size=8)
            set_text(row.cells[3],
                "No new delamination (visual + 10× magnifier) after 5 cycles. "
                "If any delamination: halt batch; run full peel test per FAI-M02. "
                "If peel force drops below 150 N/m: reject batch, re-qualify process. "
                "Note: FAI-M03 pass alone does NOT confirm long-term reliability — "
                "§3e qualification must already be on file.",
                size=8)
            print("  FAI-M03 updated to reference §3e qualification requirement")

# 1b. Add §3e heading + thermal cycling qualification table
#     Insert before SECTION 4 paragraph
s4_para = None
for p in fai.paragraphs:
    if p.text.strip().startswith("SECTION 4"):
        s4_para = p
        break
print(f"  SECTION 4 anchor: {s4_para.text[:60]!r}")

# Content to insert (in reverse order — each addprevious goes before anchor)
S3E_PARAS = [
    # (text, bold, size, color)
    ("3e.  PDMS Thermal Cycling Process Qualification  [RISK-04 — one-time, pre-production]",
     True, 11, (0x1F, 0x49, 0x7D)),
    ("This section covers the one-time process qualification test required before any "
     "production zone modules are built. It must be completed and signed off before "
     "zone module tooling is released (NP-COORD-001 gate G1-08). "
     "The test validates that the SiO₂-interlayer bonding process (NP-HW-FPC-001 §9.2) "
     "can survive the 1,800+ thermal cycles expected over device lifetime. "
     "FAI-TC02 is BLOCKING — production cannot start without a pass record on file.",
     False, 10, None),
    ("Test rationale: no published fatigue data exists for PI–PDMS bonds under "
     "LED-junction thermal cycling conditions (PDMS CTE ≈ 300 ppm/°C vs. PI ≈ 18 ppm/°C, "
     "15:1 mismatch). 200-cycle IEC 60068-2-14 qualification generates this data. "
     "If the SiO₂ interlayer process requires adjustment, it must be identified here — "
     "not after tooling is committed.",
     False, 10, None),
]

# Build the FAI-TC table first (append to end, then relocate)
fai_tc_tbl = fai.add_table(rows=1, cols=6)
fai_tc_tbl.style = "Table Grid"
HDR_BG = "1F497D"
tc_hdrs = ["Item", "Inspection / Test", "Method", "Accept Criterion", "Result", "Sign"]
tc_widths = [0.65, 2.3, 1.2, 1.55, 0.65, 0.45]
hdr_row = fai_tc_tbl.rows[0]
for i, h in enumerate(tc_hdrs):
    set_bg(hdr_row.cells[i], HDR_BG)
    r = hdr_row.cells[i].paragraphs[0].add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(8)
    hdr_row.cells[i].width = Inches(tc_widths[i])

TC_ROWS = [
    ("FAI-TC01",
     "Bonding process documentation check: confirm SiO₂ sputter parameters (75 nm, RF magnetron), "
     "O₂ plasma parameters (100 W / 30 sccm / 60 s), cure (120°C / 2 hr / 50 g/cm²), "
     "and < 10 min delay between plasma and bond contact are documented in supplier "
     "process control plan before any qualification samples are made.",
     "Process control plan review",
     "All 5 parameters documented and confirmed. Any deviation from NP-HW-FPC-001 §9.2 "
     "must be approved in writing by HW EE before qualification proceeds.",
     "", ""),
    ("FAI-TC02",
     "[BLOCKING] Full thermal cycling qualification: IEC 60068-2-14 condition Na. "
     "10 bonded PDMS–FPC coupons (same dimensions and process as production). "
     "200 cycles, −10°C to +65°C, 15 min dwell at each extreme, ramp rate ≤ 3°C/min. "
     "Inspect 5 coupons at 100 cycles (interim check); inspect all 10 at 200 cycles.",
     "IEC 60068-2-14 thermal chamber; 200 cycles; interim check at 100 cycles",
     "Zero delamination (visual + 10× magnifier) at 100-cycle interim check. "
     "Zero delamination at 200 cycles on all 10 coupons. "
     "Any delamination at any point → HALT: SiO₂ interlayer process non-conformant. "
     "Review sputter thickness, plasma parameters, and cure conditions before re-run. "
     "Production CANNOT start until this test passes.",
     "", ""),
    ("FAI-TC03",
     "Post-cycling peel test: 90° peel test on all 10 coupons immediately after FAI-TC02 "
     "thermal cycling completion. Same test method as FAI-M02 (IPC-TM-650 2.4.9).",
     "90° peel test, Instron or similar",
     "Peel force ≥ 150 N/m on all 10 coupons after 200 cycles. "
     "Target ≥ 400 N/m (pre-cycling target applies here too — cycling should not "
     "degrade peel force below 150 N/m). "
     "Record min/max/mean — document in process qualification report.",
     "", ""),
    ("FAI-TC04",
     "Post-cycling optical transmission: measure optical transmittance at 660 nm and "
     "808 nm through each of the 10 coupons after thermal cycling, using the same "
     "spectrophotometer as pre-cycling baseline measurement.",
     "Spectrophotometer, 660 nm and 808 nm",
     "Transmittance ≥ 90% of pre-cycling baseline value at both wavelengths. "
     "Any coupon below 90% → investigate for delamination-induced air gap or "
     "PDMS haze formation. Reject if < 85% absolute transmittance.",
     "", ""),
    ("FAI-TC05",
     "LED junction thermal cycling simulation (supplemental, recommended): assemble "
     "3 complete zone modules with qualified PDMS process; run 500 LED-on/LED-off "
     "thermal cycles at maximum rated current (PWM 25% duty, 180 mA peak). "
     "Monitor PD1 output each cycle — normalised to initial value.",
     "Production zone module + bench power supply; 500 LED cycles",
     "PD1 output remains within 5% of initial value after 500 cycles "
     "(no air gap forming under LED thermal cycling). "
     "This test is recommended, not mandatory for production release, "
     "but must be completed before claiming > 2-year PDMS reliability.",
     "", ""),
    ("FAI-TC06",
     "Qualification report sign-off: all FAI-TC01 through FAI-TC04 results recorded "
     "in PDMS Process Qualification Report (document number NP-QR-PDMS-001). "
     "Report must be signed by Manufacturing Eng and HW EE before production start.",
     "Document review",
     "Signed qualification report on file. Report includes: supplier name, "
     "process control plan ref, test dates, all peel force and transmittance data, "
     "thermal chamber calibration certificate. "
     "Report retained for device lifetime + 2 years (regulatory retention requirement).",
     "", ""),
]

for r_idx, row_data in enumerate(TC_ROWS):
    row = fai_tc_tbl.add_row()
    bg = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
    for c_idx, txt in enumerate(row_data):
        cell = row.cells[c_idx]
        set_bg(cell, bg)
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(8)
        if c_idx == 0:
            run.bold = True
        cell.width = Inches(tc_widths[c_idx])

# Relocate table to before SECTION 4
tbl_xml = fai_tc_tbl._tbl
s4_para._p.addprevious(tbl_xml)

# Insert spacer after table (between table and SECTION 4)
spacer = OxmlElement("w:p")
s4_para._p.addprevious(spacer)

# Insert paragraphs before the table (in reverse order)
for text, bold, size, color in reversed(S3E_PARAS):
    new_p = OxmlElement("w:p")
    tbl_xml.addprevious(new_p)
    for p_obj in fai.paragraphs:
        if p_obj._p is new_p:
            run = p_obj.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            break

print("  §3e PDMS thermal cycling qualification section inserted")

# 1c. Add RISK-04 to §9 cross-reference table (after RISK-15 or at end)
for tbl in fai.tables:
    last_risk_row = None
    for row in tbl.rows:
        if row.cells[0].text.strip().startswith("RISK-"):
            last_risk_row = row
    if last_risk_row:
        xref_widths = [0.65, 2.0, 2.35, 1.8]
        new_tr = insert_xml_row_after(
            tbl, last_risk_row._tr,
            ("RISK-04",
             "PDMS–PI CTE mismatch thermal cycling delamination",
             "FAI-M01, FAI-M02, FAI-M03, FAI-TC01–FAI-TC06",
             "MITIGATED (this FAI)"),
            xref_widths)
        for r2 in tbl.rows:
            if r2._tr is new_tr:
                set_bg(r2.cells[3], "E2EFDA")
                if r2.cells[0].paragraphs[0].runs:
                    r2.cells[0].paragraphs[0].runs[0].bold = True
                break
        print("  §9 RISK-04 cross-reference row added")
        break

fai.save(FAI_PATH)
print(f"  NP-FAI-ZM-001 saved")

# ══════════════════════════════════════════════════════════════════════════════
# 2.  NP-PROC-SUP-001 — fix SUP-C-04 peel threshold; add SUP-C-08
# ══════════════════════════════════════════════════════════════════════════════
PROC_PATH = "docs/np_proc_sup_001.docx"
proc = Document(PROC_PATH)

for tbl in proc.tables:
    for row in tbl.rows:
        cells = row.cells
        rid = cells[0].text.strip()
        if rid == "SUP-C-04":
            set_text(cells[1],
                "Batch peel test: minimum 5 coupons per production batch. "
                "90° peel test per IPC-TM-650 2.4.9. "
                "Pass criterion: peel force ≥ 150 N/m (per NP-HW-FPC-001 §9.3). "
                "Target ≥ 400 N/m. Batch released only after peel test pass. "
                "Peel test data retained with batch record.",
                size=8)
            print("  SUP-C-04 peel threshold corrected to ≥150 N/m")
        if rid == "SUP-C-07":
            # Insert SUP-C-08 after SUP-C-07
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            c_widths = [0.7, 3.2, 1.7, 0.75, 0.85]
            sup_c08 = (
                "SUP-C-08",
                "PDMS thermal cycling qualification capability: supplier must be able to "
                "provide 10 bonded PDMS–FPC coupons for IEC 60068-2-14 thermal cycling "
                "qualification (200 cycles, −10°C to +65°C) before production commences. "
                "Supplier either runs the cycling test in-house (must have calibrated thermal "
                "chamber and IEC 60068-2-14 procedure) or submits coupons to an accredited "
                "third-party test house. Test report signed by supplier quality and NeurOne "
                "HW EE required before any production FPCs are delivered. "
                "Reference: NP-FAI-ZM-001 §3e FAI-TC01–FAI-TC06.",
                "Thermal chamber calibration cert + IEC 60068-2-14 procedure, "
                "or named third-party test house with accreditation",
                "⛔ BLOCKING",
                "☐"
            )
            for c_idx, txt in enumerate(sup_c08):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(c_widths[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                if c_idx == 0:
                    b_el = OxmlElement("w:b"); rpr_el.append(b_el)
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    # Red background for BLOCKING
                    for cell in r2.cells:
                        set_bg(cell, "FFE7E7")
                    break
            print("  SUP-C-08 thermal cycling qualification requirement added")
            break

proc.save(PROC_PATH)
print(f"  NP-PROC-SUP-001 saved")

# ══════════════════════════════════════════════════════════════════════════════
# 3.  NP-COORD-001 — add G1-08 PDMS qualification gate item
# ══════════════════════════════════════════════════════════════════════════════
COORD_PATH = "docs/np_coord_001.docx"
coord = Document(COORD_PATH)

for tbl in coord.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "G1-07":
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            g108_data = [
                "G1-08",
                "G1",
                "[BLOCKING] PDMS process qualification must pass before zone module "
                "tooling is released. Required: NP-FAI-ZM-001 §3e FAI-TC02 (200-cycle "
                "IEC 60068-2-14 thermal cycling) and FAI-TC03 (post-cycle peel force "
                "≥ 150 N/m) completed and signed off in NP-QR-PDMS-001. "
                "Owner: Manufacturing Eng + HW EE. Supplier (CAT-C) must be selected "
                "and process control plan approved before qualification samples can be made. "
                "Ref: RISK-04.",
                "Mfg Eng + HW EE",
                "OPEN — CAT-C supplier not yet selected"
            ]
            widths_in = [0.65, 0.55, 3.5, 0.9, 1.2]
            for c_idx, txt in enumerate(g108_data):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[4], "FFF2CC")
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  G1-08 PDMS qualification gate item added")
            break

coord.save(COORD_PATH)
print(f"  NP-COORD-001 saved")

# ══════════════════════════════════════════════════════════════════════════════
# 4.  Risk register — RISK-04 → MITIGATED
# ══════════════════════════════════════════════════════════════════════════════
RISK_PATH = "docs/superseded/np_risk_001.docx"
risk = Document(RISK_PATH)

for tbl in risk.tables:
    for row in tbl.rows:
        cells = row.cells
        if cells[0].text.strip() == "RISK-04" and len(cells) >= 9:
            set_text(cells[1], "HIGH\n(MITIGATED)", bold=True,
                     color=(0x00, 0x70, 0x00), size=8)
            set_bg(cells[1], "E2EFDA")
            set_text(cells[5],
                "Three-tier mitigation: "
                "(1) Process — SiO₂ interlayer (75 nm RF magnetron sputter) on PI surface "
                "before O₂ plasma activation raises peel force from 7–12 N/m to "
                "174–860 N/m per literature. Specified in NP-HW-FPC-001 §9.2. "
                "(2) Production screening — FAI-M02 peel test ≥ 150 N/m per batch; "
                "FAI-M03 5-cycle accelerated screen per batch. "
                "(3) Process qualification (BLOCKING) — IEC 60068-2-14, 200 cycles "
                "−10°C to +65°C on 10 coupons before any production starts. "
                "Post-cycle peel ≥ 150 N/m required (FAI-TC02/03). "
                "Signed qualification report NP-QR-PDMS-001 must be on file. "
                "CAT-C supplier qualification per NP-PROC-SUP-001 SUP-C-08.",
                size=8)
            set_text(cells[7],
                "NP-HW-FPC-001 §9.2–9.3\n"
                "NP-FAI-ZM-001 §3d, §3e\n"
                "NP-PROC-SUP-001 SUP-C-04, SUP-C-08\n"
                "NP-COORD-001 G1-08",
                size=8)
            set_text(cells[8],
                "MITIGATED — SiO₂ process specified; 200-cycle IEC 60068-2-14 "
                "qualification required (BLOCKING) before production; "
                "NP-FAI-ZM-001 §3e FAI-TC01–TC06 added",
                bold=True, color=(0x00, 0x70, 0x00), size=8)
            set_bg(cells[8], "92D050")
            print("  RISK-04 → MITIGATED")
            break

risk.save(RISK_PATH)
print(f"  Risk register saved")
print("Done.")
