"""
Correct G2-01, G2-02, G2-05/G3-03, and G2-10 in the engineering coordination
checklist (docs/neurone_eng_coordination_checklist.docx) for the hex-tile/
socket architecture (NP-HEX-ZM-001). Each item exists in two places: a
7-column detail table (G1/G2/G3 sections) and a 5-column dashboard summary
table — both are updated to stay consistent.

G2-01 (multi-FPC cable mgmt, RISK-17/RISK-11) and G3-03 (IPX4 field-swap
test, RISK-16) are corrected the same way as their risk-register entries
(patch_hexzm_risk_corrections.py) / shell-routing doc
(patch_hexzm_shell_routing_supersession.py): NEEDS REVIEW pointing at the
open hex-tile interconnect gap, or generalized to per-tile where that's a
clean generalization (IPX4).

G2-02 (mechanical key geometry) is corrected, not retired: NP-HEX-ZM-001 §4a
still specifies a real mechanical key — universal, orientation-only, one
design for the whole socket lattice instead of 5 zone-differentiating
variants. The braille/tactile-dot/per-zone deliverables are dropped (they
existed to distinguish zones, which no longer needs distinguishing).

G2-05 (IPX4 seal decision) generalizes cleanly to per-tile language, per
NP-HEX-ZM-001 §6 ("each hex needs a perimeter gasket... a per-tile
seam-length budget is required").

G2-10 (audio zone ID firmware) is real shipped firmware
(firmware/zone_announce/) built on the retired ZONE_ID mechanism. Per commit
b96f2d4's supersession note on docs/np_fw_za_001.md, it needs porting to
UID-based detection, not a doc-only fix — flagged here, not implemented here.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COORD_PATH = "docs/neurone_eng_coordination_checklist.docx"

MARKER = "hex-tile/socket architecture (NP-HEX-ZM-001)"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def find_row(doc, row_id, ncols):
    for tbl in doc.tables:
        for row in tbl.rows:
            if row.cells[0].text.strip() == row_id and len(row.cells) == ncols:
                return row
    return None


def already_patched(doc):
    return any(MARKER in p.text for tbl in doc.tables for row in tbl.rows for c in row.cells for p in c.paragraphs)


def main():
    doc = Document(COORD_PATH)
    if already_patched(doc):
        print(f"  Already patched — skipping: {COORD_PATH}")
        return

    # ── G2-01 detail row (7 cols) ───────────────────────────────────────
    row = find_row(doc, "G2-01", 7)
    if row:
        set_cell_text(row.cells[2],
            "NEEDS REVIEW — RISK-17 / RISK-11: this item was scoped to the retired 5-zone-slot FPC "
            "routing architecture (5 discrete tails, adjacent-pair crosstalk). Superseded by the "
            "hex-tile/socket architecture (NP-HEX-ZM-001) — the ~80-socket interconnect it replaces "
            "with is real, unscoped EE work (see MECH-1/MECH-2/SMART-1 in docs/np_hex_zm_001.md §7).")
        set_cell_text(row.cells[3],
            "Cannot be completed as originally scoped (no 5 discrete zone-slot FPC tails to CAD-review). "
            "Needs re-scoping once the per-socket interconnect design exists.")
        set_cell_text(row.cells[6], "NEEDS RE-SCOPE — hex-tile socket interconnect not yet designed",
            bold=True, color=(0xBF, 0x60, 0x00))
        set_cell_bg(row.cells[6], "FFE699")
        print("  G2-01 detail row corrected")
    else:
        print("  WARNING: G2-01 detail row not found")

    # ── G2-01 dashboard row (5 cols) ────────────────────────────────────
    row = find_row(doc, "G2-01", 5)
    if row:
        set_cell_text(row.cells[3], "Multi-FPC + EEG cable separation — NEEDS RE-SCOPE (retired 5-slot architecture; NP-HEX-ZM-001 socket interconnect not yet designed)")
        set_cell_text(row.cells[4], "NEEDS RE-SCOPE", bold=True, color=(0xBF, 0x60, 0x00))
        set_cell_bg(row.cells[4], "FFE699")
        print("  G2-01 dashboard row corrected")

    # ── G2-02 detail row (7 cols) — corrected, not retired ──────────────
    row = find_row(doc, "G2-02", 7)
    if row:
        set_cell_text(row.cells[2],
            "RISK-15 (corrected for the hex-tile/socket architecture, NP-HEX-ZM-001 §4a): universal "
            "asymmetric orientation-only mechanical key, ONE design shared by every socket — not a "
            "zone-differentiating key. SMART-1 (every socket I2C/TIA-capable) removes the \"wrong "
            "socket\" hazard the 5-variant key existed to prevent, so there is nothing left to "
            "type-differentiate. Module TYPE correctness is now a software check "
            "(np_module_map_check_placement()), not a mechanical one. Braille / tactile-dot / "
            "raised-numeral per-zone deliverables are dropped — they existed to distinguish zones "
            "from each other, which is no longer a hazard.")
        set_cell_text(row.cells[3],
            "Key geometry drawing for the single universal orientation key (one design, all ~80 "
            "sockets). Written confirmation the key allows only one insertion orientation, "
            "verifiable without vision (blind-safe by construction).")
        set_cell_text(row.cells[6], "OPEN — universal orientation-key geometry required (simplified scope)",
            bold=True, color=(0xBF, 0x60, 0x00))
        print("  G2-02 detail row corrected")
    else:
        print("  WARNING: G2-02 detail row not found")

    row = find_row(doc, "G2-02", 5)
    if row:
        set_cell_text(row.cells[3], "Universal orientation-key geometry (RISK-15, corrected for hex-tile — no per-zone key variants)")
        print("  G2-02 dashboard row corrected")

    # ── G2-05 / G3-03 — generalize to per-tile ──────────────────────────
    row = find_row(doc, "G2-05", 7)
    if row:
        old = row.cells[2].text
        new = old.replace("Zone module field-replacement IPX4", "Hex-tile module field-replacement IPX4").replace(
            "zone module and shell tooling", "hex-tile module and shell tooling")
        set_cell_text(row.cells[2], new)
        old3 = row.cells[3].text
        new3 = old3.replace("10 swap cycles", "10 swap cycles per representative tile, across a sample of tile positions")
        set_cell_text(row.cells[3], new3)
        print("  G2-05 detail row generalized to per-tile")

    row = find_row(doc, "G3-03", 7)
    if row:
        old = row.cells[2].text
        new = old.replace("insert and remove zone module 10×", "insert and remove a sample of hex-tile modules 10× each")
        set_cell_text(row.cells[2], new)
        old3 = row.cells[3].text
        new3 = old3.replace("after 10 zone module swap cycles", "after 10 swap cycles per sampled tile")
        set_cell_text(row.cells[3], new3)
        print("  G3-03 detail row generalized to per-tile")

    # ── G2-10 — flag firmware porting need, do not claim it's done ─────
    row = find_row(doc, "G2-10", 7)
    if row:
        old = row.cells[2].text
        set_cell_text(row.cells[2],
            old + " NEEDS PORT: this firmware keys off the retired ZONE_ID resistor-ladder "
            "mechanism (5 parallel slot state machines). NP-HEX-ZM-001's UID-based auto-inventory "
            "(np_module_map) replaces zone-slot detection outright — the DDS audio engine and "
            "debounce algorithm are addressing-independent and remain reusable, but the slot "
            "state machine needs porting to per-socket UID-change events. See supersession note "
            "in docs/np_fw_za_001.md.")
        set_cell_text(row.cells[6],
            "PARTIAL — NP-FW-ZA-001 Rev A baselined against the retired ZONE_ID mechanism; needs "
            "porting to UID-based per-socket detection before it matches NP-HEX-ZM-001",
            bold=True, color=(0xBF, 0x60, 0x00))
        set_cell_bg(row.cells[6], "FFE699")
        print("  G2-10 detail row flagged for firmware port")

    row = find_row(doc, "G2-10", 5)
    if row:
        old = row.cells[3].text
        set_cell_text(row.cells[3], old + " NEEDS PORT to NP-HEX-ZM-001 UID-based detection (see docs/np_fw_za_001.md).")
        set_cell_text(row.cells[4], "PARTIAL — needs port off retired ZONE_ID mechanism", bold=True, color=(0xBF, 0x60, 0x00))
        set_cell_bg(row.cells[4], "FFE699")
        print("  G2-10 dashboard row flagged for firmware port")

    doc.save(COORD_PATH)
    print(f"  Coordination checklist saved: {COORD_PATH}")


if __name__ == "__main__":
    main()
