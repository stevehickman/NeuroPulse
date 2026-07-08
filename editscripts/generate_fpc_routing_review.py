"""
Generate NP-DRV-SHELL-001 Rev A:
  Zone Module FPC Routing Path — Shell Tooling Design Review
Cross-referenced from FPC spec NP-HW-FPC-001 §11.2 and risk register RISK-11.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

DATE = "2026-05-06"
DOC_NUM = "NP-DRV-SHELL-001"
REV = "Rev A"

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(doc, text, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.runs[0].font.size = Pt(10) if p.runs else None
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        rest = p.add_run(text)
        rest.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def tbl(doc, headers, rows, widths=None, header_bg="1F497D"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = t.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else "F2F2F2"
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            # Handle (text, color, bold) tuples
            if isinstance(cell_text, tuple):
                txt, color, bold = cell_text
                run = cell.paragraphs[0].add_run(txt)
                run.font.size = Pt(9)
                if color:
                    run.font.color.rgb = RGBColor(*color)
                run.bold = bold
            else:
                run = cell.paragraphs[0].add_run(str(cell_text))
                run.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def warn_box(doc, label, text):
    """Coloured callout box using a single-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "FFF2CC")
    p = cell.paragraphs[0]
    run_lbl = p.add_run(label + "  ")
    run_lbl.bold = True
    run_lbl.font.size = Pt(10)
    run_lbl.font.color.rgb = RGBColor(0xBF, 0x60, 0x00)
    run_txt = p.add_run(text)
    run_txt.font.size = Pt(10)
    doc.add_paragraph()

def req_box(doc, req_id, text):
    """Numbered requirement box."""
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    id_cell = t.rows[0].cells[0]
    txt_cell = t.rows[0].cells[1]
    set_cell_bg(id_cell, "1F497D")
    set_cell_bg(txt_cell, "EEF4FB")
    id_cell.width = Inches(1.0)
    txt_cell.width = Inches(5.5)
    run = id_cell.paragraphs[0].add_run(req_id)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    id_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = txt_cell.paragraphs[0].add_run(text)
    run2.font.size = Pt(10)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
# Document
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Cover ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NeurOne")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Zone Module FPC Routing Path\nShell Tooling Design Review")
run2.bold = True
run2.font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run(
    f"{DOC_NUM}  {REV}  |  {DATE}  |  Pre-Tooling Draft\n"
    "Companion to NP-HW-FPC-001 Rev C (Zone Module FPC Spec)\n"
    "Addresses RISK-11 in Zone Module FPC Risk Register"
)
run3.font.size = Pt(10)
run3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
doc.add_paragraph()

warn_box(doc, "STATUS:",
    "This document is a REQUIRED DESIGN REVIEW that must be completed and signed off "
    "before shell tooling first-cut. Until sign-off, all FPC routing path dimensions "
    "shown are targets — not confirmed geometry. Open items are listed in §7.")

# ── 1. Purpose and Scope ────────────────────────────────────────────────────
heading(doc, "1.  PURPOSE AND SCOPE")
body(doc,
    "This document defines the requirements for the FPC routing path between each Zone "
    "Module connector slot and the Hub PCB ZIF connector inside the NeurOne headset "
    "CFRP shell. Its purpose is to ensure shell tooling geometry enforces the minimum "
    "bend radius required for dynamic flex FPC lifetime, and to provide a structured "
    "sign-off checklist before first-cut.")
body(doc,
    "Failure to specify the routing path before tooling is cut cannot be remediated "
    "without a full shell retool. This is a one-way decision.")

tbl(doc,
    headers=["Document", "Number", "Relationship"],
    rows=[
        ("Zone Module FPC Specification", "NP-HW-FPC-001 Rev C", "Parent spec — defines FPC stackup, bend radius, and connector"),
        ("Zone Module FPC Risk Register", "NP-FPC-RISK-001", "RISK-11: shell tooling/FPC routing coordination — this doc is the mitigation"),
        ("CLAUDE.md §13.4", "Project memory", "Shell tooling listed as pending decision — must resolve before first cut"),
        ("IPC-2223D", "Industry standard", "Flexible circuit design standard — bend radius, dynamic flex requirements"),
        ("IEC 60068-2-21", "Industry standard", "Mechanical robustness — FPC lead flex fatigue test method"),
    ],
    widths=[2.3, 1.5, 2.7],
)

# ── 2. FPC Routing Path Description ─────────────────────────────────────────
heading(doc, "2.  FPC ROUTING PATH — DESCRIPTION")
body(doc,
    "The Zone Module FPC connects the zone PCB (which carries the LED array, reference "
    "photodiode, BCR421W drivers, and NTC thermistor) to the Hub PCB via a ZIF connector "
    "(Hirose FH34S-20S-0.5SH). The routing path runs from the zone slot opening in the "
    "outer CFRP shell, through the shell interior, to the Hub PCB ZIF receptacle in the "
    "control hub.")

heading(doc, "2.1  Zone Module Connector Slots", level=2)
body(doc,
    "The headset has 5 zone module slots, each housing one zone module FPC. "
    "Slot positions correspond to PBM zones:")
tbl(doc,
    headers=["Zone", "Slot position", "EEG electrode reference", "Approximate arc from hub"],
    rows=[
        ("Zone 1", "Frontal left", "Fp1/F3", "~120 mm"),
        ("Zone 2", "Frontal right", "Fp2/F4", "~120 mm"),
        ("Zone 3", "Central left", "C3", "~80 mm"),
        ("Zone 4", "Central right", "C4", "~80 mm"),
        ("Zone 5", "Parietal/occipital", "P3/P4", "~60 mm"),
    ],
    widths=[0.7, 1.4, 1.8, 1.9],
)

heading(doc, "2.2  Routing Path Segments", level=2)
body(doc,
    "Each FPC routing path has three mechanical segments that must each independently "
    "satisfy the ≥25 mm dynamic bend radius requirement:")
tbl(doc,
    headers=["Segment", "Description", "Primary constraint", "Risk if under-radius"],
    rows=[
        ("A — Exit bend",
         "FPC exits zone module PCB and turns to enter shell interior routing channel",
         "Zone module slot geometry; FR4 stiffener must not impose bend",
         "Fatigue crack at stiffener edge — catastrophic, immediate failure"),
        ("B — Interior channel",
         "FPC traverses shell interior from zone slot to hub PCB area; may curve with shell geometry",
         "Shell interior curvature — CFRP shell is not flat; routing channel must follow compound curve",
         "Repeated bend fatigue over 1,000 module swap cycles; failure may be delayed 50–200 cycles"),
        ("C — Hub entry bend",
         "FPC transitions from shell interior to horizontal plane of Hub PCB and enters ZIF connector",
         "Hub PCB height relative to shell interior; stiffener must enforce radius, not restrict it",
         "Same as Segment A — stiffener edge creates stress concentration"),
    ],
    widths=[1.1, 2.0, 1.7, 1.5],
)

# ── 3. Bend Radius Requirements ──────────────────────────────────────────────
heading(doc, "3.  BEND RADIUS REQUIREMENTS")
body(doc,
    "These requirements are derived from IPC-2223D and the FPC stackup specified in "
    "NP-HW-FPC-001 §3 (1 oz RA copper, 50 µm polyimide, total FPC thickness ~0.20 mm). "
    "All requirements apply to every segment of the routing path for all 5 zones.")

req_box(doc, "REQ-BR-01",
    "Minimum bend radius at any point along the FPC routing path: ≥25 mm (dynamic flex, "
    "actuated with every zone module insertion/removal). This is a hard lower limit — "
    "no exceptions. Shell tooling must be designed to enforce this radius by geometry, "
    "not by user care.")

req_box(doc, "REQ-BR-02",
    "Static bend radius (areas where FPC is bent once during assembly and never again): "
    "≥12.5 mm. Must be identified separately on routing path drawing and confirmed "
    "as static-only by mechanical design review.")

req_box(doc, "REQ-BR-03",
    "No bend within 5 mm of the ZIF connector tail end, FR4 stiffener edge, or any "
    "zone module PCB edge. Bending force at these transitions causes delamination of "
    "the FPC tail from the stiffener, not just copper fatigue.")

req_box(doc, "REQ-BR-04",
    "The shell interior cable management channel must prevent the FPC from taking a "
    "spontaneous bend radius smaller than 25 mm under gravity or user handling — "
    "including when the zone module is partially inserted or extracted. Channel walls "
    "must physically constrain the FPC to the required radius.")

req_box(doc, "REQ-BR-05",
    "All 5 zone routing paths must individually satisfy REQ-BR-01 through REQ-BR-04. "
    "Zones 1 and 2 (frontal, longest arc) are highest risk due to compound shell "
    "curvature. Zones 1 and 2 routing paths must be verified first in CAD review.")

body(doc,
    "Basis for ≥25 mm radius: IPC-2223D Table 4-1 recommends minimum bend radius of "
    "6× the total FPC thickness for dynamic flex (10+ cycles). At ~0.20 mm total "
    "thickness, that gives 6 × 0.20 = 1.2 mm as an absolute floor. The ≥25 mm "
    "requirement is deliberately conservative (125×) to provide margin over the "
    "1,000-cycle module swap lifetime and account for shell geometry tolerance. "
    "The 25 mm value is also what can be physically enforced by a cable channel "
    "in a headset shell interior without adding bulk.")

# ── 4. Shell Tooling Requirements ────────────────────────────────────────────
heading(doc, "4.  SHELL TOOLING REQUIREMENTS")
body(doc,
    "These requirements must be incorporated into the shell tooling specification before "
    "first-cut. They are not retrofittable — cable channel geometry is moulded into the "
    "CFRP shell interior and cannot be added post-tooling without a full shell retool.")

req_box(doc, "REQ-ST-01",
    "Each zone slot must have a dedicated FPC routing channel moulded into the CFRP "
    "shell interior. The channel must run from the zone slot ZIF lever access opening "
    "to the Hub PCB ZIF receptacle area.")

req_box(doc, "REQ-ST-02",
    "Cable channel cross-section: minimum 3 mm wide × 1.5 mm deep (FPC is 0.20 mm "
    "thick — channel depth provides clearance for FPC movement during insertion/removal "
    "without pinching). Channel width must accommodate FPC + tolerance without "
    "introducing lateral bending.")

req_box(doc, "REQ-ST-03",
    "Channel walls must be radiused at all corners. Sharp internal corners on the cable "
    "channel must be radiused ≥1 mm to prevent stress concentration on the FPC outer surface.")

req_box(doc, "REQ-ST-04",
    "The channel entry at the zone slot must align with the ZIF connector centreline to "
    "within ±0.5 mm. Misalignment greater than 0.5 mm will impose a lateral bend on "
    "the FPC tail immediately upon insertion — this is a Segment A violation.")

req_box(doc, "REQ-ST-05",
    "The routing channel must include a 'bend radius enforcer' geometry — a moulded boss "
    "or ridge on the inside of the channel at each Segment A and Segment C location, "
    "profiled to a 25 mm radius. The FPC wraps over this surface; it cannot be bent "
    "tighter than the enforcer radius by any user action.")

req_box(doc, "REQ-ST-06",
    "ZIF connector lever access opening must allow full lever actuation (180° flip) "
    "without the user's finger contacting the FPC. Minimum clearance: 8 mm above the "
    "FPC surface at the lever position.")

req_box(doc, "REQ-ST-07",
    "All five zone routing channels must be independently accessible — installing or "
    "removing Zone 1 module must not require disturbing the Zone 2 FPC.")

# ── 5. Design Review Checklist ───────────────────────────────────────────────
heading(doc, "5.  DESIGN REVIEW CHECKLIST")
body(doc,
    "This checklist must be completed in CAD review before the shell tooling specification "
    "is released to the tooling manufacturer. Each item requires a named reviewer and "
    "a pass/fail determination with supporting evidence (screenshot, measurement, or "
    "simulation result).")

tbl(doc,
    headers=["#", "Check item", "Method", "Pass criterion", "Reviewer", "Result"],
    rows=[
        ("DRC-01", "Zone 1 (frontal left) routing path: measure minimum bend radius at all three segments (A/B/C) in CAD",
         "CAD measurement", "≥25 mm at each segment", "", ""),
        ("DRC-02", "Zone 2 (frontal right) routing path: same as DRC-01",
         "CAD measurement", "≥25 mm at each segment", "", ""),
        ("DRC-03", "Zone 3 (central left) routing path: measure minimum bend radius",
         "CAD measurement", "≥25 mm", "", ""),
        ("DRC-04", "Zone 4 (central right) routing path: measure minimum bend radius",
         "CAD measurement", "≥25 mm", "", ""),
        ("DRC-05", "Zone 5 (parietal) routing path: measure minimum bend radius",
         "CAD measurement", "≥25 mm", "", ""),
        ("DRC-06", "Confirm no bend within 5 mm of ZIF connector tail end for all 5 zones",
         "CAD measurement", "Nearest bend point ≥5 mm from tail end", "", ""),
        ("DRC-07", "Cable channel cross-section: confirm ≥3 mm wide × ≥1.5 mm deep at narrowest point for all 5 channels",
         "CAD cross-section", "≥3 × 1.5 mm throughout", "", ""),
        ("DRC-08", "Bend radius enforcer geometry present at Segment A and C for all 5 zones",
         "CAD inspection", "Enforcer boss visible, radius profiled to 25 mm", "", ""),
        ("DRC-09", "Channel wall internal corners radiused ≥1 mm",
         "CAD measurement", "All interior corners ≥1 mm radius", "", ""),
        ("DRC-10", "ZIF connector lever access clearance ≥8 mm above FPC surface",
         "CAD measurement", "≥8 mm clearance at lever position, all 5 zones", "", ""),
        ("DRC-11", "Zone slot centreline alignment with cable channel centreline ≤±0.5 mm",
         "CAD GD&T analysis", "Misalignment ≤0.5 mm", "", ""),
        ("DRC-12", "Confirm Zone 1 module removal does not disturb Zone 2 FPC (and all adjacent pairs)",
         "Assembly simulation", "No cross-interference in CAD assembly", "", ""),
        ("DRC-13", "FPC length from zone module PCB to Hub PCB ZIF: confirm ≥10% slack in assembled position (slack prevents tension on connector)",
         "CAD measurement", "FPC length ≥ routing path arc × 1.10", "", ""),
        ("DRC-14", "Physical prototype (SLA or 3-DP shell section): insert/remove zone module 10× and confirm FPC does not contact channel wall edges",
         "Physical test", "No contact marks on FPC after 10 cycles", "", ""),
        ("DRC-15", "First-article production shell: insert/remove zone module 50× and inspect FPC for any deformation, whitening, or crease",
         "Physical inspection", "No deformation, whitening, or crease after 50 cycles", "", ""),
    ],
    widths=[0.55, 2.0, 1.0, 1.25, 0.9, 0.6],
)

# ── 6. FPC Length Specification ──────────────────────────────────────────────
heading(doc, "6.  FPC LENGTH AND SLACK REQUIREMENTS")
body(doc,
    "FPC length is determined by the routing path geometry. These values are targets "
    "pending CAD confirmation — final lengths must be extracted from the CAD assembly model "
    "and verified before FPC fabrication release.")

tbl(doc,
    headers=["Zone", "Estimated routing arc (mm)", "Minimum FPC length (arc × 1.10)", "Maximum FPC length (arc × 1.25)", "Status"],
    rows=[
        ("Zone 1 (frontal left)", "120 ± 10", "132 mm", "150 mm", "PENDING CAD"),
        ("Zone 2 (frontal right)", "120 ± 10", "132 mm", "150 mm", "PENDING CAD"),
        ("Zone 3 (central left)", "80 ± 8", "88 mm", "100 mm", "PENDING CAD"),
        ("Zone 4 (central right)", "80 ± 8", "88 mm", "100 mm", "PENDING CAD"),
        ("Zone 5 (parietal)", "60 ± 6", "66 mm", "75 mm", "PENDING CAD"),
    ],
    widths=[1.3, 1.6, 1.6, 1.6, 1.15],
)
body(doc,
    "10% minimum slack ensures no tensile force on the ZIF connector tail when the headset "
    "is at maximum head circumference (62 cm). 25% maximum slack prevents the FPC from "
    "pooling in the channel and creating a spontaneous tight bend.")

# ── 7. Open Items ────────────────────────────────────────────────────────────
heading(doc, "7.  OPEN ITEMS — MUST RESOLVE BEFORE TOOLING RELEASE")
warn_box(doc, "BLOCKING:",
    "None of the items below can be resolved in this document — they require a CAD model "
    "of the headset shell interior. This document exists to make the requirements explicit "
    "before the shell CAD is started, so the mechanical engineer designs the interior to "
    "meet these requirements from the first pass, not as a retrofit.")

tbl(doc,
    headers=["ID", "Item", "Who", "Needed before"],
    rows=[
        ("OI-01", "Shell interior CAD model with routing channels for all 5 zones", "Mechanical Eng", "Shell tooling spec release"),
        ("OI-02", "CAD measurement of minimum bend radius at all Segment A/B/C points for all 5 zones (DRC-01 to DRC-05)", "Mechanical Eng / HW EE", "Tooling spec release"),
        ("OI-03", "FPC length extraction from CAD assembly model for all 5 zones — update §6 table", "Mechanical Eng / HW EE", "FPC fab drawing release"),
        ("OI-04", "Bend radius enforcer geometry defined and modelled in CAD for all 5 zones", "Mechanical Eng", "Tooling spec release"),
        ("OI-05", "SLA or 3D-printed prototype of shell section (Zones 1 and 2 minimum) for physical routing verification (DRC-14)", "Mechanical Eng", "Before first-cut approval"),
        ("OI-06", "ZIF lever access opening geometry modelled and clearance verified (DRC-10 for all 5 zones)", "Mechanical Eng", "Tooling spec release"),
        ("OI-07", "Confirm FR4 stiffener length does not encroach on Segment A bend — coordinate with FPC fab drawing", "HW EE / FPC Fabricator", "FPC fab drawing release"),
    ],
    widths=[0.6, 2.8, 1.5, 1.35],
)

# ── 8. Sign-Off ──────────────────────────────────────────────────────────────
heading(doc, "8.  SIGN-OFF")
body(doc,
    "This design review document is complete when all DRC items (§5) pass and all open "
    "items (§7) are resolved. Sign-off is required from Mechanical Engineering and "
    "Hardware Electrical Engineering leads before the shell tooling specification is "
    "released to the tooling manufacturer.")

tbl(doc,
    headers=["Role", "Name", "Date", "Signature / Approval"],
    rows=[
        ("Mechanical Engineering Lead", "", "", ""),
        ("Hardware EE Lead", "", "", ""),
        ("Manufacturing Engineering", "", "", ""),
        ("Product / Programme Manager", "", "", ""),
    ],
    widths=[2.0, 1.5, 1.0, 2.2],
)

# ── 9. Revision History ──────────────────────────────────────────────────────
heading(doc, "9.  REVISION HISTORY")
tbl(doc,
    headers=["Rev", "Date", "Author", "Description"],
    rows=[
        ("A", DATE, "NeurOne Engineering", "Initial release — RISK-11 mitigation. All requirements set as targets pending CAD confirmation."),
    ],
    widths=[0.5, 1.0, 1.7, 3.55],
)

OUT = "docs/neurone_shell_fpc_routing_review.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
