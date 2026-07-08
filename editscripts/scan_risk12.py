"""Scan RISK-12 detail and all previously-specified inspection items."""
from docx import Document

risks = Document("docs/neurone_fpc_zone_module_risks_revA.docx")
tbl = risks.tables[0]

print("=== RISK-11 and RISK-12 rows ===")
for row in tbl.rows:
    rid = row.cells[0].text.strip()
    if rid in ("RISK-11", "RISK-12"):
        headers = ["ID","Severity","Title","Root Cause","Impact","Required Mitigation","Owner","Spec §","Status"]
        for h, cell in zip(headers, row.cells):
            print(f"  {h}: {cell.text.strip()!r}")
        print()

# Check procurement doc inspection table
print("=== Procurement Doc — Incoming Inspection Table ===")
proc = Document("docs/neurone_fpc_procurement_requirements.docx")
for t_idx, tbl2 in enumerate(proc.tables):
    for r_idx, row in enumerate(tbl2.rows):
        texts = [c.text.strip()[:60] for c in row.cells]
        if any(texts):
            print(f"  Table {t_idx} Row {r_idx}: {texts}")
