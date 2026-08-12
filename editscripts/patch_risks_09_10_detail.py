"""
Patch detailed sections for RISK-09 and RISK-10 in the risk register.
These sections appear to be in tables, not in body paragraphs.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RISK_PATH = "docs/superseded/np_risk_001.docx"
doc = Document(RISK_PATH)

# Print all table content to find where RISK-09 and RISK-10 are
print("=== Scanning all tables ===")
for t_idx, tbl in enumerate(doc.tables):
    for r_idx, row in enumerate(tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if "RISK-09" in text or "RISK-10" in text:
                print(f"  Table {t_idx}, Row {r_idx}, Col {c_idx}: {text[:120]!r}")

print("\n=== Scanning all paragraphs for RISK-09/10 ===")
for i, para in enumerate(doc.paragraphs):
    if "RISK-09" in para.text or "RISK-10" in para.text:
        print(f"  Para {i}: {para.text[:120]!r}")

print("\n=== Looking for 'Mitigation' in tables ===")
for t_idx, tbl in enumerate(doc.tables):
    for r_idx, row in enumerate(tbl.rows):
        for c_idx, cell in enumerate(row.cells):
            text = cell.text.strip()
            if ("Mitigation" in text or "Action" in text) and len(text) < 200:
                # Only print shorter cells (likely headers or labels)
                print(f"  Table {t_idx}, Row {r_idx}, Col {c_idx}: {text[:80]!r}")
