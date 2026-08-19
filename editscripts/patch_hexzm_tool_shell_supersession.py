"""
Mark NP-TOOL-SHELL-001 (docs/np_tool_shell_001.docx) PARTIALLY SUPERSEDED against the
hex-tile architecture, and record the per-feature disposition in the document itself.

Context. This is the one zone-module-era document the PR #223/#224 cleanup never opened.
It was flagged by inference only, from np_dt_001.md's DO-HW-06 summary line ("F-01 zone
slot anchor posts x5"). Opening it confirms F-01 and surfaces two defects a title-level
sweep could not see, because neither is visible in a feature title:

  F-03 has no cable left to carry. Its governing rule is that EEG cables run on the outer
  CFRP inner surface while the zone-module FPC bundle runs on the scalp-facing surface,
  >=2 mm of shell wall apart, achieving DRC-18's 15 mm copper separation. NP-DRV-SHELL-002
  Sec 9.1 quotes that sentence and states the trick is unavailable: EEG electrodes now live
  inside T1-B hex tiles on the L1 carrier, and Sec 5.3's conductor table records the
  replacement as "absorbing the EEG harness". The retired design's 10 separate EEG cables
  are gone, so the 8x5 mm channel has no harness to route.

  F-02 collides with a safety-governed limit. NP-TOOL-HUB-001 Rev 1 F-02 specifies the same
  hub ports with 1.0 mm dia x 0.5 mm bosses and a <=20 mm tether, where the maximum is a
  hazard control (a tethered cover must not reach the F-04 fan intake grille; FAI-HTOOL-02,
  BLOCKING). This document specifies 4.0 mm x 3.5 mm bosses and a 50 mm tether on that same
  hub shell -- 2.5x the hub document's hard maximum. Both documents are ACTIVE.

Unlike the four documents retired on 2026-08-11, this one stays in place: NP-ART-001 Sec 5
retained it deliberately, because it is the only tooling specification artifact A6 (the L0
outer bowl) has, and retiring it would leave a live artifact with no governing document.

Treatment follows the established convention (patch_hexzm_shell_routing_supersession.py and
the markdown siblings): historical content is preserved verbatim and marked, never deleted;
the banner separates retired-outright from still-reusable; and it names the gap it does not
fill. No replacement geometry is invented here -- that is MECH-1 / OI-ART-01 engineering work,
explicitly blocked until the clamp and boss geometry is fixed.

Idempotent: re-running is a no-op.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHELL_PATH = "docs/np_tool_shell_001.docx"
BANNER_MARKER = "PARTIALLY SUPERSEDED (2026-08-18)"

GREY = (0x40, 0x40, 0x40)
RED = (0xC0, 0x00, 0x00)


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def already_patched(doc):
    if any(BANNER_MARKER in p.text for p in doc.paragraphs):
        return True
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                if BANNER_MARKER in c.text:
                    return True
    return False


BLOCKS = [
    (
        "⚠ PARTIALLY SUPERSEDED (2026-08-18) — F-01 and F-03 are retired; "
        "F-02 conflicts with NP-TOOL-HUB-001; F-04 survives.",
        True,
        RED,
    ),
    (
        "This document was written against the retired five-zone-slot shell. It is NOT retired "
        "wholesale and is NOT moved to docs/superseded/ — per NP-ART-001 §5 it remains the only "
        "tooling specification artifact A6 (the L0 outer bowl) has. Nothing below is deleted; the "
        "retired geometry is retained as the record of what was specified. Disposition is per "
        "feature, and it is NOT the disposition NP-ART-001 §5 and NP-DT-001 DO-HW-06 assumed — "
        "both were written without opening this file, and both named F-01 alone.",
        False,
        GREY,
    ),
    (
        "F-01 (§2.1) — RETIRED OUTRIGHT. Five colour-coded anchor posts, one per zone slot "
        "ZM-01..ZM-05, with colours matched to NP-TOOL-ZM-001 §3.1. NP-HEX-ZM-001 §3.4 replaced the "
        "five position-unique slots with ~80 identical sockets; §4a made the tile a universal SKU "
        "with an orientation-only key and no type or zone differentiation, so there is no zone "
        "identity left to colour-code; and §5.1 moved the sockets onto the INNER bowl, while this "
        "feature specifies posts on the outer CFRP shell surface adjacent to a zone slot opening — "
        "a surface that no longer carries sockets at all. Whether unpopulated hex sockets need "
        "plugs is an open design question, not a renumbering (OI-ART-01). Consequently retired "
        "with it: SH-01..SH-05 and SH-23, FAI-CV-01..FAI-CV-04, OI-04 (five-colour bulk-material "
        "confirmation), and the F-01 half of OI-01 and OI-06. SH-23 in particular cross-checks "
        "against a zone-module mechanical key that NP-HEX-ZM-001 §4a deleted.",
        False,
        GREY,
    ),
    (
        "F-03 (§2.3) — PREMISE RETIRED; the threshold it served survives elsewhere. This is the "
        "finding the \"x5\" inference could not reach, because nothing in the feature's title is "
        "keyed to five slots. F-03 routes EEG cables from the Hub PCB to 8 discrete electrode pods "
        "in an 8x5 mm channel on the outer CFRP inner surface, separated from the zone-module FPC "
        "bundle by the shell wall to achieve DRC-18's 15 mm copper separation. NP-DRV-SHELL-002 "
        "§9.1 quotes that exact geometric rule and states it is unavailable: EEG electrodes now "
        "live inside T1-B hex tiles (NP-HEX-ZM-001 §4a) on the same L1 carrier as PBM drive "
        "current. §5.3's conductor comparison records the replacement as \"absorbing the EEG "
        "harness\" — the retired design's 10 separate EEG cables are gone, and §3.5 routes "
        "electrode analog per-cluster onto guarded lanes to the ADS1299 bank at the PAN. There is "
        "no Hub-PCB-to-electrode-pod harness for this channel to carry. Retired with it: SH-09..SH-14 "
        "and SH-22, FAI-EEG-01 and FAI-EEG-03, OI-03, and the F-03 rows of §3. What survives is not "
        "the channel but the acceptance threshold: the <5 uVpp EEG-artifact limit (FAI-EEG-02 / "
        "DRC-18c) is retained verbatim as NP-DRV-SHELL-002 SH2-DRC-16, and is now a harder test "
        "because no geometric margin backs it. Two things corroborate this independently of the "
        "present reading. cad/CAD_PARTS_LIST.md already carries EEG-ROUTE-CHANNEL as a RETIRED "
        "tombstone row, recording that the REQ-ST-01..07 molded channel features are \"deleted from "
        "shell tooling (a net simplification of the shell tool -- the complexity moves into L1 "
        "lamination)\"; and NP-RISK-002 disposes of RISK-21 on the ground that \"there are no EEG "
        "cables to route\". THIS DOCUMENT IS THE ONE THAT NEVER GOT THE MEMO -- F-03 still stands "
        "in §2 as a MANDATORY shell feature and in §5/§6 as BLOCKING checklist and FAI items, "
        "contradicting the CAD parts list and the risk file. REOPENING CONDITION, stated rather "
        "than erased: NP-HW-EEGNET-001 Rev 2 (2026-08-19, DRAFT) §7.3 records that if the EEG "
        "electrode moves off the lattice onto a geodesic strut net, RISK-21 reverts and "
        "EEG-ROUTE-CHANNEL \"returns in some form\". That is the FALLBACK option, not the lead one "
        "-- Rev 2's §1.4/§1.5 make an in-tile electrode offset the lead option specifically because "
        "it avoids cables, tail and connector entirely -- and it is gated on OI-EEGNET-01 and the "
        "blocking measurement OI-EEGNET-14. So F-03 is retired under the current design, NOT proven "
        "impossible; if that fork resolves toward the net, this feature is re-derived from scratch "
        "against L0/L1, not un-retired from the geometry below.",
        False,
        GREY,
    ),
    (
        "F-02 (§2.2) — ARCHITECTURE-INDEPENDENT, BUT IN DIRECT CONFLICT WITH NP-TOOL-HUB-001 Rev 1. "
        "Nothing about hex tiling touches it; the defect is a scope collision that post-dates this "
        "document. F-02 specifies three tethered port-cover anchors ON THE CONTROL HUB SHELL "
        "(USB-C, left and right hub accessory ports) at 4.0 mm dia x 3.5 mm boss with a 50 mm "
        "tether. NP-TOOL-HUB-001 Rev 1 F-02 — created by this document's own §9 remaining action "
        "— specifies the hub's port-cover anchors at 1.0 mm dia x 0.5 mm protrusion with a tether "
        "of <=20 mm FREE LENGTH, and that maximum is a hazard control, not a styling choice: it "
        "exists so a detached-but-tethered cover cannot reach the F-04 fan intake grille "
        "(FAI-HTOOL-02, marked BLOCKING, foreign-object-near-fan-blade hazard). This document's "
        "50 mm tether exceeds that limit by 2.5x, on the same shell, in a document that is also "
        "ACTIVE. The port counts disagree too (3 here vs 2 there). NP-TOOL-HUB-001's own "
        "back-reference is wrong in both halves — it describes this F-02 as covering \"the "
        "headset's zone-slot and lens-rim accessory ports\", but F-02 is on the hub, and lens-rim "
        "guards are out of scope per §1 and belong to NP-TOOL-LENS-001 F-07. NP-TOOL-HUB-001 is "
        "the later and BASELINED document and its tether limit is safety-derived; treat its "
        "geometry as governing until OI-ART-07 resolves ownership. Do not release F-02 geometry "
        "to a tooling manufacturer from this document.",
        False,
        GREY,
    ),
    (
        "F-04 (§2.4) — SURVIVES UNCHANGED, and remains a genuine pre-first-cut requirement. The "
        "temporal wing anchor boss for the 40 Hz mastoid vibrotactile pad is independent of how the "
        "vault is tiled; the temporal stability wing is unaffected by the two-bowl split. Its "
        "PROVISIONAL rationale is untouched — molding the boss at first cut is zero-cost and "
        "retrofitting it is $15,000-40,000 plus 6-8 weeks. SH-15..SH-19, FAI-TW-01, FAI-TW-02, "
        "OI-05, OI-07 and the F-04 rows of §3 all stand. Two consequential notes. (a) One dangling "
        "reference: §2.4's dust-cover tether is specified as \"same anchor post tether spec as "
        "F-01\", which now points at a retired feature; the referenced geometry itself survives in "
        "§3's dimension table (boss 4.0 mm dia x 3.5 mm, through-slot 2.0 x 2.5 mm, Shore 30A loop) "
        "and should be restated inline rather than by reference. (b) OI-08 tracks HOPE Phase 3 "
        "(Cognito Therapeutics, n=670) as \"expected mid-2026\", which is now past — the monitoring "
        "item is due for a status check, and this note does not supply a result.",
        False,
        GREY,
    ),
    (
        "§1 SCOPE is itself understated. It reads \"Headset outer CFRP shell and temporal stability "
        "wing\", written when the shell was one part. NP-HEX-ZM-001 §5 splits it into two nested "
        "bowls: L0 (the outer CFRP bowl carrying the complete 5-layer EMF stack) and L1 (the inner "
        "module carrier holding the sockets, cluster clamps and fluxgates). This document governs "
        "L0 only. L1 has NO tooling specification at all — NP-ART-001 A5 records it blocked on "
        "MECH-1 and OI-SHELL2-02. Three L0 features that did not exist when this was written are "
        "therefore absent from §2's \"consolidated checklist\" despite it claiming to list all "
        "molded features required in the headset shell: the four-corner clamp latches with their "
        "BeCu ground-bond fingers and Hall interlocks (NP-HEX-ZM-001 §5.4), the posterior-centre "
        "blind-mate connector boss (§5.3c), and the labyrinth overlap lip plus conductive-bead "
        "groove at the parting plane (§5.3a). All three are MECH-1, and all three are unretrofittable "
        "in exactly the sense §2 warns about.",
        False,
        GREY,
    ),
    (
        "Related Documents (§1) — three of the seven rows point at retired documents. "
        "NP-DRV-SHELL-001 (this document's stated parent design review) is SUPERSEDED, replaced by "
        "NP-DRV-SHELL-002; NP-TOOL-ZM-001 is SUPERSEDED, replaced by NP-TOOL-HEXTILE-001; and "
        "NP-FAI-ZM-001 is RETIRED, replaced per artifact by NP-FAI-001 + NP-ART-001 + "
        "NP-FAI-HUB-001. All three are indexed in docs/superseded/README.md. The CLAUDE.md section "
        "numbers cited (§7.1, §8.3, §3b) predate the Rev 33 reorganisation: interface covers are now "
        "§2.3, and the mastoid pad moved to docs/reference/accessories-roadmap.md.",
        False,
        GREY,
    ),
    (
        "STILL REUSABLE, and worth carrying into whatever replaces this document: the anchor-post + "
        "tether pattern as a loss-prevention primitive (already reused by NP-TOOL-HUB-001 F-02 and "
        "NP-TOOL-LENS-001 F-07); the boss-root stress-concentration rule (base radius >=1.0 mm on "
        "CFRP); the IPX4-at-every-boss-base sealing requirement; the draft-angle floors (>=1.5 deg "
        "on bosses, >=1 deg on interior surfaces); the §4 material specifications; and the "
        "structure of §5 itself — a named reviewer, a method and a pass criterion per row — which "
        "is the pattern NP-TOOL-HUB-001 and NP-TOOL-HEXTILE-001 both adopted.",
        False,
        GREY,
    ),
    (
        "OPEN GAP THIS NOTE DOES NOT FILL: no document yet specifies the L0 outer bowl's clamp "
        "latches, blind-mate boss, or parting-plane lip geometry, and no document specifies the L1 "
        "inner bowl's tooling at all. Both wait on MECH-1. Re-scoping or superseding this document "
        "is OI-ART-01, which NP-ART-001 §6 states cannot be completed before MECH-1 fixes the clamp "
        "and boss geometry. Until then this document is releasable for F-04 only.",
        False,
        GREY,
    ),
]


def patch_docx():
    doc = Document(SHELL_PATH)
    if already_patched(doc):
        print(f"  Already patched -- skipping: {SHELL_PATH}")
        return

    anchor = None
    for p in doc.paragraphs:
        if p.text.strip() == "1.  PURPOSE AND SCOPE":
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("Could not find '1.  PURPOSE AND SCOPE' anchor paragraph")

    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "FCE4D6")
    para = cell.paragraphs[0]

    for i, (text, bold, color) in enumerate(BLOCKS):
        if i:
            para.add_run().add_break()
            para.add_run().add_break()
        r = para.add_run(text)
        r.bold = bold
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(*color)

    anchor._p.addprevious(t._tbl)
    anchor._p.addprevious(OxmlElement("w:p"))

    # Header status cell (table 0) -- flag the partial supersession where a reader looks first.
    hdr = doc.tables[0].rows[0].cells[1]
    hp = hdr.add_paragraph()
    hr = hp.add_run(
        "⚠ PARTIALLY SUPERSEDED 2026-08-18 — see the banner at §1. "
        "Releasable for F-04 only."
    )
    hr.bold = True
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(*RED)

    # Revision history (last table) -- a content change without a revision bump is the
    # bookkeeping lapse NP-HEX-ZM-001 Rev 2 called out.
    rev = doc.tables[-1]
    row = rev.add_row()
    row.cells[0].text = "2"
    row.cells[1].text = "2026-08-18"
    row.cells[2].text = "NeurOne Systems Engineering"
    row.cells[3].text = (
        "Partial supersession against the hex-tile architecture (NP-HEX-ZM-001, "
        "NP-DRV-SHELL-002). No feature, dimension, checklist row or open item is deleted. "
        "F-01 retired outright (no zone slots, no zone colours, sockets moved to the inner "
        "bowl); F-03's premise retired (EEG electrodes are inside T1-B tiles and the EEG "
        "harness is absorbed into the cluster network -- the <5 uVpp threshold survives as "
        "NP-DRV-SHELL-002 SH2-DRC-16); F-02 architecture-independent but in conflict with "
        "NP-TOOL-HUB-001 Rev 1 F-02, whose <=20 mm tether limit is a fan-ingestion hazard "
        "control this document's 50 mm tether exceeds; F-04 survives unchanged. F-03's retirement agrees with "
        "cad/CAD_PARTS_LIST.md's existing EEG-ROUTE-CHANNEL tombstone and NP-RISK-002's RISK-21 "
        "disposition; reopening condition recorded per NP-HW-EEGNET-001 Rev 2 §7.3. Scope "
        "understated -- this governs L0 only, and three L0 features (clamp latches, "
        "blind-mate boss, parting-plane lip) are absent pending MECH-1. Revision numbering "
        "follows NP-CONV-001 Rev 3 §4.1 (integers); the historical row 'A' is Rev 1. "
        "Re-scope or supersede is OI-ART-01."
    )

    doc.save(SHELL_PATH)
    print(f"  Partial-supersession banner inserted: {SHELL_PATH}")


if __name__ == "__main__":
    patch_docx()
