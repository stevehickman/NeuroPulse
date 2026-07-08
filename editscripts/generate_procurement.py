"""Generate NeurOne FPC Zone Module Procurement Requirements document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(5)

# ── helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(10.5)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Inches(0.2)
    return p

def warn(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run('⚠  ' + text)
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
    return p

def ok(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run('✓  ' + text)
    r.italic = True; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x00, 0x77, 0x33)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(9.5)
    return p

def bullet(bold_prefix, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + '  ')
        r.bold = True; r.font.size = Pt(10.5)
    p.add_run(text).font.size = Pt(10.5)
    return p

def numbered(text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text).font.size = Pt(10.5)
    return p

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(9.5)
        c.paragraphs[0].paragraph_format.space_after = Pt(0)
    for ri, row in enumerate(rows):
        tr = t.rows[ri+1]
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            c.text = val
            for run in c.paragraphs[0].runs:
                run.font.size = Pt(9.5)
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('NeurOne')
r.bold = True; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Zone Module FPC — Procurement Requirements')
r.bold = True; r.font.size = Pt(17)
r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

for label, val in [
    ('Document ID:',    'NP-PROC-FPC-001'),
    ('Revision:',       'A  —  2026-05-06'),
    ('Status:',         'Active — must be reviewed before first purchase orders are placed'),
    ('Applies to:',     'All components for Zone Module FPC (NP-HW-FPC-001 Rev C)'),
    ('Owner:',          'Supply Chain / Hardware Engineering'),
    ('Risk register:',  'NP-HW-FPC-001 Risk Register — items RISK-01, RISK-07, RISK-08 addressed here'),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + '  '); r.bold = True; r.font.size = Pt(10)
    p.add_run(val).font.size = Pt(10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
r = p.add_run(
    'This document defines mandatory procurement specifications that must appear on '
    'purchase orders, supplier qualification letters, and incoming inspection instructions. '
    'Specifications not stated here that conflict with supplier defaults will result in '
    'non-conforming parts being delivered without notification.'
)
r.italic = True; r.font.size = Pt(10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. SCOPE
# ══════════════════════════════════════════════════════════════════
h1('1.  SCOPE AND PURPOSE')
body(
    'This document specifies procurement requirements for components used in the NeurOne '
    'Zone Module Flexible Printed Circuit (NP-HW-FPC-001). It exists because several '
    'component parameters critical to device performance and safety are not enforced by '
    'default supplier processes — they must be explicitly stated on purchase orders or '
    'they will not be controlled. Requirements are derived from risk analysis in the '
    'companion risk register (NP-HW-FPC-001 Risk Register Rev A).'
)
body(
    'All specifications in this document are MANDATORY unless marked ADVISORY. '
    'A purchase order that omits a mandatory specification is non-conforming and must '
    'not be placed without written approval from the Hardware Engineering Lead.'
)

tbl(
    headers=['Component category', 'Risk register origin', 'Section'],
    rows=[
        ['LED emitters — 660 nm and 808–830 nm', 'RISK-08 (Vf binning)', '§2'],
        ['LED driver IC — Infineon BCR421W', 'RISK-07 (variant confusion)', '§3'],
        ['Zone module connector — Hirose FH34S / JAE FF03', 'RISK-01 (wrong connector family)', '§4'],
        ['FPC fabrication — copper type and weight', 'RISK-02 (current capacity), RISK-10 (ED vs RA)', '§5'],
        ['PDMS optical window material', 'RISK-04 (CTE delamination)', '§6'],
        ['Incoming inspection requirements', 'All above', '§7'],
    ],
    widths=[2.4, 2.0, 0.7]
)

# ══════════════════════════════════════════════════════════════════
# 2. LED EMITTERS  (RISK-08 — primary item)
# ══════════════════════════════════════════════════════════════════
h1('2.  LED EMITTERS — 660 nm AND 808–830 nm')
warn(
    'RISK-08 (HIGH): Commodity LED emitters ship in Vf bins of ±0.20–0.30 V as standard. '
    'This width causes current imbalance between parallel strings even with per-string '
    'constant-current drivers, reducing dose accuracy and accelerating premature LED failure. '
    'The specifications below MUST appear on every purchase order.'
)

h2('2.1  Forward Voltage (Vf) Bin Specification')
body(
    'This is the most critical procurement specification in this document. '
    'Vf bin width directly determines current matching between LED strings. '
    'Wide bins cause current hogging even when per-string constant-current drivers are used, '
    'because the driver\'s compliance headroom is consumed unevenly across strings.'
)

tbl(
    headers=['Parameter', '660 nm LED', '808–830 nm LED'],
    rows=[
        ['Nominal forward voltage at 150 mA, 25°C', '2.0–2.2 V (confirm from datasheet)', '1.6–1.8 V (confirm from datasheet)'],
        ['MAXIMUM Vf bin width — within a single purchase order', '± 0.10 V', '± 0.10 V'],
        ['MAXIMUM Vf bin width — between purchase orders (lot-to-lot)', '± 0.15 V', '± 0.15 V'],
        ['Vf measurement condition', '150 mA DC, junction at 25°C, 10 ms pulse to avoid self-heating', 'Same'],
        ['Vf distribution data required with each lot', 'Yes — mean, σ, min, max across sampled units', 'Yes'],
        ['Minimum sample size for Vf distribution certificate', '50 units per delivery lot', '50 units per delivery lot'],
    ],
    widths=[2.8, 1.8, 1.8]
)

warn(
    'The ±0.10 V within-order bin width is tighter than commodity specification. '
    'Major LED suppliers (Seoul Semiconductor, Osram, Luminus) can meet this with '
    'explicit PO specification — but it must be stated. Without it, the supplier '
    'ships from mixed bin stock. Confirm bin availability and any price premium '
    'before first order.'
)

h3('2.1.1  Purchase Order Language — Vf Bin')
body(
    'The following text must appear verbatim in the "Special Requirements" or '
    '"Additional Specifications" section of every LED emitter purchase order:'
)
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Inches(0.4)
p.paragraph_format.right_indent = Inches(0.4)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run(
    '"All units in this order shall be from a single forward voltage (Vf) bin. '
    'Maximum Vf spread within the order: ±0.10 V, measured at [150 mA DC / 10 ms pulse / '
    'junction at 25°C]. Supplier must provide a Vf distribution certificate (mean, standard '
    'deviation, minimum, maximum) based on a minimum sample of 50 units from this lot. '
    'Orders shipped without a Vf distribution certificate will be rejected at incoming inspection."'
)
r.italic = True; r.font.size = Pt(10.5)

h2('2.2  Wavelength Bin Specification')
tbl(
    headers=['Parameter', '660 nm LED', '808–830 nm LED'],
    rows=[
        ['Peak wavelength range', '655–670 nm', '805–835 nm'],
        ['Maximum wavelength spread within order', '± 5 nm', '± 10 nm'],
        ['Measurement condition', 'At 150 mA, 25°C junction temperature', 'Same'],
        ['Wavelength certificate required', 'Yes — peak λ mean and range', 'Yes'],
    ],
    widths=[2.8, 1.8, 1.8]
)
note(
    'Wavelength tolerance is less critical than Vf for current matching, but matters '
    'for dose metering calibration. The reference photodiode calibration coefficient '
    '(irradiance-per-unit photocurrent) is wavelength-dependent. If lot-to-lot wavelength '
    'shifts exceed ±10 nm, the factory calibration coefficient must be recalculated for '
    'that lot.'
)

h2('2.3  Thermal and Package Specification')
tbl(
    headers=['Parameter', 'Requirement', 'Notes'],
    rows=[
        ['Package', 'SMD 2835 or equivalent ceramic substrate', 'Must be compatible with FPC reflow process (peak 245°C, 30 s)'],
        ['Junction-to-case thermal resistance (θjc)', '< 15 °C/W', 'Required for IEC 60601 surface temperature compliance model'],
        ['Maximum junction temperature (Tj_max)', '≥ 125°C', 'Operating headroom above 62°C throttle threshold'],
        ['Forward current — absolute maximum (If_max)', '≥ 350 mA', 'Must exceed 180 mA operating max with 2× margin'],
        ['Pulse current rating', 'Confirm ≥ 180 mA at ≤ 25% duty cycle, ≤ 1 ms pulse width', 'Required before BOM lock — see CLAUDE.md §13.4'],
        ['Rated lifetime L70', '≥ 80,000 hours at 150 mA nominal', 'L70 = time to 70% of initial luminous flux'],
        ['Moisture sensitivity level (MSL)', 'MSL 3 or better', 'Must be stored and handled per IPC/JEDEC J-STD-033'],
        ['RoHS compliance', 'Required', 'IEC 63000 / EU Directive 2011/65/EU'],
    ],
    widths=[2.5, 1.8, 2.5]
)
warn(
    'Pulse current rating at 120–180 mA for both wavelengths has not yet been '
    'confirmed from supplier datasheets. This must be verified before first LED '
    'order is placed. See CLAUDE.md §13.4 open issue.'
)

h2('2.4  Incoming Inspection — LED Emitters')
body(
    'Every delivery lot must be inspected before use in production. '
    'Incoming inspection failure results in lot quarantine and supplier notification. '
    'Do not use quarantined lots without written disposition from Hardware Engineering.'
)
tbl(
    headers=['Inspection item', 'Method', 'Sample size', 'Accept / Reject criterion'],
    rows=[
        ['Vf distribution verification', 'Measure Vf at 150 mA, 10 ms pulse, 25°C on sample units. Compare to supplier certificate.', '10 units per lot (minimum)', 'Accept: all measured Vf within ±0.10 V of lot mean stated on certificate. Reject: any unit outside ±0.15 V of lot mean, or certificate not provided.'],
        ['Peak wavelength spot-check', 'Measure λ_peak on spectrometer at 150 mA, 25°C.', '5 units per lot', 'Accept: λ_peak within stated wavelength range ± 5 nm. Reject: any unit outside range.'],
        ['Package and marking', 'Visual inspection vs. datasheet.', '100% of reel labels; 5 physical units', 'Reject: wrong part number, damaged tape, missing/illegible marking.'],
        ['MSL date code', 'Check reel label for exposure time and floor life.', '100% of reels', 'Reject: floor life expired or not stated.'],
        ['Certificate of conformance (CoC)', 'Verify CoC accompanies shipment and references PO number.', 'Per shipment', 'Reject shipment if CoC absent.'],
    ],
    widths=[1.5, 1.8, 1.2, 2.7]
)

h2('2.5  Out-of-Specification Lot Disposition')
body(
    'If incoming inspection reveals a lot that does not meet the Vf bin specification, '
    'the following disposition process applies:'
)
numbered('Quarantine the lot immediately. Label reels "HOLD — PENDING DISPOSITION". '
         'Do not allow use in production.')
numbered('Notify the supplier within 2 business days with the measured data and the '
         'specific PO specification that was not met. Request a replacement lot from '
         'a conforming bin.')
numbered('Hardware Engineering Lead assesses whether the out-of-spec lot can be '
         'used with design adjustment. Options: '
         '(a) Use if Vf spread is ≤ ±0.15 V and all strings within a zone module '
         'are sorted from the same narrow sub-bin — requires manual sorting process; '
         '(b) Use at reduced current (lower RSET value) to reduce current-hogging risk — '
         'reduces irradiance proportionally; '
         '(c) Return to supplier for replacement.')
numbered('Document disposition decision in the lot traveller. Any "use-as-is" decision '
         'requires written sign-off from Hardware Engineering Lead and Quality.')
numbered('If a supplier repeatedly ships out-of-spec lots, qualify an alternative '
         'supplier from the approved backup list (§2.6) before placing next order.')

h2('2.6  Approved Supplier List — LED Emitters')
tbl(
    headers=['Supplier', 'Status', 'Wavelengths covered', 'Contact / Notes'],
    rows=[
        ['Seoul Semiconductor', 'Primary', '660 nm + 808–830 nm', 'Confirm Vf bin capability before first order'],
        ['Osram (ams-OSRAM)', 'Primary', '660 nm + 808–830 nm', 'Confirm Vf bin capability before first order'],
        ['Luminus Devices', 'Backup', '660 nm + 808–830 nm', 'Qualify if primary suppliers cannot meet ±0.10 V bin requirement'],
    ],
    widths=[1.8, 0.8, 1.8, 2.8]
)
note(
    'No supplier is yet fully qualified. "Primary" status means the supplier has been '
    'identified as likely capable. Full qualification requires: (1) sample lot with '
    'Vf certificate; (2) incoming inspection pass; (3) prototype build and functional test. '
    'Qualification must be completed before production orders are placed.'
)

# ══════════════════════════════════════════════════════════════════
# 3. LED DRIVER IC
# ══════════════════════════════════════════════════════════════════
h1('3.  LED DRIVER IC — INFINEON BCR421W')
warn(
    'RISK-07 (MITIGATED): BCR421U (150 mA max) and BCR421W (350 mA max) are different '
    'variants with the same package (SOT-323) and similar appearance. Only BCR421W is '
    'approved. BCR421U must never be substituted, even if BCR421W is on allocation.'
)

tbl(
    headers=['Parameter', 'Required specification'],
    rows=[
        ['Approved part number', 'Infineon BCR421W (350 mA rated) ONLY'],
        ['Explicitly excluded part numbers',
         'BCR421U (150 mA max — NEVER substitute). '
         'Any other BCR42x variant requires Hardware Engineering written approval.'],
        ['Package', 'SOT-323 (SC-70), 3-pin'],
        ['Current setting', 'Via external RSET resistor — value TBD pending regulatory opinion letter. '
         'Do NOT populate RSET until RSET value is released by Hardware Engineering.'],
        ['Backup / alternate (approved)',
         'Zetex ZXLD381 (350 mA, SOT-23-3)  |  Diodes Inc. AL5809 (350 mA, SOT-23-3). '
         'Both require RSET recalculation from respective datasheets before use.'],
        ['Approved supplier', 'Infineon Technologies — authorised distributor only (Mouser, DigiKey, Arrow, Avnet)'],
        ['Counterfeit risk', 'HIGH — SOT-23/SOT-323 linear regulators are heavily counterfeited. '
         'Purchase from authorised distributor only. Request certificate of authenticity for orders > 500 units.'],
        ['RoHS compliance', 'Required'],
    ],
    widths=[2.2, 5.0]
)

h3('Purchase Order Language — BCR421W')
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Inches(0.4)
p.paragraph_format.right_indent = Inches(0.4)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run(
    '"Part number: Infineon BCR421W. Package: SOT-323. '
    'SUBSTITUTION NOT PERMITTED — BCR421U and all other BCR42x variants are explicitly '
    'excluded. If BCR421W is on allocation or unavailable, contact Hardware Engineering '
    'before placing an order for any alternative part. '
    'Purchase from authorised distributor only. Certificate of conformance required."'
)
r.italic = True; r.font.size = Pt(10.5)

# ══════════════════════════════════════════════════════════════════
# 4. ZONE MODULE CONNECTOR
# ══════════════════════════════════════════════════════════════════
h1('4.  ZONE MODULE CONNECTOR — HIROSE FH34S / JAE FF03')
ok(
    'RISK-01 (MITIGATED): Molex SlimStack removed from specification. '
    'Hirose FH34S is the correct connector family for this application.'
)

tbl(
    headers=['Parameter', 'Required specification'],
    rows=[
        ['Primary approved part number', 'Hirose FH34S-20S-0.5SH(50)  —  20-position, 0.5 mm pitch, back-flip lever ZIF, SMT'],
        ['Backup approved part number', 'JAE FF03 series  —  0.5 mm pitch, high-cycle FPC connector'],
        ['Explicitly excluded families',
         'Molex SlimStack (board-to-board — incompatible with FPC). '
         'Molex 52610 / 52745 / 52892 series (rated 20–30 cycles — below requirement). '
         'ANY 0.35 mm pitch connector (current per pin too high for power rails at 0.9 A RMS).'],
        ['Rated insertion cycles', '≥ 1,000 — MUST BE CONFIRMED FROM FULL HIROSE FH34S DATASHEET '
         'before BOM lock. Product marketing states ≥1,000 cycles; confirm in test conditions table.'],
        ['Contact current rating', '≥ 0.5 A per contact'],
        ['Pitch', '0.5 mm (NOT 0.35 mm)'],
        ['Actuation', 'Back-flip lever ZIF — zero insertion force, tool-free extraction'],
        ['Mated height', '1.0 mm maximum'],
        ['Operating temperature', '−40°C to +85°C minimum'],
        ['Approved supplier', 'Hirose Electric — authorised distributor (DigiKey, Mouser, Arrow)'],
        ['RoHS compliance', 'Required'],
    ],
    widths=[2.2, 5.0]
)

warn(
    'Before placing the first connector order: pull the complete FH34S datasheet '
    'from hirose.com and confirm that the "Endurance" or "Insertion/Withdrawal cycles" '
    'specification in the test conditions table states ≥ 1,000 cycles. '
    'If the full datasheet does not confirm this, contact Hirose applications engineering '
    'for written confirmation before placing the order.'
)

h3('Incoming Inspection — Connector')
numbered('Verify part marking against approved part number FH34S-20S-0.5SH(50).')
numbered('Verify position count (20) and pitch (0.5 mm) on 5 sample units using '
         'calibrated pin gauge.')
numbered('Perform 5 insertion/extraction cycles on 3 sample units. Verify lever '
         'actuates correctly and FPC tail seats without resistance. Measure contact '
         'resistance at pins 1 and 20 before and after: must be < 50 mΩ per contact.')
numbered('Certificate of conformance from Hirose or authorised distributor must '
         'accompany each shipment.')

# ══════════════════════════════════════════════════════════════════
# 5. FPC FABRICATION
# ══════════════════════════════════════════════════════════════════
h1('5.  FPC FABRICATION REQUIREMENTS')
warn(
    'RISK-10 (MEDIUM): FPC fabricators default to electrodeposited (ED) copper. '
    'ED copper fails under dynamic flex. This requirement MUST appear on the '
    'fabrication drawing AND in the purchase order.'
)

tbl(
    headers=['Parameter', 'Required specification', 'Default if not stated (WRONG)'],
    rows=[
        ['Copper type', 'ROLLED ANNEALED (RA) copper — mandatory', 'Electrodeposited (ED) copper'],
        ['Copper weight', '1 oz (35 µm) on all layers', 'May be substituted with 0.5 oz if not specified'],
        ['Minimum power trace width', '≥ 2.0 mm for all VLED and GND power rails — see NP-HW-FPC-001 §10.1',
         'Trace width not enforced — layout engineer\'s discretion'],
        ['Substrate', 'Polyimide (PI), 25 µm', 'May substitute polyester (PET) — not acceptable'],
        ['Adhesive', 'Acrylic or adhesiveless laminate — NO epoxy', 'Epoxy adhesive (lower thermal performance)'],
        ['Surface finish', 'ENIG on signal pads; hard gold ≥ 0.5 µm (cobalt-alloyed) on connector tail pads',
         'ENIG only — insufficient wear resistance on connector tail'],
        ['Dynamic bend radius', '≥ 25 mm at all bend zones — must be stated on drawing',
         'No bend radius enforcement'],
        ['IPC class', 'Class 2 minimum', 'Class 1'],
        ['Flammability', 'UL 94 V-0', 'Not specified'],
        ['Certificate of conformance', 'Required per lot', 'Not provided unless requested'],
    ],
    widths=[1.8, 2.6, 2.8]
)

h3('Purchase Order Language — FPC Fabrication')
p = doc.add_paragraph()
p.paragraph_format.left_indent  = Inches(0.4)
p.paragraph_format.right_indent = Inches(0.4)
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(6)
r = p.add_run(
    '"COPPER TYPE: ROLLED ANNEALED (RA) — ELECTRODEPOSITED (ED) COPPER IS NOT ACCEPTABLE '
    'AND WILL RESULT IN REJECTION OF THE ENTIRE LOT. COPPER WEIGHT: 1 OZ (35 µm) ON ALL '
    'LAYERS. MINIMUM POWER TRACE WIDTH: 2.0 mm ON ALL VLED AND GND RAILS AS SHOWN ON '
    'DRAWING. IPC CLASS 2. UL 94 V-0. CERTIFICATE OF CONFORMANCE REQUIRED."'
)
r.italic = True; r.font.size = Pt(10.5)

h3('Incoming Inspection — FPC')
numbered('Request material certification confirming RA copper from FPC fabricator. '
         'Reject lot if only ED copper certification is provided.')
numbered('Measure copper trace width on power rails using calibrated optical comparator '
         'or cross-section: minimum 2.0 mm on VLED_660, GND_660, VLED_808, GND_808 buses. '
         'Sample: 3 units per lot.')
numbered('Verify connector tail gold thickness via XRF: ≥ 0.5 µm hard gold. '
         'Sample: 2 units per lot.')
numbered('Perform 20 insertion/extraction cycles on 2 sample units through the '
         'production ZIF connector. Verify connector tail shows no delamination, '
         'pad lifting, or gold wear through to base metal.')

# ══════════════════════════════════════════════════════════════════
# 6. PDMS OPTICAL WINDOW MATERIAL
# ══════════════════════════════════════════════════════════════════
h1('6.  PDMS OPTICAL WINDOW MATERIAL')
warn(
    'RISK-04 (HIGH): PDMS–polyimide CTE mismatch (15:1) requires specific bonding '
    'process (sputtered SiO₂ interlayer). Standard PDMS from any supplier is acceptable '
    'for the base material; the bonding process is the critical variable. '
    'See NP-HW-FPC-001 §9 for full bonding process specification.'
)

tbl(
    headers=['Parameter', 'Required specification'],
    rows=[
        ['Base material', 'Polydimethylsiloxane (PDMS)'],
        ['Approved formulation', 'Dow Sylgard 184 (primary)  |  equivalent 10:1 base:curing agent formulation (backup — confirm optical transmission)'],
        ['Mixing ratio', '10:1 base:curing agent by weight — must be recorded on batch traveller'],
        ['Cure conditions', '80°C for 2 hours minimum — or 120°C for 1 hour. Do not cure at room temperature (insufficient crosslink density for thermal cycling resistance).'],
        ['Optical transmission at 660 nm (through spec thickness)', '> 92% — measure before bonding; reject batches below threshold'],
        ['Optical transmission at 808–830 nm', '> 94%'],
        ['Shore hardness', '40–50 A (confirm from formulation datasheet)'],
        ['Batch traceability', 'Lot number of base and curing agent must be recorded in batch traveller for each production run'],
        ['Storage', 'Unmixed components: store at 4–8°C, away from UV. Mixed/uncured PDMS: use within 4 hours of mixing or discard.'],
        ['Bonding process requirement', 'Sputtered SiO₂ interlayer (75 nm) on PI surface required before O₂ plasma activation. '
         'Simple O₂ plasma bond alone is NOT acceptable for production — see NP-HW-FPC-001 §9.2.'],
    ],
    widths=[2.2, 5.0]
)

note(
    'PDMS is a commodity material. Supplier qualification is less critical than for '
    'electronic components. The bonding process (§9.2 of spec) is the quality gate — '
    'peel test per batch is the acceptance criterion, not PDMS material source.'
)

# ══════════════════════════════════════════════════════════════════
# 7. INCOMING INSPECTION SUMMARY
# ══════════════════════════════════════════════════════════════════
h1('7.  INCOMING INSPECTION SUMMARY')
body(
    'The following table summarises all incoming inspection requirements across all '
    'component categories. Inspection records must be retained for a minimum of 7 years '
    'for medical device traceability.'
)
tbl(
    headers=['Component', 'Critical inspection item', 'Reject criterion', 'Sample size'],
    rows=[
        ['660 nm LED', 'Vf distribution certificate provided', 'Missing certificate → reject shipment', 'Per shipment'],
        ['660 nm LED', 'Vf spread within ±0.15 V of lot mean', 'Any unit outside → quarantine lot', '10 units / lot'],
        ['660 nm LED', 'Peak wavelength 655–670 nm', 'Any unit outside ±5 nm of range', '5 units / lot'],
        ['808–830 nm LED', 'Vf distribution certificate provided', 'Missing certificate → reject shipment', 'Per shipment'],
        ['808–830 nm LED', 'Vf spread within ±0.15 V of lot mean', 'Any unit outside → quarantine lot', '10 units / lot'],
        ['808–830 nm LED', 'Peak wavelength 805–835 nm', 'Any unit outside ±10 nm', '5 units / lot'],
        ['BCR421W', 'Part number marking confirms W variant (not U)', 'Any BCR421U found → reject entire lot', '10 units / lot'],
        ['BCR421W', 'Purchase from authorised distributor', 'Grey market source → reject', 'Per PO'],
        ['FH34S connector', 'Part number and position count (20) confirmed', 'Wrong count or part → reject', '5 units / lot'],
        ['FH34S connector', 'ZIF lever actuates and releases correctly', 'Stiff or non-actuating → quarantine', '3 units / lot'],
        ['FH34S connector', 'Contact resistance < 50 mΩ after 5 cycles', 'Any pin > 50 mΩ → quarantine lot', '3 units / lot'],
        ['FPC', 'RA copper material certificate', 'ED copper certificate or none → reject lot', 'Per lot'],
        ['FPC', 'Power trace width ≥ 2.0 mm on VLED/GND rails', 'Any rail < 2.0 mm → reject lot', '3 units / lot'],
        ['FPC', 'Hard gold on connector tail ≥ 0.5 µm (XRF)', 'Any pad < 0.4 µm → quarantine', '2 units / lot'],
        ['PDMS', 'Optical transmission ≥ 92% at 660 nm', 'Any batch < 90% → reject', '1 sample / batch'],
    ],
    widths=[1.4, 2.2, 1.8, 1.2]
)

# ══════════════════════════════════════════════════════════════════
# 8. REVISION HISTORY
# ══════════════════════════════════════════════════════════════════
h1('8.  REVISION HISTORY')
tbl(
    headers=['Rev', 'Date', 'Author', 'Description'],
    rows=[
        ['A', '2026-05-06', 'Claude / NeurOne Engineering',
         'Initial release. Addresses RISK-01 (connector), RISK-04 (PDMS), '
         'RISK-07 (BCR421W), RISK-08 (LED Vf binning), RISK-10 (RA copper) '
         'from NP-HW-FPC-001 Risk Register Rev A.'],
    ],
    widths=[0.4, 0.9, 2.3, 3.6]
)

# ── save ──────────────────────────────────────────────────────────────────────
out = '/home/user/NeurOne/docs/neurone_fpc_procurement_requirements.docx'
doc.save(out)
print(f'Saved: {out}')
