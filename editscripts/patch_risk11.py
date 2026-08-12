"""
⚠ SUPERSEDED (2026-07-28): this script's NP-DRV-SHELL-001 mitigation (5 fixed zone-slot FPC
routing paths) is retired by the hex-tile/socket architecture (NP-HEX-ZM-001) — see
editscripts/generate_fpc_routing_review.py's header note. The generic bend-radius physics
this script adds to FPC spec §11.2 (IPC-2223D basis, ≥25mm dynamic / ≥12.5mm static) remain
valid engineering constraints, but the RISK-11 mitigation claim itself needed correcting. Do
not re-run this script (its §11.2 insertion isn't idempotent and its risk-register mitigation
text is stale). The corrected RISK-11 state is applied by
editscripts/patch_hexzm_risk_corrections.py. Kept here as historical record.

Patch RISK-11:
1. Populate §11.2 in FPC spec with bend radius requirements + reference to NP-DRV-SHELL-001
2. Update risk register RISK-11 summary row and body text
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# ─────────────────────────────────────────────────────────────────────────────
# 1. Patch FPC spec §11.2
# ─────────────────────────────────────────────────────────────────────────────

SPEC_PATH = "docs/superseded/neurone_fpc_zone_module_spec_revA.docx"
spec = Document(SPEC_PATH)

# Find §11.2 heading and insert content after it
s112_idx = None
for i, para in enumerate(spec.paragraphs):
    if para.text.strip() == "11.2  Bend Radius Requirements":
        s112_idx = i
        break

print(f"§11.2 heading at paragraph {s112_idx}")

if s112_idx is not None:
    # Insert paragraphs after §11.2 heading (before §11.3)
    # We insert "before" the §11.3 paragraph to get correct order
    s113_para = spec.paragraphs[s112_idx + 1]  # this is §11.3

    CONTENT = [
        # (text, bold, is_bullet, color)
        ("[DRAWING REQUIREMENT] Minimum bend radius — dynamic flex: ≥25 mm at all points along the routing path inside the headset shell. "
         "This value is enforced by shell interior cable management channel geometry (see NP-DRV-SHELL-001). "
         "No user action, tool, or assembly step may impose a bend tighter than 25 mm on any FPC.",
         True, False, (0xC0, 0x00, 0x00)),
        ("Minimum bend radius — static flex (one-time assembly bend, not actuated during module swaps): ≥12.5 mm. "
         "Static bend locations must be identified on the routing path drawing and confirmed as static-only.",
         False, False, None),
        ("No bend within 5 mm of the ZIF connector tail end, the FR4 stiffener edge, or any zone module PCB edge. "
         "Bend at these transitions causes stiffener-to-FPC delamination, not copper fatigue.",
         False, False, None),
        ("Bend radius enforcer geometry: the shell interior cable channel must include a moulded boss or ridge "
         "at the Segment A (zone slot exit) and Segment C (Hub PCB entry) transitions, profiled to 25 mm radius, "
         "so that the FPC cannot be bent tighter by any assembly or user action.",
         False, False, None),
        ("FPC length: ≥10% slack relative to routing path arc length in the assembled headset at maximum head "
         "circumference (62 cm). Slack prevents tensile force on the ZIF connector tail. Maximum slack: 25% "
         "(excess causes spontaneous tight bend in channel).",
         False, False, None),
        ("Design review reference: Shell tooling FPC routing path must be reviewed against requirements in "
         "NP-DRV-SHELL-001 Rev 1 (Zone Module FPC Routing Path — Shell Tooling Design Review) before "
         "tooling specification release. NP-DRV-SHELL-001 §5 contains the design review checklist (15 items) "
         "that must pass before first-cut approval.",
         False, False, None),
    ]

    for (text, bold, is_bullet, color) in reversed(CONTENT):
        new_p = OxmlElement("w:p")
        s113_para._p.addprevious(new_p)
        # Find the new paragraph object
        for p_obj in spec.paragraphs:
            if p_obj._p is new_p:
                if is_bullet:
                    p_obj.style = spec.styles["List Bullet"]
                run = p_obj.add_run(text)
                run.bold = bold
                run.font.size = Pt(10)
                if color:
                    run.font.color.rgb = RGBColor(*color)
                break

    print("  §11.2 content inserted")

spec.save(SPEC_PATH)
print(f"  FPC spec saved: {SPEC_PATH}")

# ─────────────────────────────────────────────────────────────────────────────
# 2. Patch Risk Register — RISK-11
# ─────────────────────────────────────────────────────────────────────────────

RISK_PATH = "docs/superseded/neurone_fpc_zone_module_risks_revA.docx"
risk_doc = Document(RISK_PATH)

tbl = risk_doc.tables[0]
for row in tbl.rows:
    if row.cells[0].text.strip() == "RISK-11":
        # Col 1: Severity
        set_cell_text(row.cells[1], "MEDIUM\n(MITIGATED)", bold=True, color=(0x00, 0x70, 0x00))
        set_cell_bg(row.cells[1], "E2EFDA")
        # Col 5: Required Mitigation — update to reference the new doc
        set_cell_text(row.cells[5],
            "MITIGATED: Shell tooling FPC routing design review document created "
            "(NP-DRV-SHELL-001 Rev 1). Document specifies 5 routing path requirements "
            "(REQ-BR-01 to REQ-BR-05), 7 shell tooling requirements (REQ-ST-01 to REQ-ST-07), "
            "15-item design review checklist (DRC-01 to DRC-15), FPC length and slack "
            "requirements for all 5 zones, and sign-off gate before tooling release. "
            "FPC spec §11.2 updated with bend radius callouts and reference to NP-DRV-SHELL-001.")
        # Col 8: Status
        set_cell_text(row.cells[8],
            "MITIGATED — NP-DRV-SHELL-001 Rev 1 created; §11.2 populated; "
            "DRC sign-off required before first-cut",
            bold=True, color=(0x00, 0x70, 0x00))
        set_cell_bg(row.cells[8], "92D050")
        print(f"  Risk register RISK-11 updated to MITIGATED")
        break

# Also update Open Issues table (Table 1) — OI-09 references the routing path
open_issues_tbl = risk_doc.tables[1]
for row in open_issues_tbl.rows:
    if row.cells[0].text.strip() == "OI-09":
        set_cell_text(row.cells[1], "REQUIRED (addressed in NP-DRV-SHELL-001)")
        set_cell_bg(row.cells[1], "E2EFDA")
        old_txt = row.cells[2].text
        set_cell_text(row.cells[2],
            "Shell tooling FPC routing path with ≥25 mm bend radius — "
            "design review doc NP-DRV-SHELL-001 Rev 1 created. "
            "CAD model and physical prototype still required to close (see NP-DRV-SHELL-001 §7).")
        print(f"  OI-09 updated")
        break

risk_doc.save(RISK_PATH)
print(f"  Risk register saved: {RISK_PATH}")
print("Done.")
