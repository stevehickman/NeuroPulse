"""
Patch RISK-17 — Multi-FPC routing + EEG cable separation.

Changes:
  1. NP-DRV-SHELL-001: add §2.3 (multi-FPC bundle), §4b (bundle tooling reqs),
     and DRC-16 to DRC-21 to the design review checklist
  2. Risk register RISK-17 → MITIGATED with detailed section
  3. NP-COORD-001 G2-01: update deliverables to match new DRC items
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── helpers ─────────────────────────────────────────────────────────────────

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def ins_before(doc, target, text, bold=False, color=None, size=10):
    new_p = OxmlElement("w:p")
    target._p.addprevious(new_p)
    for p_obj in doc.paragraphs:
        if p_obj._p is new_p:
            run = p_obj.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p_obj

def add_drc_rows(tbl, rows):
    """Append rows to the DRC checklist table."""
    for r_data in rows:
        row = tbl.add_row()
        bg = "FFFFFF"
        widths = [0.55, 2.0, 1.0, 1.25, 0.9, 0.6]
        for c_idx, txt in enumerate(r_data):
            cell = row.cells[c_idx]
            set_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(txt))
            run.font.size = Pt(9)
            if c_idx == 0:
                run.bold = True
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)

# ─────────────────────────────────────────────────────────────────────────────
# 1.  NP-DRV-SHELL-001 — add §2.3 and new DRC items 16–21
# ─────────────────────────────────────────────────────────────────────────────

SHELL_PATH = "docs/neuropulse_shell_fpc_routing_review.docx"
shell = Document(SHELL_PATH)

# ── Insert §2.3 after §2.2 heading + table ──
# Find "3.  BEND RADIUS REQUIREMENTS" heading and insert before it
s3_para = None
for p in shell.paragraphs:
    if p.text.strip().startswith("3.  BEND RADIUS"):
        s3_para = p
        break

print(f"  §3 anchor: {s3_para.text[:60]!r}")

# §2.3 content (insert in reverse order — each ins_before prepends)
S23_CONTENT = [
    # spacer
    ("", False, None),

    # Hub PCB requirement
    ("[DRAWING REQUIREMENT] "
     "Hub PCB connector placement: all 5 ZIF receptacles (Hirose DF37NB-40DS or equivalent "
     "board-side connectors) shall be placed on the same edge of the Hub PCB, oriented so "
     "that FPC tails enter from the headset interior naturally — minimising the number of "
     "imposed bends at Segment C (Hub entry). "
     "The connector row shall be perpendicular to the dominant FPC routing direction from "
     "the headset crown toward the hub. "
     "Hub PCB connector placement is a layout constraint that must be communicated to "
     "the Hub PCB layout engineer before schematic capture.",
     True, (0xC0, 0x00, 0x00)),

    # Anchor/tie-down
    ("[DRAWING REQUIREMENT] "
     "FPC anchor and tie-down points: each FPC tail shall be anchored to the shell interior "
     "at three points along its routing path — at Segment A (zone slot exit), at the "
     "mid-point of Segment B (interior channel), and at Segment C (Hub PCB entry area). "
     "Anchors prevent FPC movement during headset use (head movement, cable tension, "
     "thermal cycling). Anchor method: moulded clip or snap-over boss in CFRP shell interior "
     "that captures the FPC without compressing it (do not pinch — no bend permitted at "
     "anchor point). Anchor boss internal radius: ≥ 5 mm (prevents local stress concentration).",
     True, (0xC0, 0x00, 0x00)),

    # EEG cable separation
    ("[DRAWING REQUIREMENT] "
     "FPC-to-EEG cable separation: zone module FPC tails carry LED drive currents up to "
     "1.8 A RMS at PWM frequencies ≥ 1 kHz. EEG electrode cables carry signals in the "
     "0.5–100 Hz band at sub-microvolt amplitude — they are the highest-impedance, "
     "most EMI-susceptible signals in the system. "
     "Minimum separation between any zone module FPC tail and any EEG electrode cable: "
     "≥ 15 mm along the full routing path inside the shell. "
     "If geometrical constraints prevent ≥ 15 mm separation at any point, a physical barrier "
     "is required: a grounded CFRP divider or 0.05 mm aluminium foil barrier bonded to the "
     "shell interior and connected to chassis GND. "
     "Separation must be verified in CAD cross-section (DRC-18) and confirmed on first-article "
     "physical inspection.",
     True, (0xC0, 0x00, 0x00)),

    # Inter-FPC separation
    ("[DRAWING REQUIREMENT] "
     "Inter-FPC separation: adjacent zone FPC tails (e.g., ZM-01 and ZM-02, both frontal) "
     "shall not contact each other along Segment B. "
     "Minimum separation between adjacent FPC surfaces: ≥ 2 mm (prevents chafing wear "
     "over 1,000+ module swap cycles and removes inductive coupling between parallel "
     "high-current paths). "
     "Each FPC shall occupy its own dedicated channel section — channels may share a "
     "common wall but must not share airspace.",
     True, (0xC0, 0x00, 0x00)),

    # Bundle cross-section
    ("Bundle cross-section analysis: five zone FPC tails, each 12 mm wide × 0.20 mm thick, "
     "must traverse the shell interior simultaneously. If routed in a common area before "
     "fanning to individual channels, the combined bundle cross-section is "
     "12 mm × ~1.5 mm (5 × 0.20 mm + 4 × 0.1 mm clearance gaps). "
     "Shell interior must provide clearance for this bundle at all common-routing areas. "
     "Individual channel sections reduce bundle to single-FPC cross-section per channel "
     "(12 mm × 0.5 mm with clearance). "
     "Shell interior CAD must show explicit bundle cross-section at the tightest point.",
     False, None),

    # §2.3 heading
    ("2.3  Multi-FPC Bundle Management and EEG Cable Separation",
     True, (0x1F, 0x49, 0x7D)),
]

for (text, bold, color) in S23_CONTENT:
    ins_before(shell, s3_para, text, bold=bold, color=color)

print("  §2.3 content inserted")

# ── Add DRC-16 to DRC-21 to the design review checklist table ──
# Find the DRC table (contains "DRC-15" as the last row)
drc_tbl = None
for tbl in shell.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "DRC-15":
            drc_tbl = tbl
            break
    if drc_tbl:
        break

print(f"  DRC table found: {drc_tbl is not None}")

if drc_tbl:
    NEW_DRC_ROWS = [
        ("DRC-16",
         "Bundle cross-section: measure total bundle width and height at the narrowest "
         "common-routing area in CAD cross-section",
         "CAD cross-section",
         "Bundle fits in available shell interior space with ≥ 2 mm clearance on all sides",
         "", ""),
        ("DRC-17",
         "Inter-FPC separation: confirm ≥ 2 mm gap between adjacent FPC surfaces "
         "(ZM-01/02, ZM-02/03, ZM-03/04, ZM-04/05) along all of Segment B for all pairs",
         "CAD measurement, all adjacent pairs",
         "≥ 2 mm separation at every point. Any pair < 1 mm → redesign channel dividers",
         "", ""),
        ("DRC-18",
         "FPC-to-EEG cable separation: overlay EEG cable routing on FPC routing in CAD; "
         "measure minimum separation at all crossing or parallel points",
         "CAD overlay measurement",
         "≥ 15 mm separation at all points. If < 15 mm anywhere: confirm barrier "
         "design in same CAD view",
         "", ""),
        ("DRC-19",
         "Anchor/tie-down points: confirm all 3 anchor points present for each of "
         "5 FPC routing paths (15 total); verify anchor boss internal radius ≥ 5 mm "
         "and no FPC compression at anchor",
         "CAD inspection, all 5 zones",
         "15 anchor points present; boss radius ≥ 5 mm; no channel narrowing at anchor",
         "", ""),
        ("DRC-20",
         "Hub PCB connector placement: confirm all 5 ZIF receptacles on same Hub PCB "
         "edge; connector row orientation verified perpendicular to dominant FPC routing "
         "direction; no FPC needs to reverse direction to reach connector",
         "Hub PCB layout + shell CAD overlay",
         "All 5 connectors on same edge; no FPC routing reversal required",
         "", ""),
        ("DRC-21",
         "Physical first-article: with all 5 zone modules installed and FPC tails routed, "
         "confirm no FPC-to-FPC contact and no FPC-to-EEG cable contact during "
         "full range of headset head-movement simulation (±30° flex, 10 cycles)",
         "Physical prototype or first-article production unit",
         "No contact at any point during motion simulation; all FPCs remain in channels",
         "", ""),
    ]
    add_drc_rows(drc_tbl, NEW_DRC_ROWS)
    print(f"  DRC-16 to DRC-21 added ({len(NEW_DRC_ROWS)} rows)")

# Update §7 open items — add OI-08 (Hub PCB connector placement)
# Find OI-07 row in the open items table and add OI-08 after it
open_items_tbl = None
for tbl in shell.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "OI-07":
            open_items_tbl = tbl
            break
    if open_items_tbl:
        break

if open_items_tbl:
    new_row = open_items_tbl.add_row()
    oi_data = ["OI-08",
               "Hub PCB connector placement confirmed for all 5 zones; layout constraint "
               "communicated to Hub PCB layout engineer",
               "HW EE / Mech Eng",
               "Before Hub PCB schematic capture"]
    oi_widths = [0.6, 2.8, 1.5, 1.35]
    for c_idx, txt in enumerate(oi_data):
        cell = new_row.cells[c_idx]
        set_bg(cell, "F5F5F5")
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        if c_idx == 0:
            run.bold = True
        if c_idx < len(oi_widths):
            cell.width = Inches(oi_widths[c_idx])
    print("  OI-08 added to NP-DRV-SHELL-001 open items")

shell.save(SHELL_PATH)
print(f"  NP-DRV-SHELL-001 saved")

# ─────────────────────────────────────────────────────────────────────────────
# 2.  Risk register — RISK-17 → MITIGATED
# ─────────────────────────────────────────────────────────────────────────────

RISK_PATH = "docs/neuropulse_fpc_zone_module_risks_revA.docx"
risk_doc = Document(RISK_PATH)

for row in risk_doc.tables[0].rows:
    if row.cells[0].text.strip() == "RISK-17":
        set_text(row.cells[1], "MEDIUM\n(MITIGATED)", bold=True,
                 color=(0x00, 0x70, 0x00))
        set_bg(row.cells[1], "E2EFDA")
        set_text(row.cells[5],
            "MITIGATED: NP-DRV-SHELL-001 Rev A extended with §2.3 "
            "(multi-FPC bundle management and EEG cable separation). "
            "New requirements: ≥ 2 mm inter-FPC separation; ≥ 15 mm FPC-to-EEG cable "
            "separation (or grounded barrier); 3 anchor/tie-down points per FPC; "
            "Hub PCB all 5 ZIF receptacles on same edge. "
            "DRC-16 to DRC-21 added to design review checklist. "
            "NP-COORD-001 G2-01 updated with specific deliverables.")
        set_text(row.cells[8],
            "MITIGATED — NP-DRV-SHELL-001 §2.3 + DRC-16–21; "
            "CAD and physical prototype sign-off required",
            bold=True, color=(0x00, 0x70, 0x00))
        set_bg(row.cells[8], "92D050")
        print("  RISK-17 → MITIGATED")
        break

# Detailed section before OPEN ISSUES
open_issues_idx = None
for i, para in enumerate(risk_doc.paragraphs):
    if "OPEN ISSUES" in para.text:
        open_issues_idx = i
        break

target = risk_doc.paragraphs[open_issues_idx]

ins_before(risk_doc, target, "")

ins_before(risk_doc, target,
    "→ Action 5 (OPEN — physical prototype): DRC-21 — motion simulation test on "
    "first-article with all 5 zone modules installed; confirm no FPC contact during "
    "±30° flex cycles. Tracked in NP-DRV-SHELL-001 §5.",
    bold=False, color=(0xC0, 0x60, 0x00), size=10)

ins_before(risk_doc, target,
    "→ Action 4 (OPEN — Hub PCB layout): Hub PCB connector placement coordination — "
    "all 5 ZIF receptacles on same edge, perpendicular to FPC routing direction. "
    "Must be communicated to Hub PCB layout engineer before schematic capture. "
    "Tracked as NP-COORD-001 G2-01 and NP-DRV-SHELL-001 OI-08.",
    bold=False, color=(0xC0, 0x60, 0x00), size=10)

ins_before(risk_doc, target,
    "→ Action 3 (OPEN — shell CAD): DRC-17/18/19/20 — inter-FPC separation, "
    "FPC-to-EEG cable separation, anchor points, and Hub PCB alignment all require "
    "a CAD model showing all 5 FPC routing paths simultaneously with EEG cable overlay. "
    "This is the primary CAD deliverable for NP-COORD-001 G2-01.",
    bold=False, color=(0xC0, 0x60, 0x00), size=10)

ins_before(risk_doc, target,
    "→ Action 2 (OPEN — shell tooling spec): Anchor boss geometry (3 per FPC, "
    "≥ 5 mm boss radius, no compression), inter-FPC channel dividers, and grounded "
    "barrier (if FPC-to-EEG gap < 15 mm) must all be in shell tooling spec before "
    "first-cut. Cannot be retrofitted.",
    bold=False, color=(0xC0, 0x60, 0x00), size=10)

ins_before(risk_doc, target,
    "→ Action 1 (DONE): NP-DRV-SHELL-001 Rev A extended with §2.3 (multi-FPC bundle "
    "management) and DRC-16 to DRC-21 (6 new design review checklist items). "
    "Requirements now specified for: bundle cross-section clearance, ≥ 2 mm inter-FPC "
    "separation, ≥ 15 mm FPC-to-EEG separation, anchor/tie-down points, "
    "and Hub PCB connector placement.",
    bold=False, color=(0x00, 0x70, 0x00), size=10)

ins_before(risk_doc, target,
    "RISK-17 extends RISK-11 (single-FPC bend radius) to the full 5-FPC system. "
    "The critical additional hazards are: (a) inductive coupling between adjacent "
    "high-current FPC tails — at 1.8 A RMS per tail, mutual inductance between "
    "parallel FPCs < 2 mm apart produces millivolt-level interference on any "
    "co-located signal line, including EEG leads; (b) EEG cable contamination — "
    "a sub-microvolt EEG signal running within 5 mm of a 1 kHz PWM power FPC "
    "will pick up > 100 µV of artifact, completely masking delta/theta band activity; "
    "(c) FPC chafing — adjacent FPCs in contact undergo relative motion of "
    "~0.5–1.5 mm on every zone module insertion/removal cycle, abrading the "
    "polyimide coverlay and eventually exposing copper.",
    bold=False, color=None, size=10)

ins_before(risk_doc, target,
    "RISK-17  |  MEDIUM (MITIGATED)  |  Multi-FPC Routing + EEG Cable Separation — §2.3 Added to NP-DRV-SHELL-001",
    bold=True, color=(0x1F, 0x49, 0x7D), size=10)

risk_doc.save(RISK_PATH)
print(f"  Risk register saved")

# ─────────────────────────────────────────────────────────────────────────────
# 3.  NP-COORD-001 — update G2-01 deliverables
# ─────────────────────────────────────────────────────────────────────────────

COORD_PATH = "docs/neuropulse_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

for tbl in coord.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) == 7 and cells[0].text.strip() == "G2-01":
            set_text(cells[3],
                "Shell CAD cross-section: all 5 FPC routing paths + EEG cable overlay "
                "simultaneously; DRC-16 to DRC-21 all verified. "
                "Specific deliverables: "
                "(a) ≥ 2 mm inter-FPC gap at all adjacent-pair cross-sections; "
                "(b) ≥ 15 mm FPC-to-EEG separation (or barrier design) at all points; "
                "(c) 15 anchor boss locations (3 per FPC) with ≥ 5 mm boss radius; "
                "(d) Hub PCB connector placement drawing showing all 5 ZIF on same edge; "
                "(e) physical first-article motion simulation test (DRC-21).",
                size=8)
            print("  G2-01 deliverables updated")
            break

coord.save(COORD_PATH)
print(f"  Coord checklist saved")
print("\nAll patches complete.")
