"""
Add mastoid LRA pad and Apple Watch sync app to additional modalities doc.
Update 40Hz vibrotactile entry to reference specific mastoid pad spec.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_section(doc, number, title, source, fields, header_bg="1F3864", status_bg=None):
    """Add a modality section: header table + body paragraphs."""
    t = doc.add_table(rows=1, cols=1)
    set_bg(t.rows[0].cells[0], header_bg)
    p = t.rows[0].cells[0].paragraphs[0]
    r = p.add_run(f"  {number}. {title}")
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if status_bg:
        set_bg(t.rows[0].cells[0], status_bg)
        p.runs[0].font.color.rgb = RGBColor(0x00, 0x40, 0x00)

    doc.add_paragraph()

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(f"Source context: {source}")
    r2.italic = True; r2.font.size = Pt(8)

    for label, text in fields:
        p3 = doc.add_paragraph()
        p3.paragraph_format.space_after = Pt(3)
        r3 = p3.add_run(f"{label}: ")
        r3.bold = True; r3.font.size = Pt(9)
        r4 = p3.add_run(text)
        r4.font.size = Pt(9)

    doc.add_paragraph()
    # Divider
    p_div = doc.add_paragraph()
    pPr = p_div._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom); pPr.append(pBdr)
    p_div.paragraph_format.space_after = Pt(6)

MOD_PATH = "docs/neurone_additional_modalities.docx"
mod = Document(MOD_PATH)

# ── Update entry 1 (40Hz vibrotactile) to reference mastoid pad spec ──────────
# Find paragraphs referencing the vibrotactile entry and append mastoid spec note
for para in mod.paragraphs:
    if "Priority" in para.text and "HOPE" in para.text and "40Hz" in para.text:
        for run in para.runs:
            if "HIGH" in run.text:
                run.text = (
                    "Priority: HIGH — await HOPE Phase 3 results (mid-2026). "
                    "Mastoid LRA pad provisionally spec'd in CLAUDE.md §3b: "
                    "LRA 8–10mm + TI DRV2605L driver, 40Hz ± 0.5Hz, mastoid placement, "
                    "$4–7 BOM/pad, $49–79 retail. Anchor boss in temporal wing tooling "
                    "at first cut (zero incremental tooling cost). Apple Watch haptic sync "
                    "is companion UX only — not a replacement for purpose-built hardware "
                    "(see marketing notes: frequency accuracy, bone coupling, regulatory)."
                )
                print("  Updated 40Hz vibrotactile priority/action note")
                break
        break

# ── Add mastoid LRA pad section ───────────────────────────────────────────────
add_section(
    mod,
    number="7",
    title="40Hz Vibrotactile — Mastoid LRA Pad (Provisional Spec)",
    source="Tsai lab MIT 2023 (40Hz vibration in mouse models); GENUS multi-sensory "
           "protocol extension; NeurOne §3b provisional spec",
    fields=[
        ("Status",
         "PROVISIONALLY SPEC'D — release contingent on HOPE Phase 3 results (mid-2026). "
         "Hardware anchor boss in temporal wing tooling at first cut (zero incremental cost)."),
        ("Evidence",
         "PRECLINICAL: Tsai lab (MIT, 2023) — 40Hz whole-body vibrotactile stimulation "
         "reduces amyloid + tau in Alzheimer's mouse models. Multi-sensory GENUS "
         "(visual + audio + tactile) more effective than any two-sense combination in "
         "preclinical. HOPE Phase 3 (n=670, results mid-2026) tests visual + audio; "
         "tactile is the next generation. No published human RCT at 40Hz tactile yet."),
        ("Actuator spec",
         "Linear resonant actuator (LRA), 8–10mm diameter (e.g. Jinlong JMC0834 or equiv). "
         "LRA preferred over ERM: precise frequency control, low harmonic distortion, "
         "flat resonance profile at 40Hz drive. Driver IC: Texas Instruments DRV2605L "
         "(I2C, open-loop mode); 40Hz ± 0.5Hz; amplitude control via gain register."),
        ("Placement rationale",
         "Mastoid process (posterior temporal / behind ear). Rationale: (1) direct bone "
         "coupling to skull — superior transmission vs wrist; (2) adjacent to temporal "
         "lobe somatosensory representations; (3) compatible with existing temporal "
         "stability wing anchor — no new shell tooling if anchor boss provisioned at "
         "first cut. Why NOT Apple Watch: Taptic Engine frequency uncharacterised for "
         "therapeutic use; watchOS API lacks raw 40Hz continuous drive; wrist has poor "
         "bone coupling; App Store prohibits disease claims. Full rationale in CLAUDE.md §15."),
        ("Output spec",
         "0.6–1.2G peak acceleration at skin surface (DRV2605L gain register configurable); "
         "40Hz ± 0.5Hz; 100% duty cycle continuous session (20 min target); 2s amplitude "
         "ramp up/down (comfort)."),
        ("Form factor",
         "30mm diameter silicone overmoulded pad; clip attachment to temporal wing anchor; "
         "3-pin pogo or JST connector to hub accessory port; Shore 30A silicone contact face."),
        ("BOM",
         "LRA $1.50–2.50 + DRV2605L $1.20–1.80 + PCB/passives $0.50 + "
         "silicone housing $0.80–1.20 = $4–7/pad. Bilateral pair: $8–14."),
        ("Retail price (projected)", "$49–79 per pad / $79–119 bilateral pair."),
        ("Power", "~80–120mW continuous from hub accessory port (existing 500mA capability)."),
        ("Firmware",
         "40Hz square-wave drive via DRV2605L open-loop; session synchronised with "
         "audio/visual 40Hz channels via hub firmware clock; amplitude ramp firmware-controlled."),
        ("Recommended action",
         "Provision temporal wing anchor boss in first-cut tooling (zero cost). "
         "Commission DRV2605L + LRA characterisation bench test. "
         "Finalise pad design post-HOPE results. Engage Tsai SAB for protocol alignment."),
    ],
    header_bg="FFF2CC",
)

# ── Add Apple Watch sync app section ──────────────────────────────────────────
add_section(
    mod,
    number="8",
    title="Apple Watch Companion Sync App (Provisional)",
    source="NeurOne §3b provisional spec; companion UX — not therapeutic delivery",
    fields=[
        ("Status",
         "PROVISIONALLY SPEC'D — software only, $0 BOM. Develop after core iOS app ships. "
         "Haptic + audio channels first; visual flicker second."),
        ("Purpose",
         "Extends NeurOne session experience to Apple Watch via three sync channels. "
         "Does NOT replace purpose-built NeurOne hardware for any therapeutic function. "
         "All therapeutic claims attach to NeurOne hardware only; Watch app is "
         "session monitoring and interface aid."),
        ("Communication",
         "BT 5.3 LE from NeurOne hub (hub already has BT radio, antennas in hub); "
         "WatchConnectivity framework via paired iPhone app; "
         "session sync over BLE GATT custom service. Sub-50ms sync latency target."),
        ("Channel 1 — Haptic sync",
         "watchOS Core Haptics delivers 40Hz pattern synchronised to hub session clock. "
         "Adds wrist somatosensory channel alongside mastoid LRA pad. "
         "Not standalone therapeutic — complement only. "
         "App displays: 'For best results, use with NeurOne vibrotactile accessory.' "
         "Note: Apple Watch haptic precision is uncharacterised for therapeutic use; "
         "this channel adds incremental sensory input, not clinical-grade stimulation."),
        ("Channel 2 — Audio sync",
         "Watch app plays binaural beats / isochronic tones / HRV breathing pacer audio "
         "through AirPods or earphones paired to Watch, synchronised to hub session. "
         "Useful when: (a) bone conduction reserved for breathing cue while earphones "
         "handle binaural beats; (b) user is away from hub speaker range; "
         "(c) user prefers AirPods over headset audio for a session."),
        ("Channel 3 — Visual sync",
         "Watch display shows: 40Hz visual flicker (reduced brightness, GENUS-compatible — "
         "brightness characterisation needed to confirm ≥100 nits at 40Hz); "
         "or EMDR left/right indicator arrow synchronised to goggle session; "
         "or session status + protocol name; or HRV biofeedback breathing ring "
         "(complementary to phone app display for wrist-glance UX)."),
        ("Additional Watch functions",
         "Session timer + haptic end-of-session alert; protocol selector for basic "
         "session start without phone; quick impedance check result notification; "
         "consumable low reminders; coherence score live feed during HRV biofeedback."),
        ("Regulatory note",
         "All Watch-delivered functions declared as session monitoring / user interface aids. "
         "No disease claims in Watch app. Therapeutic claims attach to NeurOne hardware. "
         "App Store General Wellness category. watchOS Human Interface Guidelines compliance."),
        ("BOM delta", "$0 hardware. Software development cost only."),
        ("Recommended action",
         "Add watchOS app to iOS app development roadmap (Year 1 post-launch). "
         "Prioritise: (1) session status + HRV breathing ring; (2) audio sync; "
         "(3) haptic sync; (4) 40Hz visual flicker (requires screen characterisation first)."),
    ],
    header_bg="E8F0FE",
)

mod.save(MOD_PATH)
print("Additional modalities doc saved with mastoid LRA pad (§7) and Apple Watch app (§8)")
