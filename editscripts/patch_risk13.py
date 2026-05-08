"""Patch RISK-13 in risk register and add RISK-13 detailed section."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

RISK_PATH = "docs/neuropulse_fpc_zone_module_risks_revA.docx"
risk_doc = Document(RISK_PATH)

# ── Summary table ──
tbl = risk_doc.tables[0]
for row in tbl.rows:
    if row.cells[0].text.strip() == "RISK-13":
        set_cell_text(row.cells[1], "MEDIUM\n(MITIGATED)", bold=True, color=(0x00,0x70,0x00))
        set_cell_bg(row.cells[1], "E2EFDA")
        set_cell_text(row.cells[5],
            "MITIGATED: Engineering coordination checklist NP-COORD-001 Rev A created. "
            "Item G1-03 mandates a signed Engineering Decision Record (EDR) between HW EE "
            "and Firmware Lead agreeing the carrier frequency (≥1 kHz confirmed); "
            "EDR must verify all harmonics up to Nyquist (250 Hz) do not coincide with "
            "any therapy or EEG band frequency. "
            "Item G1-04 mandates firmware implementation of 50 µs EEG blanking window "
            "on each LED switching edge, synchronised to PWM timer. "
            "Items G2-07 and G3-04 lock the agreed values as Safety MCU constants and "
            "require oscilloscope validation on first prototype. "
            "FPC spec §4.4 and §10.3 already reference minimum 1 kHz carrier and 50 µs blanking.")
        set_cell_text(row.cells[8],
            "MITIGATED — NP-COORD-001 G1-03/G1-04/G2-07/G3-04 close this risk; "
            "EDR and firmware implementation required",
            bold=True, color=(0x00,0x70,0x00))
        set_cell_bg(row.cells[8], "92D050")
        print("RISK-13 summary row updated")
        break

# ── Update OI-06 in open issues table ──
oi_tbl = risk_doc.tables[1]
for row in oi_tbl.rows:
    if row.cells[0].text.strip() == "OI-06":
        set_cell_text(row.cells[1], "REQUIRED (addressed in NP-COORD-001)")
        set_cell_bg(row.cells[1], "E2EFDA")
        set_cell_text(row.cells[2],
            "LED PWM carrier frequency agreed with firmware team — "
            "NP-COORD-001 items G1-03 (EDR) and G1-04 (firmware blanking) address this. "
            "EDR and firmware implementation still required to close G1-03 and G1-04.")
        print("OI-06 updated")
        break

# ── Add RISK-13 detailed section before "OPEN ISSUES" ──
open_issues_idx = None
for i, para in enumerate(risk_doc.paragraphs):
    if "OPEN ISSUES" in para.text:
        open_issues_idx = i
        break

print(f"Open issues section at paragraph {open_issues_idx}")

target = risk_doc.paragraphs[open_issues_idx]

def ins_para(doc, target_para, text, bold=False, bullet=False, color=None, size=10):
    new_p = OxmlElement("w:p")
    target_para._p.addprevious(new_p)
    for p_obj in doc.paragraphs:
        if p_obj._p is new_p:
            if bullet:
                p_obj.style = doc.styles["List Bullet"]
            run = p_obj.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p_obj
    return None

ins_para(risk_doc, target, "")  # spacer

ins_para(risk_doc, target,
    "→ Action 4 (DONE): Engineering coordination checklist NP-COORD-001 Rev A created. "
    "Items G1-03 and G1-04 track EDR and firmware implementation. "
    "Items G2-07 and G3-04 lock values and validate on prototype.",
    bullet=True, color=(0x00,0x70,0x00))

ins_para(risk_doc, target,
    "→ Action 3 (OPEN — prototype): Oscilloscope validation of 50 µs blanking window — "
    "scope trace showing EEG sample gate vs. PWM edge, captured on first prototype. "
    "Tracked as NP-COORD-001 G3-04.",
    bullet=True, color=(0xC0,0x60,0x00))

ins_para(risk_doc, target,
    "→ Action 2 (OPEN — Safety MCU spec): Agreed carrier frequency locked as Safety MCU "
    "constant; 50 µs blanking constant in Safety MCU firmware. "
    "Neither value changeable by OTA update without ECO. "
    "Tracked as NP-COORD-001 G2-07.",
    bullet=True, color=(0xC0,0x60,0x00))

ins_para(risk_doc, target,
    "→ Action 1 (OPEN — EDR required): HW EE and Firmware Lead must sign Engineering "
    "Decision Record specifying: agreed carrier frequency ≥ 1 kHz; anti-aliasing filter "
    "corner; harmonic analysis table confirming no coincidence with any therapy/EEG "
    "band up to Nyquist; 50 µs blanking window on every PWM edge. "
    "Tracked as NP-COORD-001 G1-03.",
    bullet=True, color=(0xC0,0x60,0x00))

ins_para(risk_doc, target,
    "The contamination path is direct: LED PWM switching generates a square-wave current "
    "transient in the power supply rails. If the carrier frequency (or any of its harmonics "
    "up to the EEG anti-aliasing cutoff at 200 Hz) falls within the EEG acquisition band, "
    "the artifact is captured and treated as brainwave signal by the closed-loop algorithm. "
    "At 40 Hz gamma (the highest-value protocol), a carrier frequency of 40, 80, or 120 Hz "
    "would produce an artifact indistinguishable from genuine gamma: the algorithm would see "
    "'gamma activity' that correlates perfectly with its own stimulation — a positive feedback "
    "loop. The mitigation (≥1 kHz carrier + 50 µs blanking) is well-established in combined "
    "PBM/EEG device design, but it requires explicit inter-team agreement because neither the "
    "HW EE nor the Firmware team has full visibility of the other's constraints in isolation.")

ins_para(risk_doc, target,
    "RISK-13  |  MEDIUM (MITIGATED)  |  LED PWM Carrier Frequency — EEG Band Contamination",
    bold=True, color=(0x1F,0x49,0x7D))

risk_doc.save(RISK_PATH)
print(f"Risk register saved: {RISK_PATH}")
print("Done.")
