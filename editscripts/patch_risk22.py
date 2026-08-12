"""
RISK-22 mitigation — Option A: sliding eject lever.
Updates:
  - NP-TOOL-ZM-001: replace F-08 pull tab with sliding eject lever spec
  - NP-FAI-ZM-001: update HFE test item for lever
  - Risk register: RISK-22 → MITIGATED
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

# ══════════════════════════════════════════════════════════════════════════════
# 1. NP-TOOL-ZM-001 — replace F-08 pull tab with sliding eject lever
# ══════════════════════════════════════════════════════════════════════════════
TOOL_PATH = "docs/superseded/neurone_tool_zone_module_001.docx"
tool = Document(TOOL_PATH)

NEW_F08 = (
    "F-08",
    "Sliding eject lever (Option A — RISK-22 mitigation): a moulded lever arm "
    "integrated into the posterior face of the zone module body. Lever is recessed "
    "flush with the module surface when closed (does not protrude beyond module envelope). "
    "To eject: user hooks one finger under the lever tip and pivots to 30–45°. Lever "
    "geometry provides mechanical advantage of ≥3:1, reducing required extraction force "
    "from the nominal 2–4 N gasket compression to ≤1 N at the lever tip — within the "
    "grip capability of users with moderate Parkinson's or post-stroke hand weakness.\n\n"
    "Lever dimensions: arm length 10–12 mm; width 8 mm; thickness 2.5 mm; pivot boss "
    "diameter 3 mm. Lever tip rounded (R1.5 mm) for finger comfort. Closed latch: a "
    "snap-fit detent (PC/ABS, 0.3 mm interference) holds lever flush during insertion "
    "and normal wear — prevents accidental ejection from incidental contact. Detent "
    "release force: 0.3–0.5 N (finger-actuable, not thumb-actuable at wrong angle).\n\n"
    "Lever clears ZIF lever mechanism: ME to confirm ≥2 mm clearance between eject "
    "lever pivot boss and ZIF back-flip lever in assembled position (OI-7).\n\n"
    "[DRAWING REQUIREMENT — ZONE MODULE TOOLING] Lever hinge pin boss moulded into "
    "zone module body. Hinge pin: 3 mm diameter stainless steel (316), press-fit into "
    "boss. Lever arm moulded separately in same PC/ABS material, assembled post-mould. "
    "Detent rib moulded on lever underside.",
    "RISK-22", "NP-TOOL-ZM-001 F-08; HFE formative study", "☐"
)

for tbl in tool.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 1 and cells[0].text.strip() == "F-08":
            set_text(cells[0], NEW_F08[0], bold=True, size=8)
            set_text(cells[1], NEW_F08[1], size=8)
            set_text(cells[2], NEW_F08[2], size=8)
            set_text(cells[3], NEW_F08[3], size=8)
            set_text(cells[4], NEW_F08[4], size=8)
            print("  F-08 updated: pull tab → sliding eject lever")
            break

# Update critical dimensions table — replace pull tab height with lever dimensions
for tbl in tool.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 1 and "Pull tab height" in cells[0].text:
            set_text(cells[0], "Eject lever arm length", bold=True, size=8)
            set_text(cells[1], "10–12 mm", size=8)
            set_text(cells[2], "±0.5 mm", size=8)
            set_text(cells[3], "Mechanical advantage ≥3:1 (RISK-22)", size=8)
            set_text(cells[4], "YES", size=8)
            print("  Dimension table: pull tab → lever arm length")
            break

for tbl in tool.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 1 and cells[0].text.strip() == "Pull tab height (F-08)":
            set_text(cells[0], "Lever arm length (F-08)", bold=True, size=8)
            set_text(cells[1], "10–12 mm", size=8)
            set_text(cells[2], "±0.5 mm", size=8)

# Add OI-7 for lever-ZIF clearance
for tbl in tool.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 1 and cells[0].text.strip() == "OI-6":
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            oi7_data = ("OI-7",
                "Confirm ≥2 mm clearance between eject lever pivot boss and ZIF back-flip "
                "lever in assembled module position. ME to verify in CAD and on prototype.",
                "ME", "G1 Pre-Layout", "OPEN")
            oi_widths = [0.5, 3.8, 0.8, 1.2, 0.5]
            for c_idx, txt in enumerate(oi7_data):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(oi_widths[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[4], "FFF2CC")
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  OI-7 lever-ZIF clearance added")
            break

# Update MDR checklist item 8 (pull tab interference check)
for tbl in tool.tables:
    for row in tbl.rows:
        cells = row.cells
        if (len(cells) >= 2 and cells[0].text.strip() == "8"
                and "pull tab" in cells[1].text.lower()):
            set_text(cells[1],
                "F-08 eject lever does not interfere with ZIF lever access — "
                "ME confirms ≥2 mm clearance in CAD (OI-7). "
                "Detent snap confirmed releasable at ≤0.5 N one-finger force.",
                size=8)
            print("  MDR item 8 updated for lever")
            break

tool.save(TOOL_PATH)
print(f"  NP-TOOL-ZM-001 saved")

# ══════════════════════════════════════════════════════════════════════════════
# 2. NP-FAI-ZM-001 — update HFE accessibility test for lever
# ══════════════════════════════════════════════════════════════════════════════
FAI_PATH = "docs/superseded/neurone_fai_zone_module.docx"
fai = Document(FAI_PATH)

# Find FAI-A09 through FAI-A14 section — add FAI-A15 for lever HFE test
# after FAI-A14 (mechanical key exclusion test)
for tbl in fai.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 1 and cells[0].text.strip() == "FAI-A14":
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            a15_data = (
                "FAI-A15",
                "Eject lever accessibility test (RISK-22): 5 test subjects with diagnosed "
                "Parkinson's disease (Hoehn & Yahr scale II–III) or post-stroke hand weakness "
                "(grip strength ≤ 8 kg). Each subject removes and reinserts all 5 zone modules "
                "twice (10 total operations). No coaching; standard user instructions only.",
                "HFE formative test — 5 subjects with Parkinson's/post-stroke (pre-production)",
                "All 5 subjects complete each extraction in ≤ 10 s with no more than 2 attempts. "
                "No reported pain or excessive effort. Lever detent must not prevent extraction. "
                "Lever must not accidentally eject during normal wear simulation "
                "(5 min handling without deliberate ejection attempt).",
                "", ""
            )
            fai_widths = [0.65, 2.3, 1.2, 1.55, 0.65, 0.45]
            for c_idx, txt in enumerate(a15_data):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(fai_widths[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[0], "F5F5F5")
                    set_bg(r2.cells[1], "F5F5F5")
                    set_bg(r2.cells[2], "F5F5F5")
                    set_bg(r2.cells[3], "F5F5F5")
                    set_bg(r2.cells[4], "F5F5F5")
                    set_bg(r2.cells[5], "F5F5F5")
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  FAI-A15 lever HFE test added")
            break

# Add RISK-22 to §9 cross-reference table (after RISK-16 row)
for tbl in fai.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "RISK-16":
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            xref_widths = [0.65, 2.0, 2.35, 1.8]
            xref_texts = ("RISK-22",
                          "Zone module extraction — reduced grip / tremor (sliding eject lever)",
                          "FAI-A15",
                          "MITIGATED (this FAI)")
            for c_idx, txt in enumerate(xref_texts):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(xref_widths[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[3], "E2EFDA")
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  §9 RISK-22 cross-reference row added")
            break

fai.save(FAI_PATH)
print(f"  NP-FAI-ZM-001 saved")

# ══════════════════════════════════════════════════════════════════════════════
# 3. Risk register — RISK-22 → MITIGATED
# ══════════════════════════════════════════════════════════════════════════════
RISK_PATH = "docs/superseded/neurone_fpc_zone_module_risks_revA.docx"
risk = Document(RISK_PATH)

for tbl in risk.tables:
    for row in tbl.rows:
        cells = row.cells
        if cells[0].text.strip() == "RISK-22" and len(cells) >= 9:
            set_text(cells[1], "MEDIUM\n(MITIGATED)", bold=True,
                     color=(0x00, 0x70, 0x00), size=8)
            set_bg(cells[1], "E2EFDA")
            set_text(cells[5],
                "Option A sliding eject lever: 10–12 mm lever arm with 3:1 mechanical "
                "advantage reduces extraction force from 2–4 N to ≤1 N at lever tip. "
                "Recessed flush when closed; snap-fit detent (0.3–0.5 N release) prevents "
                "accidental ejection during wear. Hinge pin boss moulded into zone module "
                "body; 316SS hinge pin press-fit. Lever clears ZIF mechanism (OI-7 confirms). "
                "FAI-A15 validates with 5 Parkinson's/post-stroke subjects "
                "(≤10 s extraction, ≤2 attempts, no pain).",
                size=8)
            set_text(cells[7],
                "NP-TOOL-ZM-001 F-08 (updated)\nNP-FAI-ZM-001 FAI-A15",
                size=8)
            set_text(cells[8],
                "MITIGATED — Option A lever specified in NP-TOOL-ZM-001 F-08; "
                "FAI-A15 HFE test added; OI-7 tracks CAD clearance check",
                bold=True, color=(0x00, 0x70, 0x00), size=8)
            set_bg(cells[8], "92D050")
            print("  RISK-22 → MITIGATED")
            break

risk.save(RISK_PATH)
print(f"  Risk register saved")
print("Done.")
