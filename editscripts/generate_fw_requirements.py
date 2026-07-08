import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    run = p.runs[0]
    run.bold = True
    if level == 1:
        run.font.size = Pt(13)
    else:
        run.font.size = Pt(11)
    return p

def add_normal(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    return row

# Title
title = doc.add_paragraph('NeurOne')
title.style = doc.styles['Normal']
title.runs[0].bold = True
title.runs[0].font.size = Pt(14)

subtitle = doc.add_paragraph('Zone Module Firmware Requirements')
subtitle.style = doc.styles['Normal']
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(13)

doc.add_paragraph('NP-FW-REQ-001  Rev A  |  2026-05-10  |  Pre-Layout / Pre-Tooling Draft').style = doc.styles['Normal']
doc.add_paragraph('References: NP-HW-FPC-001 Rev D · NP-COORD-001 Rev A · NP-FAI-ZM-001 Rev A · RISK-18').style = doc.styles['Normal']
doc.add_paragraph('Addresses G1-13 in NP-COORD-001: ZONE_ID firmware debounce specification for zone module detection.').style = doc.styles['Normal']

# Section 1
add_heading(doc, '1.  PURPOSE AND SCOPE')
add_normal(doc, 'This document specifies firmware requirements for zone module detection on the NeurOne headset Hub MCU. It covers the ZONE_ID resistor read algorithm (pin 18), ADC resolution requirements, debounce logic, and FAULT state behaviour. All requirements in this document are Class B (IEC 62304) for the main processor path. Any requirement that gates the ENABLE line is additionally subject to Safety MCU review (Class C).')
add_normal(doc, 'This document was created to close NP-COORD-001 G1-13. It must be baselined before any firmware for zone module detection is written.')

# Section 2
add_heading(doc, '2.  REFERENCES')
refs = [
    ('NP-HW-FPC-001 Rev D', '20-pin zone module FPC pinout; §11.4 five-layer keying; §8.4 dual-photodiode; Pin 18 = ZONE_ID; Pin 19 = PD2_CATHODE'),
    ('NP-COORD-001 Rev A', 'Engineering coordination checklist; G1-13 (this document closes it); G2-10 (audio zone ID firmware)'),
    ('NP-FAI-ZM-001 Rev A', 'FAI checklist; FAI-K04 (pin 18 contact resistance ≤30 mΩ); FAI-A15 (accessibility HFE test)'),
    ('RISK-18', 'Zone module risk register — ZONE_ID ENABLE interlock availability failure; three-layer mitigation (firmware debounce + FAI-K04 + app UX)'),
    ('CLAUDE.md §13.5', 'Locked decisions: ZONE_ID firmware debounce 3×ADC at 100ms, ≥2/3 pass; FPC 20-pin pinout locked'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Document'
t.rows[0].cells[1].text = 'Relevant content'
for ref, desc in refs:
    row = t.add_row()
    row.cells[0].text = ref
    row.cells[1].text = desc

# Section 3
add_heading(doc, '3.  ZONE_ID DETECTION — FIRMWARE REQUIREMENTS')
add_heading(doc, '3.1  Overview', level=2)
add_normal(doc, 'Each zone module contains a unique 1% 0402 SMD resistor between ZONE_ID (pin 18) and AGND (pin 17). The Hub MCU reads this resistor via ADC through a 10 kΩ pull-up to VDD (3.3 V) to identify which zone module is installed in each slot. Zone identification must occur before every session start and before ENABLE is asserted for the affected zone. ENABLE is held de-asserted (LOW) until ZONE_ID is validated.')
add_normal(doc, 'The Safety MCU owns all stimulation GPIO ENABLE lines (per NP-HW-FPC-001 §5.2 and NP-COORD-001 G1-05). The main processor reports ZONE_ID validation result to the Safety MCU over the SPI heartbeat link before ENABLE may be asserted. A FAULT from the ZONE_ID check causes the Safety MCU to hold ENABLE low regardless of any main processor request.')

add_heading(doc, '3.2  Zone ID Resistor Map (Locked)', level=2)
add_normal(doc, 'The following resistor values are locked. All resistors are 1% tolerance, 0402 package, SMD, on pin 18 of the 20-pin zone module FPC connector.')
t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
t2.rows[0].cells[0].text = 'Zone module'
t2.rows[0].cells[1].text = 'ZONE_ID resistor (pin 18)'
t2.rows[0].cells[2].text = 'Zone name'
zone_data = [
    ('ZM-01', '10 kΩ', 'Frontal Left'),
    ('ZM-02', '22 kΩ', 'Frontal Right'),
    ('ZM-03', '47 kΩ', 'Central Left'),
    ('ZM-04', '100 kΩ', 'Central Right'),
    ('ZM-05', '220 kΩ', 'Parietal / Occipital'),
]
for zm, res, name in zone_data:
    row = t2.add_row()
    row.cells[0].text = zm
    row.cells[1].text = res
    row.cells[2].text = name

add_heading(doc, '3.3  ADC Requirements', level=2)
add_normal(doc, 'FW-ADC-01 (MANDATORY): The ZONE_ID ADC channel shall use a minimum 10-bit resolution. The i.MX RT1062 and STM32G071 both provide 12-bit ADC — use 12-bit mode. 10-bit is the minimum acceptable floor; do not configure less.')
add_normal(doc, 'FW-ADC-02 (MANDATORY): The ZONE_ID pin 18 pull-up resistor is 10 kΩ ±1% to VDD (3.3 V). ADC full-scale reference = VDD. The firmware ADC driver shall not apply any additional scaling that reduces effective resolution below 10 bits for this channel.')
add_normal(doc, 'FW-ADC-03 (MANDATORY): ADC acquisition time shall be ≥10 µs per sample to allow the RC network (10 kΩ pull-up + zone module PCB capacitance, estimated ≤100 pF) to settle. With a 10 kΩ / 100 pF network, τ = 1 µs; 10 µs = 10τ (0.005% settling error) — sufficient for 12-bit resolution.')
add_normal(doc, 'FW-ADC-04 (MANDATORY): Resolution margin verification. With a 10 kΩ pull-up at VDD = 3.3 V, the ZONE_ID voltage for each module and the midpoint thresholds are:')
t3 = doc.add_table(rows=1, cols=5)
t3.style = 'Table Grid'
t3.rows[0].cells[0].text = 'Zone'
t3.rows[0].cells[1].text = 'Resistor'
t3.rows[0].cells[2].text = 'V_ZONE_ID (nom)'
t3.rows[0].cells[3].text = 'Lower threshold'
t3.rows[0].cells[4].text = 'Upper threshold'
adc_data = [
    ('ZM-01', '10 kΩ', '1.650 V', '— (none)', '1.925 V  (<14 kΩ)'),
    ('ZM-02', '22 kΩ', '2.269 V', '1.925 V  (>14 kΩ)', '2.514 V  (<32 kΩ)'),
    ('ZM-03', '47 kΩ', '2.721 V', '2.514 V  (>32 kΩ)', '2.877 V  (<68 kΩ)'),
    ('ZM-04', '100 kΩ', '3.000 V', '2.877 V  (>68 kΩ)', '3.094 V  (<150 kΩ)'),
    ('ZM-05', '220 kΩ', '3.157 V', '3.094 V  (>150 kΩ)', '— (none)'),
]
for row_data in adc_data:
    row = t3.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
add_normal(doc, 'Minimum threshold-to-nominal voltage margin: 62.5 mV (between V_ZM05 = 3.157 V and upper threshold = 3.094 V). At 12-bit resolution with VDD = 3.3 V: LSB = 0.806 mV. Minimum margin = 78 LSB. Firmware must not accept ADC readings within 5 LSB of any threshold (guard band).')
add_normal(doc, 'FW-ADC-05 (MANDATORY): The ADC thresholds above must be converted to 12-bit ADC counts in firmware using the formula: count = round(V_threshold / VDD × 4096). Threshold counts shall be defined as named constants, not magic numbers, in the firmware source. Example: ZONE_ID_THRESHOLD_ZM01_ZM02_COUNT = 2386.')

add_heading(doc, '3.4  ZONE_ID Debounce Algorithm (Locked — RISK-18 Mitigation)', level=2)
add_normal(doc, 'This algorithm is locked per CLAUDE.md §13.5. Firmware must implement it exactly as specified. No deviation without an approved Engineering Change Order (ECO).')
add_normal(doc, 'FW-DBC-01 (MANDATORY): On zone module insertion event (detected via mechanical contact closure or session start trigger), the Hub MCU firmware shall perform 3 consecutive ADC reads of pin 18 at 100 ms intervals before making any FAULT or ACCEPT decision. The 3-read sequence is mandatory — a single-sample decision is not permitted under any operating condition, including low-power or fast-start modes.')
add_normal(doc, 'FW-DBC-02 (MANDATORY): Accept criterion — a zone module is accepted if and only if ≥2 of the 3 ADC reads return a value within the valid ZONE_ID band for any single zone (per the threshold table in §3.3). All ≥2 agreeing reads must identify the same zone. A reading is "in-band" if it satisfies both the lower and upper threshold conditions for exactly one zone ID. A reading that falls between thresholds or outside all defined bands is out-of-band.')
add_normal(doc, 'FW-DBC-03 (MANDATORY): A single out-of-band reading out of 3 is tolerated without fault. This suppresses single-sample contact-bounce transients that occur during ZIF lever actuation at module insertion. Rationale: ZIF contact resistance transiently spikes during actuation; one bad reading is an expected artefact, not a failure.')
add_normal(doc, 'FW-DBC-04 (MANDATORY): FAULT state — if all 3 reads are out-of-band (no valid zone identified by ≥2 reads): assert FAULT for the affected zone, hold ENABLE low, and issue the bone conduction audio alert defined in §3.5. The FAULT condition persists until a new insertion event triggers a fresh 3-read sequence. The user must physically remove and reinsert the module to clear the FAULT — firmware shall not auto-retry without a new insertion event.')
add_normal(doc, 'FW-DBC-05 (MANDATORY): If ZONE_ID is validated (≥2/3 reads agree) but the identified zone does not match the expected zone for that physical slot: assert FAULT, hold ENABLE low, and issue the wrong-module audio alert defined in §3.5. This is the RISK-15 zone-mismatch interlock. Wrong-module FAULT is distinct from contact-fault FAULT in both alert text and app diagnostic code.')
add_normal(doc, 'FW-DBC-06 (MANDATORY): The 100 ms inter-read interval shall be implemented as a timer-based delay, not a busy-wait loop. During the debounce window (0–300 ms from insertion event), the Hub MCU may continue other tasks. ENABLE shall not be asserted for the debounced zone until the 3-read sequence completes with ACCEPT outcome.')
add_normal(doc, 'FW-DBC-07 (MANDATORY): Re-read on session start. ZONE_ID shall be re-read (full 3-read debounce sequence) at every session start, not only at insertion. A zone module may be removed and reinserted between sessions. A session may not begin until all required zones report ACCEPT from a fresh debounce sequence within the current session-start window (within 10 s of session start command).')
add_normal(doc, 'FW-DBC-08 (INFORMATIVE): Expected timing. Module insertion → first read at t=0 ms → second read at t=100 ms → third read at t=200 ms → decision at t≈200 ms → ENABLE asserted (if ACCEPT) or FAULT (if all-fail). Total latency from insertion to ACCEPT: ≈200 ms. This is imperceptible to the user and does not affect session start UX.')

add_heading(doc, '3.5  Audio Alerts (Bone Conduction)', level=2)
add_normal(doc, 'The following alert texts are locked. Bone conduction audio delivery uses the existing bone conduction transducer (Neural Audio Entrainment modality hardware — zero additional BOM). Alert firmware is in scope of NP-COORD-001 G2-10 (audio zone ID firmware).')
t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Table Grid'
t4.rows[0].cells[0].text = 'Condition'
t4.rows[0].cells[1].text = 'Alert text (spoken, bone conduction)'
t4.rows[0].cells[2].text = 'ENABLE state'
alert_data = [
    ('ACCEPT — correct zone', '"[Zone name] connected"  (e.g. "Frontal Left connected")', 'Assert ENABLE after alert'),
    ('FAULT — contact failure (all 3 reads out-of-band)', '"Zone module contact issue — remove and reinsert"', 'Hold ENABLE low'),
    ('FAULT — wrong module (ZONE_ID valid but wrong slot)', '"Wrong module — expected [expected zone name]"', 'Hold ENABLE low'),
    ('FAULT — module absent at session start', '"[Zone name] not detected — check module"', 'Hold ENABLE low'),
]
for cond, text, enable in alert_data:
    row = t4.add_row()
    row.cells[0].text = cond
    row.cells[1].text = text
    row.cells[2].text = enable
add_normal(doc, 'FW-AUD-01 (MANDATORY): Alert audio shall play through bone conduction transducer at the default session volume (user-configurable, default 60% of maximum). Alert shall not be suppressed by stealth mode (CLAUDE.md §4.7 — safety faults always fire).')
add_normal(doc, 'FW-AUD-02 (MANDATORY): Alert text shall be synthesised using the on-device TTS engine (scope: app firmware team). Alert audio files shall be pre-rendered and stored in firmware flash to ensure availability in Mode 3 Autonomous (no phone or cloud required).')

add_heading(doc, '3.6  App Diagnostic Integration', level=2)
add_normal(doc, 'FW-APP-01 (MANDATORY): Contact-fault FAULT and wrong-module FAULT shall produce distinct diagnostic codes visible in the NeurOne app. The app shall display a specific diagnostic message matching the audio alert — not a generic error. Diagnostic codes:')
add_bullet(doc, 'ZONE_FAULT_CONTACT: pin 18 unreadable after 3 reads. App message: "Zone module contact issue — remove and reinsert [zone name] module."')
add_bullet(doc, 'ZONE_FAULT_WRONG: pin 18 readable but wrong zone for this slot. App message: "Wrong module in [slot name] slot — expected [zone name]."')
add_bullet(doc, 'ZONE_FAULT_ABSENT: module not detected at session start. App message: "[Zone name] module not detected — check module is fully inserted."')
add_normal(doc, 'FW-APP-02 (MANDATORY): FAULT events shall be logged in SHDR (System Health Data Record) with: zone ID, fault type, ADC read values (all 3), timestamp (SHDR session count — no user timestamp per UHDR/SHDR boundary rules). This enables fleet-level ZIF contact wear trending (predictive maintenance, Phase 1 population average).')

# Section 4
add_heading(doc, '4.  PROCESSOR OWNERSHIP')
t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Table Grid'
t5.rows[0].cells[0].text = 'Firmware component'
t5.rows[0].cells[1].text = 'Owner processor'
t5.rows[0].cells[2].text = 'IEC 62304 class'
ownership = [
    ('ZONE_ID ADC read (3-sample debounce)', 'Hub main processor (i.MX RT1062)', 'Class B'),
    ('ZONE_ID validation result reporting to Safety MCU', 'Hub main processor → Safety MCU (SPI heartbeat)', 'Class B / Class C boundary'),
    ('ENABLE GPIO assert/de-assert on ZONE_ID outcome', 'Safety MCU (STM32G071) — owns all stimulation GPIO', 'Class C'),
    ('Bone conduction audio alert trigger', 'Hub main processor (triggers audio subsystem)', 'Class B'),
    ('App diagnostic code generation', 'Hub main processor (BLE/USB to app)', 'Class B'),
    ('SHDR FAULT event log write', 'Hub main processor (SHDR partition)', 'Class B'),
]
for comp, owner, cls in ownership:
    row = t5.add_row()
    row.cells[0].text = comp
    row.cells[1].text = owner
    row.cells[2].text = cls
add_normal(doc, 'The Safety MCU shall never assert ENABLE for any zone unless the main processor has reported ZONE_ID ACCEPT for that zone within the current session-start sequence. The Safety MCU watchdog (200 ms SPI heartbeat) does not reset ZONE_ID validation state — ZONE_ID must be re-validated at each session start explicitly.')

# Section 5
add_heading(doc, '5.  CROSS-REFERENCES AND GATE STATUS')
t6 = doc.add_table(rows=1, cols=3)
t6.style = 'Table Grid'
t6.rows[0].cells[0].text = 'Reference'
t6.rows[0].cells[1].text = 'Item'
t6.rows[0].cells[2].text = 'Status'
xrefs = [
    ('NP-COORD-001 G1-13', 'ZONE_ID firmware debounce spec written into firmware requirements document', 'CLOSED — this document'),
    ('NP-COORD-001 G2-10', 'Audio zone ID firmware — bone conduction zone name on insertion / wrong-module alert', 'OPEN — implementation required'),
    ('NP-FAI-ZM-001 FAI-K04', 'Pin 18 contact resistance ≤30 mΩ at 100 g insertion force — catches high-resistance contacts before field deployment', 'OPEN — FAI required'),
    ('RISK-18', 'ZONE_ID ENABLE interlock availability failure — 3-layer mitigation: (1) firmware debounce [this document], (2) FAI-K04, (3) app UX diagnostic', 'MITIGATED — firmware layer complete; FAI-K04 and app UX layers pending'),
    ('NP-HW-FPC-001 Rev D §11.4', 'Layer 2 electronic zone identity — ADC thresholds and ZONE_ID resistor map', 'CLOSED — hardware spec frozen'),
    ('CLAUDE.md §13.5', 'Locked decision: ZONE_ID firmware debounce 3×ADC at 100ms, ≥2/3 pass', 'LOCKED'),
]
for ref, item, status in xrefs:
    row = t6.add_row()
    row.cells[0].text = ref
    row.cells[1].text = item
    row.cells[2].text = status

# Section 6
add_heading(doc, '6.  REVISION HISTORY')
t7 = doc.add_table(rows=1, cols=4)
t7.style = 'Table Grid'
t7.rows[0].cells[0].text = 'Rev'
t7.rows[0].cells[1].text = 'Date'
t7.rows[0].cells[2].text = 'Author'
t7.rows[0].cells[3].text = 'Description'
rev_row = t7.add_row()
rev_row.cells[0].text = 'A'
rev_row.cells[1].text = '2026-05-10'
rev_row.cells[2].text = 'NeurOne Firmware'
rev_row.cells[3].text = 'Initial release. Closes NP-COORD-001 G1-13. Zone module ZONE_ID detection: ADC requirements (FW-ADC-01–05), debounce algorithm (FW-DBC-01–08), audio alerts (FW-AUD-01–02), app diagnostics (FW-APP-01–02), processor ownership table. Implements RISK-18 firmware mitigation layer.'

doc.save('/home/user/NeurOne/docs/neurone_fw_requirements_001.docx')
print('Done: neurone_fw_requirements_001.docx written')
