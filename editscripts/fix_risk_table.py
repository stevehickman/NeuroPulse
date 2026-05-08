"""
Fix risk register summary table:
- RISK-08/09/10: Status column (col 8) was not updated; Root Cause (col 3) was wrongly overwritten.
- Restore Root Cause text for RISK-08/09/10.
- Update Status (col 8) and Severity (col 1) to MITIGATED.
- Add detailed analysis sections for RISK-09 and RISK-10.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RISK_PATH = "docs/neuropulse_fpc_zone_module_risks_revA.docx"
doc = Document(RISK_PATH)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# Correct root cause texts (reconstructed from specification)
ROOT_CAUSES = {
    "RISK-08": (
        "LED emitters are sold in wide Vf bins (±0.2–0.3 V typical at rated current). "
        "Even with per-string BCR421W constant-current drivers, wide Vf spread across "
        "strings means the BCR421W compliance voltage must accommodate the widest Vf "
        "emitter — or some strings fall outside the driver's regulation range. "
        "Supply chain default: no bin restriction on PO."
    ),
    "RISK-09": (
        "SBIR Phase I proposal (§C.2) referenced 'Molex SlimStack 0.35 mm pitch' for "
        "zone module FPC connectors. 0.35 mm pitch connectors are typically rated "
        "0.3–0.5 A per contact maximum. At 0.9 A RMS per power pin (paired), "
        "0.35 mm pitch is insufficient. Spec now mandates 0.5 mm pitch (Hirose FH34S); "
        "stale reference in proposal created inconsistency between documents."
    ),
    "RISK-10": (
        "Electrodeposited (ED) copper is the industry default for FPC fabrication. "
        "Rolled Annealed (RA) copper must be explicitly called out on every purchase "
        "order and fabrication drawing. Omission results in ED copper being used "
        "without notification. ED copper fails under dynamic flex — propagating grain "
        "boundary cracks lead to open-circuit failure within 10,000–50,000 cycles."
    ),
}

STATUS_TEXTS = {
    "RISK-08": "MITIGATED — Vf bin tolerance on PO (NP-PROC-FPC-001 §2); incoming inspection §7",
    "RISK-09": "MITIGATED — SBIR proposal updated; 0.35 mm pitch eliminated; FH34S 0.5 mm confirmed",
    "RISK-10": "MITIGATED — RA copper on fab drawing (§13 note 1) and PO (NP-PROC-FPC-001 §5)",
}

SEVERITY_TEXTS = {
    "RISK-08": "HIGH\n(MITIGATED)",
    "RISK-09": "MEDIUM\n(MITIGATED)",
    "RISK-10": "MEDIUM\n(MITIGATED)",
}

# Find and update the summary table
summary_table = doc.tables[0]
changes = 0

for row in summary_table.rows:
    cells = row.cells
    if not cells or len(cells) < 9:
        continue
    risk_id = cells[0].text.strip()
    if risk_id not in ("RISK-08", "RISK-09", "RISK-10"):
        continue

    print(f"\n{risk_id}:")

    # Col 1: Severity — update to include (MITIGATED)
    set_cell_text(cells[1], SEVERITY_TEXTS[risk_id], bold=True, color=(0x00, 0x70, 0x00))
    set_cell_bg(cells[1], "E2EFDA")
    print(f"  Severity → {SEVERITY_TEXTS[risk_id]!r}")
    changes += 1

    # Col 3: Root Cause — restore correct text (was overwritten with "MITIGATED" by earlier script)
    set_cell_text(cells[3], ROOT_CAUSES[risk_id])
    print(f"  Root Cause restored")
    changes += 1

    # Col 8: Status — update to MITIGATED
    set_cell_text(cells[8], STATUS_TEXTS[risk_id], bold=True, color=(0x00, 0x70, 0x00))
    set_cell_bg(cells[8], "92D050")
    print(f"  Status → {STATUS_TEXTS[risk_id]!r}")
    changes += 1

print(f"\nTotal summary table changes: {changes}")

# ─────────────────────────────────────────────────────────────────────────────
# Add detailed analysis sections for RISK-09 and RISK-10
# (They were absent from the detailed section — only RISK-01 to RISK-08 had detail)
# ─────────────────────────────────────────────────────────────────────────────

# Find the "OPEN ISSUES" paragraph — we'll insert before it
open_issues_idx = None
for i, para in enumerate(doc.paragraphs):
    if "OPEN ISSUES" in para.text:
        open_issues_idx = i
        break

print(f"\nOpen issues section at paragraph index: {open_issues_idx}")

def add_heading(doc, text, insert_before_para):
    """Insert a paragraph before the given paragraph element."""
    new_para = OxmlElement("w:p")
    insert_before_para._p.addprevious(new_para)
    # Find the new para and style it
    for p in doc.paragraphs:
        if p._p is new_para:
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            return p
    return None

def add_body_para(doc, text, insert_before_para, bullet=False, color=None):
    """Insert a body paragraph before the given paragraph element."""
    new_para = OxmlElement("w:p")
    insert_before_para._p.addprevious(new_para)
    for p in doc.paragraphs:
        if p._p is new_para:
            run = p.add_run(text)
            run.font.size = Pt(10)
            if color:
                run.font.color.rgb = RGBColor(*color)
            if bullet:
                p.style = doc.styles['List Bullet']
            return p
    return None

if open_issues_idx is not None:
    target_para = doc.paragraphs[open_issues_idx]

    # ── RISK-10 detail (added second = appears first since we insert before) ──
    # Actually inserting "before" means the last thing we add appears first.
    # So add RISK-10 first, then RISK-09 — they'll appear in correct order.

    # Add a blank spacer
    add_body_para(doc, "", target_para)

    # RISK-10 detail
    add_body_para(doc, (
        "→ Action 3 (DONE): FPC spec §13 Fab Note 1 marked as DRAWING REQUIREMENT — fabricator "
        "must acknowledge RA copper requirement on order confirmation."
    ), target_para, bullet=True, color=(0x00, 0x70, 0x00))

    add_body_para(doc, (
        "→ Action 2 (DONE): Procurement document NP-PROC-FPC-001 §5 contains verbatim PO "
        "language: 'Copper foil shall be Rolled Annealed (RA) per IPC-4204/11 Type I. "
        "Electrodeposited (ED) copper is NOT acceptable for dynamic flex.' Incoming "
        "inspection §7 line item 7 requires copper foil certification."
    ), target_para, bullet=True, color=(0x00, 0x70, 0x00))

    add_body_para(doc, (
        "→ Action 1 (DONE): FPC spec §3.1 and §13 Fab Note 1 updated to mandate RA copper "
        "explicitly. Note is designated a DRAWING REQUIREMENT — not a suggestion."
    ), target_para, bullet=True, color=(0x00, 0x70, 0x00))

    add_body_para(doc, (
        "RA copper (per IPC-4204/11 Type I) has elongation at break of 20–30% versus "
        "3–8% for ED copper. In dynamic flex applications, ED copper grain boundaries "
        "propagate cracks after 10,000–50,000 flex cycles. At 1,000–3,000 session cycles "
        "per year (cable-connected daily use), ED copper may fail within 4–10 years — "
        "within the device warranty period. RA copper has documented 10M+ cycle life "
        "in dynamic flex applications when routed at ≥25 mm radius. The fabricator "
        "will not substitute RA for ED without being told — ED is cheaper and the "
        "default. This is purely a purchase order and drawing discipline issue."
    ), target_para)

    add_heading(doc, "RISK-10  |  MEDIUM (MITIGATED)  |  ED Copper Default — RA Copper Mandated on Drawing and PO", target_para)

    # Add blank spacer between RISK-09 and RISK-10
    add_body_para(doc, "", target_para)

    # ── RISK-09 detail ──
    add_body_para(doc, (
        "→ Action 3 (DONE): Procurement document NP-PROC-FPC-001 §4 explicitly prohibits "
        "Molex SlimStack and 0.35 mm pitch. Hirose FH34S-20S-0.5SH or JAE FF03 only. "
        "Substitute prohibition language included."
    ), target_para, bullet=True, color=(0x00, 0x70, 0x00))

    add_body_para(doc, (
        "→ Action 2 (DONE): SBIR Phase I proposal updated — all references to 'Molex "
        "SlimStack' replaced with 'Hirose FH34S-20S-0.5SH'; all references to "
        "'0.35 mm pitch' replaced with '0.5 mm pitch'. Three paragraph-level changes made."
    ), target_para, bullet=True, color=(0x00, 0x70, 0x00))

    add_body_para(doc, (
        "→ Action 1 (DONE): Connector spec updated to Hirose FH34S-20S-0.5SH (0.5 mm pitch, "
        "back-flip lever ZIF, ≥1,000 cycles) in FPC spec §5, CLAUDE.md §7.1, and "
        "risk RISK-01 mitigation. 0.35 mm pitch is explicitly excluded."
    ), target_para, bullet=True, color=(0x00, 0x70, 0x00))

    add_body_para(doc, (
        "The SBIR Phase I application (§C.2) listed 'Molex SlimStack 0.35 mm pitch FPC connector' "
        "as part of the design freeze deliverable. This was an early-draft reference that "
        "predated the RISK-01 resolution. At 0.35 mm pitch, contact current ratings are "
        "typically 0.3–0.5 A — insufficient for 0.9 A RMS power pins (paired). "
        "The correct connector (Hirose FH34S, 0.5 mm pitch) is rated 0.5 A per contact; "
        "power pins are paired (2 × 0.5 A = 1.0 A, margin above 0.9 A RMS). "
        "Leaving the old reference in the proposal created a document inconsistency "
        "that could undermine credibility with reviewers who know FPC specifications."
    ), target_para)

    add_heading(doc, "RISK-09  |  MEDIUM (MITIGATED)  |  SBIR Proposal 0.35 mm Pitch Reference — Corrected to 0.5 mm", target_para)

    # Add a "DETAILED ANALYSIS — MEDIUM RISKS (MITIGATED)" section heading
    add_body_para(doc, "", target_para)
    for p in list(doc.paragraphs):
        if "OPEN ISSUES" in p.text and p is target_para:
            break
    add_heading(doc, "DETAILED ANALYSIS — MEDIUM RISKS (MITIGATED)", target_para)
    print("  Added RISK-09 and RISK-10 detailed sections")

doc.save(RISK_PATH)
print(f"\nRisk register saved: {RISK_PATH}")
print("Done.")
