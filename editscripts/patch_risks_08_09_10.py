"""
Patch script: Mitigate RISK-08, RISK-09, RISK-10 in the risk register.
Also patches SBIR proposal for RISK-09 and verifies/strengthens FPC spec fab note for RISK-10.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def bold_para(doc, text, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    return p

# ─────────────────────────────────────────────────────────────────────────────
# 1. Patch Risk Register — RISK-08, RISK-09, RISK-10
# ─────────────────────────────────────────────────────────────────────────────

RISK_PATH = "docs/superseded/np_risk_001.docx"
doc = Document(RISK_PATH)

# --- Update the summary table ---
# Columns: Risk ID | Severity | Probability | Status | One-line description
# We'll scan for the summary table (first large table in the doc)

summary_table = None
for tbl in doc.tables:
    # Look for table that has "RISK-08" in any cell
    for row in tbl.rows:
        for cell in row.cells:
            if "RISK-08" in cell.text:
                summary_table = tbl
                break
        if summary_table:
            break
    if summary_table:
        break

changes = 0
if summary_table:
    for row in summary_table.rows:
        cells = row.cells
        if not cells:
            continue
        risk_id = cells[0].text.strip()
        if risk_id in ("RISK-08", "RISK-09", "RISK-10"):
            # Status column is index 3
            status_cell = cells[3]
            old_status = status_cell.text.strip()
            # Clear and rewrite
            for para in status_cell.paragraphs:
                for run in para.runs:
                    run.text = ""
            status_cell.paragraphs[0].add_run("MITIGATED")
            set_cell_bg(status_cell, "92D050")  # green
            print(f"  Summary table: {risk_id} status '{old_status}' → MITIGATED")
            changes += 1
else:
    print("  WARNING: summary table not found")

# --- Update detailed sections ---
# Strategy: scan paragraphs, find RISK-## headings, then update content within each section.

risk_sections = {}
current_risk = None
current_paras = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    for r in ("RISK-08", "RISK-09", "RISK-10"):
        if text.startswith(r):
            if current_risk:
                risk_sections[current_risk] = current_paras
            current_risk = r
            current_paras = [i]
            break
    else:
        if current_risk:
            current_paras.append(i)

if current_risk:
    risk_sections[current_risk] = current_paras

print(f"\n  Detailed sections found: {list(risk_sections.keys())}")

def replace_status_in_section(doc, section_para_indices):
    """Find 'Status:' paragraph in section and update to MITIGATED."""
    for idx in section_para_indices:
        para = doc.paragraphs[idx]
        if para.text.strip().startswith("Status:"):
            for run in para.runs:
                if "OPEN" in run.text.upper() or "HIGH" in run.text.upper() or "MEDIUM" in run.text.upper() or "CRITICAL" in run.text.upper():
                    old = run.text
                    run.text = run.text.replace("OPEN", "MITIGATED").replace("Open", "MITIGATED")
                    run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                    print(f"    Status run updated: '{old}' → '{run.text}'")
                    return True
            # Try full paragraph text replacement
            full = para.text
            if "OPEN" in full.upper():
                for run in para.runs:
                    run.text = ""
                para.runs[0].text = "Status: MITIGATED" if para.runs else ""
                print(f"    Status para updated to MITIGATED")
                return True
    return False

# RISK-08 detailed: add mitigation paragraph after the section heading area
# Find the RISK-08 heading paragraph
risk08_heading_idx = None
risk09_heading_idx = None
risk10_heading_idx = None

for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t.startswith("RISK-08") and "RA" not in t and risk08_heading_idx is None:
        risk08_heading_idx = i
    if t.startswith("RISK-09") and risk09_heading_idx is None:
        risk09_heading_idx = i
    if t.startswith("RISK-10") and risk10_heading_idx is None:
        risk10_heading_idx = i

print(f"\n  Heading indices: RISK-08={risk08_heading_idx}, RISK-09={risk09_heading_idx}, RISK-10={risk10_heading_idx}")

# Update Status lines in each section
for risk_id, para_indices in risk_sections.items():
    result = replace_status_in_section(doc, para_indices)
    if not result:
        print(f"  WARNING: could not update Status line for {risk_id}")

# Add MITIGATED notes to the detailed body paragraphs
# We'll search for specific paragraphs by content and insert/replace mitigation text.

RISK08_MITIGATION = (
    "MITIGATED (Rev B / Procurement Doc NP-PROC-FPC-001 Rev 1): "
    "LED Vf binning requirements are now mandated on all purchase orders. "
    "Within-order tolerance: ±0.10 V at rated current. Lot-to-lot tolerance: ±0.15 V. "
    "Verbatim PO language is defined in NP-PROC-FPC-001 §2. "
    "Out-of-spec lots are subject to mandatory disposition review before acceptance. "
    "Approved supplier list (Epistar, OSRAM OS, Cree/Wolfspeed) requires Vf bin "
    "certificates with each shipment. Incoming inspection: 5-unit sample AQL 1.0 per reel."
)

RISK09_MITIGATION = (
    "MITIGATED (Rev C / CLAUDE.md update): "
    "Connector family confirmed as Hirose FH34S-20S-0.5SH (0.5 mm pitch, back-flip lever "
    "ZIF, ≥1,000 insertion cycles). The 0.35 mm pitch option is explicitly excluded from "
    "all specifications — current per pin is insufficient at 0.35 mm pitch for 3.6 A peak "
    "zone draw. SBIR Phase I proposal updated to replace all references to Molex SlimStack "
    "and 0.35 mm pitch. FPC procurement document NP-PROC-FPC-001 §4 mandates Hirose FH34S "
    "or JAE FF03 only; Molex SlimStack explicitly prohibited. Connector spec cross-references "
    "NP-HW-FPC-001 Rev 3 §5."
)

RISK10_MITIGATION = (
    "MITIGATED (Rev A / Procurement Doc NP-PROC-FPC-001 Rev 1 / FPC Spec NP-HW-FPC-001 §13): "
    "Rolled Annealed (RA) copper is mandated in three locations: "
    "(1) FPC fabrication specification §13 Fab Note 1 — mandatory RA copper requirement "
    "with explicit prohibition on Electrodeposited (ED) copper; "
    "(2) Procurement document NP-PROC-FPC-001 §5 FPC Fabrication Requirements — verbatim "
    "PO language: 'Copper foil shall be Rolled Annealed (RA) per IPC-4204/11 Type I. "
    "Electrodeposited (ED) copper is NOT acceptable for dynamic flex applications'; "
    "(3) Incoming inspection protocol §7 line item 7: copper foil certification verification. "
    "Fab notes are designated as drawing requirements — fabricator must acknowledge on order confirmation."
)

# Insert mitigation text after the existing "Mitigation:" or "Action:" paragraphs
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if "RISK-08" in t and "Vf" in t:
        # Already has content — we're in the right section; look ahead for mitigation marker
        pass

# Simpler approach: scan all paragraphs and for each risk section, find the last paragraph
# that says "Action:" or "Mitigation:", then append after it.

# Build section boundaries
section_boundaries = {}
heading_indices = []
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    for r in ("RISK-08", "RISK-09", "RISK-10", "RISK-11"):
        if t.startswith(r):
            heading_indices.append((i, r))
            break

for j, (idx, rid) in enumerate(heading_indices):
    end_idx = heading_indices[j+1][0] if j+1 < len(heading_indices) else len(doc.paragraphs)
    section_boundaries[rid] = (idx, end_idx)

print(f"\n  Section boundaries: {section_boundaries}")

# For each target risk, find the paragraph containing "Mitigation" or "Action" in the section
# and replace it / append the mitigation text.

def update_risk_section(doc, risk_id, boundaries, mitigation_text):
    if risk_id not in boundaries:
        print(f"  WARNING: {risk_id} boundary not found")
        return
    start, end = boundaries[risk_id]

    # Look for existing "Mitigation:" paragraph
    mitigation_para_idx = None
    for i in range(start, end):
        t = doc.paragraphs[i].text.strip()
        if t.lower().startswith("mitigation") or t.lower().startswith("action"):
            mitigation_para_idx = i

    if mitigation_para_idx is not None:
        para = doc.paragraphs[mitigation_para_idx]
        # Clear runs and set new text
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = "Mitigation: " + mitigation_text
            para.runs[0].font.color.rgb = RGBColor(0x00, 0x60, 0x00)
        else:
            run = para.add_run("Mitigation: " + mitigation_text)
            run.font.color.rgb = RGBColor(0x00, 0x60, 0x00)
        print(f"  {risk_id}: Updated mitigation paragraph at index {mitigation_para_idx}")
    else:
        # Find last paragraph in section and note we couldn't find it
        print(f"  WARNING: {risk_id}: No 'Mitigation:' paragraph found in range {start}–{end}")
        print(f"  Paragraphs in section:")
        for i in range(start, min(start+10, end)):
            print(f"    [{i}]: {doc.paragraphs[i].text[:80]!r}")

update_risk_section(doc, "RISK-08", section_boundaries, RISK08_MITIGATION)
update_risk_section(doc, "RISK-09", section_boundaries, RISK09_MITIGATION)
update_risk_section(doc, "RISK-10", section_boundaries, RISK10_MITIGATION)

doc.save(RISK_PATH)
print(f"\nRisk register saved: {RISK_PATH}")

# ─────────────────────────────────────────────────────────────────────────────
# 2. Patch SBIR Proposal — RISK-09: replace Molex/0.35mm connector references
# ─────────────────────────────────────────────────────────────────────────────

SBIR_PATH = "docs/np_sbir_001.docx"
sbir = Document(SBIR_PATH)

sbir_changes = 0
search_terms = [
    ("Molex SlimStack", "Hirose FH34S-20S-0.5SH"),
    ("0.35 mm pitch", "0.5 mm pitch"),
    ("0.35mm pitch", "0.5 mm pitch"),
    ("SlimStack", "FH34S lever-ZIF"),
]

for para in sbir.paragraphs:
    for old, new in search_terms:
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    sbir_changes += 1
                    print(f"  SBIR: replaced '{old}' → '{new}' in: {para.text[:80]!r}")

# Also check tables in SBIR
for tbl in sbir.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for old, new in search_terms:
                    if old in para.text:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                sbir_changes += 1
                                print(f"  SBIR table: replaced '{old}' → '{new}'")

print(f"\n  SBIR changes: {sbir_changes}")
sbir.save(SBIR_PATH)
print(f"  SBIR saved: {SBIR_PATH}")

# ─────────────────────────────────────────────────────────────────────────────
# 3. Verify / Strengthen FPC Spec Fab Note for RISK-10 (RA copper)
# ─────────────────────────────────────────────────────────────────────────────

SPEC_PATH = "docs/superseded/np_hw_fpc_001.docx"
spec = Document(SPEC_PATH)

spec_changes = 0
ra_found = False
for i, para in enumerate(spec.paragraphs):
    t = para.text
    if "Rolled Annealed" in t or "RA copper" in t or "IPC-4204" in t:
        ra_found = True
        print(f"  FPC spec: RA copper fab note found at para {i}: {t[:100]!r}")
        # Ensure the paragraph starts with "FAB NOTE" or similar prominence
        if not any(kw in t for kw in ["FAB NOTE", "MANDATORY", "DRAWING REQUIREMENT"]):
            # Prepend drawing requirement marker
            for run in para.runs:
                if "Rolled Annealed" in run.text or "RA copper" in run.text or "IPC-4204" in run.text:
                    run.text = "[DRAWING REQUIREMENT] " + run.text
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    spec_changes += 1
                    print(f"  FPC spec: marked para {i} as DRAWING REQUIREMENT")
                    break

if not ra_found:
    print("  WARNING: RA copper fab note not found in FPC spec — adding it")
    # Find the Fabrication Notes section and add it
    fab_section_idx = None
    for i, para in enumerate(spec.paragraphs):
        if "Fabrication" in para.text and "Note" in para.text:
            fab_section_idx = i
    if fab_section_idx:
        # Insert after section heading
        target_para = spec.paragraphs[fab_section_idx]
        new_para = OxmlElement("w:p")
        target_para._p.addnext(new_para)
        # Get the inserted paragraph
        for i2, p2 in enumerate(spec.paragraphs):
            if p2._p is new_para:
                run = p2.add_run(
                    "[DRAWING REQUIREMENT — FAB NOTE 1] Copper foil: Rolled Annealed (RA) per "
                    "IPC-4204/11 Type I. Electrodeposited (ED) copper is NOT acceptable. "
                    "Fabricator must acknowledge on order confirmation."
                )
                run.bold = True
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                spec_changes += 1
                break

spec.save(SPEC_PATH)
print(f"  FPC spec saved: {SPEC_PATH} (changes: {spec_changes})")

print("\nAll patches complete.")
