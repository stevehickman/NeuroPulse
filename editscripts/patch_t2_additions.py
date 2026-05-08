"""
Add sLORETA-guided HD-tDCS and cervical VNS to T2 stack:
  1. Design brief Rev 4 — completed decisions bullets
  2. Additional modalities doc — mark both as ADDED
  3. FAI checklist — HD-tDCS verification items (FAI-HD01–HD04) and
     cervical VNS items (FAI-CV01–CV03)
  4. Coordination checklist — G3-07 (HD-tDCS electrode spec) and
     G3-08 (cervical VNS safety interlock)
  5. Risk register — RISK-25 (cervical VNS cardiac proximity)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, size=8, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def insert_para_after(anchor_p, text, size=9):
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    new_rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2)))
    new_rpr.append(sz); new_r.append(new_rpr)
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve"); new_t.text = text
    new_r.append(new_t); new_p.append(new_r)
    anchor_p.addnext(new_p)
    return new_p

def add_row_after(anchor_tr, row_data, col_widths_in):
    new_tr = OxmlElement("w:tr")
    for c_idx, txt in enumerate(row_data):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        w_el = OxmlElement("w:tcW")
        w_el.set(qn("w:w"), str(int(col_widths_in[c_idx] * 1440)))
        w_el.set(qn("w:type"), "dxa")
        tc_pr.append(w_el); tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr_el = OxmlElement("w:rPr")
        sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
        rpr_el.append(sz_el); r_el.append(rpr_el)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
        r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
        new_tr.append(tc)
    anchor_tr.addnext(new_tr)
    return new_tr

# ══════════════════════════════════════════════════════════════════════════════
# 1. Design Brief Rev 4 — completed decisions
# ══════════════════════════════════════════════════════════════════════════════
BRIEF_PATH = "docs/neuropulse_brief_r4.docx"
brief = Document(BRIEF_PATH)

# Find the HRV biofeedback paragraph we added previously and insert after it
hrv_para = None
for para in brief.paragraphs:
    if "HRV biofeedback protocol" in para.text and "resonance" in para.text.lower():
        hrv_para = para
        break

if hrv_para:
    p1 = insert_para_after(
        hrv_para._p,
        "sLORETA-guided HD-tDCS (T2): 4×1 ring montage (center anode + 4 return cathodes); "
        "Ag/AgCl 3.5mm dual-rated electrodes in T2 qEEG cap; 16-ch tACS driver provides "
        "independent current sources (no additional stimulation hardware); sLORETA source "
        "map → MNI target → automatic 10-20 electrode mapping → personalised current "
        "distribution; 40µC/cm² safety MCU limit; ≤2mA/electrode; ~3–5× spatial focality "
        "vs standard tDCS. Evidence: Jog/UCLA 2025 (n=71, personalised HD-tDCS); "
        "BRIGhTMIND 2024 (n=255, personalised targeting outperforms fixed F3).",
        size=9
    )
    insert_para_after(
        p1,
        "Cervical VNS (T2 accessory): neck-worn tcVNS module stimulating vagus trunk at "
        "carotid sheath; higher vagal activation than auricular branch; safety MCU owns "
        "enable with cardiac monitor interlock (auto-cutoff if HR changes >15 BPM in 5s); "
        "510(k) predicate = electroCore gammaCore K163334 (cluster headache) + K173323 "
        "(migraine); gel pad consumable; +$35–55 BOM. Indications: cluster headache, "
        "migraine, depression adjunct, post-stroke rehabilitation.",
        size=9
    )
    print("  Brief: sLORETA HD-tDCS and cervical VNS bullets added")
else:
    print("  WARNING: HRV anchor paragraph not found in brief")

brief.save(BRIEF_PATH)
print("  Design brief saved")

# ══════════════════════════════════════════════════════════════════════════════
# 2. Additional Modalities Doc — mark both as ADDED
# ══════════════════════════════════════════════════════════════════════════════
MOD_PATH = "docs/neuropulse_additional_modalities.docx"
mod = Document(MOD_PATH)

marks = {
    "sLORETA-Guided HD-tDCS": "ADDED TO T2 STACK",
    "Cervical": "ADDED TO T2 STACK",
}

for tbl in mod.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for key, label in marks.items():
                if key in cell.text and label not in cell.text:
                    orig_text = cell.paragraphs[0].text
                    for run in cell.paragraphs[0].runs:
                        run.text = ""
                    run = cell.paragraphs[0].add_run(f"✓ {label} — " + orig_text.lstrip("0123456789. "))
                    run.bold = True; run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                    set_bg(cell, "E2EFDA")
                    print(f"  Modalities doc: '{key}' marked {label}")

mod.save(MOD_PATH)
print("  Additional modalities doc saved")

# ══════════════════════════════════════════════════════════════════════════════
# 3. FAI checklist — add HD-tDCS and cervical VNS test items
# ══════════════════════════════════════════════════════════════════════════════
FAI_PATH = "docs/neuropulse_fai_zone_module.docx"
fai = Document(FAI_PATH)

# Find last HRV item (FAI-H04) as insertion anchor
fai_h04_tr = None
anchor_tbl = None
for tbl in fai.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "FAI-H04":
            fai_h04_tr = row._tr
            anchor_tbl = tbl
            break
    if fai_h04_tr is not None:
        break

if fai_h04_tr is not None:
    fai_widths = [0.65, 2.3, 1.2, 1.55, 0.65, 0.45]

    HD_ITEMS = [
        ("FAI-HD01",
         "sLORETA source localisation accuracy: T2 qEEG session (21-ch, eyes-closed resting "
         "state, 5 min). sLORETA algorithm computes cortical source map. Verify: source "
         "reconstruction completes within 10s; DLPFC peak activation localises to within "
         "±15mm MNI coordinates of expected region (validated against published normative "
         "qEEG atlas). Bench test with 5 healthy adults.",
         "Computational + bench verification (5 subjects)",
         "sLORETA output within 10s. DLPFC localisation error ≤15mm MNI vs atlas. "
         "Consistent across 5 subjects (σ ≤ 8mm inter-subject).",
         "", ""),
        ("FAI-HD02",
         "HD-tDCS electrode mapping: verify firmware correctly translates sLORETA MNI "
         "target coordinates to 10-20 electrode positions and generates 4×1 montage "
         "configuration. Test with 5 simulated target locations (F3/Fz/Cz/P3/Oz). "
         "Confirm current distribution computed correctly for each montage.",
         "Firmware + computational verification",
         "Correct 10-20 position identified for all 5 test targets. 4×1 montage "
         "configuration matches manual calculation ±5%. Current sum across all electrodes "
         "= 0 (charge-balanced) in all montages.",
         "", ""),
        ("FAI-HD03",
         "HD-tDCS focality vs standard tDCS: bench measurement using gel phantom. "
         "Compare electric field distribution for (a) standard 2-electrode 2mA tDCS and "
         "(b) 4×1 ring HD-tDCS 2mA at same scalp location. Measure with embedded "
         "electrode array (or finite element model validated against phantom). "
         "Confirm HD-tDCS half-maximum field width ≤ 3× narrower than standard.",
         "Phantom / FEM bench test",
         "HD-tDCS FWHM ≤ 3cm vs standard tDCS FWHM ≥ 8cm in gel phantom "
         "(consistent with Bikson lab published data). Peak field strength within "
         "safety envelope.",
         "", ""),
        ("FAI-HD04",
         "Simultaneous EEG + HD-tDCS signal quality: 21-ch EEG recorded during and after "
         "1mA HD-tDCS (3 min on, 2 min post). Verify: (a) EEG artefact during stimulation "
         "is rejected correctly by artefact filter; (b) post-stimulation EEG SNR recovers "
         "to ≥95% of pre-stimulation baseline within 30s (electrode polarisation recovery); "
         "(c) no electrode damage or impedance drift >5kΩ after 10 consecutive sessions. "
         "3 bench subjects.",
         "3 subjects, signal quality + electrode durability",
         "EEG recovers ≥95% SNR within 30s post-stimulation. Impedance drift ≤5kΩ "
         "after 10 sessions. No visible electrode damage on inspection.",
         "", ""),
    ]

    CV_ITEMS = [
        ("FAI-CV01",
         "Cervical VNS electrode placement verification: confirm correct gel electrode "
         "placement on skin overlying sternocleidomastoid muscle / carotid sheath using "
         "anatomical landmarks. 5 subjects (3 male, 2 female, BMI range 20–32). "
         "Verify impedance <5kΩ at placement; no stimulation over carotid artery directly "
         "(ultrasound-guided confirmation in bench study).",
         "Anatomical bench verification, 5 subjects",
         "Impedance <5kΩ at correct placement in all 5 subjects. Ultrasound confirms "
         "electrode not directly over carotid bifurcation. Placement guide in IFU "
         "confirmed adequate by all 5 subjects.",
         "", ""),
        ("FAI-CV02",
         "Cardiac safety interlock: connect cervical VNS to simulated artefact ECG "
         "waveform generator. Deliver 1mA cervical stimulation. Inject simulated HR "
         "change of +16 BPM at t=3s. Verify: safety MCU detects HR change >15 BPM "
         "within ≤500ms and cuts stimulation enable line. No stimulation resumes "
         "without user confirmation.",
         "Bench safety interlock test (simulated ECG generator)",
         "Stimulation cutoff within ≤500ms of HR change >15 BPM threshold. "
         "Stimulation does not resume automatically. Log entry created in SHDR "
         "with timestamp and HR value at cutoff.",
         "", ""),
        ("FAI-CV03",
         "Cervical VNS tolerability: 10-minute session at 0.5mA (low) and 1.0mA "
         "(standard), bilateral cervical placement. 5 healthy adult subjects. "
         "Collect: subject comfort rating (NRS 0–10), carotid palpation result "
         "before/after, ECG rhythm pre/during/post. No vasovagal response, no "
         "carotid bruit change, no arrhythmia.",
         "5 subjects, tolerability + cardiac monitoring",
         "NRS comfort ≤4/10 at 1.0mA in all subjects. No vasovagal episode. "
         "No cardiac arrhythmia on ECG during or post-stimulation. "
         "No carotid bruit change on auscultation.",
         "", ""),
    ]

    prev_tr = fai_h04_tr
    for items, color in [(HD_ITEMS, "F0F0FF"), (CV_ITEMS, "FFF0F0")]:
        for item in items:
            new_tr = OxmlElement("w:tr")
            for c_idx, txt in enumerate(item):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(fai_widths[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            prev_tr.addnext(new_tr)
            for r2 in anchor_tbl.rows:
                if r2._tr is new_tr:
                    for ci in range(6):
                        set_bg(r2.cells[ci], color)
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            prev_tr = new_tr
            print(f"  FAI: {item[0]} inserted")
else:
    print("  WARNING: FAI-H04 not found — FAI items not inserted")

fai.save(FAI_PATH)
print("  FAI checklist saved")

# ══════════════════════════════════════════════════════════════════════════════
# 4. Coordination checklist — G3-07 and G3-08
# ══════════════════════════════════════════════════════════════════════════════
COORD_PATH = "docs/neuropulse_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

# Find G3-06 (last G3 item) — add G3-07 and G3-08 after it
g306_rows = []
for tbl in coord.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "G3-06":
            g306_rows.append((tbl, row))

new_g3_items = [
    ("G3-07", "G3", "EE + FW",
     "sLORETA-guided HD-tDCS T2 specification: (1) T2 qEEG electrode dual-rated spec "
     "(Ag/AgCl 3.5mm, simultaneous EEG+stimulation); (2) sLORETA-to-10-20 mapping "
     "algorithm validated (FAI-HD01); (3) 4×1 montage current distribution verified "
     "(FAI-HD02–HD03); (4) simultaneous EEG+tDCS signal quality (FAI-HD04).",
     "OPEN — T2 hardware specification"),
    ("G3-08", "G3", "EE + ME + Regulatory",
     "Cervical VNS T2 accessory specification: (1) electrode placement guide in IFU "
     "(FAI-CV01); (2) cardiac safety interlock bench verified (FAI-CV02); (3) "
     "tolerability study complete (FAI-CV03); (4) 510(k) pre-submission strategy "
     "aligned with FDA (predicate: electroCore K163334 + K173323).",
     "OPEN — T2 accessory specification"),
]

widths = [0.65, 0.55, 1.3, 2.8, 1.4]
for tbl, row in g306_rows:
    prev_tr = row._tr
    for item in new_g3_items:
        new_tr = add_row_after(prev_tr, item, widths)
        for r2 in tbl.rows:
            if r2._tr is new_tr:
                set_bg(r2.cells[4], "FFF2CC")
                if r2.cells[0].paragraphs[0].runs:
                    r2.cells[0].paragraphs[0].runs[0].bold = True
                break
        prev_tr = new_tr
        print(f"  Coord: {item[0]} added in table with G3-06")

coord.save(COORD_PATH)
print("  Coordination checklist saved")

# ══════════════════════════════════════════════════════════════════════════════
# 5. Risk register — RISK-25 cervical VNS cardiac proximity
# ══════════════════════════════════════════════════════════════════════════════
RISK_PATH = "docs/neuropulse_fpc_zone_module_risks_revA.docx"
risk = Document(RISK_PATH)

# Find last risk row (RISK-24) to insert after
risk24_tr = None
risk_tbl = None
for tbl in risk.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "RISK-24":
            risk24_tr = row._tr
            risk_tbl = tbl
            break
    if risk24_tr is not None:
        break

if risk24_tr is not None:
    # Determine column count from header
    n_cols = len(risk_tbl.rows[0].cells)
    risk_widths = [0.55, 0.7, 0.6, 0.8, 0.7, 1.9, 0.55, 1.0, 1.3][:n_cols]
    # Pad if needed
    while len(risk_widths) < n_cols:
        risk_widths.append(0.7)

    r25_data = [
        "RISK-25",
        "LOW\n(MITIGATED)",
        "T2 cervical VNS proximity to carotid artery and vagus trunk",
        "Cervical tcVNS electrode placed over sternocleidomastoid in proximity to "
        "carotid artery; inadvertent carotid sinus stimulation could trigger baroreceptor "
        "reflex → bradycardia / hypotension / vasovagal syncope",
        "MEDIUM\n(before mitigation)",
        "Option A — safety MCU cardiac interlock: PPG or ECG rhythm monitor feeds "
        "safety MCU; stimulation enable cut if HR changes >15 BPM within 5s of onset; "
        "auto-cutoff, requires user confirmation to resume. "
        "IFU specifies electrode placement landmarks to avoid carotid bifurcation directly. "
        "Pre-stimulation impedance check <5kΩ confirms correct skin contact. "
        "Bench study (FAI-CV01–CV02) validates placement and interlock. "
        "Regulatory precedent: electroCore gammaCore (K163334) cleared for same anatomy.",
        "T2 ONLY\n(T1 = auricular\nonly)",
        "FAI-CV02 (cardiac interlock bench test)\nFAI-CV03 (tolerability, 5 subjects)\n"
        "510(k) predicate K163334",
        "MITIGATED — safety MCU cardiac interlock; IFU placement guide; "
        "FAI-CV01–CV03 BLOCKING before T2 cervical VNS launch",
    ]

    # Trim or pad to match column count
    r25_data = r25_data[:n_cols]
    while len(r25_data) < n_cols:
        r25_data.append("")

    new_tr = add_row_after(risk24_tr, r25_data, risk_widths[:n_cols])
    for r2 in risk_tbl.rows:
        if r2._tr is new_tr:
            set_bg(r2.cells[1], "E2EFDA")
            if r2.cells[1].paragraphs[0].runs:
                r2.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x00,0x70,0x00)
                r2.cells[1].paragraphs[0].runs[0].bold = True
            if len(r2.cells) > 8:
                set_bg(r2.cells[8], "92D050")
            break
    print("  Risk register: RISK-25 (cervical VNS cardiac) added")
else:
    print("  WARNING: RISK-24 not found in risk register")

risk.save(RISK_PATH)
print("  Risk register saved")

print("\nDone.")
