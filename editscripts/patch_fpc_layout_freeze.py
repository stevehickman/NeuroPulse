"""
Issue #10 — FPC layout freeze + PD2 aperture position for mould spec.

Changes:
  NP-HW-FPC-001  Rev 4 — pin 19 notes locked, §8.4 XY position specified,
                          doc status changed to LAYOUT FROZEN
  NP-TOOL-ZM-001        — F-04 updated with XY coords, critical dims added,
                          OI-3 closed
  NP-DRV-SHELL-001      — DRC-23 note updated (FPC layout now available for
                          CAD overlay check at G2)
  NP-COORD-001          — G1-15 added (layout freeze CLOSED)

PD2 centre position (locked):
  X = 33.0 mm, Y = 39.0 mm from module reference corner (lower-left of LED
  array zone, scalp-facing view). This is the geometric centre of the 66 × 78 mm
  LED array, co-located in XY with PD1 (opposite face) for direct ratio
  calculation without spatial correction.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ──────────────────────────────────────────────────────────────────

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def append_row(table, cells_text, widths_inches, bg=None, bold_col0=True):
    """Append a new row to *table* with the given cell texts and widths."""
    new_tr = OxmlElement("w:tr")
    table.rows[-1]._tr.addnext(new_tr)   # insert after last row
    for idx, (txt, w) in enumerate(zip(cells_text, widths_inches)):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        w_el = OxmlElement("w:tcW")
        w_el.set(qn("w:w"), str(int(w * 1440)))
        w_el.set(qn("w:type"), "dxa")
        tc_pr.append(w_el)
        tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr_el = OxmlElement("w:rPr")
        sz_el = OxmlElement("w:sz")
        sz_el.set(qn("w:val"), "16")
        rpr_el.append(sz_el)
        r_el.append(rpr_el)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve")
        t_el.text = txt
        r_el.append(t_el)
        p_el.append(r_el)
        tc.append(p_el)
        new_tr.append(tc)
    # retrieve the row python-docx object that wraps new_tr and apply formatting
    for row in table.rows:
        if row._tr is new_tr:
            if bg:
                for c in row.cells:
                    set_bg(c, bg)
            if bold_col0 and row.cells[0].paragraphs[0].runs:
                row.cells[0].paragraphs[0].runs[0].bold = True
            return row
    return None


# ══════════════════════════════════════════════════════════════════════════════
# 1. NP-HW-FPC-001 — pin 19 + §8.4 position + doc freeze (Rev D)
# ══════════════════════════════════════════════════════════════════════════════
FPC_PATH = "docs/superseded/neurone_fpc_zone_module_spec_revA.docx"
fpc = Document(FPC_PATH)

# 1a. Update pin 19 Notes cell (Table 7, row where col 0 = '19')
pinout_tbl = fpc.tables[7]
for row in pinout_tbl.rows:
    if row.cells[0].text.strip() == "19":
        set_text(row.cells[4],
            "PD2 centre at X = 33.0 mm, Y = 39.0 mm from module reference corner "
            "(lower-left of LED array zone, scalp-facing view). This is the geometric "
            "centre of the 66 × 78 mm LED array, co-located in XY with PD1 on the "
            "opposite face. FPC contact pad: 1.6 mm diameter annular ring, scalp-facing "
            "copper layer, hard gold ≥ 0.5 µm. Aperture in scalp-facing PDMS window: "
            "1.2 mm diameter, centred on pad (F-04, NP-TOOL-ZM-001). "
            "POSITION LOCKED — change control required.",
            size=8)
        print("  Pin 19 Notes updated with locked XY position")
        break

# 1b. Update §8.4 PD2 spec table (Table 10) — add Position row
pd2_tbl = fpc.tables[10]

# Check if Position row already exists
has_position = any("Position" in row.cells[0].text for row in pd2_tbl.rows)
if not has_position:
    # Insert after header (after row 0)
    # Find "Quantity per zone" row (row 1) and add Position after it
    for row in pd2_tbl.rows:
        if "Quantity per zone" in row.cells[0].text:
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            pos_cells = (
                "Position (LOCKED)",
                "X = 33.0 ± 0.2 mm, Y = 39.0 ± 0.2 mm from module reference corner "
                "(lower-left of LED array zone, scalp-facing view). Geometric centre of "
                "66 × 78 mm LED array. Co-located in XY with PD1 (§8.1–8.3) for direct "
                "PD1/PD2 ratio calculation without spatial correction factor. "
                "FPC contact pad: 1.6 mm annular ring, scalp-facing copper layer. "
                "F-04 aperture in zone module mould centred on this coordinate."
            )
            widths = [1.4, 5.1]
            for idx, (txt, w) in enumerate(zip(pos_cells, widths)):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(w * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el)
                tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz")
                sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el)
                r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve")
                t_el.text = txt
                r_el.append(t_el)
                p_el.append(r_el)
                tc.append(p_el)
                new_tr.append(tc)
            # Format new row
            for r2 in pd2_tbl.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[0], "E2EFDA")
                    set_bg(r2.cells[1], "E2EFDA")
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  §8.4 PD2 spec table: Position (LOCKED) row added")
            break

# 1c. Update document header paragraphs (Revision, Status, Date)
for para in fpc.paragraphs:
    if para.text.startswith("Revision:"):
        for run in para.runs:
            run.text = ""
        para.runs[0].text = "Revision:  D (Layout Frozen — Released for Mould Specification)"
        print("  Revision header updated to Rev 4")
        break

for para in fpc.paragraphs:
    if para.text.startswith("Status:"):
        for run in para.runs:
            run.text = ""
        para.runs[0].text = (
            "Status:  LAYOUT FROZEN — Released for mould F-04 specification. "
            "Change control required for any pin position, component position, "
            "or footprint change."
        )
        print("  Status updated to LAYOUT FROZEN")
        break

for para in fpc.paragraphs:
    if para.text.startswith("Date:"):
        for run in para.runs:
            run.text = ""
        para.runs[0].text = "Date:  2026-05-09"
        print("  Date updated to 2026-05-09")
        break

# 1d. Add Rev 4 entry to revision history table (Table 17)
rev_tbl = fpc.tables[17]
rev_d = (
    "D",
    "2026-05-09",
    "Claude / NeurOne Engineering",
    "Issue #10 — FPC layout frozen. Pin 19 (PD2_CATHODE) position locked: "
    "X = 33.0 mm, Y = 39.0 mm from module reference corner (geometric centre of "
    "LED array). §8.4 updated with PD2 Position row. Document status changed to "
    "LAYOUT FROZEN. Enables NP-TOOL-ZM-001 F-04 specification and mould design "
    "review. DRC-18 routing strategy per NP-DRV-SHELL-001 §2.4 confirmed; "
    "CAD overlay measurement to proceed at G2 gate."
)
rev_widths = [0.4, 0.8, 1.6, 3.7]
append_row(rev_tbl, rev_d, rev_widths)
print("  Rev 4 entry added to revision history")

fpc.save(FPC_PATH)
print(f"  NP-HW-FPC-001 saved (Rev D, LAYOUT FROZEN)")


# ══════════════════════════════════════════════════════════════════════════════
# 2. NP-TOOL-ZM-001 — F-04 XY coords, critical dims, close OI-3
# ══════════════════════════════════════════════════════════════════════════════
TOOL_PATH = "docs/superseded/neurone_tool_zone_module_001.docx"
tool = Document(TOOL_PATH)

# 2a. Update F-04 description (Table 1)
features_tbl = tool.tables[1]
for row in features_tbl.rows:
    if row.cells[0].text.strip() == "F-04":
        set_text(row.cells[1],
            "RISK-14 Option B PD2 aperture: 1.2 mm circular aperture in scalp-facing "
            "PDMS window, positioned at X = 33.0 mm, Y = 39.0 mm from module reference "
            "corner (lower-left of LED array zone, scalp-facing view). This is the "
            "geometric centre of the 66 × 78 mm LED array — co-located in XY with PD1 "
            "on the LED-facing surface. Aperture tolerance: ±0.2 mm on both axes. "
            "Aperture wall: angled 15° inward (funnel entry) to maximise backscattered "
            "light collection from scalp tissue. Aperture must be free of flash and "
            "burr (Ra ≤ 1.6 µm interior surface). "
            "POSITION LOCKED — NP-HW-FPC-001 Rev 4 §8.4 — 2026-05-09.",
            size=8)
        set_text(row.cells[4], "☐", size=8)
        print("  F-04 description updated with locked XY position")
        break

# 2b. Add PD2 XY dimension rows to critical dimensions table (Table 3)
dims_tbl = tool.tables[3]
# Check if already present
has_pd2_x = any("PD2 aperture centre X" in row.cells[0].text for row in dims_tbl.rows)
if not has_pd2_x:
    # Find the existing "PD2 aperture diameter" row and add X and Y after it
    for row in dims_tbl.rows:
        if "PD2 aperture diameter" in row.cells[0].text:
            dim_widths = [2.2, 0.7, 0.9, 2.1, 0.6]
            # Insert X row
            new_tr_x = OxmlElement("w:tr")
            row._tr.addnext(new_tr_x)
            x_row_data = (
                "PD2 aperture centre X (F-04)",
                "33.0 mm",
                "±0.2 mm",
                "PD2 optical alignment — match FPC contact pad centre (RISK-14)",
                "YES"
            )
            for idx, (txt, w) in enumerate(zip(x_row_data, dim_widths)):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(w * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el)
                tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz")
                sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el)
                r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve")
                t_el.text = txt
                r_el.append(t_el)
                p_el.append(r_el)
                tc.append(p_el)
                new_tr_x.append(tc)
            # Insert Y row after X row
            new_tr_y = OxmlElement("w:tr")
            new_tr_x.addnext(new_tr_y)
            y_row_data = (
                "PD2 aperture centre Y (F-04)",
                "39.0 mm",
                "±0.2 mm",
                "PD2 optical alignment — match FPC contact pad centre (RISK-14)",
                "YES"
            )
            for idx, (txt, w) in enumerate(zip(y_row_data, dim_widths)):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(w * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el)
                tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz")
                sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el)
                r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve")
                t_el.text = txt
                r_el.append(t_el)
                p_el.append(r_el)
                tc.append(p_el)
                new_tr_y.append(tc)
            # Apply row formatting
            for r2 in dims_tbl.rows:
                if r2._tr is new_tr_x or r2._tr is new_tr_y:
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
            print("  Critical dims: PD2 aperture centre X and Y rows added")
            break

# 2c. Close OI-3 (Table 6)
oi_tbl = tool.tables[6]
for row in oi_tbl.rows:
    if row.cells[0].text.strip() == "OI-3":
        set_text(row.cells[4],
            "CLOSED — NP-HW-FPC-001 Rev 4 frozen 2026-05-09. "
            "PD2 at X = 33.0 mm, Y = 39.0 mm from module reference corner. "
            "See NP-HW-FPC-001 Rev 4 §8.4 and pin 19 notes. "
            "F-04 position handed to mould designer.",
            bold=True, color=(0x00, 0x70, 0x00), size=8)
        set_bg(row.cells[4], "E2EFDA")
        print("  OI-3 closed")
        break

# 2d. Update mould design review checklist item for F-04 (Table 4)
# Item 4 should reference F-04 and include the XY co-ord in the Notes
mdr_tbl = tool.tables[4]
for row in mdr_tbl.rows:
    cells = row.cells
    if (len(cells) >= 2 and cells[0].text.strip() == "4"
            and ("F-04" in cells[1].text or "PD2" in cells[1].text or "aperture" in cells[1].text.lower())):
        # Update Notes column if present, else update requirement column
        if len(cells) >= 4:
            set_text(cells[3],
                "F-04 aperture centre: X = 33.0 ± 0.2 mm, Y = 39.0 ± 0.2 mm "
                "from module reference corner. Confirm in mould drawing. "
                "Position locked per NP-HW-FPC-001 Rev 4 §8.4 (2026-05-09).",
                size=8)
        print("  MDR checklist item 4 (F-04) notes updated")
        break

tool.save(TOOL_PATH)
print(f"  NP-TOOL-ZM-001 saved")


# ══════════════════════════════════════════════════════════════════════════════
# 3. NP-DRV-SHELL-001 — DRC-23 note update
# ══════════════════════════════════════════════════════════════════════════════
SHELL_PATH = "docs/superseded/neurone_shell_fpc_routing_review.docx"
shell = Document(SHELL_PATH)

for tbl in shell.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 1 and cells[0].text.strip() == "DRC-23":
            # Update owner/status cell (usually last) to note FPC layout now available
            last_c = cells[-1]
            cur = last_c.text.strip()
            if "OPEN" in cur or not cur:
                set_text(last_c,
                    "OPEN — FPC layout frozen (NP-HW-FPC-001 Rev 4, 2026-05-09). "
                    "EEG routing strategy per §2.4.2 confirmed (dedicated 8 × 5 mm "
                    "channel, opposite CFRP surface from FPC bundle). "
                    "CAD overlay measurement (DRC-18a) to proceed at G2 gate when "
                    "shell CAD model is available (OI-09, G2-11).",
                    size=8)
                set_bg(last_c, "FFF2CC")
                print("  DRC-23 note updated (FPC layout now available)")
            break

shell.save(SHELL_PATH)
print(f"  NP-DRV-SHELL-001 saved")


# ══════════════════════════════════════════════════════════════════════════════
# 4. NP-COORD-001 — add G1-15 layout freeze CLOSED
# ══════════════════════════════════════════════════════════════════════════════
COORD_PATH = "docs/neurone_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

# Find the summary table (Table 15 based on scan — 35 rows, 5 cols) and
# add G1-15 after the G1-14 row
SUMMARY_TBL_IDX = None
for i, tbl in enumerate(coord.tables):
    if (len(tbl.columns) == 5
            and any(row.cells[0].text.strip() == "G1-14" for row in tbl.rows)):
        SUMMARY_TBL_IDX = i
        break

if SUMMARY_TBL_IDX is not None:
    summ_tbl = coord.tables[SUMMARY_TBL_IDX]
    # Check if G1-15 already exists
    has_g115 = any(row.cells[0].text.strip() == "G1-15" for row in summ_tbl.rows)
    if not has_g115:
        for row in summ_tbl.rows:
            if row.cells[0].text.strip() == "G1-14":
                new_tr = OxmlElement("w:tr")
                row._tr.addnext(new_tr)
                g115_data = (
                    "G1-15",
                    "G1",
                    "HW EE",
                    "FPC layout freeze + PD2 aperture position (Issue #10): "
                    "NP-HW-FPC-001 Rev 4 released with PD2 at X = 33.0 mm, "
                    "Y = 39.0 mm. F-04 in NP-TOOL-ZM-001 updated. OI-3 closed. "
                    "Enables mould design review (NP-TOOL-ZM-001 §5). "
                    "DRC-18 CAD verification to follow at G2 (G2-11).",
                    "CLOSED — 2026-05-09"
                )
                g115_widths = [0.55, 0.4, 0.9, 4.5, 1.1]
                for idx, (txt, w) in enumerate(zip(g115_data, g115_widths)):
                    tc = OxmlElement("w:tc")
                    tc_pr = OxmlElement("w:tcPr")
                    w_el = OxmlElement("w:tcW")
                    w_el.set(qn("w:w"), str(int(w * 1440)))
                    w_el.set(qn("w:type"), "dxa")
                    tc_pr.append(w_el)
                    tc.append(tc_pr)
                    p_el = OxmlElement("w:p")
                    r_el = OxmlElement("w:r")
                    rpr_el = OxmlElement("w:rPr")
                    sz_el = OxmlElement("w:sz")
                    sz_el.set(qn("w:val"), "16")
                    rpr_el.append(sz_el)
                    r_el.append(rpr_el)
                    t_el = OxmlElement("w:t")
                    t_el.set(qn("xml:space"), "preserve")
                    t_el.text = txt
                    r_el.append(t_el)
                    p_el.append(r_el)
                    tc.append(p_el)
                    new_tr.append(tc)
                # Format new row
                for r2 in summ_tbl.rows:
                    if r2._tr is new_tr:
                        set_bg(r2.cells[4], "E2EFDA")
                        if r2.cells[0].paragraphs[0].runs:
                            r2.cells[0].paragraphs[0].runs[0].bold = True
                        break
                print("  G1-15 layout freeze row added to summary table")
                break

coord.save(COORD_PATH)
print(f"  NP-COORD-001 saved")

print("\nDone — all documents updated for Issue #10 FPC layout freeze.")
