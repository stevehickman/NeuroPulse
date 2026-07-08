"""
Add HRV biofeedback protocol to:
  1. Design brief Rev 4 — VNS+HRV completed-decision bullet + pending item removed
  2. Additional modalities doc — mark HRV biofeedback as ADDED/LOCKED
  3. FAI checklist — add HRV biofeedback test items (FAI-H01 through FAI-H04)
  4. Coordination checklist — add G2-12 (HRV biofeedback firmware spec)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, size=8, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def insert_para_after(anchor_p, text, size=9, bold=False):
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    new_rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2)))
    new_rpr.append(sz)
    if bold:
        b_el = OxmlElement("w:b")
        new_rpr.append(b_el)
    new_r.append(new_rpr)
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = text
    new_r.append(new_t)
    new_p.append(new_r)
    anchor_p.addnext(new_p)
    return new_p

def add_row_after(anchor_tr, row_data, col_widths_in, bold_col0=True):
    new_tr = OxmlElement("w:tr")
    for c_idx, txt in enumerate(row_data):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        w_el = OxmlElement("w:tcW")
        w_el.set(qn("w:w"), str(int(col_widths_in[c_idx] * 1440)))
        w_el.set(qn("w:type"), "dxa")
        tc_pr.append(w_el); tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr_el = OxmlElement("w:rPr")
        sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
        rpr_el.append(sz_el)
        if c_idx == 0 and bold_col0:
            b_el = OxmlElement("w:b"); rpr_el.append(b_el)
        r_el.append(rpr_el)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
        r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
        new_tr.append(tc)
    anchor_tr.addnext(new_tr)
    return new_tr

# ══════════════════════════════════════════════════════════════════════════════
# 1. Design Brief Rev 4
# ══════════════════════════════════════════════════════════════════════════════
BRIEF_PATH = "docs/neurone_brief_r4.docx"
brief = Document(BRIEF_PATH)

# Add to completed decisions section (after para 70 — auricular VNS clip)
hrv_inserted = False
for pi, para in enumerate(brief.paragraphs):
    if "Auricular VNS clip = HRV PPG" in para.text:
        insert_para_after(
            para._p,
            "HRV biofeedback protocol (software only, $0 BOM): resonance frequency "
            "breathing pacer (6 breaths/min default, personalised per user); real-time "
            "coherence score (LF peak / LF+HF power); four protocols: standalone coherence "
            "training, HRV+taVNS synchronised (stimulation timed to inspiration phase), "
            "HRV+EEG dual biofeedback, HRV+PBM simultaneous (replicates 2025 multi-modal RCT). "
            "Bone conduction delivers breathing cue. UHDR: HRV time series + session logs.",
            size=9
        )
        hrv_inserted = True
        print("  Brief: HRV biofeedback note appended after VNS/HRV bullet")
        break

if not hrv_inserted:
    print("  WARNING: VNS/HRV paragraph not found in brief")

brief.save(BRIEF_PATH)
print(f"  Design brief saved")

# ══════════════════════════════════════════════════════════════════════════════
# 2. Additional Modalities Doc — mark HRV biofeedback as ADDED/LOCKED
# ══════════════════════════════════════════════════════════════════════════════
MOD_PATH = "docs/neurone_additional_modalities.docx"
mod = Document(MOD_PATH)

for tbl in mod.tables:
    for row in tbl.rows:
        full = ' '.join(c.text for c in row.cells)
        if "HRV" in full and "biofeedback" in full.lower():
            # Find the header cell and mark it
            for cell in row.cells:
                if "HRV Biofeedback" in cell.text:
                    orig = cell.paragraphs[0]
                    orig_text = orig.text
                    for run in orig.runs:
                        run.text = ""
                    run = orig.add_run("✓ ADDED TO T1 STACK — " + orig_text.replace("4. ", ""))
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                    set_bg(cell, "E2EFDA")
                    print("  Modalities doc: HRV biofeedback marked as ADDED")
                    break

# Also scan paragraphs for the priority/action text and update
for para in mod.paragraphs:
    if "HIGH — no hardware cost" in para.text and "HRV" in para.text:
        for run in para.runs:
            if "HIGH — no hardware cost" in run.text:
                run.text = run.text.replace(
                    "HIGH — no hardware cost",
                    "ADDED TO T1 STACK — was HIGH priority, now implemented"
                )
                run.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                print("  Modalities doc: priority note updated")
                break

mod.save(MOD_PATH)
print(f"  Additional modalities doc saved")

# ══════════════════════════════════════════════════════════════════════════════
# 3. FAI checklist — add HRV biofeedback test section
# ══════════════════════════════════════════════════════════════════════════════
FAI_PATH = "docs/neurone_fai_zone_module.docx"
fai = Document(FAI_PATH)

# Find FAI-A15 (last accessibility item) or last item in accessibility section
# and append HRV biofeedback items after it
fai_a15_tr = None
for tbl in fai.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "FAI-A15":
            fai_a15_tr = row._tr
            anchor_tbl = tbl
            print(f"  FAI: found FAI-A15, will add HRV items after")
            break
    if fai_a15_tr:
        break

if fai_a15_tr:
    # Determine column widths from existing table structure
    fai_widths = [0.65, 2.3, 1.2, 1.55, 0.65, 0.45]
    HRV_ITEMS = [
        ("FAI-H01",
         "HRV coherence baseline (standalone): subject performs 5-minute guided breathing "
         "session at 6 breaths/min. App displays real-time coherence score. Verify: PPG "
         "R-R interval successfully detected from auricular clip (>95% beat detection); "
         "coherence score updates at ≤2s interval; bone conduction audio breathing cue "
         "audible at ≥60dB at ear.",
         "App + firmware verification",
         "Beat detection ≥95% (validated against ECG reference in bench test). Coherence "
         "score range 0–10 correctly calculated. Audio cue timing: ±0.1s of target.",
         "", ""),
        ("FAI-H02",
         "HRV+taVNS synchronisation: 10-minute session, HRV+taVNS synchronised protocol. "
         "Verify firmware correctly identifies inspiration phase from PPG waveform and "
         "triggers taVNS pulse within ±150ms of peak inspiration for ≥80% of respiratory "
         "cycles. Subject: 3 healthy adults.",
         "3 subjects, bench + human verification",
         "Inspiration-triggered stimulation latency ≤150ms in ≥80% of cycles. "
         "No taVNS stimulation during expiration-only phases. No subject reports "
         "unexpected sensation from mis-timed stimulation.",
         "", ""),
        ("FAI-H03",
         "HRV+EEG dual biofeedback: 10-minute session, both coherence score and EEG "
         "alpha/theta band power displayed simultaneously. Verify: both metrics update "
         "independently without interference; EEG signal quality unchanged when taVNS "
         "is running concurrently (signal-to-noise check).",
         "App + signal quality verification",
         "Both metrics update at ≤2s intervals. EEG SNR during concurrent taVNS "
         "≥90% of baseline (taVNS artefact <5µV RMS after filtering).",
         "", ""),
        ("FAI-H04",
         "HRV+PBM multi-modal session: 15-minute combined session (PBM zone 1–5 at "
         "standard dose + HRV coherence training simultaneously). Verify: PBM dosimetry "
         "unaffected by concurrent HRV session; HRV coherence scores in expected range "
         "(≥5.0 for subjects trained in resonance breathing); no thermal interaction "
         "between PBM NTC monitoring and PPG signal processing.",
         "3 subjects with prior HRV biofeedback experience",
         "PBM J/cm² dose within ±5% of target. HRV coherence ≥5.0 in all 3 subjects "
         "(indicates usable biofeedback loop). NTC temperature reporting unaffected.",
         "", ""),
    ]

    prev_tr = fai_a15_tr
    for item in HRV_ITEMS:
        new_tr = OxmlElement("w:tr")
        for c_idx, txt in enumerate(item):
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            w_el = OxmlElement("w:tcW")
            w_el.set(qn("w:w"), str(int(fai_widths[c_idx] * 1440)))
            w_el.set(qn("w:type"), "dxa")
            tc_pr.append(w_el); tc.append(tc_pr)
            p_el = OxmlElement("w:p")
            r_el = OxmlElement("w:r")
            rpr_el = OxmlElement("w:rPr")
            sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
            rpr_el.append(sz_el); r_el.append(rpr_el)
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
            r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
            new_tr.append(tc)
        prev_tr.addnext(new_tr)
        # Style: grey background for HRV items
        for r2 in anchor_tbl.rows:
            if r2._tr is new_tr:
                for ci in range(6):
                    set_bg(r2.cells[ci], "F0F7FF")
                if r2.cells[0].paragraphs[0].runs:
                    r2.cells[0].paragraphs[0].runs[0].bold = True
                break
        prev_tr = new_tr
        print(f"  FAI: {item[0]} inserted")

fai.save(FAI_PATH)
print(f"  FAI checklist saved")

# ══════════════════════════════════════════════════════════════════════════════
# 4. Coordination checklist — add G2-12 (HRV biofeedback firmware spec)
# ══════════════════════════════════════════════════════════════════════════════
COORD_PATH = "docs/neurone_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

# Find G2-11 in the tables and add G2-12 after it
for tbl in coord.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "G2-11":
            g212_data = (
                "G2-12", "G2", "FW + HW EE",
                "HRV biofeedback firmware spec: resonance breathing detection, coherence "
                "algorithm (LF peak / LF+HF), inspiration-phase taVNS trigger (±150ms), "
                "dual EEG+HRV display protocol, bone conduction breathing cue timing. "
                "FAI-H01 through FAI-H04 pass before G2 gate sign-off.",
                "OPEN — firmware specification to be written"
            )
            widths = [0.65, 0.55, 1.3, 2.8, 1.4]
            new_tr = add_row_after(row._tr, g212_data, widths)
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[4], "FFF2CC")
                    break
            print("  Coord checklist: G2-12 added after G2-11")
            break

# Also add G2-12 to dashboard (Table 15)
# Find Table 15 or whichever table is the dashboard
for tbl in coord.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "G2-11":
            # Check if this looks like dashboard (5 cols)
            if len(row.cells) == 5:
                # This is the dashboard - G2-11 is already here, add G2-12 after
                g212_dash = (
                    "G2-12", "G2", "FW + HW EE",
                    "HRV biofeedback firmware spec — coherence algo, taVNS timing, dual EEG+HRV display",
                    "OPEN — spec to be written (FAI-H01–H04)"
                )
                new_tr = add_row_after(row._tr, g212_dash, widths)
                for r2 in tbl.rows:
                    if r2._tr is new_tr:
                        set_bg(r2.cells[4], "FFF2CC")
                        break
                print("  Coord dashboard: G2-12 added")
                break

coord.save(COORD_PATH)
print(f"  Coordination checklist saved")

print("\nDone.")
