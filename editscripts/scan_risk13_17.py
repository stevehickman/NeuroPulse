"""Scan RISK-13 through RISK-17 and all open issues for engineering coordination items."""
from docx import Document

risks = Document("docs/neurone_fpc_zone_module_risks_revA.docx")

print("=== RISK-13 through RISK-17 (summary table) ===")
tbl = risks.tables[0]
for row in tbl.rows:
    rid = row.cells[0].text.strip()
    if rid in ("RISK-13","RISK-14","RISK-15","RISK-16","RISK-17"):
        headers = ["ID","Severity","Title","Root Cause","Impact","Required Mitigation","Owner","Spec §","Status"]
        for h, cell in zip(headers, row.cells):
            print(f"  {h}: {cell.text.strip()!r}")
        print()

print("=== Open Issues table ===")
oi_tbl = risks.tables[1]
for row in oi_tbl.rows:
    texts = [c.text.strip() for c in row.cells]
    if any(texts):
        print(f"  {texts}")

print()
print("=== Detailed paragraphs for RISK-13 onward ===")
for i, para in enumerate(risks.paragraphs):
    t = para.text.strip()
    if t and (any(f"RISK-{n}" in t for n in range(13,18)) or "DETAILED ANALYSIS" in t):
        print(f"  [{i}] {t[:120]!r}")
        for j in range(i+1, min(i+12, len(risks.paragraphs))):
            p2 = risks.paragraphs[j]
            if p2.text.strip():
                print(f"  [{j}] {p2.text.strip()[:120]!r}")
            if any(f"RISK-{n}" in p2.text for n in range(13,18)) and j > i+1:
                break
