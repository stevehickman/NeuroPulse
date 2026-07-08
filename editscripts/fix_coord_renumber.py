"""
Renumber coord checklist items to avoid ID conflicts:
  G1-07 (ZONE_ID debounce, Table 6) → G1-13
  G1-08 (PDMS qual, Table 6)        → G1-14
Add G1-13 and G1-14 to dashboard (Table 15) after G1-12.
Update cross-references in risk register and FPC spec.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

# ── 1. Coord checklist: renumber in Table 6 ───────────────────────────────────
COORD_PATH = "docs/neurone_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

# Table 6: find the two items in the right section
# G1-07 (Table 6) = ZONE_ID debounce → G1-13
# G1-08 (Table 6) = PDMS qual        → G1-14
tbl6 = coord.tables[6]
for row in tbl6.rows:
    cells = row.cells
    rid = cells[0].text.strip()
    if rid == "G1-07" and len(cells) >= 3 and "debounce" in cells[2].text.lower():
        set_text(cells[0], "G1-13", bold=True, size=8)
        print("  Table 6: G1-07 (ZONE_ID debounce) → G1-13")
    elif rid == "G1-08" and len(cells) >= 3 and "PDMS" in cells[2].text:
        set_text(cells[0], "G1-14", bold=True, size=8)
        print("  Table 6: G1-08 (PDMS qual) → G1-14")

# ── 2. Dashboard (Table 15): add G1-13 and G1-14 after G1-12 ─────────────────
tbl15 = coord.tables[15]
g112_tr = None
for row in tbl15.rows:
    if row.cells[0].text.strip() == "G1-12":
        g112_tr = row._tr
        break

if g112_tr:
    NEW_ROWS = [
        ("G1-14", "G1", "Mfg Eng + HW EE",
         "PDMS thermal cycling qual (FAI-TC02) BLOCKING before production",
         "OPEN — CAT-C supplier not yet selected"),
        ("G1-13", "G1", "EE",
         "ZONE_ID firmware debounce spec (3×ADC / RISK-18)",
         "OPEN — firmware req to be written"),
    ]
    widths_in = [0.65, 0.55, 1.3, 2.8, 1.4]
    for row_data in NEW_ROWS:   # reversed order so G1-13 ends up above G1-14
        new_tr = OxmlElement("w:tr")
        for c_idx, txt in enumerate(row_data):
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            w_el = OxmlElement("w:tcW")
            w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
            w_el.set(qn("w:type"), "dxa")
            tc_pr.append(w_el); tc.append(tc_pr)
            p_el = OxmlElement("w:p")
            r_el = OxmlElement("w:r")
            rpr_el = OxmlElement("w:rPr")
            sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
            rpr_el.append(sz_el); r_el.append(rpr_el)
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
            r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
            new_tr.append(tc)
        g112_tr.addnext(new_tr)
        for r2 in tbl15.rows:
            if r2._tr is new_tr:
                set_bg(r2.cells[4], "FFF2CC")
                if r2.cells[0].paragraphs[0].runs:
                    r2.cells[0].paragraphs[0].runs[0].bold = True
                break
        print(f"  Dashboard: added {row_data[0]}")

# ── 3. Verify final ID counts ─────────────────────────────────────────────────
print("\nFinal item counts (G1-07 through G1-14):")
for check_id in ("G1-07", "G1-08", "G1-13", "G1-14"):
    count = sum(1 for tbl in coord.tables
                for row in tbl.rows
                if row.cells[0].text.strip() == check_id)
    print(f"  {check_id}: {count}")

coord.save(COORD_PATH)
print(f"\nNP-COORD-001 saved")

# ── 4. Update risk register cross-refs ────────────────────────────────────────
# RISK-18 mitigation spec ref: G2-07 → keep (was correct), but update "G1-07"
# reference if it appears in RISK-18 or RISK-04 entries
RISK_PATH = "docs/neurone_fpc_zone_module_risks_revA.docx"
risk = Document(RISK_PATH)

for tbl in risk.tables:
    for row in tbl.rows:
        cells = row.cells
        rid = cells[0].text.strip()
        if rid == "RISK-18" and len(cells) >= 8:
            existing = cells[7].text
            if "G1-07" in existing:
                set_text(cells[7], existing.replace("G1-07", "G1-13"), size=8)
                print("  Risk register RISK-18 spec ref: G1-07 → G1-13")
        if rid == "RISK-04" and len(cells) >= 8:
            existing = cells[7].text
            if "G1-08" in existing:
                set_text(cells[7], existing.replace("G1-08", "G1-14"), size=8)
                print("  Risk register RISK-04 spec ref: G1-08 → G1-14")

risk.save(RISK_PATH)
print(f"Risk register saved")

# ── 5. Update FPC spec ZONE_ID paragraph ─────────────────────────────────────
FPC_PATH = "docs/neurone_fpc_zone_module_spec_revA.docx"
fpc = Document(FPC_PATH)
changed = 0
for p in fpc.paragraphs:
    if "G1-07" in p.text or "G2-07" in p.text:
        for run in p.runs:
            if "G1-07" in run.text:
                run.text = run.text.replace("G1-07", "G1-13")
                changed += 1
if changed:
    fpc.save(FPC_PATH)
    print(f"FPC spec: replaced G1-07 → G1-13 ({changed} occurrence(s))")
else:
    print("FPC spec: no G1-07 references found")

print("Done.")
