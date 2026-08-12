"""Get full content of §11.2 in FPC spec."""
from docx import Document

spec = Document("docs/superseded/np_hw_fpc_001.docx")
for i, para in enumerate(spec.paragraphs):
    t = para.text.strip()
    if "11.2" in t:
        print(f"[{i}] {t!r}")
        for j in range(i+1, min(i+8, len(spec.paragraphs))):
            p2 = spec.paragraphs[j]
            if p2.text.strip():
                print(f"[{j}] {p2.text.strip()!r}")
            if "11.3" in p2.text:
                break
