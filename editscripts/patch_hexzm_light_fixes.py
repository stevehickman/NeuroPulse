"""
Small, mechanical terminology fixes for the hex-tile/socket architecture
(NP-HEX-ZM-001) that don't require re-scoping or a supersession note —
just "zone module" -> "hex-tile module" generalization where the underlying
engineering concept (per-module gasket, per-module dose metering) carries
over unchanged, per docs/np_hex_zm_001.md §6 ("each hex needs a perimeter
gasket... a per-tile seam-length budget is required").

1. RISK-16 (docs/superseded/neurone_fpc_zone_module_risks_revA.docx) — IPX4
   field-replacement seal. The self-sealing compression gasket approach is
   architecture-independent; only the noun needs updating.

2. docs/neurone_clinical_trials_strategy.docx — "5 independently addressable
   zones" asserted a retired, fixed slot count as a marketing differentiator.
   The real differentiator (individually addressable modules with real-time
   per-module dose metering) is still true and is in fact stronger under the
   hex-tile lattice (~80 addressable positions, not 5) — but this script does
   not assert a specific socket count here, since the active-surface boundary
   (gate ACT-1) and REG-1 registration are still open per
   docs/np_hex_zm_001.md §7. "600 LEDs" is also dropped: CLAUDE.md §3
   modality 1 states LED count scales with how many T1-A tiles are populated
   in a given build, not a fixed number.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RISK_PATH = "docs/superseded/neurone_fpc_zone_module_risks_revA.docx"
CLINICAL_PATH = "docs/neurone_clinical_trials_strategy.docx"

OLD_CLINICAL = (
    "Device equivalence: Vielight RCTs used Vielight devices at specific irradiance, "
    "duty cycle, and LED count. NeurOne (600 LEDs, dual-wavelength, 400 mW/cm² "
    "pulsed, 5 independently addressable zones) has no published comparator. Any "
    "marketing claim implying NeurOne produces outcomes 'at least as good as Vielight' "
    "requires head-to-head data or regulatory counsel opinion."
)
NEW_CLINICAL = (
    "Device equivalence: Vielight RCTs used Vielight devices at specific irradiance, "
    "duty cycle, and LED count. NeurOne (dual-wavelength, 400 mW/cm² pulsed, "
    "individually addressable hex-tile PBM modules across the socket lattice with "
    "real-time per-module J/cm² dose metering) has no published comparator. Any "
    "marketing claim implying NeurOne produces outcomes 'at least as good as Vielight' "
    "requires head-to-head data or regulatory counsel opinion."
)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def fix_risk16():
    doc = Document(RISK_PATH)
    tbl = doc.tables[0]
    for row in tbl.rows:
        if row.cells[0].text.strip() == "RISK-16":
            if "hex-tile module" in row.cells[3].text:
                print("  RISK-16 already patched — skipping")
                return
            for idx in (2, 3, 4, 5):
                old = row.cells[idx].text
                new = old.replace("zone module", "hex-tile module").replace("Zone module", "Hex-tile module")
                if new != old:
                    set_cell_text(row.cells[idx], new)
            print("  RISK-16 terminology generalized to hex-tile module")
            break
    doc.save(RISK_PATH)


def fix_clinical():
    doc = Document(CLINICAL_PATH)
    changed = False
    for p in doc.paragraphs:
        if OLD_CLINICAL in p.text:
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = NEW_CLINICAL
            else:
                p.add_run(NEW_CLINICAL)
            changed = True
            print("  Clinical trials doc: '5 independently addressable zones' claim corrected")
            break
    if not changed:
        print("  Clinical trials doc: target text not found (already patched?)")
        return
    doc.save(CLINICAL_PATH)


if __name__ == "__main__":
    fix_risk16()
    fix_clinical()
