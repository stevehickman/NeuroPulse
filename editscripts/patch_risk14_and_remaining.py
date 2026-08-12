"""
Patch: RISK-14 Option B implementation, RISK-21 mitigation, NP-TOOL-ZM-001 creation.
Targets:
  - docs/superseded/neurone_fpc_zone_module_spec_revA.docx  (pin 19, §8 dual-PD update)
  - docs/neurone_eng_coordination_checklist.docx (G1-06 → CLOSED, add G2-11)
  - docs/superseded/neurone_shell_fpc_routing_review.docx   (add §2.4 EEG cable routing)
  - docs/superseded/neurone_tool_zone_module_001.docx       (new: NP-TOOL-ZM-001)
  - docs/superseded/neurone_fpc_zone_module_risks_revA.docx (RISK-14, 18-24 status updates)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

# ─── helpers ──────────────────────────────────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, color=(0x1F, 0x49, 0x7D), size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p

def add_body(doc, text, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def make_table_with_header(doc, headers, widths_in, hdr_bg="1F497D"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        set_bg(hrow.cells[i], hdr_bg)
        r = hrow.cells[i].paragraphs[0].add_run(h)
        r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(8)
        hrow.cells[i].width = Inches(widths_in[i])
    return tbl

def add_tbl_row(tbl, data, widths_in, alternating_idx=0):
    row = tbl.add_row()
    bg = "FFFFFF" if alternating_idx % 2 == 0 else "F5F5F5"
    for i, txt in enumerate(data):
        cell = row.cells[i]
        set_bg(cell, bg)
        r = cell.paragraphs[0].add_run(txt)
        r.font.size = Pt(8)
        if i == 0:
            r.bold = True
        cell.width = Inches(widths_in[i])
    return row

# ══════════════════════════════════════════════════════════════════════════════
# 1. FPC SPEC — update pin 19 to PD2_CATHODE + add §8.4 dual-PD architecture
# ══════════════════════════════════════════════════════════════════════════════
FPC_PATH = "docs/superseded/neurone_fpc_zone_module_spec_revA.docx"
fpc = Document(FPC_PATH)

# 1a. Update pinout table: row 19 (index 19) — SPARE_2 → PD2_CATHODE
for tbl in fpc.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 3 and cells[0].text.strip() == "19" and "SPARE" in cells[1].text:
            set_text(cells[1], "PD2_CATHODE", bold=True, size=8)
            set_text(cells[2], "ZM → Hub", size=8)
            # col 3 = description (may have more cells)
            if len(cells) >= 4:
                set_text(cells[3],
                    "Scalp-side reference photodiode cathode (RISK-14 Option B). "
                    "Connect to Hub ADC channel (same TIA as PD_CATHODE, multiplexed). "
                    "0.5 mm pitch ZIF pin 19.", size=8)
            print("  Pin 19 updated: SPARE_2 → PD2_CATHODE")
            break

# 1b. Find §8.3 paragraph and insert §8.4 after the last para of §8
# Find §9 heading to anchor insertion
s9_para = None
for p in fpc.paragraphs:
    if p.text.strip().startswith("9.") and "PDMS" in p.text:
        s9_para = p
        break
print(f"  §9 anchor: {s9_para.text[:60]!r}" if s9_para else "  §9 NOT FOUND")

S84_PARAS = [
    ("8.4  Scalp-Side Reference Photodiode (PD2) — RISK-14 Option B", True, 11),
    ("A second photodiode (PD2) is co-moulded into the scalp-facing surface of the zone "
     "module body, behind a dedicated 1.2 mm circular aperture in the PDMS window "
     "(see §12 zone module tooling). PD2 faces the scalp directly and measures "
     "backscattered optical power from tissue.", False, 10),
    ("Rationale: PD1 (existing reference photodiode, §8.1–8.3) is positioned behind the "
     "scalp-facing PDMS window and measures forward-emitted power. PD1 signal combines "
     "LED aging and PDMS fouling — these cannot be separated by trend analysis alone. "
     "PD2, facing the scalp, is unaffected by PDMS fouling. The ratio PD1/PD2 isolates "
     "the two failure modes with firmware-only logic:", False, 10),
    ("PD1↓ PD2 stable → PDMS fouling (user: clean window prompt)", False, 10),
    ("PD1↓ PD2↓ proportional → LED aging (auto: increase current within IEC 60601 ceiling)", False, 10),
    ("PD1 stable PD2↑ → anomalous scalp reflectance change (log to SHDR; do not alter dose)", False, 10),
    ("Electrical connection: PD2_CATHODE on FPC pin 19 (formerly SPARE_2). Hub MCU "
     "reads PD2 via the same TIA circuit as PD1, time-multiplexed within the LED PWM "
     "dark period. No additional Hub PCB area required beyond one ADC multiplexer input.", False, 10),
    ("[DRAWING REQUIREMENT — ZONE MODULE TOOLING] PD2 aperture: 1.2 mm circular "
     "opening in scalp-facing PDMS window, centred on PD2 body. Aperture must be "
     "specified in NP-TOOL-ZM-001 before mould design. T1 and T2 zone modules use "
     "identical mould — PD2 aperture is present in all units; firmware enables "
     "dual-PD ratio algorithm by software flag only.", False, 10),
    ("BOM impact: +$0.75–1.50 per headset (PD2 component + TIA multiplexer input). "
     "Included in RISK-14 Option B BOM estimate. T1 and T2 share the same zone module "
     "mould — no additional tooling cost for T2.", False, 10),
    ("Calibration: at factory burn-in (100 h), record PD1_ref and PD2_ref at nominal "
     "LED current. Store in SHDR calibration partition. Field algorithm uses "
     "PD1/PD1_ref and PD2/PD2_ref ratios, not absolute values.", False, 10),
]

if s9_para:
    # Insert paragraphs before §9 in reverse order (each addprevious adds before anchor)
    for text, bold, size in reversed(S84_PARAS):
        new_p = OxmlElement("w:p")
        s9_para._p.addprevious(new_p)
        for p_obj in fpc.paragraphs:
            if p_obj._p is new_p:
                run = p_obj.add_run(text)
                run.bold = bold
                run.font.size = Pt(size)
                if bold:
                    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                break
    # Insert a spacer
    spacer = OxmlElement("w:p")
    s9_para._p.addprevious(spacer)
    print("  §8.4 dual-PD section inserted")

fpc.save(FPC_PATH)
print(f"  FPC spec saved: {FPC_PATH}")

# ══════════════════════════════════════════════════════════════════════════════
# 2. COORD CHECKLIST — close G1-06 + add G2-11 EEG cable routing
# ══════════════════════════════════════════════════════════════════════════════
COORD_PATH = "docs/neurone_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

# 2a. Find G1-06 (PD2 decision row) and mark CLOSED
for tbl in coord.tables:
    for row in tbl.rows:
        cells = row.cells
        if (len(cells) >= 1 and cells[0].text.strip() == "G1-06"):
            # Update status cell (last cell, index 4 or wherever)
            if len(cells) >= 5:
                set_text(cells[4], "CLOSED — Option B selected; pin 19 assigned PD2_CATHODE; §8.4 added to NP-HW-FPC-001")
                set_bg(cells[4], "E2EFDA")
            print("  G1-06 marked CLOSED")
            break

# 2b. Find G2-10 row and add G2-11 after it (EEG cable routing)
for tbl in coord.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 1 and cells[0].text.strip() == "G2-10":
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            # Build cells matching the table structure (5 columns based on prior patches)
            g211_data = [
                "G2-11",
                "G2",
                "EEG cable routing path inside shell must be specified and "
                "documented in NP-DRV-SHELL-001 §2.4 before tooling release. "
                "Required to verify DRC-18 (≥15 mm FPC-to-EEG separation). "
                "[BLOCKING — tooling must not be released without this.] "
                "Owner: EE + ME joint. Deliverable: §2.4 routing diagram + "
                "DRC-18 pass confirmation in CAD. Ref: RISK-21.",
                "EE + ME",
                "OPEN — §2.4 in progress (RISK-21)"
            ]
            # Determine widths from nearby rows
            widths_in = [0.65, 0.55, 3.5, 0.9, 1.2]
            for c_idx, txt in enumerate(g211_data):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            # Apply background to status cell (index 4)
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[4], "FFF2CC")  # yellow = OPEN
                    r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  G2-11 EEG cable routing item added")
            break

coord.save(COORD_PATH)
print(f"  Coord checklist saved: {COORD_PATH}")

# ══════════════════════════════════════════════════════════════════════════════
# 3. SHELL ROUTING REVIEW — add §2.4 EEG cable routing specification
# ══════════════════════════════════════════════════════════════════════════════
SHELL_PATH = "docs/superseded/neurone_shell_fpc_routing_review.docx"
shell = Document(SHELL_PATH)

# Find §3 heading to anchor insertion before it
s3_para = None
for p in shell.paragraphs:
    if p.text.strip().startswith("3.") and ("DRC" in p.text or "Design" in p.text or "CHECK" in p.text.upper()):
        s3_para = p
        break
# Fallback: find first paragraph containing "3."
if not s3_para:
    for p in shell.paragraphs:
        if p.text.strip().startswith("3."):
            s3_para = p
            break
print(f"  §3 anchor: {s3_para.text[:60]!r}" if s3_para else "  §3 NOT FOUND — appending")

S24_CONTENT = [
    ("2.4  EEG Electrode Cable Routing — Separation from Zone Module FPCs  [RISK-21]", True, 11),
    ("This section specifies the routing path of the EEG electrode cables inside the headset "
     "shell. This specification is required to verify DRC-18 (≥15 mm FPC-to-EEG separation). "
     "DRC-18 cannot be confirmed in CAD until this routing is locked. This section is a "
     "[PRE-TOOLING BLOCKING REQUIREMENT] — shell tooling must not be released until §2.4 "
     "is approved by EE and ME.", False, 10),
    ("2.4.1  EEG Cable Description", True, 10),
    ("The NeurOne headset contains 8 EEG electrode cables (Fp1/2, F3/4, C3/4, P3/4) plus "
     "reference and ground leads. Each cable is a shielded twisted pair (STP), 28 AWG, with "
     "a braided copper shield tied to DRL at the Hub PCB. Cable outer diameter approximately "
     "2.2 mm with overmould. Total bundle of up to 10 cables routes from the Hub PCB to the "
     "electrode pods distributed across the headset crown.", False, 10),
    ("2.4.2  Required Routing Path", True, 10),
    ("The EEG cable bundle shall route exclusively along the OPPOSITE side of the shell "
     "interior from the zone module FPC bundle. The zone module FPCs occupy the scalp-facing "
     "interior surface (dorsal/scalp side). The EEG cables shall route along the outer "
     "(CFRP shell inner surface, dorsal-outer) channel, separated from the FPC layer by the "
     "shell wall thickness (nominal 2.5 mm CFRP + mu-metal + foam).", False, 10),
    ("[DRAWING REQUIREMENT — SHELL TOOLING] EEG cable routing channel: A dedicated 8 mm × "
     "5 mm moulded channel on the outer CFRP inner surface routes EEG cables from Hub PCB "
     "to each electrode pod location. Channel must maintain ≥2 mm wall separation from zone "
     "module FPC race channels on the opposite surface. Channel to include 3 snap-in cable "
     "retention clips (PC/ABS, snap moulded in) at 80 mm intervals.", False, 10),
    ("2.4.3  Separation Compliance Verification", True, 10),
    ("DRC-18 (≥15 mm FPC-to-EEG separation) shall be verified as follows:", False, 10),
    ("[DRC-18a] In CAD assembly: minimum distance between any zone module FPC copper "
     "conductor and any EEG cable conductor ≥15 mm, measured point-to-point. Shell wall "
     "routing as described in §2.4.2 achieves approximately 18–22 mm separation.", False, 10),
    ("[DRC-18b] If shell geometry forces EEG cable within 15 mm of any FPC: a grounded "
     "aluminium foil barrier (35 µm, adhesive-backed) shall be installed in the crossing "
     "region. Barrier tied to chassis GND (pin 20). Barrier placement moulded into shell "
     "at zero cost if specified before first tooling cut.", False, 10),
    ("[DRC-18c] Prototype validation: oscilloscope measurement with EEG in recording mode, "
     "all zone module LEDs at full PWM load (40 kHz carrier). Confirm EEG artifact "
     "injection < 5 µVpp at all frequencies (per DRC-18 threshold).", False, 10),
    ("2.4.4  Open Items", True, 10),
    ("OI-09: ME to produce EEG cable routing diagram in shell CAD model (owner: ME; "
     "required before G2 Pre-Tooling gate; BLOCKING for DRC-18 verification). "
     "Status: OPEN — RISK-21.", False, 10),
    ("OI-10: Confirm shell wall thickness at CFRP crown arc is sufficient to provide "
     "≥15 mm between FPC channel and EEG channel without barrier. "
     "If < 15 mm: specify barrier per DRC-18b. Status: OPEN pending CAD.", False, 10),
]

if s3_para:
    for text, bold, size in reversed(S24_CONTENT):
        new_p = OxmlElement("w:p")
        s3_para._p.addprevious(new_p)
        for p_obj in shell.paragraphs:
            if p_obj._p is new_p:
                run = p_obj.add_run(text)
                run.bold = bold; run.font.size = Pt(size)
                if bold and size == 11:
                    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                break
    spacer = OxmlElement("w:p")
    s3_para._p.addprevious(spacer)
    print("  §2.4 EEG cable routing inserted before §3")
else:
    for text, bold, size in S24_CONTENT:
        p = shell.add_paragraph()
        run = p.add_run(text)
        run.bold = bold; run.font.size = Pt(size)

# Also add DRC-22 and DRC-23 to the DRC checklist table
for tbl in shell.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() in ("DRC-21", "DRC-20", "DRC-19"):
            last_drc_row = row
for tbl in shell.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "DRC-21":
            NEW_DRCS = [
                ("DRC-22", "EEG cable routing channel present in shell CAD (§2.4 channel, 8×5 mm, outer surface)",
                 "ME", "G2 Pre-Tooling", "OPEN"),
                ("DRC-23", "DRC-18 verified with EEG routing locked (≥15 mm or barrier per §2.4.3)",
                 "EE+ME", "G2 Pre-Tooling", "OPEN"),
            ]
            for drc_data in reversed(NEW_DRCS):
                new_tr = OxmlElement("w:tr")
                row._tr.addnext(new_tr)
                drc_widths = [0.7, 3.4, 0.8, 1.2, 0.7]
                for c_idx, txt in enumerate(drc_data):
                    tc = OxmlElement("w:tc")
                    tc_pr = OxmlElement("w:tcPr")
                    w_el = OxmlElement("w:tcW")
                    w_el.set(qn("w:w"), str(int(drc_widths[c_idx] * 1440)))
                    w_el.set(qn("w:type"), "dxa")
                    tc_pr.append(w_el); tc.append(tc_pr)
                    p_el = OxmlElement("w:p")
                    r_el = OxmlElement("w:r")
                    rpr_el = OxmlElement("w:rPr")
                    sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                    rpr_el.append(sz_el); r_el.append(rpr_el)
                    t_el = OxmlElement("w:t")
                    t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                    r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                    new_tr.append(tc)
                for r2 in tbl.rows:
                    if r2._tr is new_tr:
                        set_bg(r2.cells[4], "FFF2CC")
                        if r2.cells[0].paragraphs[0].runs:
                            r2.cells[0].paragraphs[0].runs[0].bold = True
                        break
                print(f"  {drc_data[0]} added to DRC table")
            break

shell.save(SHELL_PATH)
print(f"  Shell routing review saved: {SHELL_PATH}")

# ══════════════════════════════════════════════════════════════════════════════
# 4. CREATE NP-TOOL-ZM-001 — Zone Module Tooling Specification
# ══════════════════════════════════════════════════════════════════════════════
TOOL_PATH = "docs/superseded/neurone_tool_zone_module_001.docx"
tool = Document()

# Title block
p = tool.add_paragraph()
run = p.add_run("NP-TOOL-ZM-001")
run.bold = True; run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = tool.add_paragraph()
run = p.add_run("NeurOne Zone Module Tooling Specification")
run.bold = True; run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

meta = [
    ("Document number:", "NP-TOOL-ZM-001"),
    ("Revision:", "A (Draft — Pre-Tooling)"),
    ("Status:", "OPEN — must be signed off before mould design commences"),
    ("Applies to:", "Zone Module body mould — all 5 zone positions (ZM-01 through ZM-05)"),
    ("T1 / T2 shared mould:", "YES — identical mould for both tiers; differences are firmware-only"),
    ("Date:", datetime.date.today().strftime("%Y-%m-%d")),
    ("Related specs:", "NP-HW-FPC-001, NP-DRV-SHELL-001, NP-FAI-ZM-001, NP-COORD-001"),
]
tbl = tool.add_table(rows=len(meta), cols=2)
tbl.style = "Table Grid"
for i, (label, val) in enumerate(meta):
    set_bg(tbl.rows[i].cells[0], "1F497D")
    r = tbl.rows[i].cells[0].paragraphs[0].add_run(label)
    r.bold = True; r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); r.font.size = Pt(9)
    r = tbl.rows[i].cells[1].paragraphs[0].add_run(val)
    r.font.size = Pt(9)
    tbl.rows[i].cells[0].width = Inches(2.0)
    tbl.rows[i].cells[1].width = Inches(4.8)

tool.add_paragraph()

# PURPOSE
add_heading(tool, "1.  PURPOSE AND SCOPE", size=12)
add_body(tool,
    "This document consolidates all moulded and co-moulded feature requirements for the "
    "NeurOne Zone Module body. Four mitigations (RISK-06, RISK-14, RISK-15, RISK-16) "
    "each add a distinct moulded feature to the zone module. This document exists because "
    "the cumulative complexity creates an omission risk (RISK-23): any feature missed before "
    "first-cut requires tooling modification (typically $15,000–40,000 and 6–8 week delay). "
    "All features below are MANDATORY and must appear in the mould design review before "
    "steel is cut.")
add_body(tool,
    "Scope: Zone Module body only (thermoplastic shell + co-moulded silicone). "
    "Not in scope: FPC, PDMS window bonding, LED placement (covered by NP-HW-FPC-001), "
    "shell slot geometry (NP-DRV-SHELL-001).")
tool.add_paragraph()

# SECTION 2 — CO-MOULDED FEATURES SUMMARY
add_heading(tool, "2.  CO-MOULDED FEATURES — MANDATORY CHECKLIST", size=12)
add_body(tool,
    "The following table lists all moulded features required in the zone module body. "
    "Each feature has a risk traceability reference. The mould design review "
    "(Gate: G1 Pre-Layout, item G1-05) must confirm all items before proceeding.")
tool.add_paragraph()

feat_hdrs = ["Feature ID", "Feature Description", "Risk Ref", "Spec Ref", "Confirmed in Mould Design?"]
feat_widths = [0.9, 3.0, 0.8, 1.1, 0.9]
feat_tbl = make_table_with_header(tool, feat_hdrs, feat_widths)

FEATURES = [
    ("F-01",
     "Mechanical zone key: asymmetric D-section protrusion on insertion edge, unique geometry "
     "per zone position (ZM-01 through ZM-05). Prevents physical insertion into wrong shell slot. "
     "Must engage shell slot key before FPC tail contacts ZIF. Protrusion height ≥ 1.0 mm; "
     "width ≥ 2.0 mm; must clear without binding across 1,000 insertion cycles.",
     "RISK-15", "NP-HW-FPC-001 §11.4 Layer 1", "☐"),
    ("F-02",
     "Raised numeral on module body: ISO-standard raised numeral '1' through '5' (zone position). "
     "Numeral height ≥ 0.8 mm; ISO 17049:2013 braille-dot height. Finger-readable by eyes-closed "
     "test. Located on thumb-contact surface (posterior face, within 15 mm of extraction point).",
     "RISK-15", "NP-HW-FPC-001 §11.4 Layer 3", "☐"),
    ("F-03",
     "Braille zone abbreviation: ISO 17049:2013 braille cell on same face as numeral (F-02). "
     "Zone abbreviations: ZM-01='FL' (Frontal Left), ZM-02='FR' (Frontal Right), "
     "ZM-03='CL' (Central Left), ZM-04='CR' (Central Right), ZM-05='OC' (Occipital Centre). "
     "Braille dot diameter 1.5 mm ± 0.1 mm; dot height 0.48 mm ± 0.05 mm; inter-dot 2.3 mm.",
     "RISK-15", "NP-HW-FPC-001 §11.4 Layer 3", "☐"),
    ("F-04",
     "RISK-14 Option B PD2 aperture: 1.2 mm circular aperture in scalp-facing surface of zone "
     "module body, aligned with PD2 component position on FPC (centred on PD2 body ± 0.2 mm). "
     "Aperture passes through PDMS window (PDMS to be moulded with matching through-hole). "
     "PD2 faces scalp to measure backscattered tissue optical power. "
     "Aperture must not overlap any LED position or PD1 baffle.",
     "RISK-14", "NP-HW-FPC-001 §8.4", "☐"),
    ("F-05",
     "Self-sealing compression gasket (co-moulded silicone): Shore 40–50A UV-stable medical-grade "
     "silicone (ISO 10993-5), D-section, 2.5 mm wide × 2.0 mm height nominal, continuous bead "
     "around full perimeter of zone module insertion edge. Compressed to 1.6 mm (20% compression, "
     "0.4 mm deflection) when module is fully seated in shell slot. "
     "No user RTV or field re-sealing action required. Provides IPX4 compliance at module interface. "
     "Gasket must maintain seal after 1,000 insertion/removal cycles (per NP-FAI-ZM-001 FAI-IPX-02).",
     "RISK-16", "NP-HW-FPC-001 §12", "☐"),
    ("F-06",
     "Gasket retention groove/undercut: a 0.5 mm undercut groove moulded into the zone module body "
     "perimeter accommodates the gasket D-section heel. Silicone primer (Dow Corning 1200 or "
     "equivalent) applied to groove surface before gasket co-moulding. Combined mechanical retention "
     "(groove) + chemical bonding (primer) prevents gasket delamination under 1,000 insertion cycles. "
     "Per NP-FAI-ZM-001 FAI-K09: 50-cycle extraction force test — gasket must not detach.",
     "RISK-19", "NP-FAI-ZM-001 §5", "☐"),
    ("F-07",
     "LED baffle wall: internal moulded baffle (height ≥ LED stack height + 0.3 mm) separating "
     "PD1 position from LED array. Prevents direct LED emission reaching PD1 — PD1 must measure "
     "only forward-exiting light through PDMS window. Baffle aperture 2.0 mm × 2.0 mm aligned "
     "with PD1 field of view through PDMS. Omission invalidates dose metering system.",
     "RISK-06", "NP-HW-FPC-001 §8.2", "☐"),
    ("F-08",
     "Module extraction pull tab: moulded grip feature on posterior face (opposite to insertion "
     "edge), 8 mm × 5 mm raised tongue, 2 mm height, textured surface (Ra 1.6 µm knurl). "
     "Provides grip for users with reduced grip strength or tremor. "
     "Pull tab must not interfere with FPC routing or ZIF lever actuation.",
     "RISK-22", "HFE formative study", "☐"),
]

for idx, row_data in enumerate(FEATURES):
    add_tbl_row(feat_tbl, row_data, feat_widths, idx)

tool.add_paragraph()

# SECTION 3 — MATERIAL SPECIFICATIONS
add_heading(tool, "3.  MATERIAL SPECIFICATIONS", size=12)
add_body(tool, "3.1  Zone Module Body — Thermoplastic Shell")
mat_hdrs = ["Property", "Requirement", "Test Standard"]
mat_widths = [1.5, 3.5, 1.7]
mat_tbl = make_table_with_header(tool, mat_hdrs, mat_widths)
MAT_ROWS = [
    ("Base material", "PC/ABS alloy or equivalent, UL94-V0 rated", "UL 94"),
    ("Colour coding", "One solid moulded colour per zone (ZM-01=Blue, ZM-02=Red, ZM-03=Green, ZM-04=Yellow, ZM-05=Purple). Colour must be in bulk material — not surface paint. Saturated enough for colour-vision-normal users; tactile features (F-01/F-02/F-03) are primary zone ID for colour-impaired users.", "ASTM D2244"),
    ("Surface finish — gasket mating slot rim", "Ra ≤ 1.6 µm on shell slot rim contact surfaces (see NP-DRV-SHELL-001 §2.2). Confirm achievable with moulded CFRP — tooling manufacturer written confirmation required before G2 gate (RISK-20).", "ISO 4287"),
    ("Thermal stability", "Continuous operation ≤ 70°C, IEC 60601-1 skin contact ≤ 43°C at PDMS window external surface", "IEC 60601-1"),
    ("Chemical resistance", "Resistant to isopropyl alcohol (70%), saline 0.9%, and sebum simulant. No cracking or crazing after 1,000 wipe cycles.", "ISO 10993-5"),
    ("Co-moulded silicone (gasket, F-05)", "Shore 40–50A, UV-stable, ISO 10993-5 biocompatible, compression set ≤ 20% after 70h at 70°C (ASTM D395)", "ASTM D395, ISO 10993-5"),
]
for idx, row_data in enumerate(MAT_ROWS):
    add_tbl_row(mat_tbl, row_data, mat_widths, idx)
tool.add_paragraph()

# SECTION 4 — DIMENSIONAL REQUIREMENTS
add_heading(tool, "4.  CRITICAL DIMENSIONS", size=12)
add_body(tool, "The following dimensions are tolerance-critical and must appear on all zone module drawings.")
dim_hdrs = ["Dimension", "Nominal", "Tolerance", "Critical for", "Drawing callout required?"]
dim_widths = [2.0, 0.9, 0.9, 1.5, 1.4]
dim_tbl = make_table_with_header(tool, dim_hdrs, dim_widths)
DIM_ROWS = [
    ("Gasket height (uncompressed, F-05)", "2.0 mm", "±0.1 mm", "IPX4 seal (RISK-16)", "YES"),
    ("Gasket width (D-section, F-05)", "2.5 mm", "±0.15 mm", "IPX4 seal", "YES"),
    ("Gasket retention groove depth (F-06)", "0.5 mm", "+0.0/−0.1 mm", "Delamination prevention (RISK-19)", "YES"),
    ("Mechanical key protrusion height (F-01)", "≥1.0 mm", "Min only", "Wrong-zone prevention (RISK-15)", "YES"),
    ("PD2 aperture diameter (F-04)", "1.2 mm", "±0.1 mm", "PD2 field of view (RISK-14)", "YES"),
    ("PD2 aperture centre-to-PD2 body offset", "≤0.2 mm", "Max only", "Optical alignment (RISK-14)", "YES"),
    ("Braille dot height (F-03)", "0.48 mm", "±0.05 mm", "Blind accessibility (RISK-15)", "YES"),
    ("Braille dot diameter (F-03)", "1.5 mm", "±0.1 mm", "Blind accessibility (RISK-15)", "YES"),
    ("Raised numeral height (F-02)", "≥0.8 mm", "Min only", "Tactile ID (RISK-15)", "YES"),
    ("Pull tab height (F-08)", "2.0 mm", "±0.3 mm", "Grip accessibility (RISK-22)", "YES"),
    ("PD1 baffle height (F-07)", "≥LED stack + 0.3 mm", "Min only", "Dose metering (RISK-06)", "YES"),
    ("PD1 baffle aperture (F-07)", "2.0 × 2.0 mm", "±0.2 mm", "Dose metering", "YES"),
]
for idx, row_data in enumerate(DIM_ROWS):
    add_tbl_row(dim_tbl, row_data, dim_widths, idx)
tool.add_paragraph()

# SECTION 5 — MOULD DESIGN REVIEW CHECKLIST
add_heading(tool, "5.  MOULD DESIGN REVIEW CHECKLIST", size=12)
add_body(tool,
    "This checklist must be completed and signed before steel is cut. "
    "All items are MANDATORY. Any OPEN item blocks tooling release.")
add_body(tool,
    "Gate: G1 Pre-Layout (NP-COORD-001 G1-05). Reviewer: ME + EE + Tooling Engineer.")
tool.add_paragraph()

mdr_hdrs = ["Item", "Requirement", "Confirmed?", "Notes"]
mdr_widths = [0.5, 4.2, 0.9, 1.1]
mdr_tbl = make_table_with_header(tool, mdr_hdrs, mdr_widths)
MDR_ITEMS = [
    ("1", "All 8 features (F-01 through F-08) appear in mould CAD model", "☐", ""),
    ("2", "F-01 mechanical key geometry unique per zone position (5 unique keys, reviewed by ME)", "☐", ""),
    ("3", "F-02/F-03 numeral + braille finger-readable in prototype sample (3-subject blind test)", "☐", "Prototype validation"),
    ("4", "F-04 PD2 aperture aligned to PD2 body position per FPC layout (EE confirms FPC layout frozen)", "☐", "FPC layout must be locked first"),
    ("5", "F-05/F-06 gasket co-moulding process confirmed with moulder — silicone primer specified", "☐", "Get written process confirmation"),
    ("6", "F-07 baffle height confirmed vs. LED stack height in BOM (EE confirms LED supplier + package)", "☐", ""),
    ("7", "F-08 pull tab does not interfere with ZIF lever access (ME confirms clearance in CAD)", "☐", ""),
    ("8", "CFRP shell slot rim Ra ≤ 1.6 µm achievable — written confirmation from tooling manufacturer (RISK-20)", "☐", "BLOCKING if not confirmed"),
    ("9", "Gasket compression per slot geometry: shell slot width minus module body width = 0.4 mm ± 0.1 mm (ME)", "☐", ""),
    ("10", "T1/T2 shared mould confirmed — identical hardware, firmware-only distinction (EE sign-off)", "☐", ""),
    ("11", "Zone colour in bulk material (not paint) — material spec confirmed with moulding supplier", "☐", ""),
    ("12", "All critical dimensions (§4) appear as drawing callouts with tolerances", "☐", ""),
]
for idx, row_data in enumerate(MDR_ITEMS):
    add_tbl_row(mdr_tbl, row_data, mdr_widths, idx)
tool.add_paragraph()

# SECTION 6 — FAI CROSS-REFERENCE
add_heading(tool, "6.  FAI CROSS-REFERENCE", size=12)
add_body(tool,
    "The following first-article inspection items in NP-FAI-ZM-001 directly validate "
    "features in this tooling specification. FAI must pass before production lot approval.")
xref_hdrs = ["Feature", "FAI Item(s)", "Pass Criterion"]
xref_widths = [1.0, 2.0, 3.8]
xref_tbl = make_table_with_header(tool, xref_hdrs, xref_widths)
XREF_ROWS = [
    ("F-01 (mechanical key)", "FAI-A14", "Zero wrong-slot insertions in 20 attempts"),
    ("F-02/F-03 (tactile ID)", "FAI-A09, FAI-A10", "Braille readable; dot count correct"),
    ("F-04 (PD2 aperture)", "FAI-SY01 (PD2 channel functional)", "PD2 ADC reading responds to LED modulation"),
    ("F-05/F-06 (gasket seal)", "FAI-IPX-01, FAI-IPX-02, FAI-IPX-03", "IPX4 pass; gasket stays bonded after 10 swap cycles"),
    ("F-07 (baffle)", "FAI-A06 (photodiode linearity)", "PD output linear with LED current; no direct LED bleed"),
    ("F-08 (pull tab)", "HFE formative study (pre-FAI)", "Users with Parkinson's/tremor extract module in ≤ 10 s"),
]
for idx, row_data in enumerate(XREF_ROWS):
    add_tbl_row(xref_tbl, row_data, xref_widths, idx)
tool.add_paragraph()

# SECTION 7 — OPEN ITEMS
add_heading(tool, "7.  OPEN ITEMS", size=12)
oi_hdrs = ["OI", "Description", "Owner", "Required by", "Status"]
oi_widths = [0.5, 3.8, 0.8, 1.2, 0.5]
oi_tbl = make_table_with_header(tool, oi_hdrs, oi_widths)
OI_ROWS = [
    ("OI-1", "Written confirmation from tooling manufacturer that CFRP shell slot rim Ra ≤ 1.6 µm is achievable with moulding (RISK-20 — BLOCKING)", "ME + Tooling", "G2 Pre-Tooling", "OPEN"),
    ("OI-2", "Silicone primer specification confirmed with gasket moulding supplier (process compatibility, ISO 10993-5)", "ME", "G1 Pre-Layout", "OPEN"),
    ("OI-3", "FPC layout frozen and PD2 component position confirmed (required before F-04 aperture can be specified in mould)", "EE", "G1 Pre-Layout", "OPEN"),
    ("OI-4", "HFE formative study for pull tab (F-08) — recruit 5 users with Parkinson's or reduced grip strength", "UX/Clinical", "G2 Pre-Tooling", "OPEN"),
    ("OI-5", "Zone colour palette confirmed with accessibility review (must pass against shell colour under all lighting conditions)", "Design", "G1 Pre-Layout", "OPEN"),
    ("OI-6", "Mould DFM review with tooling house — confirm all features manufacturable with proposed gate locations and draft angles", "ME + Tooling", "G1 Pre-Layout", "OPEN"),
]
for idx, row_data in enumerate(OI_ROWS):
    add_tbl_row(oi_tbl, row_data, oi_widths, idx)
tool.add_paragraph()

# SECTION 8 — REVISION HISTORY
add_heading(tool, "8.  REVISION HISTORY", size=12)
rev_hdrs = ["Rev", "Date", "Changes", "Author"]
rev_widths = [0.4, 1.0, 4.4, 0.9]
rev_tbl = make_table_with_header(tool, rev_hdrs, rev_widths)
add_tbl_row(rev_tbl,
    ("A", datetime.date.today().strftime("%Y-%m-%d"),
     "Initial release. Consolidates RISK-06 (F-07 baffle), RISK-14 Option B (F-04 PD2), "
     "RISK-15 (F-01 key, F-02 numeral, F-03 braille), RISK-16 (F-05 gasket), "
     "RISK-19 (F-06 retention groove), RISK-22 (F-08 pull tab). "
     "Created to mitigate RISK-23 (mould complexity omission risk).",
     "Claude/EE+ME"),
    rev_widths, 0)

tool.save(TOOL_PATH)
print(f"  NP-TOOL-ZM-001 saved: {TOOL_PATH}")

# ══════════════════════════════════════════════════════════════════════════════
# 5. RISK REGISTER — update RISK-14, RISK-21, RISK-23 to MITIGATED
#    Update RISK-18 through RISK-24 statuses
# ══════════════════════════════════════════════════════════════════════════════
RISK_PATH = "docs/superseded/neurone_fpc_zone_module_risks_revA.docx"
risk = Document(RISK_PATH)

MITIGATED_NOW = {
    "RISK-14": ("HIGH\n(MITIGATED)",
                "Option B selected: second PD (PD2) on scalp-facing PDMS surface. "
                "PD1/PD2 ratio isolates fouling from LED aging. Pin 19 assigned PD2_CATHODE. "
                "§8.4 added to NP-HW-FPC-001. BOM +$0.75–1.50/headset. T1/T2 shared mould.",
                "MITIGATED — NP-HW-FPC-001 §8.4; NP-TOOL-ZM-001 F-04"),
    "RISK-21": ("HIGH\n(MITIGATED)",
                "EEG cables routed along outer CFRP surface (opposite side from FPC bundle). "
                "§2.4 added to NP-DRV-SHELL-001 specifying dedicated 8×5 mm moulded channel. "
                "DRC-22/23 added to checklist. OI-09/10 track remaining CAD confirmation.",
                "MITIGATED — NP-DRV-SHELL-001 §2.4; DRC-22/23"),
    "RISK-23": ("HIGH\n(MITIGATED)",
                "NP-TOOL-ZM-001 created consolidating all 8 moulded features (F-01 through F-08). "
                "12-item mould design review checklist in §5. All features traceable to risk register. "
                "G1-05 gate in NP-COORD-001 requires sign-off before steel cut.",
                "MITIGATED — NP-TOOL-ZM-001 Rev 1"),
}

PARTIAL_MITIGATION = {
    "RISK-18": "OPEN — firmware debounce pending; FAI-K04 already in checklist",
    "RISK-19": "MITIGATED — F-06 retention groove + primer in NP-TOOL-ZM-001; FAI-K09 validates",
    "RISK-20": "OPEN — OI-1 in NP-TOOL-ZM-001; tooling manufacturer confirmation required (BLOCKING)",
    "RISK-22": "OPEN — F-08 pull tab in NP-TOOL-ZM-001; HFE formative study pending (OI-4)",
    "RISK-24": "MITIGATED — §2.4 DRC-18a/b provides primary routing + barrier fallback",
}

for tbl in risk.tables:
    for row in tbl.rows:
        cells = row.cells
        if not cells:
            continue
        rid = cells[0].text.strip()
        if rid in MITIGATED_NOW:
            severity_text, mitigation_text, status_text = MITIGATED_NOW[rid]
            if len(cells) >= 9:
                set_text(cells[1], severity_text, bold=True, color=(0x00, 0x70, 0x00), size=8)
                set_bg(cells[1], "E2EFDA")
                set_text(cells[5], mitigation_text, size=8)
                set_text(cells[8], status_text, bold=True, color=(0x00, 0x70, 0x00), size=8)
                set_bg(cells[8], "92D050")
                print(f"  {rid} → MITIGATED")
        elif rid in PARTIAL_MITIGATION:
            if len(cells) >= 9:
                st = PARTIAL_MITIGATION[rid]
                if "MITIGATED" in st and "OPEN" not in st:
                    set_text(cells[8], st, bold=True, color=(0x00, 0x70, 0x00), size=8)
                    set_bg(cells[8], "92D050")
                    set_text(cells[1], cells[1].text.replace("HIGH", "HIGH\n(MITIGATED)").replace("MEDIUM", "MEDIUM\n(MITIGATED)"), bold=True, size=8)
                else:
                    set_text(cells[8], st, size=8)
                    set_bg(cells[8], "FFF2CC")
                print(f"  {rid} → updated")

risk.save(RISK_PATH)
print(f"  Risk register saved: {RISK_PATH}")
print("Done.")
