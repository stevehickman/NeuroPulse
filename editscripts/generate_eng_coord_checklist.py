"""
⚠ PARTIALLY SUPERSEDED (2026-07-28): this generator's G2-01 (multi-FPC/5 zone tails), G2-02
(5-variant zone-differentiating mechanical key), and the G1-13/G2-10 ZONE_ID-based firmware
items describe the retired 5-zone-slot architecture — see docs/np_hex_zm_001.md. Do NOT
re-run: docs/neurone_eng_coordination_checklist.docx has since been extended by many later
patch_*.py scripts (fix_coord_duplicates, patch_risk16_cont, patch_t2_additions, etc.), so a
from-scratch regeneration here would discard that work. The corrected G2-01/G2-02/G2-05/
G3-03/G2-10 state is applied in place by editscripts/patch_hexzm_eng_coord_corrections.py.
Most of this document (G1, G3, and the rest of G2) is unrelated to zone-module architecture
and remains current. Kept here as the historical base generator.

Generate NP-COORD-001 Rev 1:
  Zone Module FPC — Engineering Coordination Checklist
Covers RISK-13 (primary) and consolidates all cross-discipline coordination items
from RISK-01 through RISK-17 and OI-01 through OI-13.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATE = "2026-05-06"
DOC_NUM = "NP-COORD-001"
REV = "Rev A"

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def body(doc, text, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    if p.runs:
        p.runs[0].font.size = Pt(10)
    return p

def callout(doc, label, text, bg="FFF2CC", label_rgb=(0xBF, 0x60, 0x00)):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + "  ")
    r1.bold = True; r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(*label_rgb)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()

def gate_banner(doc, gate_id, gate_name, description):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "1F497D")
    p = cell.paragraphs[0]
    r1 = p.add_run(f"GATE {gate_id}  —  {gate_name}  |  ")
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2 = p.add_run(description)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    doc.add_paragraph()

def coord_table(doc, rows, section_bg="E8F0FB"):
    """
    Coordination item table.
    rows: list of (item_id, discipline_pair, description, deliverable, gate, risk_refs, status)
    """
    headers = ["ID", "Disciplines", "Coordination Item", "Required Deliverable", "Gate", "Risk(s)", "Status"]
    widths  = [0.55,  1.15,          2.25,                1.85,                  0.6,    0.65,       0.7]

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"

    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hdr_row.cells[i], "1F497D")
        run = hdr_row.cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8)

    STATUS_COLORS = {
        "DONE":      "92D050",
        "OPEN":      "FFE699",
        "BLOCKING":  "FF0000",
        "REQUIRED":  "FFE699",
        "ADVISORY":  "D9E1F2",
        "MITIGATED": "92D050",
    }

    for r_idx, row_data in enumerate(rows):
        item_id, disciplines, desc, deliverable, gate, risk_refs, status = row_data
        row = t.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"

        status_key = status.split()[0].upper()
        status_bg = STATUS_COLORS.get(status_key, "FFFFFF")

        for c_idx, txt in enumerate([item_id, disciplines, desc, deliverable, gate, risk_refs, status]):
            cell = row.cells[c_idx]
            use_bg = status_bg if c_idx == 6 else bg
            set_cell_bg(cell, use_bg)
            run = cell.paragraphs[0].add_run(str(txt))
            run.font.size = Pt(8)
            if c_idx == 0:
                run.bold = True
            if c_idx == 6 and "BLOCKING" in status.upper():
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_bg(cell, "C00000")

    for i, w in enumerate(widths):
        for row in t.rows:
            row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t

# ─────────────────────────────────────────────────────────────────────────────
# Document
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()
sec = doc.sections[0]
sec.left_margin = Inches(0.75)
sec.right_margin = Inches(0.75)
sec.top_margin = Inches(0.85)
sec.bottom_margin = Inches(0.85)

# ── Cover ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NeurOne")
r.bold = True; r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Zone Module FPC\nEngineering Coordination Checklist")
r2.bold = True; r2.font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(
    f"{DOC_NUM}  {REV}  |  {DATE}  |  Pre-Layout / Pre-Tooling Draft\n"
    "References: NP-HW-FPC-001 Rev 3 · NP-PROC-FPC-001 Rev 1 · "
    "NP-DRV-SHELL-001 Rev 1 · NP-FAI-ZM-001 Rev 1\n"
    "Addresses RISK-13 (primary) and all cross-discipline coordination items "
    "from RISK-01 through RISK-17"
)
r3.font.size = Pt(9)
r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

callout(doc, "PURPOSE:",
    "This document tracks every cross-discipline engineering coordination item "
    "required before PCB layout release, shell tooling release, and production "
    "release of the Zone Module FPC assembly. Items are organised by gate — "
    "all items in a gate must be DONE before that gate can be declared open. "
    "Unresolved items at a gate are production-blocking.",
    bg="EEF4FB", label_rgb=(0x1F, 0x49, 0x7D))

# ── 1. Gate Structure ──
heading(doc, "1.  GATE STRUCTURE")
body(doc, "Three sequential gates govern the Zone Module FPC development programme:")

t_gates = doc.add_table(rows=4, cols=4)
t_gates.style = "Table Grid"
set_cell_bg(t_gates.rows[0].cells[0], "1F497D")
set_cell_bg(t_gates.rows[0].cells[1], "1F497D")
set_cell_bg(t_gates.rows[0].cells[2], "1F497D")
set_cell_bg(t_gates.rows[0].cells[3], "1F497D")
for i, h in enumerate(["Gate", "Name", "Opens", "Blocked items if not open"]):
    run = t_gates.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(9)

gate_rows = [
    ("G1", "Pre-Layout", "PCB layout may begin",
     "All G1 items open → layout cannot start → no prototype → no FAI → no production"),
    ("G2", "Pre-Tooling", "Shell tooling spec may be released to tooling manufacturer",
     "All G2 items open → tooling cut without provisions → retrofit impossible"),
    ("G3", "Pre-Production", "Production build may begin (after FAI pass)",
     "All G3 items open → production units non-conformant"),
]
for r_idx, (g, name, opens, consequence) in enumerate(gate_rows):
    row = t_gates.rows[r_idx + 1]
    bg = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
    for c_idx, txt in enumerate([g, name, opens, consequence]):
        set_cell_bg(row.cells[c_idx], bg)
        run = row.cells[c_idx].paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        if c_idx == 0: run.bold = True
for row in t_gates.rows:
    row.cells[0].width = Inches(0.45)
    row.cells[1].width = Inches(1.0)
    row.cells[2].width = Inches(2.0)
    row.cells[3].width = Inches(3.35)
doc.add_paragraph()

# ── 2. GATE G1 — Pre-Layout ──
heading(doc, "2.  GATE G1 — PRE-LAYOUT COORDINATION")
gate_banner(doc, "G1", "PRE-LAYOUT",
    "All items below must be DONE before PCB layout begins. "
    "Layout decisions are irreversible — RSET values, trace widths, connector "
    "footprints, and signal routing all depend on these agreements.")

# ── G1-A: HW EE ↔ Regulatory Counsel ──
heading(doc, "G1-A  |  HW EE  ↔  Regulatory Counsel  (External)", level=2)
callout(doc, "RISK-03:", "Regulatory opinion on 400 mW/cm² peak pulsed irradiance determines approved LED drive current. "
    "RSET resistor values and string topology cannot be finalised until opinion is received. "
    "Hardware is designed for 180 mA worst-case; all other outcomes (lower) are accommodated by RSET R-ladder.",
    bg="FFEEEE", label_rgb=(0xC0,0x00,0x00))

coord_table(doc, [
    ("G1-01",
     "CEO / Regulatory Counsel",
     "Commission regulatory opinion letter: 400 mW/cm² peak pulsed at ≤25% duty cycle — FDA General Wellness Policy pathway confirmation. "
     "Also assess Vielight comparison claim under FTC implied claim doctrine.",
     "Signed regulatory opinion letter from specialist counsel (PBM / digital health). "
     "Identifies any irradiance ceiling restriction.",
     "G1",
     "RISK-03",
     "BLOCKING"),
    ("G1-02",
     "HW EE ← Regulatory Counsel",
     "Once opinion received: set RSET values for approved current. "
     "If approved current < 180 mA, select RSET from R-ladder per §4.4 calculation. "
     "Update BOM and layout net list.",
     "Updated BOM with confirmed RSET values. Layout net list sign-off.",
     "G1",
     "RISK-03",
     "OPEN — awaiting G1-01"),
])

# ── G1-B: HW EE ↔ Firmware ──
heading(doc, "G1-B  |  HW EE  ↔  Firmware Lead", level=2)
callout(doc, "RISK-13 (primary):",
    "LED PWM carrier frequency must be agreed between HW EE and Firmware before layout. "
    "The carrier frequency drives PWM circuit component selection (filter caps, gate driver). "
    "EEG sample-rate anti-aliasing filter is at 200 Hz; any carrier ≥ 400 Hz aliases to DC and is invisible. "
    "1 kHz minimum is specified in NP-HW-FPC-001 §4.4. Additionally, a 50 µs EEG blanking window "
    "on each LED switching edge must be implemented in firmware to suppress switching transients. "
    "Neither the carrier frequency nor the blanking window can be changed after layout is cut.",
    bg="FFF2CC", label_rgb=(0xBF,0x60,0x00))

coord_table(doc, [
    ("G1-03",
     "HW EE + Firmware Lead",
     "RISK-13 — Formally agree LED PWM carrier frequency. "
     "Candidate: 1 kHz (verified: aliases to DC through 200 Hz anti-alias filter; "
     "does not alias into 40 Hz gamma, 10 Hz alpha, or any BES/tACS therapy band up to 40 Hz; "
     "not a harmonic of any 0.5–100 Hz visual entrainment frequency). "
     "Agreement must be documented — informal Slack agreement is not sufficient.",
     "Signed engineering decision record (EDR) stating: agreed carrier = X kHz; "
     "anti-alias filter corner frequency = Y Hz; alias frequency at carrier = Z Hz (verify < 0.5 Hz); "
     "harmonics checked against all therapy frequencies up to 100 Hz.",
     "G1",
     "RISK-13",
     "OPEN — EDR required"),
    ("G1-04",
     "Firmware Lead",
     "RISK-13 — Implement 50 µs EEG blanking window on each LED switching edge. "
     "Blanking suppresses switching transients from entering EEG ADC during transition. "
     "Blanking must be synchronised to PWM timer output — not a fixed delay. "
     "Confirm firmware implements blanking and provide test evidence (scope trace showing "
     "EEG input held during blanking window).",
     "Firmware implementation note documenting blanking mechanism + oscilloscope trace "
     "showing EEG sample gate vs PWM edge, captured in first prototype validation.",
     "G1",
     "RISK-13",
     "OPEN — firmware implementation"),
    ("G1-05",
     "HW EE + Firmware Lead",
     "RISK-02 — Firmware must enforce ≤25% LED duty cycle as inviolable hardware limit. "
     "Safety MCU owns this gate — app cannot override. "
     "25% duty cycle is the design basis for all power trace sizing (RMS current = peak × √0.25). "
     "If firmware allows >25% duty cycle, power traces will be undersized.",
     "Safety MCU firmware requirement documented: 'LED duty cycle shall not exceed 25% "
     "under any operating condition. App-requested duty cycles > 25% shall be silently capped.' "
     "Requirement in Safety MCU specification (IEC 62304 Class C).",
     "G1",
     "RISK-02",
     "OPEN — Safety MCU spec required"),
    ("G1-06",
     "HW EE + Firmware Lead",
     "RISK-14 — Decide: single-photodiode two-measurement protocol vs. second photodiode "
     "on scalp side of PDMS for LED aging / PDMS fouling discrimination. "
     "This decision affects Hub PCB routing (does second PD need additional ADC channel?) "
     "and zone module mould geometry (second PD requires aperture in different location). "
     "Must decide before layout — cannot add second PD pin after layout is cut.",
     "Engineering decision record: chosen approach (two-measurement protocol or second PD); "
     "if second PD: revised zone module BOM, additional pin allocation in 20-pin connector, "
     "revised §8 of NP-HW-FPC-001; if two-measurement: firmware algorithm spec.",
     "G1",
     "RISK-14",
     "OPEN — design decision required"),
])

# ── G1-C: HW EE ↔ Supply Chain ──
heading(doc, "G1-C  |  HW EE  ↔  Supply Chain", level=2)

coord_table(doc, [
    ("G1-07",
     "HW EE + Supply Chain",
     "OI-03 (BLOCKING) — Confirm LED emitter part numbers at rated pulse current. "
     "Seoul Semiconductor and Osram (ams-OSRAM) candidates for 660 nm and 808–830 nm. "
     "Required data: Vf at 180 mA / 25% DC / 25°C; θjc; package footprint (must be "
     "compatible with FPC reflow, peak 245°C, 30 s); L70 rating at 150 mA; MSL. "
     "Vf value drives RSET calculation. Package footprint drives PCB courtyard.",
     "Approved part numbers for 660 nm and 808–830 nm LED emitters with datasheet "
     "values for Vf, θjc, and footprint. Entered in BOM before layout.",
     "G1",
     "RISK-03 / OI-03",
     "BLOCKING"),
    ("G1-08",
     "HW EE",
     "OI-02 — Pull Hirose FH34S-20S-0.5SH full datasheet; confirm 'endurance: 1,000 cycles' "
     "in test conditions. Confirm JAE FF03 availability as backup. "
     "Cycle rating must appear in the datasheet — not just the marketing brief.",
     "Datasheet on file with 1,000-cycle rating confirmed. "
     "If FH34S datasheet does not state 1,000 cycles, escalate to connector vendor for test data.",
     "G1",
     "RISK-01 / OI-02",
     "OPEN — datasheet pull required"),
    ("G1-09",
     "HW EE + Supply Chain",
     "OI-05 — Confirm LED Vf bin tolerance (±0.10 V within-order) is achievable by "
     "Seoul Semiconductor and Osram for first production order. "
     "Some vendors offer bin-controlled lots as a premium service; others cannot. "
     "If primary suppliers cannot meet ±0.10 V, qualify Luminus Devices as primary.",
     "Written confirmation from ≥1 approved LED supplier that ±0.10 V bin control "
     "is available for first production order quantity.",
     "G1",
     "RISK-08 / OI-05",
     "OPEN — supplier confirmation required"),
    ("G1-10",
     "HW EE",
     "OI-04 — Update BOM: BCR421W in all LED driver positions. Confirm no BCR421U on any "
     "BOM revision. One-line change but must be done and version-controlled before layout.",
     "BOM revision with BCR421W confirmed in all driver positions. "
     "BCR421U must not appear on any released BOM.",
     "G1",
     "RISK-07 / OI-04",
     "OPEN — BOM update required"),
])

# ── G1-D: HW EE ↔ Optical / Mech Eng ──
heading(doc, "G1-D  |  HW EE  ↔  Optical Engineering  +  Mechanical Engineering", level=2)

coord_table(doc, [
    ("G1-11",
     "HW EE + Optical Eng + Mech Eng",
     "OI-07 — Photodiode part number and baffle geometry must be defined before zone "
     "module mould geometry is finalised. Baffle: 1 mm aperture, 2–3 mm height, black POM "
     "or anodised Al (per §8.2). Baffle housing is moulded into zone module body — "
     "cannot be added post-tooling. "
     "Also decide: ADPD174GGI (integrated TIA + optical isolation) vs. discrete TIA "
     "(per OI-13). Integrated part changes Hub PCB footprint.",
     "Photodiode part number confirmed. Baffle geometry drawing in mech spec. "
     "ADPD174GGI vs discrete decision documented. Zone module mould drawing updated.",
     "G1",
     "RISK-06 / OI-07 / OI-13",
     "OPEN — mould geometry blocking"),
    ("G1-12",
     "HW EE + Manufacturing Eng",
     "OI-08 — Confirm SiO₂ interlayer sputtering (RISK-04 PDMS delamination mitigation) "
     "is available at an identified contract process house before FPC tooling is released. "
     "If process house cannot accommodate sputtering in production workflow, "
     "evaluate APTES silane coupling agent as alternative (lower bond strength but "
     "may be sufficient — requires qualification data).",
     "Named process house with confirmed SiO₂ sputtering capability (RF magnetron, "
     "75 nm on PI). Lead time and NRE cost confirmed. "
     "If APTES alternative chosen: qualification bond strength data (≥50 N/m peel) on file.",
     "G1",
     "RISK-04 / OI-08",
     "OPEN — process house identification required"),
])

# ── 3. GATE G2 — Pre-Tooling ──
heading(doc, "3.  GATE G2 — PRE-TOOLING COORDINATION")
gate_banner(doc, "G2", "PRE-TOOLING",
    "All items below must be DONE before shell tooling spec is released to manufacturer. "
    "Shell tooling geometry is irreversible — features not specified before first-cut "
    "cannot be retrofitted.")

# ── G2-A: Mechanical ↔ HW EE ──
heading(doc, "G2-A  |  Mechanical Engineering  ↔  HW EE", level=2)

coord_table(doc, [
    ("G2-01",
     "Mech Eng + HW EE",
     "RISK-17 / RISK-11 — Multi-FPC cable management inside shell: "
     "5 zone FPC tails run simultaneously from zone slots to Hub PCB connector. "
     "Shell interior must separate zone FPCs from EEG electrode cables (to prevent "
     "RF coupling between LED power current and EEG signal leads). "
     "Hub PCB connector placement must minimise routing complexity for longest FPC runs "
     "(Zone 1 and 2, frontal, ~120 mm arc). "
     "All routing channels addressed in NP-DRV-SHELL-001; multi-FPC bundle management "
     "requires additional coordination.",
     "Shell CAD cross-section showing all 5 FPC channels + EEG cable routing simultaneously. "
     "Confirmed separation ≥ 10 mm between LED power FPC tails and EEG signal cables at "
     "all routing path segments. Hub PCB connector placement confirmed.",
     "G2",
     "RISK-17 / RISK-11",
     "OPEN — shell CAD required"),
    ("G2-02",
     "Mech Eng",
     "RISK-15 — Unique asymmetric mechanical key geometry for each of 5 zone slots. "
     "Key profiles must be defined such that: "
     "(a) no two zones have the same key; (b) key is asymmetric (prevents 180° mis-rotation); "
     "(c) key tolerances allow smooth insertion by users with limited dexterity. "
     "Key geometry must be in shell tooling spec before first-cut. "
     "NP-FAI-ZM-001 FAI-A01–A08 verifies alignment post-tooling.",
     "Key geometry drawings for all 5 zone slots: tab/notch dimensions, tolerances "
     "(recommend ±0.1 mm on key tab, ±0.15 mm on slot notch for smooth insertion). "
     "Written confirmation that no two zone key profiles are interchangeable.",
     "G2",
     "RISK-15 / OI-10",
     "OPEN — key geometry required"),
    ("G2-03",
     "Mech Eng + HW EE",
     "RISK-12 — Zone module stiffener datum features (reference holes or edges) must be "
     "defined and added to FPC fabrication drawing before tooling release. "
     "Datum features are required for CMM fixturing during FAI (NP-FAI-ZM-001 FAI-A07). "
     "Coordinate: HW EE places datum features on FPC drawing; Mech Eng confirms "
     "they are compatible with CMM fixture design.",
     "FPC fabrication drawing revision with datum features (position ±0.1 mm). "
     "CMM fixture concept confirmed by Mech Eng.",
     "G2",
     "RISK-12",
     "OPEN — FPC drawing revision required"),
    ("G2-04",
     "Mech Eng + HW EE",
     "NP-DRV-SHELL-001 sign-off — all 8 design review checklist items (DRC-01 to DRC-15) "
     "that can be verified in CAD must be completed before tooling spec release. "
     "Physical prototype items (DRC-14) required before first-cut approval (not tooling spec release). "
     "See NP-DRV-SHELL-001 §5 for full checklist and §8 for sign-off.",
     "Completed NP-DRV-SHELL-001 §5 checklist with all CAD-verifiable items passed. "
     "All sign-off boxes in §8 completed.",
     "G2",
     "RISK-11",
     "OPEN — DRC sign-off required"),
    ("G2-05",
     "Mech Eng + UX",
     "RISK-16 — Zone module field-replacement IPX4 seal design decision: "
     "self-sealing silicone gasket (preferred) vs. user-applied RTV syringe. "
     "Gasket geometry must be in zone module and shell tooling spec before first-cut. "
     "If RTV syringe approach chosen: define included accessories (syringe in box), "
     "document re-sealing procedure, add re-sealing step to zone module swap user guide.",
     "Engineering decision record: chosen approach. "
     "If gasket: gasket geometry drawing in mech spec, material spec (Shore 40A silicone), "
     "IPX4 test protocol after 10 swap cycles. "
     "If RTV: syringe BOM entry, procedure documentation.",
     "G2",
     "RISK-16 / OI-11",
     "OPEN — design decision required"),
])

# ── G2-B: Firmware ↔ Safety MCU ──
heading(doc, "G2-B  |  Firmware Lead  ↔  Safety MCU Team", level=2)

coord_table(doc, [
    ("G2-06",
     "Firmware + Safety MCU",
     "RISK-03 — Once regulatory opinion received (G1-01) and RSET values confirmed (G1-02): "
     "firmware must enforce approved peak current as a hard limit in Safety MCU. "
     "Safety MCU owns all stimulation GPIO enable lines. Approved current ceiling must be "
     "a Safety MCU constant, not a firmware parameter (app cannot override).",
     "Safety MCU firmware constant: MAX_LED_CURRENT_MA = approved value. "
     "Unit test demonstrating Safety MCU rejects any main processor request exceeding limit. "
     "In Safety MCU IEC 62304 Class C documentation.",
     "G2",
     "RISK-03",
     "OPEN — awaiting G1-01"),
    ("G2-07",
     "Firmware + Safety MCU",
     "RISK-13 — PWM carrier frequency lock: once G1-03 EDR is signed, the agreed carrier "
     "frequency becomes a Safety MCU constant and a locked firmware parameter. "
     "Neither the carrier frequency nor the EEG blanking window timing may be changed "
     "via OTA update without a new hardware revision (because both affect PCB component "
     "values and trace routing).",
     "Safety MCU firmware constant: LED_PWM_CARRIER_HZ = agreed value. "
     "Blanking window constant: EEG_BLANK_US = 50 (or agreed value). "
     "Change-controlled — requires ECO to modify.",
     "G2",
     "RISK-13",
     "OPEN — awaiting G1-03"),
])

# ── G2-C: Firmware ↔ Systems Eng ──
heading(doc, "G2-C  |  Firmware / Systems Eng  ↔  HW EE", level=2)

coord_table(doc, [
    ("G2-08",
     "Firmware + Systems Eng",
     "RISK-14 — If two-measurement protocol chosen (G1-06): "
     "firmware algorithm for LED aging vs. PDMS fouling discrimination must be "
     "designed before first prototype, so that first-article validation can test "
     "the full dose metering claim end-to-end. "
     "Algorithm inputs: session measurement, NTC temperature (stable LED condition), "
     "factory baseline, time-series trend from SHDR. "
     "Algorithm must run on-device (Mode 3 autonomous — no cloud required).",
     "Algorithm specification document: inputs, outputs, decision thresholds, "
     "SHDR storage format for trend data. Reviewed by Systems Eng and HW EE. "
     "Implemented in prototype firmware before FAI.",
     "G2",
     "RISK-14",
     "OPEN — awaiting G1-06"),
    ("G2-09",
     "Firmware + HW EE",
     "RISK-13 — PWM carrier harmonic analysis: document that agreed carrier frequency "
     "and all harmonics up to Nyquist (250 Hz) do not coincide with: "
     "(a) any BES/tACS therapy frequency (0.5–40 Hz); "
     "(b) EEG neurofeedback bands (delta 0.5–4, theta 4–8, alpha 8–13, beta 13–30, gamma 30–100 Hz); "
     "(c) visual entrainment frequencies (0.5–100 Hz); "
     "(d) VNS stimulation frequencies (1–25 Hz). "
     "This is a one-time calculation but must be on record.",
     "Harmonic analysis table: carrier = X kHz; harmonics at 2X, 3X, ... up to 250 Hz; "
     "confirmed no coincidence with any therapy frequency. "
     "Signed by HW EE and Firmware Lead.",
     "G2",
     "RISK-13",
     "OPEN — calculation required"),
])

# ── 4. GATE G3 — Pre-Production ──
heading(doc, "4.  GATE G3 — PRE-PRODUCTION COORDINATION")
gate_banner(doc, "G3", "PRE-PRODUCTION",
    "All items below must be DONE before production release. "
    "Some items can run in parallel with prototype build; none may be waived.")

coord_table(doc, [
    ("G3-01",
     "Manufacturing Eng + HW EE",
     "RISK-04 — PDMS bonding process qualification: complete thermal cycling test "
     "(IEC 60068-2-14, 200 cycles −40 to +85°C) on production-representative PDMS/SiO₂/PI "
     "bond coupons. Pass criterion: no delamination; peel force ≥ 50 N/m after cycling. "
     "This is a process qualification, not a design change — process house must run it.",
     "Thermal cycling qualification report with pass/fail per IEC 60068-2-14. "
     "Minimum 10 coupons. Report on file as manufacturing process qualification record.",
     "G3",
     "RISK-04",
     "OPEN — process house required first"),
    ("G3-02",
     "Mech Eng",
     "RISK-11 — Physical prototype (SLA or 3D-printed shell section, Zones 1 and 2 minimum): "
     "insert/remove zone module 10× and confirm no FPC contact with channel wall edges. "
     "Required by NP-DRV-SHELL-001 DRC-14 before first-cut approval.",
     "DRC-14 sign-off: no contact marks on FPC after 10 cycles on physical prototype. "
     "Photo evidence filed. NP-DRV-SHELL-001 DRC-14 marked PASS.",
     "G3",
     "RISK-11",
     "OPEN — prototype required"),
    ("G3-03",
     "Mech Eng",
     "RISK-16 — IPX4 validation after field replacement: insert and remove zone module "
     "10× (simulating user field replacement), then conduct IPX4 splash test. "
     "Required regardless of seal approach chosen (gasket or RTV).",
     "IPX4 test report: pass/fail per IEC 60529 after 10 zone module swap cycles. "
     "Test conducted on production-representative unit with production seal approach.",
     "G3",
     "RISK-16",
     "OPEN — design decision (G2-05) required first"),
    ("G3-04",
     "HW EE + Firmware",
     "RISK-13 — First prototype validation: capture oscilloscope trace showing "
     "EEG input vs. LED PWM edge to confirm 50 µs blanking is active and synchronised. "
     "Also verify no EEG artifact visible during LED switching in scope trace.",
     "Scope trace showing: (1) PWM edge; (2) EEG sample gate going low during blanking "
     "window; (3) EEG ADC output flat during blanking. Filed as prototype validation record.",
     "G3",
     "RISK-13",
     "OPEN — prototype required"),
    ("G3-05",
     "Manufacturing Eng + Supply Chain",
     "All RISK-08 incoming inspection procedures (NP-PROC-FPC-001 §7) must be exercised "
     "on first production component lots before assembly begins. "
     "Any reject at incoming inspection triggers disposition review before assembly proceeds.",
     "Completed incoming inspection records for first production lots of: "
     "660 nm LED, 808 nm LED, BCR421W, FH34S connector, FPC laminate, PDMS. "
     "All records in quality management system.",
     "G3",
     "RISK-08 / RISK-07 / RISK-01 / RISK-10",
     "OPEN — component orders required first"),
    ("G3-06",
     "Manufacturing Eng + HW EE + Mech Eng",
     "NP-FAI-ZM-001 Rev 1 — complete all 50+ FAI items before production release. "
     "All BLOCKING items (FAI-D01, FAI-C01, FAI-K01, FAI-F01, FAI-A01–A05) must pass. "
     "Non-conformances in §7 must be dispositioned with Engineering sign-off.",
     "Completed NP-FAI-ZM-001 checklist with all items PASS or formally dispositioned. "
     "Sign-off by Manufacturing Eng, HW EE, Quality.",
     "G3",
     "RISK-12 + all",
     "OPEN — prototype required first"),
])

# ── 5. Summary Status Dashboard ──
heading(doc, "5.  SUMMARY STATUS DASHBOARD")
body(doc, "Current status of all coordination items. Updated as items close.")

dash_headers = ["ID", "Gate", "Owner", "Item (abbreviated)", "Status"]
dash_widths = [0.6, 0.5, 1.35, 4.0, 0.85]

t_dash = doc.add_table(rows=1, cols=len(dash_headers))
t_dash.style = "Table Grid"
hdr_dash = t_dash.rows[0]
for i, h in enumerate(dash_headers):
    set_cell_bg(hdr_dash.cells[i], "1F497D")
    run = hdr_dash.cells[i].paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(8)

STATUS_COLORS_DASH = {
    "DONE": ("92D050", (0x00,0x50,0x00)),
    "BLOCKING": ("C00000", (0xFF,0xFF,0xFF)),
    "OPEN": ("FFE699", (0x60,0x40,0x00)),
}

dash_items = [
    ("G1-01","G1","CEO / Reg Counsel","Regulatory opinion letter — 400 mW/cm² irradiance","BLOCKING"),
    ("G1-02","G1","HW EE","Set RSET values per approved current","OPEN — awaiting G1-01"),
    ("G1-03","G1","HW EE + FW Lead","LED PWM carrier frequency EDR (RISK-13)","OPEN — EDR required"),
    ("G1-04","G1","Firmware Lead","50 µs EEG blanking window implementation (RISK-13)","OPEN"),
    ("G1-05","G1","HW EE + Safety MCU","≤25% duty cycle hard limit in Safety MCU spec","OPEN"),
    ("G1-06","G1","HW EE + FW + Systems","Photodiode: 2-measurement protocol vs second PD (RISK-14)","OPEN — design decision"),
    ("G1-07","G1","HW EE + Supply Chain","LED emitter part numbers confirmed (OI-03)","BLOCKING"),
    ("G1-08","G1","HW EE","Hirose FH34S 1,000-cycle rating confirmed from datasheet","OPEN"),
    ("G1-09","G1","HW EE + Supply Chain","LED Vf ±0.10 V bin control confirmed by supplier","OPEN"),
    ("G1-10","G1","HW EE","BOM updated to BCR421W (no BCR421U)","OPEN"),
    ("G1-11","G1","HW EE + Optical + Mech","Photodiode part, baffle geometry, PD integration decision","OPEN — mould blocking"),
    ("G1-12","G1","HW EE + Mfg","SiO₂ sputtering process house confirmed (RISK-04)","OPEN"),
    ("G2-01","G2","Mech + HW EE","Multi-FPC + EEG cable separation in shell","OPEN"),
    ("G2-02","G2","Mech Eng","Mechanical key geometry all 5 zones (RISK-15)","OPEN"),
    ("G2-03","G2","Mech + HW EE","Stiffener datum features on FPC drawing (RISK-12)","OPEN"),
    ("G2-04","G2","Mech + HW EE","NP-DRV-SHELL-001 DRC sign-off — CAD items","OPEN"),
    ("G2-05","G2","Mech + UX","IPX4 field-replacement seal design decision (RISK-16)","OPEN"),
    ("G2-06","G2","FW + Safety MCU","Approved current ceiling in Safety MCU constant","OPEN — awaiting G1-01"),
    ("G2-07","G2","FW + Safety MCU","PWM carrier and blanking locked as Safety MCU constants","OPEN — awaiting G1-03"),
    ("G2-08","G2","FW + Systems","PDMS fouling algorithm spec (if two-measurement chosen)","OPEN — awaiting G1-06"),
    ("G2-09","G2","FW + HW EE","PWM harmonic analysis — no therapy frequency coincidence","OPEN"),
    ("G3-01","G3","Mfg + HW EE","PDMS thermal cycling qualification 200 cycles","OPEN"),
    ("G3-02","G3","Mech Eng","Physical prototype FPC routing path test (DRC-14)","OPEN"),
    ("G3-03","G3","Mech Eng","IPX4 test after 10 field replacement cycles","OPEN — awaiting G2-05"),
    ("G3-04","G3","HW EE + FW","PWM blanking oscilloscope validation on prototype","OPEN"),
    ("G3-05","G3","Mfg + Supply Chain","Incoming inspection all first production component lots","OPEN"),
    ("G3-06","G3","Mfg + HW EE + Mech","NP-FAI-ZM-001 full FAI pass","OPEN"),
]

for r_idx, (item_id, gate, owner, item, status) in enumerate(dash_items):
    row = t_dash.add_row()
    bg = "FFFFFF" if r_idx % 2 == 0 else "F8F8F8"
    status_key = status.split()[0].upper()
    s_bg, s_fg = STATUS_COLORS_DASH.get(status_key, ("FFE699", (0x60,0x40,0x00)))
    for c_idx, txt in enumerate([item_id, gate, owner, item, status]):
        cell = row.cells[c_idx]
        set_cell_bg(cell, s_bg if c_idx == 4 else bg)
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(8)
        if c_idx == 0: run.bold = True
        if c_idx == 4:
            run.font.color.rgb = RGBColor(*s_fg)
            run.bold = ("BLOCKING" in status.upper())

for row in t_dash.rows:
    for i, w in enumerate(dash_widths):
        row.cells[i].width = Inches(w)
doc.add_paragraph()

body(doc,
    f"Status at {DATE}: 2 BLOCKING items (G1-01 regulatory opinion, G1-07 LED part numbers). "
    "All other items OPEN. Gate G1 cannot be declared open until BLOCKING items are resolved. "
    "No layout, tooling, or production work may begin until BLOCKING items are closed.",
    space_after=8)

# ── 6. Revision History ──
heading(doc, "6.  REVISION HISTORY")
t_rev = doc.add_table(rows=2, cols=4)
t_rev.style = "Table Grid"
for i, h in enumerate(["Rev", "Date", "Author", "Description"]):
    set_cell_bg(t_rev.rows[0].cells[i], "1F497D")
    run = t_rev.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(8)
for c_idx, txt in enumerate(["A", DATE, "NeurOne Engineering",
    "Initial release. RISK-13 (primary) — LED PWM carrier frequency / EEG coordination. "
    "Consolidates all cross-discipline coordination items from RISK-01 through RISK-17 "
    "and OI-01 through OI-13 into 3-gate structure. 27 coordination items total."]):
    set_cell_bg(t_rev.rows[1].cells[c_idx], "F5F5F5")
    run = t_rev.rows[1].cells[c_idx].paragraphs[0].add_run(txt)
    run.font.size = Pt(8)
for row in t_rev.rows:
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(0.9)
    row.cells[2].width = Inches(1.5)
    row.cells[3].width = Inches(4.0)

OUT = "docs/neurone_eng_coordination_checklist.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
