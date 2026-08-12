"""
Correct RISK-11, RISK-15, and RISK-17 in the zone module risk register
(docs/superseded/np_risk_001.docx) for the hex-tile/socket
architecture (NP-HEX-ZM-001), following PR #223's commits and the
supersession-note convention it established for the markdown docs.

This register is NOT globally superseded — most of its 25 risks are still
current — so unlike the FPC spec (patch_hexzm_spec_supersession.py) this is a
targeted per-row correction, not a document-level banner.

RISK-11 / RISK-17: both point to NP-DRV-SHELL-001 (the shell FPC routing
review), whose entire premise — 5 fixed zone-slot FPC routing paths with
measured arc lengths — is retired. The hex-tile interconnect for ~80 sockets
is real, unfinished EE work (see MECH-1/MECH-2/SMART-1 residuals in
docs/np_hex_zm_001.md §7) — not something to invent numbers for here.
Per CLAUDE.md/memory convention (least-asserting default for safety-adjacent
claims), these are downgraded from MITIGATED to NEEDS REVIEW rather than left
falsely asserting a closed mitigation.

RISK-15: the root cause itself (fixed-position, color-coded zone modules that
can be inserted into the wrong slot) no longer exists. Modules are
type-agnostic hex tiles with a universal orientation-only mechanical key
(docs/np_hex_zm_001.md §4a) — there is no "wrong zone" left to insert into.
This is a genuine simplification, not a downgrade: rewritten to describe why
the risk is now structurally eliminated rather than mitigated by added
hardware/firmware layers.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RISK_PATH = "docs/superseded/np_risk_001.docx"

MARKER = "hex-tile/socket architecture (NP-HEX-ZM-001)"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def find_row(tbl, risk_id):
    for row in tbl.rows:
        if row.cells[0].text.strip() == risk_id:
            return row
    return None


def already_patched(doc):
    return any(MARKER in p.text for tbl in doc.tables for row in tbl.rows for c in row.cells for p in c.paragraphs)


def main():
    doc = Document(RISK_PATH)
    if already_patched(doc):
        print(f"  Already patched — skipping: {RISK_PATH}")
        return

    tbl = doc.tables[0]

    # ── RISK-11 ──────────────────────────────────────────────────────────
    row = find_row(tbl, "RISK-11")
    if row:
        set_cell_text(row.cells[1], "MEDIUM\n(NEEDS REVIEW)", bold=True, color=(0xBF, 0x60, 0x00))
        set_cell_bg(row.cells[1], "FFF2CC")
        set_cell_text(row.cells[5],
            "NEEDS REVIEW: the cited mitigation (NP-DRV-SHELL-001 Rev 1/B — 5 fixed zone-slot FPC "
            "routing paths, REQ-BR-01..05) was scoped to the retired position-unique 5-slot zone "
            "module architecture, itself superseded by the hex-tile/socket architecture "
            "(NP-HEX-ZM-001, 2026-07-15). The ~80-socket interconnect that replaces it is real, "
            "unfinished EE work — no document yet specifies per-socket FPC/cable routing, bend "
            "radius, or channel geometry for the hex-tile lattice (see MECH-1/MECH-2/SMART-1 "
            "residuals in docs/np_hex_zm_001.md §7). The generic bend-radius physics (IPC-2223D "
            "basis, ≥25mm dynamic / ≥12.5mm static) remain valid engineering constraints and should "
            "carry forward into whatever per-socket interconnect design replaces NP-DRV-SHELL-001; "
            "the specific 5-zone routing paths and arc-length table do not.", size=9)
        set_cell_text(row.cells[8],
            "NEEDS REVIEW — mitigation was scoped to the retired 5-slot architecture; hex-tile "
            "socket interconnect design (~80 sockets) not yet started",
            bold=True, color=(0xBF, 0x60, 0x00), size=9)
        set_cell_bg(row.cells[8], "FFE699")
        print("  RISK-11 corrected")
    else:
        print("  WARNING: RISK-11 row not found")

    # ── RISK-15 ──────────────────────────────────────────────────────────
    row = find_row(tbl, "RISK-15")
    if row:
        set_cell_text(row.cells[2], "Zone keying risk structurally eliminated by the hex-tile/socket "
            "architecture (NP-HEX-ZM-001)", bold=False, size=9)
        set_cell_text(row.cells[3],
            "SUPERSEDED ROOT CAUSE (was: color-coded zone modules, each valid in exactly one of 5 "
            "fixed slots — a wrong-slot insertion was possible and needed color-independent keying). "
            "Under NP-HEX-ZM-001, all modules are type-agnostic hex tiles — \"any type (T1-A/B/C) "
            "inserts into any socket — same size, same mount, orientation-only key (no type keying)\" "
            "(docs/np_hex_zm_001.md §4a). There is no zone-specific module and no wrong slot to "
            "insert one into: the fixed color-per-zone identity this risk assumed does not exist.",
            size=9)
        set_cell_text(row.cells[4],
            "N/A — eliminated by design. The residual hazard (if any) is orientation, not zone "
            "identity: a module inserted backwards. This is prevented purely mechanically (the "
            "universal asymmetric orientation key), independent of vision, for every user including "
            "blind users, with no firmware, reading, or user attention required.", size=9)
        set_cell_text(row.cells[5],
            "MITIGATED BY ELIMINATION: the universal orientation-only mechanical key "
            "(docs/np_hex_zm_001.md §4a) is passive and blind-safe by construction — it requires no "
            "firmware, no reading, no color vision, and no zone-specific tooling. Module TYPE "
            "correctness (not zone identity) is separately checked in software: "
            "np_module_map_check_placement() verifies the live UID-reported inventory against each "
            "protocol's required module map and guides the user to fix any mismatch via the app "
            "(docs/np_hex_zm_001.md §4a, \"Protocol ↔ module-map matching\"). "
            "RETIRED (do not carry forward): the 5-layer keying system (mechanical zone-differentiating "
            "key + ZONE_ID pin-18 resistor ladder + per-zone braille/tactile markers + bone-conduction "
            "zone-name announcement) — all five layers existed only to solve wrong-slot insertion, "
            "which no longer exists as a hazard. The bone-conduction announcement firmware "
            "(NP-FW-ZA-001, tracked by NP-COORD-001 G2-10) is real shipped firmware built on the "
            "retired ZONE_ID mechanism and needs porting to UID-based detection, not new keying "
            "hardware — see docs/superseded/np_fw_za_001.md.", size=9)
        set_cell_text(row.cells[8],
            "MITIGATED — risk eliminated by the hex-tile/socket architecture (NP-HEX-ZM-001); "
            "orientation key + software placement-check replace the retired 5-layer zone keying "
            "system. NP-FW-ZA-001 firmware port to UID-based detection tracked separately.",
            bold=True, color=(0x00, 0x70, 0x00), size=9)
        set_cell_bg(row.cells[8], "92D050")
        print("  RISK-15 corrected")
    else:
        print("  WARNING: RISK-15 row not found")

    # ── RISK-17 ──────────────────────────────────────────────────────────
    row = find_row(tbl, "RISK-17")
    if row:
        set_cell_text(row.cells[1], "MEDIUM\n(NEEDS REVIEW)", bold=True, color=(0xBF, 0x60, 0x00))
        set_cell_bg(row.cells[1], "FFF2CC")
        set_cell_text(row.cells[5],
            "NEEDS REVIEW: the cited mitigation (NP-DRV-SHELL-001 §2.3 — multi-FPC bundle management "
            "and inter-FPC separation across 5 adjacent zone-slot tails) was scoped to the retired "
            "5-zone-slot routing architecture, superseded by the hex-tile/socket architecture "
            "(NP-HEX-ZM-001, 2026-07-15). At ~80 sockets the adjacent-tail crosstalk/bend-radius "
            "analysis needs to be redone against whatever per-socket interconnect design replaces "
            "the 5 discrete zone-slot FPC tails — that design does not exist yet (see "
            "MECH-1/MECH-2/SMART-1 in docs/np_hex_zm_001.md §7). Not carried forward without "
            "re-derivation: the specific inter-FPC separation distances and pairwise DRC checks, "
            "which were sized for 5 large tails, not ~80 small ones.", size=9)
        set_cell_text(row.cells[8],
            "NEEDS REVIEW — mitigation was scoped to the retired 5-slot architecture; hex-tile "
            "socket interconnect design (~80 sockets) not yet started",
            bold=True, color=(0xBF, 0x60, 0x00), size=9)
        set_cell_bg(row.cells[8], "FFE699")
        print("  RISK-17 corrected")
    else:
        print("  WARNING: RISK-17 row not found")

    doc.save(RISK_PATH)
    print(f"  Risk register saved: {RISK_PATH}")


if __name__ == "__main__":
    main()
