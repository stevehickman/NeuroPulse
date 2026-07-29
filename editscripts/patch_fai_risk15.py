"""
⚠ SUPERSEDED (2026-07-28): the FAI-A09..A14 items this script adds test the retired 5-layer
zone-keying system (mechanical zone-differentiating key + ZONE_ID resistor ladder + per-zone
braille/tactile markers) — see editscripts/patch_risk15.py's header note for why that system
no longer exists under the hex-tile/socket architecture (NP-HEX-ZM-001). Do not re-run. The
corrected FAI-A01..A15 state (including A09-A14 rewritten for the universal orientation key
and the UID-based placement-check software gate) is applied by
editscripts/patch_hexzm_fai_corrections.py. Kept here as historical record.

Patch NP-FAI-ZM-001: add RISK-15 accessibility verification items to §4a
and cross-reference table in §9.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

FAI_PATH = "docs/neurone_fai_zone_module.docx"
fai = Document(FAI_PATH)

# ── Find the FAI-A08 row and append new rows after it ──
target_tbl = None
target_row_idx = None
for ti, tbl in enumerate(fai.tables):
    for ri, row in enumerate(tbl.rows):
        if row.cells[0].text.strip() == "FAI-A08":
            target_tbl = tbl
            target_row_idx = ri
            break
    if target_tbl:
        break

print(f"FAI-A08 found in table {ti}, row {target_row_idx}")

NEW_ROWS = [
    ("FAI-A09",
     "Zone module tactile identification: verify raised numeral (1–5) and braille "
     "zone abbreviation are present and readable on all 5 zone module bodies",
     "Physical inspection; finger-reading test by 3 test subjects with eyes closed",
     "Numeral present on all 5 modules; braille readable by each test subject "
     "(confirmed match to zone position). Any absent or unreadable → reject.",
     "", ""),
    ("FAI-A10",
     "Shell tactile dot patterns: verify N raised dots adjacent to each zone slot "
     "for zones 1–5; confirm pattern is finger-findable from the slot edge",
     "Physical inspection; blind-condition finger test by 3 test subjects",
     "Correct dot count at each slot; dots reachable within 5 s by test subjects "
     "without visual reference",
     "", ""),
    ("FAI-A11",
     "ZONE_ID electronic verification: read ADC voltage at pin 18 for all 5 zones "
     "and confirm correct resistor value per zone (ZM-01=10k, 02=22k, 03=47k, "
     "04=100k, 05=220k ±5% ADC tolerance)",
     "Hub MCU diagnostic mode (or bench ADC measurement)",
     "All 5 zones read correct ID within ±5% band. Any zone outside band → reject module lot.",
     "", ""),
    ("FAI-A12",
     "[BLOCKING] Wrong-zone ENABLE block: insert ZM-02 module into ZM-01 slot; "
     "confirm firmware asserts FAULT, blocks ENABLE, and issues audio alert via "
     "bone conduction",
     "Live firmware test",
     "ENABLE held low within 200 ms of module insertion. Bone conduction announces "
     "'Wrong module' within 3 s. Session cannot start.",
     "", ""),
    ("FAI-A13",
     "Correct-zone audio confirmation: insert all 5 modules correctly; confirm "
     "bone conduction announces correct zone name for each zone during headset "
     "initialisation sequence",
     "Live firmware test; listen to bone conduction output",
     "All 5 zone names announced correctly and in sequence. Audio intelligible at "
     "normal ambient noise level (≤60 dB background).",
     "", ""),
    ("FAI-A14",
     "Mechanical key exclusion test: attempt to insert each zone module into each "
     "wrong slot (ZM-01 → slots 2–5, etc.); confirm mechanical key prevents insertion "
     "in all incorrect slot/module combinations",
     "Physical manual test — 20 wrong-slot attempts per unit",
     "Zero incorrect insertions possible. Module must be fully stopped by mechanical "
     "key before FPC tail contacts ZIF. Any mis-insertion possible → halt, escalate.",
     "", ""),
]

if target_tbl and target_row_idx is not None:
    # Add new rows after FAI-A08
    widths = [0.65, 2.3, 1.2, 1.55, 0.65, 0.45]

    for new_row_data in NEW_ROWS:
        # Insert row after target_row_idx
        ref_row = target_tbl.rows[target_row_idx]
        new_tr = OxmlElement("w:tr")
        ref_row._tr.addnext(new_tr)

        # Find the new row object
        new_row = None
        for row in target_tbl.rows:
            if row._tr is new_tr:
                new_row = row
                break

        if new_row is None:
            # Append at end
            new_row = target_tbl.add_row()

        # Add cells
        for c_idx, txt in enumerate(new_row_data):
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            w_el = OxmlElement("w:tcW")
            w_el.set(qn("w:w"), str(int(widths[c_idx] * 1440)))
            w_el.set(qn("w:type"), "dxa")
            tc_pr.append(w_el)
            tc.append(tc_pr)
            p_el = OxmlElement("w:p")
            r_el = OxmlElement("w:r")
            rpr_el = OxmlElement("w:rPr")
            sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
            rpr_el.append(sz_el)
            r_el.append(rpr_el)
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve")
            t_el.text = txt
            r_el.append(t_el)
            p_el.append(r_el)
            tc.append(p_el)
            new_row._tr.append(tc)

            set_cell_bg(new_row.cells[c_idx], "F5F5F5")

        target_row_idx += 1  # advance so next row inserts after the one we just added
        print(f"  Added {new_row_data[0]}")

# ── Update §9 cross-reference table — add RISK-15 row ──
for tbl in fai.tables:
    # Look for table with RISK-12 row
    for row in tbl.rows:
        if "RISK-12" in row.cells[0].text:
            # Add RISK-15 row after RISK-12
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)

            xref_widths = [0.65, 2.0, 2.35, 1.8]
            xref_data = ("RISK-15",
                         "Zone keying — colour-blind + blind-accessible",
                         "FAI-A09, FAI-A10, FAI-A11, FAI-A12, FAI-A13, FAI-A14",
                         "MITIGATED (this FAI)")

            for c_idx, txt in enumerate(xref_data):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(xref_widths[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el)
                tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el)
                r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve")
                t_el.text = txt
                r_el.append(t_el)
                p_el.append(r_el)
                tc.append(p_el)
                new_tr.append(tc)

            # Find cells and apply background
            new_row_obj = None
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    new_row_obj = r2
                    break
            if new_row_obj:
                set_cell_bg(new_row_obj.cells[3], "E2EFDA")
                new_row_obj.cells[0].paragraphs[0].runs[0].bold = True

            print("  §9 cross-reference RISK-15 row added")
            break

fai.save(FAI_PATH)
print(f"FAI checklist saved: {FAI_PATH}")
