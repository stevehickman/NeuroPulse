"""
Generate NeuroPulse Design Brief Revision 4.
Copies neuropulse_brief_r3.docx and applies all changes from the current
engineering session: dual-PD, five-layer keying, gasket, eject lever, PDMS
bonding, EEG routing, multi-FPC mgmt, new risks, new documents.
"""
import copy, datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ───────────────────────────────────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def append_xml_row(tbl, data, widths_in=None, bold_col0=True, bg=None):
    new_tr = OxmlElement("w:tr")
    for c_idx, txt in enumerate(data):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        if widths_in:
            w_el = OxmlElement("w:tcW")
            w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
            w_el.set(qn("w:type"), "dxa")
            tc_pr.append(w_el)
        tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr_el = OxmlElement("w:rPr")
        if bold_col0 and c_idx == 0:
            b_el = OxmlElement("w:b"); rpr_el.append(b_el)
        sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "18")
        rpr_el.append(sz_el); r_el.append(rpr_el)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
        r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
        new_tr.append(tc)
    tbl._tbl.append(new_tr)
    if bg:
        for r2 in tbl.rows:
            if r2._tr is new_tr:
                for cell in r2.cells:
                    set_bg(cell, bg)
                break
    return new_tr

def insert_xml_row_after(tbl, anchor_tr, data, widths_in=None, bold_col0=True, bg=None):
    new_tr = OxmlElement("w:tr")
    for c_idx, txt in enumerate(data):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        if widths_in:
            w_el = OxmlElement("w:tcW")
            w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
            w_el.set(qn("w:type"), "dxa")
            tc_pr.append(w_el)
        tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr_el = OxmlElement("w:rPr")
        if bold_col0 and c_idx == 0:
            b_el = OxmlElement("w:b"); rpr_el.append(b_el)
        sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "18")
        rpr_el.append(sz_el); r_el.append(rpr_el)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
        r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
        new_tr.append(tc)
    anchor_tr.addnext(new_tr)
    if bg:
        for r2 in tbl.rows:
            if r2._tr is new_tr:
                for cell in r2.cells:
                    set_bg(cell, bg)
                break
    return new_tr

# ══════════════════════════════════════════════════════════════════════════════
# Load Rev 3 as base
# ══════════════════════════════════════════════════════════════════════════════
doc = Document("docs/neuropulse_brief_r3.docx")

# ── 1. Update title / date / revision ─────────────────────────────────────────
for p in doc.paragraphs:
    if "Revision 3" in p.text and "Master Design Brief" in p.text:
        for run in p.runs:
            run.text = run.text.replace("Revision 3", "Revision 4")
        print("  Title updated → Revision 4")
    if "Revision 3 —" in p.text and "incorporates" in p.text:
        for run in p.runs:
            if "Revision 3" in run.text:
                run.text = run.text.replace(
                    "Revision 3 — incorporates all decisions through current session",
                    f"{datetime.date.today().strftime('%B %-d, %Y')}  ·  "
                    "Revision 4 — zone module engineering package complete; "
                    "24-risk register; dual-PD, five-layer keying, gasket, eject lever, "
                    "PDMS qualification, EEG routing")
        print("  Date/subtitle updated")
    if "May 4, 2026" in p.text or "Revision 3 —" in p.text:
        for run in p.runs:
            if "May 4, 2026" in run.text:
                run.text = run.text.replace(
                    "May 4, 2026  ·  Revision 3 — incorporates all decisions through current session",
                    f"{datetime.date.today().strftime('%B %-d, %Y')}  ·  "
                    "Revision 4 — zone module engineering package complete")

# ── 2. Add Rev 4 row to revision summary table (Table 0) ─────────────────────
rev_tbl = doc.tables[0]
append_xml_row(rev_tbl,
    ("Zone module engineering",
     "Dual photodiode PD1/PD2 (RISK-14 Option B). Five-layer zone keying — mechanical key "
     "+ ZONE_ID resistor + braille/numeral + tactile dots + bone conduction audio (RISK-15). "
     "Self-sealing co-moulded silicone gasket, no user RTV (RISK-16). Sliding eject lever "
     "3:1 MA ≤1N force (RISK-22). Multi-FPC bundle separation ≥2mm / EEG ≥15mm (RISK-17). "
     "EEG cable routing channel moulded into outer CFRP surface (RISK-21). PDMS SiO₂ "
     "interlayer bonding + 200-cycle IEC 60068-2-14 thermal cycling qualification "
     "BLOCKING before production (RISK-04). 24-risk register: 22 MITIGATED, 2 OPEN "
     "(RISK-03 regulatory opinion; RISK-20 CFRP Ra confirmation — both external). "
     "New documents: NP-HW-FPC-001, NP-PROC-FPC-001, NP-DRV-SHELL-001, NP-FAI-ZM-001, "
     "NP-COORD-001, NP-TOOL-ZM-001, NP-PROC-SUP-001."),
    bg="EBF3FB")
print("  Rev summary table row added")

# ── 3. Update durability table (Table 6) — add new session rows ───────────────
dur_tbl = doc.tables[6]

# Find the last row to anchor after
last_row = dur_tbl.rows[-1]

NEW_DUR_ROWS = [
    ("Second scalp-side photodiode PD2 per zone (RISK-14 Option B)",
     "PRE-TOOLING CRITICAL",
     "+$0.75–1.50/headset total",
     "Co-located on scalp-facing PDMS surface (pin 19 = PD2_CATHODE). PD1/PD2 ratio "
     "in firmware isolates PDMS fouling (PD1↓, PD2 stable) from LED aging (both↓ "
     "proportional). Eliminates service ambiguity single-PD could not resolve. "
     "T1 and T2 use identical zone module mould — firmware flag only."),
    ("Five-layer zone module keying (RISK-15)",
     "PRE-TOOLING CRITICAL",
     "Included in tooling",
     "Layer 1: asymmetric mechanical key per zone (physically prevents mis-insertion). "
     "Layer 2: ZONE_ID resistor (ZM-01=10kΩ → ZM-05=220kΩ, 1%, pin 18; Hub ADC reads "
     "before ENABLE). Layer 3: ISO 17049 braille zone abbreviation + raised numeral on "
     "module body. Layer 4: N tactile dots on shell adjacent to each zone slot. "
     "Layer 5: bone conduction audio confirms zone name on insertion. "
     "Covers colour-blind AND blind users. ZONE_ID firmware: 3×ADC at 100ms, "
     "≥2/3 pass required (RISK-18 debounce)."),
    ("Self-sealing co-moulded silicone gasket per zone module (RISK-16 Option A)",
     "PRE-TOOLING CRITICAL",
     "+$0.30–0.60/module",
     "Shore 40–50A UV-stable medical-grade silicone (ISO 10993-5), D-section 2.5mm × "
     "2.0mm, co-moulded with mechanical retention groove + silicone primer. Provides "
     "IPX4 seal at zone module interface with no user RTV or re-sealing. "
     "20% compression (0.4mm deflection) when seated. FAI-IPX-02 BLOCKING: "
     "IPX4 must pass after 10 field swap cycles before production."),
    ("Zone module sliding eject lever (RISK-22 Option A)",
     "PRE-TOOLING CRITICAL",
     "+$0.40/module",
     "10–12mm lever arm, 3:1 mechanical advantage, reduces extraction force from "
     "2–4 N to ≤1 N at lever tip. Recessed flush when closed; snap-fit detent "
     "(0.3–0.5 N release) prevents accidental ejection. 316SS hinge pin press-fit "
     "into moulded boss. FAI-A15 HFE test: 5 subjects with Parkinson's "
     "(Hoehn & Yahr II–III) or post-stroke weakness; ≤10 s extraction, ≤2 attempts."),
    ("Multi-FPC bundle management inside shell (RISK-17)",
     "PRE-TOOLING CRITICAL",
     "$0 (routing discipline)",
     "≥2mm separation between adjacent zone module FPCs (DRC-16). ≥15mm separation "
     "between FPC bundle and EEG electrode cables, or grounded 35µm Al foil barrier "
     "(DRC-18). 3 anchor bosses per FPC at ≥80mm intervals. All 5 Hub ZIF connectors "
     "on same PCB edge. Specified in NP-DRV-SHELL-001 §2.3."),
    ("EEG electrode cable routing channel in shell (RISK-21)",
     "PRE-TOOLING CRITICAL",
     "$0 (tooling)",
     "Dedicated 8×5mm moulded channel on outer CFRP inner surface, opposite side "
     "from FPC bundle. Provides ≥15mm separation from zone module FPCs. "
     "3 snap-in cable retention clips at 80mm intervals. Must be in shell tooling "
     "spec before first cut. Specified in NP-DRV-SHELL-001 §2.4. "
     "OI-09: ME to produce routing diagram before G2 gate."),
    ("PDMS–PI SiO₂ interlayer bonding process (RISK-04)",
     "PRE-TOOLING CRITICAL",
     "Process cost (CAT-C supplier)",
     "75 nm RF magnetron sputtered SiO₂ on PI coverlay + O₂ plasma activation (100 W / "
     "30 sccm / 60 s) of both surfaces. Bond within 10 min of plasma. Cure 120°C / 2 hr "
     "/ 50 g/cm². Achieves 174–860 N/m peel force (vs 7–12 N/m without interlayer). "
     "200-cycle IEC 60068-2-14 thermal cycling qualification (−10°C/+65°C) BLOCKING "
     "before any production FPCs. Batch peel ≥150 N/m per lot. NP-FAI-ZM-001 §3e."),
]

for row_data in NEW_DUR_ROWS:
    tr = append_xml_row(dur_tbl, row_data)
    # colour-code status cell
    for r2 in dur_tbl.rows:
        if r2._tr is tr:
            set_bg(r2.cells[1], "FCE4D6")  # orange = pre-tooling critical
            if r2.cells[0].paragraphs[0].runs:
                r2.cells[0].paragraphs[0].runs[0].bold = True
            break
    print(f"  Durability row added: {row_data[0][:55]}")

# ── 4. Update risks table (Table 8) — add RISK-20 and update RISK-04 ──────────
risk_tbl = doc.tables[8]

# Update RISK-04 row if present, add RISK-20 row, ensure existing rows current
risk04_tr = None
for row in risk_tbl.rows:
    if "PDMS" in row.cells[1].text and "delamination" in row.cells[1].text:
        risk04_tr = row._tr
        set_text(row.cells[0], "HIGH\n(MITIGATED)", bold=True,
                 color=(0x00, 0x70, 0x00), size=9)
        set_bg(row.cells[0], "E2EFDA")
        set_text(row.cells[2],
            "MITIGATED — SiO₂ 75nm interlayer bonding process specified "
            "(NP-HW-FPC-001 §9.2). 200-cycle IEC 60068-2-14 qualification "
            "BLOCKING before production (NP-FAI-ZM-001 FAI-TC02). "
            "CAT-C supplier selection required per NP-PROC-SUP-001 SUP-C-08. "
            "Batch peel ≥150 N/m per production lot (FAI-M02).", size=9)
        print("  Table 8: RISK-04 row updated to MITIGATED")
        break

# Add new HIGH rows for RISK-20 and RISK-22 (if not already present)
# Find last HIGH row or just append
append_xml_row(risk_tbl,
    ("HIGH (OPEN)",
     "CFRP shell slot rim surface finish (Ra ≤ 1.6 µm) — RISK-20\n"
     "Moulded CFRP typically achieves Ra 1.6–3.2 µm. Gasket seal (RISK-16) "
     "requires Ra ≤ 1.6 µm at slot rim. Gap between achievable and required "
     "finish is unconfirmed.",
     "BLOCKING — tooling cannot be released until CFRP tooling supplier provides "
     "written confirmation (letter + Ra measurement data from representative coupon) "
     "that Ra ≤ 1.6 µm is achievable. If Ra > 1.6 µm: escalate to ME + EE; "
     "gasket geometry or shell design must be revised. "
     "Tracked in NP-PROC-SUP-001 SUP-M-07/SUP-B-01 and §9."),
    bg="FCE4D6")
print("  Table 8: RISK-20 row added")

# ── 5. Update §8.1 Pending decisions ─────────────────────────────────────────
# Find the pending decisions bullet paragraphs and replace them
pending_start = None
pending_end = None
for i, p in enumerate(doc.paragraphs):
    if "8.1" in p.text and "Pending" in p.text:
        pending_start = i
    if pending_start and "8.2" in p.text and "Completed" in p.text:
        pending_end = i
        break

PENDING_ITEMS = [
    "Product name trademark search and clearance (US/EU/CA/AU) — before ANY external conversation",
    "400 mW/cm² regulatory opinion letter (RISK-03 — BLOCKING for all public material referencing irradiance)",
    "CFRP shell slot rim Ra ≤ 1.6 µm — written confirmation from tooling supplier with Ra measurement data (RISK-20 — BLOCKING for tooling release). Supplier qualification item: NP-PROC-SUP-001 SUP-M-07.",
    "PDMS CAT-C supplier selection + IEC 60068-2-14 200-cycle thermal cycling qualification passed (RISK-04 — BLOCKING before any production FPCs built). NP-FAI-ZM-001 FAI-TC02 must be on file.",
    "Zone module FPC layout freeze — PD2 component position must be locked before mould aperture (F-04 in NP-TOOL-ZM-001) can be specified",
    "Zone module mould design review — NP-TOOL-ZM-001 §5 checklist (all 8 features F-01 through F-08) signed off before steel cut",
    "ZONE_ID firmware debounce specification written into firmware requirements document (3×ADC at 100ms, ≥2/3 pass — RISK-18)",
    "HFE formative study for sliding eject lever — 5 subjects with Parkinson's/post-stroke grip weakness (NP-FAI-ZM-001 FAI-A15, NP-TOOL-ZM-001 OI-4)",
    "EEG cable routing path in shell CAD model (NP-DRV-SHELL-001 §2.4 OI-09 — required for DRC-18 verification)",
    "Hub tooling: intranasal probe dock + anchor posts for all interface covers + large-radius Boa cable channel + tool-free fan",
    "Shell tooling: anchor posts (5 zone + 3 port positions, colour-coded) + EEG cable routing channel (§2.4)",
    "Lens tooling: sliding rail + N42 magnet positions (≥1mm polymer wall) + AgNW spec + hard coat + EC driver contacts",
    "Goggle arm tooling: anchor hook for lens rim guard tether",
    "eMMC partition architecture + UHDR/SHDR separate encryption — before first firmware line",
    "Dual-bank OTA bootloader — must be in bootloader from first firmware line",
    "Partner optician network contract (S3 Rx programme)",
    "SBIR Phase I application — initiate at company formation",
    "First researcher contact: Rashidi-Ranjbar, then Jog, then Naeser",
    "SAB formation — Tsai outreach first (617-324-0305, SAB role only)",
    "LED emitter pulse current rating verification (660nm + 808–830nm at 120–180mA)",
]

if pending_start:
    # Replace existing pending bullet paragraphs
    bullets_to_remove = []
    for i in range(pending_start + 1, pending_end):
        p = doc.paragraphs[i]
        if p.text.strip() and p.text.strip() not in ("8.1 Pending decisions — action required before first tooling cut",):
            bullets_to_remove.append(p._p)

    for p_el in bullets_to_remove:
        p_el.getparent().remove(p_el)
    print(f"  Removed {len(bullets_to_remove)} old pending items")

    # Insert new pending items before the §8.2 heading
    s82_para = doc.paragraphs[pending_end - len(bullets_to_remove)]
    for item in reversed(PENDING_ITEMS):
        new_p = OxmlElement("w:p")
        s82_para._p.addprevious(new_p)
        for p_obj in doc.paragraphs:
            if p_obj._p is new_p:
                run = p_obj.add_run("• " + item)
                run.font.size = Pt(10)
                break
    print(f"  Inserted {len(PENDING_ITEMS)} updated pending items")

# ── 6. Update §8.2 Completed decisions — add new locked items ────────────────
NEW_COMPLETED = [
    "FPC 20-pin pinout locked: pin 18 = ZONE_ID (ZM-01=10kΩ → ZM-05=220kΩ, 1% 0402); pin 19 = PD2_CATHODE (scalp-side reference photodiode)",
    "Dual photodiode (RISK-14 Option B): PD2 on scalp-facing PDMS surface; PD1/PD2 ratio separates fouling from LED aging; T1 and T2 use identical zone module mould",
    "Five-layer zone keying (RISK-15): mechanical key + ZONE_ID resistor + braille/numeral + tactile shell dots + bone conduction audio; works for colour-blind and blind users",
    "Self-sealing gasket (RISK-16 Option A): co-moulded Shore 40–50A silicone, no user RTV, IPX4 rated after 10 field swap cycles",
    "Sliding eject lever (RISK-22 Option A): 10–12mm arm, 3:1 mechanical advantage, ≤1N extraction force; accessibility target Parkinson's Hoehn & Yahr II–III",
    "ZONE_ID firmware debounce: 3× ADC reads at 100ms intervals, ≥2/3 must pass before FAULT (RISK-18)",
    "PDMS bonding process: SiO₂ 75nm interlayer + O₂ plasma, ≥150 N/m peel; 200-cycle IEC 60068-2-14 qualification BLOCKING before production (RISK-04)",
    "EEG cable routing: dedicated 8×5mm channel on outer CFRP surface, ≥15mm separation from FPC bundle (RISK-21)",
    "Multi-FPC bundle management: ≥2mm inter-FPC separation; ≥15mm FPC-to-EEG cable; 3 anchor bosses per FPC; all 5 Hub ZIF connectors on same PCB edge (RISK-17)",
    "Risk register complete: 24 risks (RISK-01 through RISK-24); 22 MITIGATED; 2 OPEN (RISK-03 regulatory opinion; RISK-20 CFRP Ra — both external, pending supplier/counsel)",
    "Engineering documentation package complete: NP-HW-FPC-001, NP-PROC-FPC-001, NP-DRV-SHELL-001, NP-FAI-ZM-001, NP-COORD-001, NP-TOOL-ZM-001, NP-PROC-SUP-001",
]

# Find last completed decision paragraph and insert after it
last_completed_p = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Design brief Revision"):
        last_completed_p = p
        break

if last_completed_p:
    for item in reversed(NEW_COMPLETED):
        new_p = OxmlElement("w:p")
        last_completed_p._p.addprevious(new_p)
        for p_obj in doc.paragraphs:
            if p_obj._p is new_p:
                run = p_obj.add_run("• " + item)
                run.font.size = Pt(10)
                break
    print(f"  Inserted {len(NEW_COMPLETED)} new completed decisions")

# ── 7. Update final document reference bullet ─────────────────────────────────
for p in doc.paragraphs:
    if p.text.strip().startswith("Design brief Revision 1"):
        for run in p.runs:
            if "Revision 3" in run.text:
                run.text = run.text.replace(
                    "and this Revision 3",
                    "Revision 3, and this Revision 4")
        print("  Final doc reference updated")

# ── Save as Rev 4 ─────────────────────────────────────────────────────────────
OUT_PATH = "docs/neuropulse_brief_r4.docx"
doc.save(OUT_PATH)
print(f"\nSaved: {OUT_PATH}")
print("Done.")
