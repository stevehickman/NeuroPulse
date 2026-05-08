"""Scan full structure of risk register document."""
from docx import Document

RISK_PATH = "docs/neuropulse_fpc_zone_module_risks_revA.docx"
doc = Document(RISK_PATH)

print(f"Tables: {len(doc.tables)}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print()

# Print all table info
for t_idx, tbl in enumerate(doc.tables):
    print(f"Table {t_idx}: {len(tbl.rows)} rows × {len(tbl.columns)} cols")
    for r_idx, row in enumerate(tbl.rows):
        row_texts = [cell.text.strip()[:60] for cell in row.cells]
        if any(row_texts):
            print(f"  Row {r_idx}: {row_texts}")
    print()

# Print all paragraphs
print("Paragraphs:")
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"  [{i}] style={para.style.name!r}: {para.text[:100]!r}")
