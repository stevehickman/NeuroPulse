"""
NP-DB-005 §4 — replace the pre-hex BOM / COGS / GM% columns with the NP-COST-001 re-derivation.

Context
-------
NP-DB-005 Rev 5 §4 reproduced CLAUDE.md §2.1's cost columns behind an explicit caveat. The caveat
was a stopgap. NP-COST-001 Rev 1 (2026-08-16) re-derives those columns against the hex-tile
architecture and this script carries the result in.

Retail was LOCKED when this script was written and was UNLOCKED by principal direction on
2026-08-16 (NP-COST-001 Rev 2 §8). **The retail guard below is retained deliberately.** Unlocking
the constraint did not change any price -- the six configurations keep the prices currently in
force while the pricing decision is taken separately -- so the guard still does useful work: it
detects a price that has drifted without a decision behind it. When new prices ARE set, this script
is the place to carry them in, and the guard becomes the record of what they were changed from.

Why a patch script and not a generator: there is no `scripts/create_np_db_005.py` in this
repository (checked 2026-08-16) -- NP-DB-005 has no generator of record. The established repo
convention for editing a .docx of record is a `editscripts/patch_*.py` script, which is what this
is. See NP-COST-001 and the CLAUDE.md Rev 38 change note.

Idempotent: re-running is a no-op once the new figures are in place.

    python3 editscripts/patch_db005_bom_recost.py
"""

import sys
from docx import Document
from docx.shared import Pt, RGBColor

DOC = "docs/np_db_005.docx"

# config -> (BOM, COGS, GM%).  Retail is deliberately absent: it is locked.
# Source: NP-COST-001 Rev 2 §4.
NEW = {
    "Core — EEG only": ("$360–423", "$554–650", "−23% to −45%"),
    "Home Lite":       ("$642–705", "$896–984", "−50% to −64%"),
    "Home Standard ★": ("$897–959", "$1,196–1,278", "−41% to −51%"),
    "Home Premium":    ("$952–1,014", "$1,287–1,371", "−7% to −14%"),
    "Pro Entry":       ("$1,463–1,525", "$2,398–2,500", "+50% to +52%"),
    "Pro Full":        ("$2,136–2,198", "$3,728–3,836", "+73%"),
}

# Retail is UNLOCKED but unchanged.  Asserted, never written -- see module docstring.
LOCKED_RETAIL = {
    "Core — EEG only": "$449",
    "Home Lite": "$599",
    "Home Standard ★": "$849",
    "Home Premium": "$1,199",
    "Pro Entry": "$4,999",
    "Pro Full": "$13,999",
}

COL_CONFIG, COL_BOM, COL_COGS, COL_RETAIL, COL_GM = 0, 1, 2, 3, 4

CAVEAT = (
    "⚠ BOM / COGS / GM% re-derived 2026-08-16 against the hex-tile architecture "
    "(NP-COST-001 Rev 2 §4). RETAIL PRICING IS UNLOCKED (principal, 2026-08-16) -- but NO PRICE "
    "IS SET: the prices in this table are the ones currently in force, and every configuration "
    "is loss-making at them. NP-COST-001 §8 publishes the implied ladder -- break-even Home "
    "Standard ~$1,196-1,278, target margin $1,869-1,997 (2.20-2.35x the current $849); Core "
    "$955-1,121; Home Lite $1,445-1,587; Home Premium $2,475-2,637. Setting a price is a "
    "separate commercial decision and OI-HEXTILE-06 must be decided first (OI-COST-10), else "
    "the price is set against a cost that decision invalidates. Every cost figure is a "
    "FLOOR: it excludes term U (the emitter-count delta net of the retired hub-side drive "
    "stage), which is uncosted in both directions because OI-HEXTILE-02 has selected no "
    "660/808 nm emitter -- so OI-HUB-C08 remains OPEN and cannot be closed from the record. "
    "All four T1 configurations are gross-margin negative at the floor. Three inputs are "
    "assumptions, not decisions: per-configuration tile population (OI-COST-01), full-L1-per-"
    "configuration (OI-COST-04), per-configuration BOM->COGS multiplier (OI-COST-05). Socket "
    "count ~80 is PROVISIONAL pending REG-1/ACT-1. Read NP-COST-001 before quoting any figure."
)


def set_cell(cell, text, bold=False, color=None):
    """Replace a cell's text, preserving the first run's formatting."""
    para = cell.paragraphs[0]
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)
    if para.runs:
        first, rest = para.runs[0], para.runs[1:]
        for r in rest:
            r._element.getparent().remove(r._element)
        first.text = text
        run = first
    else:
        run = para.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


def find_config_table(doc):
    for table in doc.tables:
        header = [c.text.strip() for c in table.rows[0].cells]
        if len(header) > COL_GM and header[COL_BOM] == "BOM" and header[COL_COGS] == "COGS":
            return table
    return None


def main():
    doc = Document(DOC)
    table = find_config_table(doc)
    if table is None:
        sys.exit("FAIL: NP-DB-005 configuration table not found -- document structure changed.")

    seen, changed = set(), 0
    for row in table.rows[1:]:
        name = row.cells[COL_CONFIG].text.strip()
        if name not in NEW:
            sys.exit(f"FAIL: unexpected configuration row {name!r} -- refusing to guess.")
        seen.add(name)

        # Retail is locked: verify, never write.
        retail = row.cells[COL_RETAIL].text.strip()
        if retail != LOCKED_RETAIL[name]:
            sys.exit(f"FAIL: retail for {name!r} is {retail!r}, expected "
                     f"{LOCKED_RETAIL[name]!r}. Retail is locked -- aborting.")

        bom, cogs, gm = NEW[name]
        if (row.cells[COL_BOM].text.strip() == bom
                and row.cells[COL_COGS].text.strip() == cogs
                and row.cells[COL_GM].text.strip() == gm):
            continue  # already patched

        set_cell(row.cells[COL_BOM], bom)
        set_cell(row.cells[COL_COGS], cogs)
        # Negative margin is rendered red + bold: it must not read as an ordinary number.
        negative = gm.startswith("−")
        set_cell(row.cells[COL_GM], gm, bold=True,
                 color=RGBColor(0xC0, 0x00, 0x00) if negative else None)
        changed += 1

    missing = set(NEW) - seen
    if missing:
        sys.exit(f"FAIL: configurations absent from the table: {sorted(missing)}")

    # Caveat paragraph, immediately after the table. Replaced in place if already present.
    marker = "BOM / COGS / GM% re-derived"
    existing = next((p for p in doc.paragraphs if marker in p.text), None)
    if existing is not None:
        set_cell_para = existing
        for r in set_cell_para.runs[1:]:
            r._element.getparent().remove(r._element)
        if set_cell_para.runs:
            set_cell_para.runs[0].text = CAVEAT
    else:
        para = doc.add_paragraph()
        table._element.addnext(para._element)
        run = para.add_run(CAVEAT)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.save(DOC)
    print(f"OK: {DOC} -- {changed} configuration row(s) recosted, "
          f"{len(LOCKED_RETAIL)} locked retail price(s) verified unchanged.")


if __name__ == "__main__":
    main()
