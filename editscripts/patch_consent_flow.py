"""
Update consent decision flow with:
1. On-device anonymization architecture requirement (UHDR section)
2. Irreversibility notice + forward-effectiveness guarantee (L3 blanket consent)
3. Per-project consent workflow — add forward-effectiveness clarification
4. New §5.3 "Research data anonymization architecture" requirement
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def append_text(cell, text, bold=False, size=9):
    """Append text to existing content in a cell (new paragraph)."""
    para = cell.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

def add_row_after(anchor_tr, row_data, col_widths_in, status_col=None, status_color=None):
    new_tr = OxmlElement("w:tr")
    for c_idx, txt in enumerate(row_data):
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        w_el = OxmlElement("w:tcW")
        w_el.set(qn("w:w"), str(int(col_widths_in[c_idx] * 1440)))
        w_el.set(qn("w:type"), "dxa")
        tc_pr.append(w_el); tc.append(tc_pr)
        p_el = OxmlElement("w:p")
        r_el = OxmlElement("w:r")
        rpr_el = OxmlElement("w:rPr")
        sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
        rpr_el.append(sz_el); r_el.append(rpr_el)
        t_el = OxmlElement("w:t")
        t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
        r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
        new_tr.append(tc)
    anchor_tr.addnext(new_tr)
    return new_tr

BRIEF_PATH = "docs/neuropulse_brief_r4.docx"
brief = Document(BRIEF_PATH)

# ── Locate tables ─────────────────────────────────────────────────────────────
# Table 2: UHDR/SHDR  — find by scanning for "UHDR" in first column
# Table 3: Consent tiers — find by scanning for "Monitor" in first column
# Table 4: A priori consent — find by scanning for "L1" in first column

uhdr_tbl = None
consent_tier_tbl = None
apriori_tbl = None

for ti, tbl in enumerate(brief.tables):
    first_col_texts = [r.cells[0].text.strip() for r in tbl.rows]
    if any("UHDR" in t for t in first_col_texts) and uhdr_tbl is None:
        uhdr_tbl = (ti, tbl)
        print(f"UHDR table: Table {ti}")
    if any("Monitor" in t for t in first_col_texts) and consent_tier_tbl is None:
        consent_tier_tbl = (ti, tbl)
        print(f"Consent tier table: Table {ti}")
    if any("L1" in t for t in first_col_texts) and apriori_tbl is None:
        apriori_tbl = (ti, tbl)
        print(f"A priori consent table: Table {ti}")

# ── 1. UHDR table — append on-device anonymization note to UHDR storage row ──
if uhdr_tbl:
    tbl = uhdr_tbl[1]
    for row in tbl.rows:
        if row.cells[0].text.strip() == "UHDR":
            # Find the storage or definition cell — typically column 1 (long text)
            if len(row.cells) >= 2:
                existing = row.cells[1].text
                if "biometric" in existing or "AES" in existing or "backup" in existing:
                    append_text(
                        row.cells[1],
                        "\nResearch data flow: All anonymization of UHDR for research "
                        "purposes occurs on-device (within the NeuroPulse app), generating "
                        "a fresh extract directly from the AES-256 encrypted source partition "
                        "for each approved study. NeuroPulse infrastructure receives only the "
                        "pre-anonymized extract. NeuroPulse never decrypts, accesses, or "
                        "temporarily holds raw UHDR at any point in the research data pipeline. "
                        "No persistent per-user anonymized data store is maintained on "
                        "NeuroPulse servers.",
                        size=9
                    )
                    print("  UHDR row: on-device anonymization note appended")
                    break

# ── 2. A priori consent table — update L3 blanket consent "If yes" cell ──────
if apriori_tbl:
    tbl = apriori_tbl[1]
    for row in tbl.rows:
        label = row.cells[0].text.strip()
        if label.startswith("L3"):
            # "If yes" is column 2 (0-indexed)
            if len(row.cells) >= 4:
                if_yes_cell = row.cells[2]
                existing = if_yes_cell.text
                if "blanket" in existing.lower() or "pre-approve" in existing.lower() or "k≥10" in existing.lower():
                    append_text(
                        if_yes_cell,
                        "\nIrreversibility notice (displayed at consent time): "
                        "\"Once your anonymised data has been included in a published study, "
                        "it cannot be individually withdrawn from that dataset — this is a "
                        "fundamental property of k-anonymised aggregate data. However, because "
                        "NeuroPulse anonymises your data fresh from your device for each new "
                        "study, withdrawing consent immediately and permanently stops any further "
                        "data flowing to any future dataset — including data from sessions that "
                        "occurred before your withdrawal. Your historical sessions remain on your "
                        "device under your sole control.\"",
                        size=9
                    )
                    print("  L3 blanket consent: irreversibility + forward-effectiveness notice appended")
                    break

# ── 3. A priori consent table — update L4 row to reference the notice ─────────
if apriori_tbl:
    tbl = apriori_tbl[1]
    for row in tbl.rows:
        label = row.cells[0].text.strip()
        if label.startswith("L4"):
            if len(row.cells) >= 4:
                brand_cell = row.cells[3]  # Brand ambassador mechanism column
                existing = brand_cell.text
                if "Results" in existing or "results" in existing:
                    append_text(
                        brand_cell,
                        "\nResults notification closes the loop for all enrolled participants — "
                        "including null results. Users who withdrew consent still receive results "
                        "for studies they previously participated in (notification-only, no new data).",
                        size=9
                    )
                    print("  L4 results row: withdrawal-participant results note appended")
                    break

# ── 4. Consent tier table — update Research tier to note on-device anon ───────
if consent_tier_tbl:
    tbl = consent_tier_tbl[1]
    for row in tbl.rows:
        if "Research" in row.cells[0].text:
            if len(row.cells) >= 4:
                # Last column or UHDR elements column
                for ci, cell in enumerate(row.cells):
                    if "k-anon" in cell.text.lower() or "k≥10" in cell.text or "anonymis" in cell.text.lower():
                        append_text(
                            cell,
                            "\nAnonymization occurs on-device. NeuroPulse receives "
                            "only pre-anonymized extracts per study.",
                            size=9
                        )
                        print(f"  Research tier row col {ci}: on-device anon note appended")
                        break
            break

# ── 5. Per-project consent workflow — insert forward-effectiveness note ────────
# Find the per-project workflow numbered list paragraphs by scanning all paragraphs
# Looking for: "Per-project contact workflow" or "Patient decision" step

per_project_para_idx = None
for pi, para in enumerate(brief.paragraphs):
    if "Per-project contact workflow" in para.text or "per-project consent" in para.text.lower():
        per_project_para_idx = pi
        print(f"  Per-project workflow para at index {pi}: '{para.text[:60]}'")
        break

# Insert after the numbered steps — find "Patient decision" step 4
withdrawal_inserted = False
for pi, para in enumerate(brief.paragraphs):
    if ("Results notification closes loop" in para.text
            or "Results notification" in para.text
            or ("closes loop" in para.text and "opted in" in para.text)):
        # Insert a new paragraph immediately after
        new_p = OxmlElement("w:p")
        new_r = OxmlElement("w:r")
        new_rpr = OxmlElement("w:rPr")
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "18")  # 9pt
        new_rpr.append(sz); new_r.append(new_rpr)
        new_t = OxmlElement("w:t")
        new_t.text = (
            "Consent withdrawal effect: If a participant withdraws consent after "
            "step 5, all future data flows to all studies are immediately blocked. "
            "Because anonymization occurs on-device fresh per study, withdrawal is "
            "fully effective for all data periods — including sessions that occurred "
            "before the withdrawal date. Already-published study datasets cannot be "
            "modified (irreversibility notice given at step 4); no new data flows ever."
        )
        new_r.append(new_t); new_p.append(new_r)
        para._p.addnext(new_p)
        withdrawal_inserted = True
        print(f"  Per-project withdrawal note inserted after para {pi}")
        break

if not withdrawal_inserted:
    print("  WARNING: Per-project workflow para not found — withdrawal note not inserted")

# ── 6. Add §5.3 Research Data Anonymization Architecture section ──────────────
# Find §5.2 Predictive maintenance heading and insert §5.3 after its last paragraph
predictive_para_idx = None
for pi, para in enumerate(brief.paragraphs):
    if "Predictive maintenance" in para.text and ("5.2" in para.text or "SHDR" in para.text):
        predictive_para_idx = pi
        print(f"  §5.2 para at index {pi}: '{para.text[:60]}'")

# Alternative: find §6 Clinical Consent heading and insert §5.3 before it
consent_heading_idx = None
for pi, para in enumerate(brief.paragraphs):
    if "Clinical Consent" in para.text and ("6" in para.text or "CONSENT" in para.text.upper()):
        consent_heading_idx = pi
        print(f"  §6 consent heading at index {pi}: '{para.text[:60]}'")
        break

insert_after_idx = None
if consent_heading_idx and consent_heading_idx > 0:
    insert_after_idx = consent_heading_idx - 1

if insert_after_idx is not None:
    anchor_p = brief.paragraphs[insert_after_idx]._p

    # Heading paragraph
    h_p = OxmlElement("w:p")
    h_pPr = OxmlElement("w:pPr")
    h_pStyle = OxmlElement("w:pStyle"); h_pStyle.set(qn("w:val"), "Heading2")
    h_pPr.append(h_pStyle); h_p.append(h_pPr)
    h_r = OxmlElement("w:r")
    h_rPr = OxmlElement("w:rPr")
    h_sz = OxmlElement("w:sz"); h_sz.set(qn("w:val"), "24")
    h_rPr.append(h_sz); h_r.append(h_rPr)
    h_t = OxmlElement("w:t")
    h_t.text = "5.3 Research data anonymization architecture (locked)"
    h_r.append(h_t); h_p.append(h_r)
    anchor_p.addnext(h_p)

    paras = [
        ("Architecture requirement",
         "All anonymization of UHDR data for research purposes must occur on-device, "
         "within the NeuroPulse app, before any data leaves the device. The process for "
         "each approved study is:"),
        ("Step 1 — Study approval signal",
         "NeuroPulse server sends the device a signed study descriptor: study ID, "
         "anonymization parameters (k-value ≥10, suppression rules, date-rounding interval), "
         "and approved UHDR element list. Study descriptor is cryptographically signed by "
         "NeuroPulse (same signing infrastructure as session protocols)."),
        ("Step 2 — On-device extract generation",
         "The app reads the encrypted UHDR partition (decryption uses only the user biometric-"
         "derived key — app is the only process with transient access during an active session). "
         "Anonymization transformations are applied in-app: (a) k-anonymity grouping, "
         "(b) date/time rounding to ≥1-week intervals, (c) all direct identifiers removed, "
         "(d) quasi-identifier suppression per study descriptor. The anonymized extract is "
         "signed with the study ID and device hardware attestation token."),
        ("Step 3 — Extract transmission",
         "Only the pre-anonymized, signed extract is transmitted to NeuroPulse research "
         "infrastructure. Raw UHDR never leaves the device. NeuroPulse servers receive and "
         "store only the extract. No persistent per-user anonymized data store is maintained; "
         "each study receives its own extract. Extracts are keyed to study ID and device ID only."),
        ("Step 4 — No server-side re-identification path",
         "NeuroPulse infrastructure cannot re-identify users from extracts: no raw UHDR, "
         "no biometric key, no linkage table. The device ID maps only to the SHDR warranty owner "
         "record (shipping address + warranty dates) — SHDR contains no user biology. "
         "Researchers access only aggregated study datasets with no device ID fields."),
        ("Consent withdrawal effect",
         "Because anonymization is performed on-device, per-study, on-demand: withdrawing "
         "consent permanently blocks the device from processing any future study descriptors. "
         "No further extracts are generated or transmitted. This withdrawal is effective for "
         "all data periods — including UHDR from sessions that occurred before the withdrawal "
         "date — because the source data remains on the device and the extraction step cannot "
         "be triggered remotely without device participation. Already-transmitted extracts "
         "included in published datasets cannot be individually removed (k-anonymized aggregates "
         "cannot be individually withdrawn — irreversibility notice is given at consent time)."),
        ("Audit trail",
         "Study ID, study descriptor hash, extract transmission timestamp, and extract size "
         "(byte count, not content) are logged in SHDR. User can inspect which studies their "
         "device has contributed to at any time via the app. This log is the user's proof of "
         "participation and is never shared with researchers."),
    ]

    # Insert in reverse order so addnext preserves sequence
    prev_p = h_p
    for label, body in paras:
        body_p = OxmlElement("w:p")
        body_r = OxmlElement("w:r")
        body_rPr = OxmlElement("w:rPr")
        body_sz = OxmlElement("w:sz"); body_sz.set(qn("w:val"), "18")
        body_rPr.append(body_sz); body_r.append(body_rPr)
        body_t = OxmlElement("w:t")
        body_t.set(qn("xml:space"), "preserve")
        body_t.text = f"{label}: {body}"
        body_r.append(body_t); body_p.append(body_r)
        prev_p.addnext(body_p)
        prev_p = body_p

    print("  §5.3 Research data anonymization architecture inserted")
else:
    print("  WARNING: Could not locate §6 heading for §5.3 insertion")

brief.save(BRIEF_PATH)
print(f"\nDesign Brief Rev 4 saved with consent flow updates")
