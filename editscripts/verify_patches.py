"""Quick verification of all patches."""
from docx import Document

print("=== SBIR Proposal — connector references ===")
sbir = Document("docs/np_sbir_001.docx")
for para in sbir.paragraphs:
    if "SlimStack" in para.text or "0.35" in para.text or "FH34S" in para.text or "Hirose" in para.text:
        print(f"  {para.text[:120]!r}")

print("\n=== FPC Spec — fab note RA copper ===")
spec = Document("docs/superseded/np_hw_fpc_001.docx")
for i, para in enumerate(spec.paragraphs):
    if "Rolled Annealed" in para.text or "DRAWING REQUIREMENT" in para.text or "IPC-4204" in para.text:
        print(f"  [{i}] {para.text[:120]!r}")

print("\n=== Risk Register — RISK-08/09/10 Summary Table ===")
risks = Document("docs/superseded/np_risk_001.docx")
tbl = risks.tables[0]
for row in tbl.rows:
    rid = row.cells[0].text.strip()
    if rid in ("RISK-08", "RISK-09", "RISK-10"):
        print(f"  {rid}:")
        print(f"    Severity: {row.cells[1].text.strip()!r}")
        print(f"    Root Cause: {row.cells[3].text.strip()[:80]!r}")
        print(f"    Status: {row.cells[8].text.strip()!r}")

print("\n=== Risk Register — RISK-09/10 detail paragraphs ===")
for para in risks.paragraphs:
    if "RISK-09" in para.text or "RISK-10" in para.text or "MEDIUM (MITIGATED)" in para.text:
        print(f"  {para.text[:100]!r}")

print("\n=== Procurement Doc — connector exclusions ===")
proc = Document("docs/np_proc_fpc_001.docx")
for para in proc.paragraphs:
    if "SlimStack" in para.text or "0.35" in para.text or "FH34S" in para.text or "explicitly" in para.text.lower():
        print(f"  {para.text[:120]!r}")
for tbl in proc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            if "SlimStack" in cell.text or "0.35" in cell.text:
                print(f"  Table: {cell.text[:100]!r}")
