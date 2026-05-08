"""Continue patch_risk16: fix coord dashboard + add FAI IPX4 section."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def ins_before(doc, target, text, bold=False, color=None, size=10):
    new_p = OxmlElement("w:p")
    target._p.addprevious(new_p)
    for p_obj in doc.paragraphs:
        if p_obj._p is new_p:
            run = p_obj.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p_obj
    return None

# ── Coord checklist: fix dashboard G2-05 ──
COORD_PATH = "docs/neuropulse_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)
for tbl in coord.tables:
    for row in tbl.rows:
        cells = row.cells
        if (len(cells) == 5
                and cells[0].text.strip() == "G2-05"
                and cells[1].text.strip() == "G2"):
            set_text(cells[4], "OPEN — tooling spec; FAI test required")
            set_bg(cells[4], "E2EFDA")
            print("  Dashboard G2-05 updated")
            break
coord.save(COORD_PATH)
print(f"  Coord checklist saved")

# ── FAI: add §4c Sealing and IPX4 Verification ──
FAI_PATH = "docs/neuropulse_fai_zone_module.docx"
fai = Document(FAI_PATH)

# Find SECTION 5 heading paragraph
s5_para = None
for p in fai.paragraphs:
    if "SECTION 5" in p.text:
        s5_para = p
        break
print(f"  SECTION 5 anchor: {s5_para.text[:60]!r}")

# Build the IPX4 table, append to end of document, then relocate via XML
fai_tbl = fai.add_table(rows=1, cols=6)
fai_tbl.style = "Table Grid"

col_headers = ["Item", "Inspection / Test", "Method", "Accept Criterion", "Result", "Sign"]
widths_in   = [0.65,   2.3,                1.2,      1.55,               0.65,     0.45]
HDR_BG = "1F497D"

hdr_row = fai_tbl.rows[0]
for i, h in enumerate(col_headers):
    set_bg(hdr_row.cells[i], HDR_BG)
    run = hdr_row.cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(8)
    hdr_row.cells[i].width = Inches(widths_in[i])

FAI_IPX_ROWS = [
    ("FAI-IPX-01",
     "Initial IPX4 baseline: assembled headset, all 5 zone modules factory-installed. "
     "IEC 60529 IPX4 water spray from all directions, 10 minutes.",
     "IEC 60529 IPX4 nozzle, 10 min",
     "No water ingress to Hub PCB, any connector, or zone module PCB. "
     "EEG impedance all channels < 10 kΩ after test (functional check).",
     "", ""),
    ("FAI-IPX-02",
     "[BLOCKING] Field-replacement IPX4 re-test: remove and re-insert all 5 zone "
     "modules 10 times each (50 total swaps, no tools, no RTV, user-pace); "
     "then repeat full IEC 60529 IPX4 spray test without any re-sealing.",
     "10 swap cycles per zone then IEC 60529 IPX4",
     "Same pass criterion as FAI-IPX-01. Zero water ingress after 10 field swaps. "
     "Any ingress → HALT production; gasket geometry or shell rim tolerance "
     "non-conformant. Escalate to Mechanical Engineering immediately.",
     "", ""),
    ("FAI-IPX-03",
     "Gasket visual and dimensional check after 10 swap cycles: inspect all 5 zone "
     "module gaskets under 10× magnification; measure gasket cross-section height "
     "on 2 modules with calibrated comparator.",
     "10× magnifier; calibrated comparator (2 modules)",
     "No tears, cuts, or embedded debris. Compression set < 10%: "
     "gasket height ≥ 1.80 mm (vs. 2.00 mm nominal) after cycle release.",
     "", ""),
    ("FAI-IPX-04",
     "Gasket shelf-life / ageing qualification (first-article qualification, "
     "not every lot): store 5 sealed zone modules at 40°C / 75% RH for 90 days "
     "(equivalent to ~2 years at ambient per Arrhenius). Perform IPX4 test and "
     "gasket dimension check at end of storage.",
     "IEC 60068-2-66 damp heat storage; then IPX4 + dimensional check",
     "IPX4 pass. Gasket dimensions within 5% of nominal. "
     "Compression set < 15% after ageing. "
     "Qualification pass required before first production run.",
     "", ""),
]

for r_idx, row_data in enumerate(FAI_IPX_ROWS):
    row = fai_tbl.add_row()
    bg = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
    for c_idx, txt in enumerate(row_data):
        cell = row.cells[c_idx]
        set_bg(cell, bg)
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(8)
        if c_idx == 0:
            run.bold = True
        cell.width = Inches(widths_in[c_idx])

# Move the table XML to just before SECTION 5
tbl_xml = fai_tbl._tbl
s5_para._p.addprevious(tbl_xml)

# Insert a spacer paragraph after the table (now between table and SECTION 5)
spacer_p = OxmlElement("w:p")
s5_para._p.addprevious(spacer_p)

# Insert the section heading before the table
heading_p = OxmlElement("w:p")
tbl_xml.addprevious(heading_p)
for p_obj in fai.paragraphs:
    if p_obj._p is heading_p:
        run = p_obj.add_run(
            "4c.  Sealing and IPX4 Verification  [RISK-16 — Option A gasket]")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        break

# Insert body note before the heading
note_p = OxmlElement("w:p")
heading_p.addprevious(note_p)
for p_obj in fai.paragraphs:
    if p_obj._p is note_p:
        run = p_obj.add_run(
            "Self-sealing co-moulded silicone gasket — no user action, no RTV, "
            "no re-sealing procedure. FAI-IPX-02 is BLOCKING.")
        run.font.size = Pt(10)
        break

print("  §4c IPX4 section inserted before SECTION 5")

# Add RISK-16 to §9 cross-reference table
for tbl2 in fai.tables:
    for row in tbl2.rows:
        if row.cells[0].text.strip() == "RISK-15":
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            xref_widths = [0.65, 2.0, 2.35, 1.8]
            xref_texts = ("RISK-16",
                          "IPX4 field-replacement — Option A gasket",
                          "FAI-IPX-01, FAI-IPX-02, FAI-IPX-03, FAI-IPX-04",
                          "MITIGATED (this FAI)")
            for c_idx, txt in enumerate(xref_texts):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(xref_widths[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            for r2 in tbl2.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[3], "E2EFDA")
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  §9 cross-reference RISK-16 row added")
            break

fai.save(FAI_PATH)
print(f"  FAI checklist saved: {FAI_PATH}")
print("Done.")
