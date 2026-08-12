"""
NOTE (2026-07-28): the self-sealing compression gasket approach this script describes is
architecture-independent and still valid under the hex-tile/socket architecture
(NP-HEX-ZM-001 §6: "each hex needs a perimeter gasket... a per-tile seam-length budget is
required") — this script is NOT retired the way its RISK-11/15/17 siblings are. Only its
"zone module" terminology was stale; editscripts/patch_hexzm_light_fixes.py generalizes the
RISK-16 risk-register text to "hex-tile module" in place. Do not re-run this script itself
(not idempotent against the now-further-patched documents); its FPC-spec §12 gasket
dimensions and FAI-IPX-01..04 test protocol remain the current guidance, just described with
the old noun in the FPC spec / FAI doc bodies this script wrote directly.

Patch RISK-16 — Option A: self-sealing compression gasket.

Changes:
  1. FPC spec §12: replace RTV bead text with gasket spec
  2. Risk register RISK-16 → MITIGATED; add detailed section
  3. NP-COORD-001 G2-05: update to Option A decided / gasket spec
  4. NP-FAI-ZM-001: add §4c sealing/IPX4 sub-section + post-cycle IPX4 items
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── helpers ────────────────────────────────────────────────────────────────

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=9):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def ins_before(doc, target, text, bold=False, bullet=False, color=None, size=10):
    """Insert a paragraph immediately before `target` paragraph."""
    new_p = OxmlElement("w:p")
    target._p.addprevious(new_p)
    for p_obj in doc.paragraphs:
        if p_obj._p is new_p:
            if bullet:
                try:
                    p_obj.style = doc.styles["List Bullet"]
                except Exception:
                    pass
            run = p_obj.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor(*color)
            return p_obj
    return None

def append_table_row(tbl, cols_data, widths_in, bg_even="FFFFFF", status_col=None):
    """Append a row to an existing table."""
    new_row = tbl.add_row()
    for c_idx, txt in enumerate(cols_data):
        cell = new_row.cells[c_idx]
        set_bg(cell, bg_even)
        run = cell.paragraphs[0].add_run(str(txt))
        run.font.size = Pt(8)
        if c_idx == 0:
            run.bold = True
        if status_col is not None and c_idx == status_col:
            set_bg(cell, "FFE699")
    if widths_in:
        for i, w in enumerate(widths_in):
            if i < len(new_row.cells):
                new_row.cells[i].width = Inches(w)
    return new_row

# ─────────────────────────────────────────────────────────────────────────────
# 1.  FPC spec §12 — replace RTV bead with Option A gasket spec
# ─────────────────────────────────────────────────────────────────────────────

SPEC_PATH = "docs/superseded/neurone_fpc_zone_module_spec_revA.docx"
spec = Document(SPEC_PATH)

GASKET_SPEC = (
    "[DRAWING REQUIREMENT — ZONE MODULE TOOLING]  "
    "Zone module edge seal — self-sealing compression gasket (Option A, selected): "
    "A co-moulded Shore 40–50A UV-stable medical-grade silicone (ISO 10993-5) "
    "lip seal runs continuously around the full perimeter of the zone module face, "
    "1 mm inset from the PDMS window edge. "
    "Profile: D-section, 2.5 mm wide × 2.0 mm height nominal; "
    "compressed to 1.6 mm (20% compression, 0.4 mm deflection) at full insertion. "
    "Compression force contribution: ~2–4 N (additive to mechanical key engagement force). "
    "The gasket is co-moulded with the zone module body — it is permanent and requires no "
    "user attention on any reinsertion. "
    "On every zone module removal and reinsertion, the gasket re-seats automatically "
    "against the shell slot rim. No RTV, no syringe, no user procedure. "
    "IPX4 compliance is maintained through the device lifetime for all user-replaceable "
    "zone module swaps. "
    "The shell slot rim contact surface must be smooth (Ra ≤ 1.6 µm) and free of "
    "draft-angle steps that would prevent full gasket seating. "
    "Specify in shell tooling: slot rim width ≥ 2.5 mm to provide full gasket bearing surface. "
    "BOM delta vs. RTV bead approach: +$0.30–0.60 per zone module "
    "(co-moulded gasket NRE absorbed if specified before first tooling cut)."
)

s12_changes = 0
for i, para in enumerate(spec.paragraphs):
    t = para.text.strip()
    # Replace the RTV bead paragraph
    if t.startswith("Zone module edge seal: silicone RTV bead"):
        for run in para.runs:
            run.text = ""
        run = para.add_run(GASKET_SPEC)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        s12_changes += 1
        print(f"  §12 RTV bead paragraph replaced at para [{i}]")
        break

# Also update the IPX4 test row in the test requirements table
for ti, tbl in enumerate(spec.tables):
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 4 and "IPX4 splash test" in cells[0].text:
            set_text(cells[3],
                "Qualification on assembled headset — PLUS field-replacement IPX4 test: "
                "repeat after 10 zone module swap cycles per NP-FAI-ZM-001 FAI-IPX-02.",
                size=9)
            s12_changes += 1
            print(f"  IPX4 test table row updated in Table {ti}")
            break

spec.save(SPEC_PATH)
print(f"  FPC spec saved ({s12_changes} changes)")

# ─────────────────────────────────────────────────────────────────────────────
# 2.  Risk register — RISK-16 → MITIGATED + detailed section
# ─────────────────────────────────────────────────────────────────────────────

RISK_PATH = "docs/superseded/neurone_fpc_zone_module_risks_revA.docx"
risk_doc = Document(RISK_PATH)

for row in risk_doc.tables[0].rows:
    if row.cells[0].text.strip() == "RISK-16":
        set_text(row.cells[1], "LOW\n(MITIGATED)", bold=True,
                 color=(0x00, 0x70, 0x00))
        set_bg(row.cells[1], "E2EFDA")
        set_text(row.cells[5],
            "MITIGATED — Option A selected: self-sealing compression gasket. "
            "Co-moulded Shore 40–50A silicone lip seal on zone module perimeter; "
            "20% compression (0.4 mm deflection) against shell slot rim on every insertion. "
            "No user action, no RTV syringe, no procedure. "
            "IPX4 maintained through all field replacements. "
            "Gasket spec added to FPC spec §12. "
            "IPX4 re-test after 10 swap cycles required (NP-FAI-ZM-001 §4c).")
        set_text(row.cells[8],
            "MITIGATED — Option A gasket spec in §12; "
            "IPX4 field-replacement test in NP-FAI-ZM-001 FAI-IPX-01 to FAI-IPX-04",
            bold=True, color=(0x00, 0x70, 0x00))
        set_bg(row.cells[8], "92D050")
        print("  RISK-16 summary row → MITIGATED")
        break

# Detailed section before OPEN ISSUES
open_issues_idx = None
for i, para in enumerate(risk_doc.paragraphs):
    if "OPEN ISSUES" in para.text:
        open_issues_idx = i
        break

target = risk_doc.paragraphs[open_issues_idx]

ins_before(risk_doc, target, "")

ins_before(risk_doc, target,
    "→ Action 3 (OPEN — FAI): IPX4 field-replacement test (NP-FAI-ZM-001 FAI-IPX-02): "
    "10 swap cycles then IEC 60529 IPX4 splash test. Must pass before production release.",
    bullet=True, color=(0xC0, 0x60, 0x00))

ins_before(risk_doc, target,
    "→ Action 2 (OPEN — tooling spec): Shell slot rim: smooth contact surface Ra ≤ 1.6 µm, "
    "rim width ≥ 2.5 mm, no draft-angle steps. Must be in shell tooling spec before first-cut.",
    bullet=True, color=(0xC0, 0x60, 0x00))

ins_before(risk_doc, target,
    "→ Action 1 (OPEN — zone module tooling spec): Co-moulded gasket spec (Shore 40–50A, "
    "D-section 2.5 mm × 2.0 mm, 20% compression at insertion) in zone module mould drawing. "
    "Must be specified before first-cut — cannot be retrofitted.",
    bullet=True, color=(0xC0, 0x60, 0x00))

ins_before(risk_doc, target,
    "Option A (self-sealing gasket) was selected over Option B (user-applied RTV syringe) "
    "because silicone RTV requires accurate bead application around a complex perimeter — "
    "a task prone to gaps, excess material fouling the PDMS window, and non-compliance "
    "when users skip the step. The gasket approach removes user behaviour from the "
    "IPX4 compliance path entirely. The co-moulded gasket adds ~$0.30–0.60 BOM "
    "per zone module ($1.50–3.00 per headset for 5 zones) vs. an RTV syringe "
    "accessory that would cost ~$1.50–3.00 per headset anyway — comparable cost, "
    "guaranteed compliance vs. user-dependent compliance.")

ins_before(risk_doc, target,
    "RISK-16  |  LOW (MITIGATED)  |  IPX4 Field-Replacement Seal — Option A Gasket Selected",
    bold=True, color=(0x1F, 0x49, 0x7D))

# Update OI-11
for row in risk_doc.tables[1].rows:
    if row.cells[0].text.strip() == "OI-11":
        set_text(row.cells[1], "DONE (decision made: Option A gasket)")
        set_bg(row.cells[1], "E2EFDA")
        set_text(row.cells[2],
            "Zone module field-replacement seal — Option A selected: "
            "self-sealing co-moulded silicone gasket. "
            "Gasket spec in FPC spec §12. Shell slot rim spec required in shell tooling. "
            "FAI field-replacement IPX4 test required (NP-FAI-ZM-001 FAI-IPX-01 to FAI-IPX-04).")
        print("  OI-11 updated")
        break

risk_doc.save(RISK_PATH)
print(f"  Risk register saved")

# ─────────────────────────────────────────────────────────────────────────────
# 3.  NP-COORD-001 — update G2-05
# ─────────────────────────────────────────────────────────────────────────────

COORD_PATH = "docs/neurone_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

for tbl in coord.tables:
    for row in tbl.rows:
        if row.cells[0].text.strip() == "G2-05":
            cells = row.cells
            # Description (col 2)
            set_text(cells[2],
                "RISK-16 — Zone module field-replacement IPX4 seal. "
                "DECISION MADE: Option A — self-sealing co-moulded silicone compression "
                "gasket on zone module perimeter (Shore 40–50A, D-section 2.5×2.0 mm, "
                "20% compression). No RTV syringe, no user procedure. "
                "Two tooling coordination items remain: "
                "(a) gasket geometry in zone module mould drawing before first-cut; "
                "(b) shell slot rim: smooth Ra ≤ 1.6 µm, width ≥ 2.5 mm, no draft steps.",
                size=8)
            # Deliverable (col 3)
            set_text(cells[3],
                "Zone module mould drawing with gasket co-mould spec. "
                "Shell tooling spec: slot rim surface finish Ra ≤ 1.6 µm, width ≥ 2.5 mm. "
                "FAI IPX4 field-replacement test (NP-FAI-ZM-001 FAI-IPX-01–FAI-IPX-04) pass.",
                size=8)
            # Status (col 6)
            set_text(cells[6], "OPEN — tooling spec items remain", size=8)
            set_bg(cells[6], "E2EFDA")
            print(f"  NP-COORD-001 G2-05 updated")
            break

# Update dashboard row for G2-05 (5-col dashboard table: ID|Gate|Owner|Item|Status)
for tbl in coord.tables:
    for row in tbl.rows:
        cells = row.cells
        if (len(cells) == 5
                and cells[0].text.strip() == "G2-05"
                and cells[1].text.strip() == "G2"):
            set_text(cells[4], "OPEN — tooling spec; FAI test required", size=8)
            set_bg(cells[4], "E2EFDA")
            print("  Dashboard G2-05 updated")
            break

coord.save(COORD_PATH)
print(f"  Coord checklist saved")

# ─────────────────────────────────────────────────────────────────────────────
# 4.  NP-FAI-ZM-001 — add §4c Sealing and IPX4 Verification
# ─────────────────────────────────────────────────────────────────────────────

FAI_PATH = "docs/superseded/neurone_fai_zone_module.docx"
fai = Document(FAI_PATH)

# We need to insert a new sub-section heading "4c. Sealing and IPX4 Verification"
# and a table with items FAI-IPX-01 to FAI-IPX-04.
# Best insertion point: after the §4b table (which ends before SECTION 5 heading).

# Find SECTION 5 paragraph index
s5_para = None
for i, p in enumerate(fai.paragraphs):
    if "SECTION 5" in p.text:
        s5_para = p
        break

print(f"  SECTION 5 anchor: {s5_para.text[:60]!r}")

# Insert heading and body paragraph before SECTION 5
ins_before(fai, s5_para, "")  # spacer

# Build the table XML inline using python-docx (simpler: add a table to the doc
# then move its XML to the correct position)

# Step 1: add table at end of document
fai_tbl = fai.add_table(rows=1, cols=6)
fai_tbl.style = "Table Grid"

HDR_BG = "1F497D"
col_headers = ["Item", "Inspection / Test", "Method", "Accept Criterion", "Result", "Sign"]
widths_in   = [0.65,   2.3,                 1.2,      1.55,               0.65,     0.45]

hdr_row = fai_tbl.rows[0]
for i, h in enumerate(col_headers):
    set_bg(hdr_row.cells[i], HDR_BG)
    run = hdr_row.cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(8)

FAI_IPX_ROWS = [
    ("FAI-IPX-01",
     "Initial IPX4 baseline: assembled headset (all 5 zone modules installed "
     "at factory) — IEC 60529 IPX4 water spray test, 10 min from all directions",
     "IEC 60529 IPX4 spray nozzle",
     "No water ingress to Hub PCB, any connector, or any zone module PCB. "
     "Electrical function unchanged after test (EEG impedance all channels < 10 kΩ).",
     "", ""),
    ("FAI-IPX-02",
     "[BLOCKING] Field-replacement IPX4 re-test: remove and re-insert each zone "
     "module 10× (simulating 10 user field swaps); then repeat full IEC 60529 "
     "IPX4 spray test without factory re-sealing",
     "10 swap cycles per zone (50 total), then IEC 60529 IPX4",
     "Same pass criterion as FAI-IPX-01. No water ingress after 10 field swaps. "
     "If any ingress: halt production — gasket geometry or shell rim tolerance "
     "non-conformant; escalate to Mechanical Engineering.",
     "", ""),
    ("FAI-IPX-03",
     "Gasket visual inspection after 10 swap cycles: inspect all 5 zone module "
     "gaskets under 10× magnification for tears, compression set, or deformation "
     "that would prevent re-sealing",
     "10× magnifier, all 5 modules post-cycle",
     "No tears, cuts, or permanent deformation > 10% of cross-section. "
     "Gasket returns to ≥ 90% of original height after removal (compression set < 10%).",
     "", ""),
    ("FAI-IPX-04",
     "Gasket shelf-life check: inspect a factory-sealed zone module stored 12 months "
     "at 25°C / 50% RH; perform IPX4 test and gasket dimensional check",
     "Accelerated ageing (IEC 60068-2-66) or 12-month real-time sample; "
     "then IPX4 test",
     "IPX4 pass after storage period. Gasket dimensions within 5% of nominal. "
     "(First-article qualification only — not required for every production lot.)",
     "", ""),
]

for r_idx, (item_id, desc, method, criterion, result, sign) in enumerate(FAI_IPX_ROWS):
    row = fai_tbl.add_row()
    bg = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
    for c_idx, txt in enumerate([item_id, desc, method, criterion, result, sign]):
        cell = row.cells[c_idx]
        set_bg(cell, bg)
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(8)
        if c_idx == 0:
            run.bold = True

for row in fai_tbl.rows:
    for i, w in enumerate(widths_in):
        if i < len(row.cells):
            row.cells[i].width = Inches(w)

# Step 2: build sub-section heading paragraph to insert before SECTION 5
# First insert the table by moving its XML before s5_para
tbl_xml = fai_tbl._tbl
s5_para._p.addprevious(tbl_xml)

# Insert heading before the (now-placed) table
ins_before(fai, s5_para, "")  # spacer after table — insert before s5
# Actually we need to insert the heading before the table; the table is now
# at s5_para position - 1. Find the paragraph just before s5_para which should
# be the spacer we just added, and insert heading before that.
# Simpler: we already inserted a spacer para before s5_para; now we need the
# heading to be before that spacer. So insert heading before s5_para again
# (which puts it between the last spacer and the table).
# The order from top: ... §4b content ... [heading] [table] [spacer] [SECTION 5]
# That's slightly wrong order. Let me fix by inserting the heading before the table XML.

# Insert heading before table
heading_p = OxmlElement("w:p")
tbl_xml.addprevious(heading_p)
for p_obj in fai.paragraphs:
    if p_obj._p is heading_p:
        run = p_obj.add_run(
            "4c. Sealing and IPX4 Verification (Option A gasket)  [RISK-16]")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        break

# Insert body note before the heading
note_p = OxmlElement("w:p")
heading_p.addprevious(note_p)
for p_obj in fai.paragraphs:
    if p_obj._p is note_p:
        run = p_obj.add_run(
            "All items use the self-sealing co-moulded gasket (Option A). "
            "No RTV or user re-sealing procedure. "
            "FAI-IPX-02 is BLOCKING — production cannot release without field-replacement IPX4 pass.")
        run.font.size = Pt(10)
        break

# Step 3: add RISK-16 to §9 cross-reference
for tbl2 in fai.tables:
    for row in tbl2.rows:
        if "RISK-15" in row.cells[0].text:
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            xref_widths = [0.65, 2.0, 2.35, 1.8]
            xref_data = ("RISK-16",
                         "IPX4 field-replacement seal — Option A gasket",
                         "FAI-IPX-01, FAI-IPX-02, FAI-IPX-03, FAI-IPX-04",
                         "MITIGATED (this FAI)")
            for c_idx, txt in enumerate(xref_data):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(xref_widths[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el)
                tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el)
                r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve")
                t_el.text = txt
                r_el.append(t_el)
                p_el.append(r_el)
                tc.append(p_el)
                new_tr.append(tc)
            # Find and color
            for r2 in tbl2.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[3], "E2EFDA")
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  §9 cross-reference RISK-16 row added")
            break

fai.save(FAI_PATH)
print(f"  FAI checklist saved: {FAI_PATH}")
print("\nAll patches complete.")
