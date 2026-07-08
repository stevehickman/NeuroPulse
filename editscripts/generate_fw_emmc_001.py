"""
Generate NP-FW-EMMC-001 Rev A
eMMC Partition Architecture and Storage Encryption Specification
Addresses CLAUDE.md §13.4 pending decisions:
  - eMMC partition architecture + separate UHDR/SHDR encryption (before any storage code)
  - Dual-bank OTA bootloader (must be from first firmware line)
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(13) if level == 1 else Pt(11)
    return p


def add_normal(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


# ── Title ──────────────────────────────────────────────────────────────────────
title = doc.add_paragraph('NeurOne')
title.style = doc.styles['Normal']
title.runs[0].bold = True
title.runs[0].font.size = Pt(14)

subtitle = doc.add_paragraph('eMMC Partition Architecture and Storage Encryption Specification')
subtitle.style = doc.styles['Normal']
subtitle.runs[0].bold = True
subtitle.runs[0].font.size = Pt(13)

doc.add_paragraph(
    'NP-FW-EMMC-001  Rev A  |  2026-05-11  |  Pre-Storage-Code Baseline'
).style = doc.styles['Normal']
doc.add_paragraph(
    'References: CLAUDE.md §4.1 · §5.1 · §5.3 · §7.1 · §7.2  ·  NP-COORD-001 Rev A.3'
).style = doc.styles['Normal']
doc.add_paragraph(
    'Addresses CLAUDE.md §13.4 pending decisions: (1) eMMC partition architecture + '
    'separate UHDR/SHDR encryption in firmware specification before any storage code is '
    'written; (2) Dual-bank OTA bootloader from first firmware line. '
    'This document must be baselined and reviewed before any eMMC driver, filesystem '
    'mount, or data-write firmware code is authored.'
).style = doc.styles['Normal']


# ── §1 Purpose and Scope ───────────────────────────────────────────────────────
add_heading(doc, '1.  PURPOSE AND SCOPE')
add_normal(doc,
    'This document specifies the eMMC partition layout, filesystem configuration, '
    'encryption architecture, key management, session data classification, backup '
    'protocol, and write-endurance monitoring for the NeurOne Hub. It governs '
    'all firmware that reads from or writes to the 8 GB industrial eMMC.')
add_normal(doc,
    'All requirements herein are MANDATORY unless explicitly marked INFORMATIVE. '
    'No storage firmware code may be authored until this document is baselined. '
    'Deviations require an approved Engineering Change Order (ECO).')
add_normal(doc,
    'IEC 62304 classification: Main processor storage firmware — Class B. '
    'Safety MCU has no direct eMMC access; all safety-relevant data classification '
    'is enforced by the main processor (i.MX RT1062) before any write occurs.')


# ── §2 References ──────────────────────────────────────────────────────────────
add_heading(doc, '2.  REFERENCES')
refs = [
    ('CLAUDE.md §4.1', 'Hardware specification: 8 GB industrial eMMC, SLC cache, 30,000+ P/E cycles, LittleFS, write-protected firmware partition, separate UHDR/SHDR partitions from first firmware line'),
    ('CLAUDE.md §5.1', 'UHDR/SHDR definitions, UHDR AES-256 biometric-derived key, SHDR separate encryption, boundary resolution rules, automated nightly UHDR backup'),
    ('CLAUDE.md §5.2', 'Predictive maintenance via SHDR fleet telemetry; write endurance monitoring'),
    ('CLAUDE.md §5.3', 'Research data anonymization architecture: on-device, per-study, fresh per request; NeurOne never holds raw UHDR'),
    ('CLAUDE.md §7.1', 'Industrial eMMC + LittleFS locked design change; dual-bank OTA bootloader must be in bootloader from first firmware line'),
    ('NP-COORD-001 Rev A.3', 'Engineering coordination checklist; gate structure; summary status dashboard'),
    ('NP-FW-REQ-001 Rev A', 'Zone module firmware requirements; ZONE_ID debounce; SHDR FAULT event logging format'),
    ('IEEE 1619-2007', 'XTS-AES mode for storage encryption (sector-level, no ciphertext expansion)'),
    ('RFC 9106', 'Argon2id — memory-hard key derivation function for UHDR key (resists GPU brute-force of biometric/PIN inputs)'),
    ('RFC 5869', 'HKDF — HMAC-based key derivation for SHDR device manufacturing key'),
    ('JEDEC JESD84-B51', 'eMMC 5.1 standard — hardware boot partitions, write protection groups, reliable write'),
]
t = doc.add_table(rows=1, cols=2)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Document / Standard'
t.rows[0].cells[1].text = 'Relevant content'
for ref, desc in refs:
    row = t.add_row()
    row.cells[0].text = ref
    row.cells[1].text = desc


# ── §3 eMMC Hardware Specification ────────────────────────────────────────────
add_heading(doc, '3.  EMMC HARDWARE SPECIFICATION')
add_normal(doc,
    'EMMC-HW-01 (MANDATORY): Device shall use an industrial-grade eMMC component '
    'meeting all of the following: total user capacity ≥ 8 GiB; SLC write cache '
    'enabled; endurance rating ≥ 30,000 P/E cycles on SLC cache blocks; operating '
    'temperature range −40°C to +85°C (industrial grade); JEDEC eMMC 5.1 compliant; '
    'supports hardware write protection groups (WP_GRP_SIZE per datasheet).')
add_normal(doc,
    'EMMC-HW-02 (MANDATORY): The eMMC component shall expose two hardware boot '
    'partitions (Boot Partition 1 and Boot Partition 2) each ≥ 4 MiB, as specified '
    'by JEDEC JESD84-B51 §7.3. These partitions are outside the user data area and '
    'are used exclusively for bootloader images (§8.1).')
add_normal(doc,
    'EMMC-HW-03 (MANDATORY): The i.MX RT1062 eMMC interface shall operate at HS200 '
    '(200 MHz DDR) or HS400 (400 MHz DDR) if supported by the selected eMMC component. '
    'Minimum acceptable: HS52 (52 MHz SDR). Bus width: 8-bit (eMMC x8). Interface '
    'parameters locked in bootloader — do not change via OTA.')
add_normal(doc,
    'EMMC-HW-04 (MANDATORY): eMMC write endurance counter (device-internal wear '
    'indicator if available, e.g., EXT_CSD[268] DEVICE_LIFE_TIME_EST_TYP_A/B) shall '
    'be read at every power-on and logged to the SHDR partition. See §14.')
add_normal(doc,
    'EMMC-HW-05 (INFORMATIVE): Candidate components as of 2026-05-11: Micron MTFC8GAKAJCN '
    '(8 GB, HS200, SLC cache, −40°C to +85°C, 30,000 P/E); Kingston EMMC08G-S100-A08 '
    '(8 GB, HS200, industrial). Final part selection requires procurement sign-off per '
    'NP-PROC-SUP-001. This specification is part-agnostic — any component meeting '
    'EMMC-HW-01 through EMMC-HW-04 is acceptable.')


# ── §4 Partition Layout ────────────────────────────────────────────────────────
add_heading(doc, '4.  PARTITION LAYOUT')
add_heading(doc, '4.1  eMMC Hardware Boot Partitions (Outside User Area)', level=2)
add_normal(doc,
    'The eMMC hardware boot partitions are accessed by setting BOOT_PARTITION_ENABLE '
    'in EXT_CSD[179]. They are NOT part of the GPT user area and do not consume user '
    'area capacity. The i.MX RT1062 boot ROM loads from Boot Partition 1 on cold boot.')

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
for cell, hdr in zip(t2.rows[0].cells, ['Partition', 'Size', 'Contents', 'Write protection']):
    cell.text = hdr
boot_parts = [
    ('eMMC Boot Partition 1', '4 MiB (hardware minimum)', 'Primary bootloader image (IVT + DCD + bootloader binary). i.MX RT1062 boot ROM entry point.', 'Write-protected via EXT_CSD[173] BOOT_WP. Cleared only during deliberate bootloader update sequence.'),
    ('eMMC Boot Partition 2', '4 MiB (hardware minimum)', 'Backup bootloader image — identical content to Boot Partition 1 after each successful update. Used as fallback if Boot Partition 1 is corrupt.', 'Write-protected via EXT_CSD[173] BOOT_WP. Updated in sync with Boot Partition 1.'),
]
for cols in boot_parts:
    row = t2.add_row()
    for cell, val in zip(row.cells, cols):
        cell.text = val

add_heading(doc, '4.2  User Area Partition Table (GPT)', level=2)
add_normal(doc,
    'EMMC-PT-01 (MANDATORY): The eMMC user area shall use a GUID Partition Table (GPT) '
    'layout. GPT is required for: (a) partition integrity verification via CRC32 header '
    'checksum; (b) redundant GPT header at LBA −1 (end of device); (c) named partition '
    'entries enabling firmware to locate partitions by GUID rather than fixed LBA offset. '
    'MBR-style partition tables are not permitted.')
add_normal(doc,
    'EMMC-PT-02 (MANDATORY): All partition boundaries shall align to eMMC erase block '
    'boundaries. Minimum alignment: 4 MiB (4,096 KiB). This ensures LittleFS block '
    'operations align with eMMC erase units and prevents write amplification.')
add_normal(doc,
    'EMMC-PT-03 (MANDATORY): The following partition table is locked. Total user area: '
    '8,192 MiB. Partition layout must be written to eMMC at manufacturing time before '
    'any firmware is installed. No partition may be resized, merged, or relocated without '
    'an approved ECO and full re-flash of the device.')

t3 = doc.add_table(rows=1, cols=6)
t3.style = 'Table Grid'
for cell, hdr in zip(t3.rows[0].cells,
                     ['#', 'Partition name', 'Size (MiB)', 'Filesystem', 'Encryption', 'Write protection']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
partitions = [
    ('0', 'GPT header + protective MBR', '1', '—', 'None', 'Write-protected (GPT integrity)'),
    ('1', 'Firmware Bank A', '128', 'Raw binary image', 'None (firmware is authenticated via cryptographic signature)', 'eMMC WP group locked after production programming. Write-enabled only during OTA Bank A update (deliberate).'),
    ('2', 'Firmware Bank B', '128', 'Raw binary image', 'None (signature verified before boot flag flip)', 'Write-enabled during OTA download. Write-protected after boot flag flip to Bank B.'),
    ('3', 'Safety MCU Firmware Staging', '4', 'Raw binary image', 'None (signed image)', 'Write-enabled during explicit Safety MCU update only. Safety MCU firmware update requires deliberate user action — never automated.'),
    ('4', 'Config / Calibration', '16', 'LittleFS', 'AES-256-XTS, device manufacturing key (see §7)', 'Write-protected after factory programming; updated only via authenticated firmware command.'),
    ('5', 'SHDR', '512', 'LittleFS', 'AES-256-XTS, device manufacturing key (see §7)', 'No hardware WP (SHDR appended continuously during device operation).'),
    ('6', 'UHDR', '6,903', 'LittleFS', 'AES-256-XTS, user biometric-derived key (see §6)', 'No hardware WP (UHDR appended continuously during sessions). NeurOne NEVER holds the decryption key.'),
    ('7', 'Scratch', '500', 'None (raw blocks, zeroed at boot)', 'None (Scratch contains no persistent user or health data — cleared on every boot)', 'No hardware WP.'),
]
for cols in partitions:
    row = t3.add_row()
    for cell, val in zip(row.cells, cols):
        cell.text = val

add_normal(doc, 'Partition size verification: 1 + 128 + 128 + 4 + 16 + 512 + 6,903 + 500 = 8,192 MiB ✓')

add_heading(doc, '4.3  UHDR/SHDR Boundary Enforcement', level=2)
add_normal(doc,
    'EMMC-PT-04 (MANDATORY): The physical partition boundary between UHDR (partition 6) '
    'and SHDR (partition 5) enforces the UHDR/SHDR data separation defined in CLAUDE.md §5.1 '
    'at the hardware level. No firmware code path may write SHDR-classified data to the '
    'UHDR partition or UHDR-classified data to the SHDR partition. Cross-partition writes '
    'are a Class B firmware defect and shall be detected by unit test.')
add_normal(doc,
    'EMMC-PT-05 (MANDATORY): All data elements must be classified as UHDR or SHDR before '
    'any write is issued. The classification decision is made at the data-source layer '
    '(sensor driver or session manager), not at the storage layer. Storage layer receives '
    'pre-classified data with a destination tag (UHDR | SHDR) and must reject any write '
    'that does not carry a valid tag. See §12 for the full classification table.')


# ── §5 LittleFS Configuration ─────────────────────────────────────────────────
add_heading(doc, '5.  LITTLEFS CONFIGURATION')
add_heading(doc, '5.1  Why LittleFS', level=2)
add_normal(doc,
    'LittleFS is selected per CLAUDE.md §4.1 for the following properties: '
    '(a) power-loss resilience — every write is atomic; no filesystem corruption on '
    'unexpected power removal during session; '
    '(b) wear leveling — dynamic wear leveling distributes writes across all erase blocks, '
    'maximising SLC cache endurance; '
    '(c) embedded footprint — ≈4 KB RAM per mounted filesystem instance; runs on '
    'i.MX RT1062 without OS; '
    '(d) block device abstraction — cleanly separates eMMC driver from filesystem logic; '
    'encryption layer (§6, §7) inserted between LittleFS block device callbacks and '
    'eMMC driver without modifying LittleFS source.')

add_heading(doc, '5.2  LittleFS Instance Parameters', level=2)
add_normal(doc,
    'EMMC-FS-01 (MANDATORY): Each encrypted partition (UHDR, SHDR, Config) shall mount '
    'a separate LittleFS instance. The instances are independent — a filesystem error in '
    'one partition does not affect the other. Scratch partition is NOT mounted as '
    'LittleFS — it is accessed as raw blocks.')

t4 = doc.add_table(rows=1, cols=4)
t4.style = 'Table Grid'
for cell, hdr in zip(t4.rows[0].cells,
                     ['Parameter', 'UHDR', 'SHDR', 'Config']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
lfs_params = [
    ('block_size', '4,096 bytes (4 KiB)', '4,096 bytes (4 KiB)', '4,096 bytes (4 KiB)'),
    ('block_count', '1,767,168 (6,903 MiB ÷ 4 KiB)', '131,072 (512 MiB ÷ 4 KiB)', '4,096 (16 MiB ÷ 4 KiB)'),
    ('read_size', '256 bytes', '256 bytes', '256 bytes'),
    ('prog_size', '256 bytes', '256 bytes', '256 bytes'),
    ('lookahead_size', '512 bytes', '256 bytes', '64 bytes'),
    ('cache_size', '4,096 bytes', '4,096 bytes', '512 bytes'),
    ('block_cycles', '500 (LittleFS wear leveling trigger)', '500', '200'),
    ('file_max', '2,147,483,647 (LFS_FILE_MAX)', '2,147,483,647', '65,536'),
    ('attr_max', '1,022 bytes', '1,022 bytes', '256 bytes'),
    ('name_max', '255 bytes', '255 bytes', '64 bytes'),
    ('metadata_max', '4,096 bytes', '4,096 bytes', '4,096 bytes'),
]
for row_data in lfs_params:
    row = t4.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

add_normal(doc,
    'EMMC-FS-02 (MANDATORY): LittleFS block_size must equal the eMMC erase block size '
    '(4,096 bytes minimum for eMMC 5.1 compatible components). Verify block_size matches '
    'the selected eMMC component\'s ERASE_GRP_SIZE from EXT_CSD before production '
    'programming. If erase block > 4,096 bytes, increase block_size to match and '
    'recalculate block_count for each partition.')
add_normal(doc,
    'EMMC-FS-03 (MANDATORY): LittleFS shall be compiled with LFS_NO_MALLOC defined '
    '(static allocation only). Dynamic memory allocation in the storage path is '
    'prohibited on the i.MX RT1062 to prevent heap fragmentation during long sessions.')
add_normal(doc,
    'EMMC-FS-04 (MANDATORY): All three LittleFS instances (UHDR, SHDR, Config) shall '
    'be mounted at system startup before any session may begin. A mount failure on any '
    'instance is a FAULT — session start is blocked and a diagnostic is logged to the '
    'SHDR partition (if SHDR mount succeeded) or to UART/USB-C debug log if SHDR mount '
    'also failed.')


# ── §6 UHDR Partition Specification ───────────────────────────────────────────
add_heading(doc, '6.  UHDR PARTITION SPECIFICATION')
add_heading(doc, '6.1  Overview and Access Control', level=2)
add_normal(doc,
    'The UHDR partition (User Health Data Record) stores all data that reveals '
    'information about the user as a person — EEG waveforms, HRV time series, PPG '
    'optical signal, neurofeedback scores, session timestamps, protocol parameters, '
    'closed-loop adaptation events, PBM dose per zone, user-entered outcome logs, '
    'eye-open/closed state, and all other elements listed in CLAUDE.md §5.1 UHDR '
    'definition.')
add_normal(doc,
    'EMMC-UHDR-01 (MANDATORY): NeurOne infrastructure shall NEVER hold, request, '
    'receive, or process raw UHDR data. The AES-256-XTS decryption key is derived '
    'from the user\'s biometric/PIN input and never leaves the device in plaintext. '
    'This is a product design invariant — no support, engineering, research, or '
    'regulatory workflow may bypass it. The firmware must make it technically '
    'impossible to transmit raw UHDR data over any interface (USB-C, BT, Wi-Fi) '
    'without user-initiated decryption on-device.')
add_normal(doc,
    'EMMC-UHDR-02 (MANDATORY): Clinician access to UHDR elements is provided only '
    'via the clinical consent engine (CLAUDE.md §6). The consent engine operates '
    'on-device: it reads UHDR (using the user\'s decrypted key, which requires user '
    'authentication), extracts the consented elements, and transmits only those '
    'elements — never raw UHDR — to the clinician platform. The extraction and '
    'transmission happen within a single authenticated session; the key is zeroed '
    'from SRAM immediately after the extraction completes.')

add_heading(doc, '6.2  Encryption: AES-256-XTS', level=2)
add_normal(doc,
    'EMMC-UHDR-03 (MANDATORY): The UHDR partition shall be encrypted using AES-256-XTS '
    '(IEEE 1619-2007). XTS is selected for storage because: (a) it operates on '
    'fixed-size data units (sectors) without ciphertext expansion — ciphertext is the '
    'same size as plaintext, preserving LittleFS block geometry; (b) each sector is '
    'encrypted with a unique tweak derived from its logical block address, preventing '
    'copy-and-paste attacks between sectors; (c) no authentication tag — authentication '
    'is provided separately via LittleFS CRC and file-level HMAC-SHA256 on session '
    'records (§6.5).')
add_normal(doc,
    'EMMC-UHDR-04 (MANDATORY): AES-256-XTS key material is 512 bits (two independent '
    '256-bit keys: K1 for AES-XTS encryption, K2 for tweak encryption). Both keys are '
    'derived from the same Argon2id output (§6.3) by splitting the 64-byte output: '
    'K1 = output[0:32], K2 = output[32:64]. Neither K1 nor K2 is stored on the device '
    'in plaintext.')
add_normal(doc,
    'EMMC-UHDR-05 (MANDATORY): The XTS data unit size shall be 512 bytes, matching the '
    'LittleFS read_size and prog_size. The XTS tweak for each 512-byte unit is: '
    'tweak = (UHDR_partition_LBA_start + block_offset) as a 128-bit little-endian integer. '
    'This is the standard XTS tweak construction per IEEE 1619-2007 §5.1.')
add_normal(doc,
    'EMMC-UHDR-06 (MANDATORY): AES-256-XTS shall be implemented using the i.MX RT1062 '
    'DCP (Data Co-Processor) hardware accelerator for AES-128. Since DCP supports '
    'AES-128 only, AES-256-XTS shall be implemented in software using the hardware '
    'AES-128 primitive for the underlying cipher (AES-256 with two-key 128-bit halves). '
    'Alternatively, use a vetted software AES-256 library (e.g., mbedTLS 3.x AES-XTS '
    'mode) if hardware acceleration for full AES-256 becomes available. No custom '
    'cryptographic implementations.')

add_heading(doc, '6.3  Key Derivation: Argon2id', level=2)
add_normal(doc,
    'EMMC-UHDR-07 (MANDATORY): The UHDR encryption key (K1 || K2, 64 bytes) shall be '
    'derived using Argon2id (RFC 9106) with the following locked parameters:')
t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Table Grid'
for cell, hdr in zip(t5.rows[0].cells, ['Parameter', 'Value', 'Rationale']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
argon_params = [
    ('Algorithm', 'Argon2id', 'Hybrid of Argon2i (side-channel resistance) and Argon2d (GPU resistance). RFC 9106 recommended for password hashing.'),
    ('Memory (m)', '65,536 KiB (64 MiB)', 'Fills i.MX RT1062 external LPSDR4 (32 MB) + internal SRAM during KDF; renders GPU/ASIC parallelism economically infeasible. Runs once at session start on 32 MB LPSDR4.'),
    ('Iterations (t)', '4', 'RFC 9106 minimum recommended for interactive use. Increases time cost while memory cost dominates.'),
    ('Parallelism (p)', '1', 'Single-threaded on Cortex-M7 (FreeRTOS single-core). Must be 1.'),
    ('Output length', '64 bytes (512 bits)', 'Provides K1 (bytes 0–31) and K2 (bytes 32–63) for AES-256-XTS.'),
    ('Password input', 'SHA-256(biometric_template) OR SHA-256(PIN_string), 32 bytes', 'Biometric or PIN is hashed to fixed-length input before Argon2id. The raw biometric is never the direct Argon2id input. Minimum PIN: 6 digits or 4-word passphrase.'),
    ('Salt', '32 bytes, random, generated by TRNG at first device setup', 'Stored in Config partition (partition 4) in plaintext. Salt is device-unique — same PIN on two devices produces different keys. Salt is NOT secret.'),
    ('Associated data (ad)', 'b"NeurOne-UHDR-v1" (18 bytes)', 'Domain separator — prevents key reuse between firmware versions or future key types.'),
    ('Secret (secret)', 'None (empty)', 'Argon2id secret parameter not used; all key strength comes from memory+time cost.'),
    ('Key derivation timing', 'Once per session unlock (Mode 1 / Mode 2 app connection)', 'Key held in SRAM for session duration, zeroed on session end or power-down.'),
    ('KDF execution location', 'i.MX RT1062 main processor, LPSDR4 working memory', 'Never executed on Safety MCU. Safety MCU has no access to UHDR key material.'),
]
for row_data in argon_params:
    row = t5.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

add_normal(doc,
    'EMMC-UHDR-08 (MANDATORY): Expected Argon2id execution time on i.MX RT1062 @ 600 MHz '
    'with 64 MiB memory: approximately 3–8 seconds. This is acceptable — key derivation '
    'occurs once per session unlock, not per write. A progress indicator shall be displayed '
    'in the app during KDF execution. In Mode 3 Autonomous (no app), KDF executes silently; '
    'status LED breathes during derivation.')
add_normal(doc,
    'EMMC-UHDR-09 (MANDATORY): Key verification. After Argon2id derivation, the firmware '
    'shall attempt to read a known sentinel record from the UHDR partition header block '
    'using the derived key. If decryption succeeds and the sentinel magic bytes match '
    '(b"UHDR_v1" + 25 bytes of zeros = 32 bytes total, written at first setup), the key '
    'is accepted. If the sentinel does not match, the key is wrong — zero K1/K2 from SRAM '
    'immediately and report authentication failure to the app. Three consecutive failures '
    'impose a 30-second lockout (doubles with each subsequent failure: 30s, 60s, 120s). '
    'Lockout state is stored in Config partition (SHDR key, so NeurOne can read if '
    'needed for support — lockout count is device health data, not user health data).')

add_heading(doc, '6.4  Key Storage and Lifecycle', level=2)
add_normal(doc,
    'EMMC-UHDR-10 (MANDATORY): The UHDR key (K1, K2) exists in plaintext only in '
    'i.MX RT1062 SRAM (non-cached OCRAM region) during an active session. Key lifecycle:')
add_bullet(doc, 'Key generation: Argon2id executed at session unlock (app authentication or Mode 3 pre-programmed session start). Working memory allocated in LPSDR4, zeroed immediately after KDF completes.')
add_bullet(doc, 'Key active: K1, K2 stored in a dedicated 128-byte SRAM key cache region (non-paged, non-DMA-accessible from external peripherals). Key cache region is marked non-cacheable in MPU configuration.')
add_bullet(doc, 'Key zeroing triggers: (a) session end (user-initiated or timeout); (b) power-down or reset; (c) authentication failure (§6.3 sentinel mismatch); (d) 200 ms SPI heartbeat timeout (Safety MCU watchdog triggers all-stimulation cutoff; main processor must also zero UHDR key on watchdog event).')
add_bullet(doc, 'Key never written to any persistent storage in plaintext. Key never transmitted over USB-C, BT, or Wi-Fi in any form.')
add_bullet(doc, 'Key rotation: Key rotation (re-encryption of entire UHDR partition with a new Argon2id-derived key) is a user-initiated operation triggered from the app security settings. Re-encryption uses the Scratch partition as intermediate staging. Full re-encryption of 6,903 MiB at eMMC HS200 speeds takes approximately 60–180 seconds; app displays progress. Old key is zeroed before new key is stored in SRAM cache.')
add_normal(doc,
    'EMMC-UHDR-11 (MANDATORY): The UHDR salt (stored in Config partition) shall never '
    'be transmitted to NeurOne servers. The salt is device-local. If a user factory-'
    'resets the device, a new salt is generated (TRNG), the UHDR partition is re-'
    'initialised (all data lost — user warned with two confirmation steps), and the '
    'old salt is overwritten 3× with TRNG data before being replaced.')

add_heading(doc, '6.5  Session Record Format and Integrity', level=2)
add_normal(doc,
    'EMMC-UHDR-12 (MANDATORY): Each session produces one UHDR session record file, '
    'named by session counter (unsigned 64-bit integer, zero-padded to 20 digits). '
    'File format (all fields encrypted at rest via AES-256-XTS):')
t6 = doc.add_table(rows=1, cols=3)
t6.style = 'Table Grid'
for cell, hdr in zip(t6.rows[0].cells, ['Field', 'Size', 'Content']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
session_fields = [
    ('Magic', '8 bytes', 'b"NPUHDR01" — version marker'),
    ('Session counter', '8 bytes (uint64, little-endian)', 'Monotonic counter, incremented per session, stored in Config partition. SHDR stores unsigned session count only (no user biology).'),
    ('Protocol hash', '32 bytes (SHA-256)', 'SHA-256 of the signed session protocol descriptor. Allows verification that recorded data matches the intended protocol.'),
    ('Data length', '8 bytes (uint64)', 'Total byte length of the data payload that follows.'),
    ('Data payload', 'Variable', 'EDF+ formatted EEG + HRV + PPG + neurofeedback time series. EDF+ header followed by signal data per standard. Enables import into EEGLab, BrainAnalyzer, MNE-Python without proprietary tools.'),
    ('Adaptation log', 'Variable', 'JSON array of closed-loop adaptation events: {timestamp_sample, channel, old_value, new_value, trigger_metric}. Sample-indexed (not wall-clock) to avoid UHDR/SHDR boundary issues.'),
    ('Outcome log', 'Variable', 'User-entered text/numeric entries (optional). UTF-8, max 4,096 bytes per entry.'),
    ('Record HMAC', '32 bytes (HMAC-SHA256)', 'HMAC-SHA256(key=K1, data=all preceding fields). Provides authenticated integrity check separate from XTS encryption. K1 is the same UHDR encryption key — dual use of K1 for both XTS and HMAC is acceptable here because AES-XTS and HMAC-SHA256 use K1 in distinct cryptographic contexts.'),
]
for row_data in session_fields:
    row = t6.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

add_normal(doc,
    'EMMC-UHDR-13 (MANDATORY): UHDR directory structure under the LittleFS mount point:')
add_bullet(doc, '/uhdr/sessions/ — session record files (one per session, named by counter)')
add_bullet(doc, '/uhdr/outcomes/ — standalone user outcome log entries (not tied to sessions)')
add_bullet(doc, '/uhdr/backup_state/ — incremental backup manifest (last backed-up session counter, checksum)')
add_bullet(doc, '/uhdr/_header — sentinel record for key verification (§6.3) + salt echo for integrity check')


# ── §7 SHDR Partition Specification ───────────────────────────────────────────
add_heading(doc, '7.  SHDR PARTITION SPECIFICATION')
add_heading(doc, '7.1  Overview and Access Control', level=2)
add_normal(doc,
    'The SHDR partition (System Health Data Record) stores data that reveals the '
    'condition of the device — not the biology of the user. Contents are defined in '
    'CLAUDE.md §5.1 and include: LED output ratio per zone, NTC temperature profiles, '
    'EMF shielding attenuation ratio, unsigned session count, consumable counts, '
    'USB-C insertion counter, PD negotiation log, impact events (between sessions only), '
    'fan RPM, supercapacitor cycles, firmware version history, OTA log, accessory '
    'authentication pass/fail, calibration coefficient history.')
add_normal(doc,
    'EMMC-SHDR-01 (MANDATORY): NeurOne may read SHDR data for warranty, predictive '
    'maintenance, and fleet analytics. SHDR is transmitted to NeurOne fleet database '
    'on USB-C connect (with warranty consent). The SHDR encryption key is held by '
    'NeurOne (§7.3). This is intentional — SHDR contains no user biology, only '
    'device health metrics linked to device ID and warranty owner ID.')
add_normal(doc,
    'EMMC-SHDR-02 (MANDATORY): Any data element that fails the UHDR/SHDR boundary test '
    '("does this tell us about the device\'s condition, with nothing that reveals user '
    'biology?") must be written to UHDR, not SHDR. Borderline cases are resolved per '
    'CLAUDE.md §5.1 boundary case resolution rule: when in doubt → UHDR.')

add_heading(doc, '7.2  Encryption: AES-256-XTS', level=2)
add_normal(doc,
    'EMMC-SHDR-03 (MANDATORY): The SHDR partition shall be encrypted using AES-256-XTS '
    '(IEEE 1619-2007) with the same XTS parameters as UHDR (§6.2): 512-byte data units, '
    'LBA-derived tweak. The SHDR key is distinct from the UHDR key — the same XTS '
    'parameters with different key material do not allow cross-partition decryption.')
add_normal(doc,
    'EMMC-SHDR-04 (MANDATORY): SHDR key material is 64 bytes (K1_shdr || K2_shdr), '
    'both derived from the device manufacturing key (§7.3). The SHDR key is static '
    'for the device lifetime (no rotation). Key is loaded from secure element at boot '
    'and held in SRAM key cache alongside the UHDR key.')

add_heading(doc, '7.3  Key Derivation: Device Manufacturing Key', level=2)
add_normal(doc,
    'EMMC-SHDR-05 (MANDATORY): The SHDR encryption key is derived at manufacturing '
    'time using HKDF-SHA256 (RFC 5869):')
add_bullet(doc, 'IKM (input key material): 32-byte NeurOne manufacturing root key (MRK). MRK is a fleet-wide secret held in NeurOne secure key management system (HSM). MRK is NEVER written to any individual device.')
add_bullet(doc, 'Salt: device serial number (32-byte SHA-256 hash of the eMMC CID register, which is unique per component).')
add_bullet(doc, 'Info: b"NeurOne-SHDR-v1" (18 bytes) — domain separator.')
add_bullet(doc, 'Output length: 64 bytes → K1_shdr = output[0:32], K2_shdr = output[32:64].')
add_bullet(doc, 'The 64-byte SHDR key is injected into the device Config partition (partition 4) during manufacturing, encrypted under the i.MX RT1062 SNVS OTPMK (One-Time Programmable Master Key, 256-bit key fused into OTP at manufacturing — NXP HAB/SNVS). NeurOne can re-derive the SHDR key for any device using MRK + device serial number — enabling SHDR telemetry decryption for fleet analytics and warranty support.')
add_normal(doc,
    'EMMC-SHDR-06 (MANDATORY): The i.MX RT1062 SNVS OTPMK shall be fused at '
    'manufacturing using NXP\'s HAB (High Assurance Boot) key provisioning flow. '
    'OTPMK fusing is irreversible. The OTPMK is used to decrypt the SHDR key stored '
    'in Config partition — it is not the SHDR key itself. This means: even if an '
    'attacker reads the Config partition, they cannot recover the SHDR key without '
    'the device OTPMK (which is hardware-fused and not readable via software).')
add_normal(doc,
    'EMMC-SHDR-07 (MANDATORY): SHDR key lifecycle: loaded from Config partition at '
    'boot (OTPMK-decrypted), stored in SRAM key cache region for device power-on '
    'duration. Zeroed on power-down. SHDR key is available from the moment Config '
    'partition is mounted — SHDR can receive FAULT events and startup logs before '
    'the user unlocks UHDR.')

add_heading(doc, '7.4  SHDR File Structure', level=2)
add_normal(doc,
    'EMMC-SHDR-08 (MANDATORY): SHDR directory structure under the LittleFS mount point:')
add_bullet(doc, '/shdr/telemetry/ — periodic device telemetry snapshots (one file per eMMC session count checkpoint, e.g., every 100 sessions)')
add_bullet(doc, '/shdr/faults/ — FAULT event log files (zone faults, stimulation cutoffs, watchdog events, mount errors)')
add_bullet(doc, '/shdr/calibration/ — factory calibration coefficients + ADS1299 self-calibration history + NTC cross-calibration history')
add_bullet(doc, '/shdr/ota/ — OTA update log (version installed, bank used, install timestamp as eMMC session count)')
add_bullet(doc, '/shdr/consumables/ — consumable session counts (intranasal sleeves, hydrogel tips, VNS pads, audio mesh, interface covers)')
add_bullet(doc, '/shdr/research/ — audit log of research study participation (study ID, descriptor hash, extract byte count, eMMC session count — never UHDR content)')
add_bullet(doc, '/shdr/shielding/ — EMF attenuation measurement history per fleet SHDR monitoring (CLAUDE.md §4.3)')
add_bullet(doc, '/shdr/supercap/ — supercapacitor cycle count, NTC aging estimate, hub thermistor history')

add_normal(doc,
    'EMMC-SHDR-09 (MANDATORY): All SHDR timestamps are stored as the unsigned eMMC '
    'session count integer (not wall-clock time), per the UHDR/SHDR boundary rule in '
    'CLAUDE.md §5.1: "device session count (unsigned integer) → SHDR; session '
    'timestamps → UHDR". Session count is incremented in Config partition at session '
    'start; it is the only time-like reference in SHDR.')


# ── §8 Firmware Partitions and Dual-Bank OTA ─────────────────────────────────
add_heading(doc, '8.  FIRMWARE PARTITIONS AND DUAL-BANK OTA')
add_heading(doc, '8.1  Bootloader (eMMC Hardware Boot Partitions)', level=2)
add_normal(doc,
    'EMMC-BL-01 (MANDATORY): The bootloader image shall reside in eMMC Boot Partition 1 '
    '(primary) and eMMC Boot Partition 2 (backup). The i.MX RT1062 boot ROM reads from '
    'eMMC Boot Partition 1 by default. Bootloader source code is version-controlled and '
    'changes require a separate ECO and re-qualification — bootloader updates are '
    'infrequent by design.')
add_normal(doc,
    'EMMC-BL-02 (MANDATORY): Bootloader responsibilities (in order):')
add_bullet(doc, '(1) Configure i.MX RT1062 clocks, DDR interface, SNVS, MPU.')
add_bullet(doc, '(2) Read boot bank flag from SNVS general purpose register (SNVS_LPGPR0): 0x00 = boot Bank A (partition 1); 0x01 = boot Bank B (partition 2).')
add_bullet(doc, '(3) Verify the target bank firmware image signature (Ed25519 signature over SHA-256 of firmware binary, signed with NeurOne firmware signing key). Reject unsigned or invalid images.')
add_bullet(doc, '(4) If Bank A fails verification: attempt Bank B. If Bank B also fails: enter USB-C DFU recovery mode (§8.5). Never boot an unverified image.')
add_bullet(doc, '(5) Load verified firmware image from target bank into i.MX RT1062 TCM + LPSDR4, jump to firmware entry point.')
add_bullet(doc, '(6) Log boot event (selected bank, verification outcome, boot count) to SHDR /shdr/ota/ via a minimal SHDR writer that runs before LittleFS is fully mounted — uses raw block writes to a reserved SHDR journal area, replayed by LittleFS on next full mount.')
add_normal(doc,
    'EMMC-BL-03 (MANDATORY): Bootloader size shall not exceed 256 KiB (within the '
    '4 MiB boot partition). Remaining boot partition space is reserved for future '
    'bootloader expansion. Bootloader is compiled as position-independent code to '
    'allow SRAM relocation if required.')

add_heading(doc, '8.2  Firmware Bank A (Primary, Write-Protected)', level=2)
add_normal(doc,
    'EMMC-FW-01 (MANDATORY): Firmware Bank A (partition 1, 128 MiB) contains the '
    'production firmware image. After manufacturing programming, Bank A is write-'
    'protected using the eMMC hardware write protection mechanism (EXT_CSD[171] '
    'WP_GRP_ENABLE + EXT_CSD[28] WRITE_PROT register). Write protection covers the '
    'entire Bank A erase group range.')
add_normal(doc,
    'EMMC-FW-02 (MANDATORY): Bank A write protection shall be verified by firmware '
    'at every boot: attempt a single-byte write to LBA 0 of Bank A and confirm it '
    'is rejected by the eMMC. If write protection is not active on Bank A (possible '
    'if manufacturing step was skipped), log a SHDR FAULT and alert via status LED '
    '(red blink). This does NOT block operation — a device with Bank A unprotected '
    'is functional but is flagged for manufacturing review.')

add_heading(doc, '8.3  OTA Update Sequence (Bank A → Bank B or Bank B → Bank A)', level=2)
add_normal(doc,
    'EMMC-OTA-01 (MANDATORY): OTA update sequence — all steps are MANDATORY, none '
    'may be skipped by firmware or by app command:')
t7 = doc.add_table(rows=1, cols=3)
t7.style = 'Table Grid'
for cell, hdr in zip(t7.rows[0].cells, ['Step', 'Action', 'Failure handling']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
ota_steps = [
    ('1', 'App or Mode 4 download delivers signed OTA package (firmware binary + Ed25519 signature + manifest) via USB-C to Scratch partition.', 'Transmission error → retry. Scratch partition write failure → abort OTA, log SHDR fault, retain current firmware.'),
    ('2', 'Main processor verifies Ed25519 signature over SHA-256 of firmware binary using NeurOne firmware public key (stored in Firmware Bank A firmware, not in bootloader — bootloader uses its own embedded copy). If signature invalid: delete OTA package from Scratch, log SHDR fault, abort.', 'Invalid signature → abort. Never flash an unverified image.'),
    ('3', 'Clear write protection on inactive bank (if Bank A is active, target is Bank B — Bank B write protection is not set by default; if Bank B is active, target Bank A — temporarily clear Bank A WP for this operation only).', 'WP clear failure → abort OTA.'),
    ('4', 'Write firmware binary to inactive bank in 512-byte aligned chunks. After each 4 MiB block, verify written data by reading back and comparing SHA-256. On mismatch: erase block and retry up to 3×, then abort.', 'Write or verify failure after 3 retries → abort OTA, log fault. Device continues on current bank.'),
    ('5', 'After full image written and verified: compute SHA-256 of entire inactive bank image, compare to manifest SHA-256. On mismatch: erase inactive bank, abort OTA.', 'SHA-256 mismatch → abort, erase inactive bank.'),
    ('6', 'Update boot bank flag in SNVS_LPGPR0 to point to newly written bank.', 'SNVS write failure → abort (device boots current bank next reset).'),
    ('7', 'Log OTA completion to SHDR /shdr/ota/: new version, target bank, install eMMC session count.', 'SHDR log failure is non-fatal — OTA proceeds.'),
    ('8', 'Reboot. Bootloader selects new bank, verifies signature, boots new firmware.', 'Bootloader verification failure on new bank → bootloader falls back to old bank, logs error, resets boot flag. Device recovers automatically.'),
    ('9', 'New firmware performs post-OTA validation: self-test, SHDR mount, Safety MCU SPI ping. If any validation fails: request bootloader revert to previous bank on next reset (write revert flag to SNVS_LPGPR1), schedule reboot.', 'Auto-revert ensures device always boots a working image.'),
]
for row_data in ota_steps:
    row = t7.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

add_normal(doc,
    'EMMC-OTA-02 (MANDATORY): At no point during OTA update may a session be '
    'in progress. Firmware shall reject OTA package delivery while a session is '
    'active. OTA must complete before the next session — the device is usable between '
    'OTA download and the reboot step (steps 1–7 may complete while device is connected '
    'to USB-C power but not in active session).')
add_normal(doc,
    'EMMC-OTA-03 (MANDATORY): The Safety MCU firmware (partition 3) shall NEVER be '
    'updated automatically via OTA. Safety MCU firmware update requires: '
    '(a) explicit user confirmation in app ("Update safety processor firmware — '
    'this requires your device to be connected to USB-C power"); '
    '(b) the update image in Scratch partition to pass Ed25519 signature verification; '
    '(c) main processor to flash Safety MCU via SPI, one 256-byte page at a time, '
    'with readback verification per page; '
    '(d) Safety MCU to perform self-test after update and report result to main '
    'processor before being armed. If any step fails, Safety MCU retains previous '
    'firmware and reports FAULT to app.')

add_heading(doc, '8.4  USB-C DFU Recovery Mode', level=2)
add_normal(doc,
    'EMMC-DFU-01 (MANDATORY): USB-C DFU (Device Firmware Upgrade) recovery mode shall '
    'be implemented in the bootloader and accessible without any firmware being present '
    'in either bank. DFU entry conditions:')
add_bullet(doc, 'Both Bank A and Bank B fail Ed25519 signature verification on boot.')
add_bullet(doc, 'User holds the physical reset button (hardware input, separate from app command) for 10 seconds during power-on.')
add_bullet(doc, 'App issues deliberate DFU-mode command (requires app authentication + double-confirmation).')
add_normal(doc,
    'EMMC-DFU-02 (MANDATORY): In DFU mode, the device presents as a USB DFU class '
    'device (USB DFU 1.1 specification). The NeurOne recovery tool (macOS, Windows, '
    'Linux) can flash a new firmware image to Bank A. DFU mode does NOT expose UHDR or '
    'SHDR partitions — only Firmware Bank A is writable in DFU mode. UHDR remains '
    'encrypted and inaccessible during DFU. DFU-flashed images must pass the same '
    'Ed25519 signature verification as OTA images.')


# ── §9 Config / Calibration Partition ─────────────────────────────────────────
add_heading(doc, '9.  CONFIG / CALIBRATION PARTITION')
add_normal(doc,
    'EMMC-CFG-01 (MANDATORY): The Config partition (partition 4, 16 MiB, LittleFS, '
    'AES-256-XTS with device manufacturing key) stores:')
add_bullet(doc, 'UHDR Argon2id salt (32 bytes, TRNG-generated at first setup).')
add_bullet(doc, 'SHDR AES-256-XTS key (K1_shdr || K2_shdr, 64 bytes, encrypted under SNVS OTPMK).')
add_bullet(doc, 'Device serial number (SHA-256 of eMMC CID, 32 bytes).')
add_bullet(doc, 'eMMC session counter (unsigned 64-bit integer, incremented at each session start — the only time-like SHDR reference).')
add_bullet(doc, 'UHDR authentication lockout state (consecutive failure count, lockout expiry in session-count units).')
add_bullet(doc, 'Bootloader version history (last 10 versions, session-count-indexed).')
add_bullet(doc, 'Factory calibration coefficients: PBM photodiode baseline per zone (PD1_baseline[5], PD2_baseline[5]); ADS1299 gain/offset per channel (8 channels); NTC thermistor coefficients per sensor (headset zones + hub).')
add_bullet(doc, 'USB-C PD negotiation profile (last-known charger capability, for power management optimisation).')
add_bullet(doc, 'BT/Wi-Fi pairing credentials (encrypted, device-scoped — not user identity).')
add_normal(doc,
    'EMMC-CFG-02 (MANDATORY): Config partition is write-protected after factory '
    'programming of all read-only fields (salt, SHDR key, serial number, calibration). '
    'Only the session counter, lockout state, and last-known PD profile may be written '
    'during normal operation. These fields occupy a dedicated Config partition journal '
    'area (raw-write, no LittleFS overhead for high-frequency updates).')


# ── §10 Scratch Partition ─────────────────────────────────────────────────────
add_heading(doc, '10.  SCRATCH PARTITION')
add_normal(doc,
    'EMMC-SCR-01 (MANDATORY): The Scratch partition (partition 7, 500 MiB) is '
    'unencrypted, unmounted (no filesystem), and zeroed on every boot before any '
    'other partition is mounted. It MUST NOT contain persistent data. Uses:')
add_bullet(doc, 'OTA firmware download staging (step 1 of §8.3 OTA sequence).')
add_bullet(doc, 'UHDR key rotation intermediate staging (§6.4): re-encrypted UHDR blocks staged here before being written back to UHDR partition.')
add_bullet(doc, 'On-device research anonymization workspace (§15): Argon2id output and anonymized extract assembly happen in Scratch; cleared immediately after transmission.')
add_bullet(doc, 'EDF+ export staging: when user initiates UHDR export via USB-C, decrypted and formatted EDF+ data is staged in Scratch before transmission, then immediately zeroed.')
add_normal(doc,
    'EMMC-SCR-02 (MANDATORY): Scratch zero-on-boot is a firm requirement. The '
    'bootloader shall issue eMMC erase commands covering the full Scratch partition '
    'LBA range before handing off to firmware. Erase operation uses eMMC CMD38 '
    '(ERASE) with ERASE_GROUP_START + ERASE_GROUP_END set to the Scratch partition '
    'boundaries. On eMMC with sanitize support (EXT_CSD[187] SANITIZE_START), '
    'firmware shall issue SANITIZE after erase to prevent recovery from erased blocks.')
add_normal(doc,
    'EMMC-SCR-03 (MANDATORY): If the Scratch erase fails at boot (eMMC ERASE command '
    'returns error): the device shall log the failure to SHDR and refuse to mount '
    'UHDR until the Scratch erase succeeds or the user explicitly bypasses via app '
    '(double-confirmation). This ensures research anonymization workspace is always '
    'clean.')


# ── §11 Boot Sequence ─────────────────────────────────────────────────────────
add_heading(doc, '11.  BOOT SEQUENCE')
add_normal(doc,
    'EMMC-BOOT-01 (MANDATORY): Complete boot sequence with storage subsystem '
    'initialisation:')
t8 = doc.add_table(rows=1, cols=3)
t8.style = 'Table Grid'
for cell, hdr in zip(t8.rows[0].cells, ['Phase', 'Action', 'FAULT handling']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
boot_seq = [
    ('B-1  Bootloader', 'Configure clocks, MPU, SNVS. Read boot bank flag. Verify firmware signature. Load and jump to firmware.', 'Both banks fail → USB-C DFU mode.'),
    ('B-2  Scratch erase', 'Firmware issues eMMC ERASE + optional SANITIZE over full Scratch partition before any other eMMC operation.', 'Erase fail → log SHDR fault (raw journal write), block UHDR mount until resolved.'),
    ('B-3  Config mount', 'Mount Config partition LittleFS (SHDR key from SNVS OTPMK). Load SHDR key (K1_shdr, K2_shdr). Load UHDR salt. Load session counter. Load factory calibration.', 'Config mount fail → FAULT. No session possible. Red LED blink. USB-C DFU available.'),
    ('B-4  SHDR mount', 'Mount SHDR partition LittleFS (K1_shdr, K2_shdr). Log boot event: bank used, firmware version, boot count.', 'SHDR mount fail → log to UART/USB-C debug. Device operational but SHDR telemetry lost until remount succeeds.'),
    ('B-5  eMMC health', 'Read EXT_CSD wear indicator. Verify Bank A write protection active. Log both to SHDR /shdr/telemetry/.', 'WP not active → SHDR FAULT. Not blocking.'),
    ('B-6  Safety MCU init', 'Main processor sends SPI heartbeat to Safety MCU. Safety MCU acknowledges and confirms all stimulation ENABLE lines are de-asserted.', 'No SPI heartbeat → FAULT. No session possible.'),
    ('B-7  Standby', 'System enters standby (Mode 1 / Mode 3 awaiting session start). UHDR is NOT mounted until user authentication (§6.3). BT/Wi-Fi available for app connection.', '—'),
    ('B-8  Session unlock', 'User authenticates via app (biometric / PIN). Argon2id executes, UHDR key derived, UHDR sentinel verified. UHDR LittleFS mounted.', 'KDF failure / sentinel mismatch → authentication failure, lockout timer starts.'),
    ('B-9  Session start', 'Session counter incremented in Config. New UHDR session record file created. Safety MCU armed (ENABLE lines may now be asserted per zone ID validation).', 'UHDR write fail at session start → abort session, log SHDR fault.'),
]
for row_data in boot_seq:
    row = t8.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val


# ── §12 Session Data Classification ───────────────────────────────────────────
add_heading(doc, '12.  SESSION DATA CLASSIFICATION (UHDR vs SHDR)')
add_normal(doc,
    'EMMC-CLS-01 (MANDATORY): Every data element produced during a session must be '
    'classified before any write is issued. The following table is the authoritative '
    'classification for all sensor outputs. This table is derived from CLAUDE.md §5.1 '
    'and resolves all boundary cases listed there. Classification changes require ECO.')
t9 = doc.add_table(rows=1, cols=4)
t9.style = 'Table Grid'
for cell, hdr in zip(t9.rows[0].cells,
                     ['Data element', 'Classification', 'Destination partition', 'Rationale']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
classification = [
    ('EEG waveforms (all channels, raw ADC samples)', 'UHDR', 'UHDR /sessions/', 'Reveals neurological state of person.'),
    ('EEG impedance raw values (per electrode, per sample)', 'UHDR', 'UHDR /sessions/', 'Skin impedance reveals physiological state (hydration, skin condition).'),
    ('EEG impedance trend slope (slope of impedance over session)', 'SHDR', 'SHDR /telemetry/', 'Electrode wear trend — device condition, no user biology.'),
    ('HRV time series (R-R intervals, RMSSD, LF/HF, coherence score)', 'UHDR', 'UHDR /sessions/', 'Reveals cardiac and autonomic state of person.'),
    ('PPG optical signal (raw and processed)', 'UHDR', 'UHDR /sessions/', 'Reveals cardiovascular physiology.'),
    ('Neurofeedback performance score (per session)', 'UHDR', 'UHDR /sessions/', 'Reveals cognitive performance.'),
    ('PBM dose per zone (J/cm², session total)', 'UHDR', 'UHDR /sessions/', 'Reveals photobiomodulation treatment received by person.'),
    ('PBM LED output ratio PD1/PD2 (per zone, per session)', 'SHDR', 'SHDR /telemetry/', 'LED aging and PDMS fouling trend — device condition. Ratio not absolute dose.'),
    ('NTC temperature raw values (during active session)', 'UHDR', 'UHDR /sessions/', 'Scalp-contact temperature reveals physiological proximity information.'),
    ('NTC temperature profile (between sessions, ambient trending)', 'SHDR', 'SHDR /telemetry/', 'Between-session thermal profile — device condition, not user biology.'),
    ('Session start/end wall-clock timestamps', 'UHDR', 'UHDR /sessions/', 'Reveals when person used device — behavioural/lifestyle data.'),
    ('Session count (unsigned integer, no timestamps)', 'SHDR', 'SHDR /telemetry/', 'Device usage count only. Explicitly listed as SHDR in CLAUDE.md §5.1.'),
    ('Protocol parameters used (stimulation waveforms, frequencies, durations)', 'UHDR', 'UHDR /sessions/', 'Reveals what therapeutic protocol person chose — health-related.'),
    ('Closed-loop adaptation events (frequency adjustments, trigger metrics)', 'UHDR', 'UHDR /sessions/', 'Reveals how person\'s brain responded to stimulation.'),
    ('User-entered symptom / outcome logs', 'UHDR', 'UHDR /outcomes/', 'Explicit health narrative entered by user.'),
    ('IR eye-open/closed state (during session)', 'UHDR', 'UHDR /sessions/', 'Explicitly listed as UHDR in CLAUDE.md §5.1.'),
    ('Safety interlock log (goggle halt events, photoparoxysmal detection)', 'SHDR', 'SHDR /faults/', 'Device safety event — device condition, not user biology. Explicitly listed as SHDR in CLAUDE.md §5.1.'),
    ('Accelerometer data during active session', 'UHDR', 'UHDR /sessions/', 'Motion during session can reveal user activity/context. CLAUDE.md §5.1 boundary resolution.'),
    ('Impact events (g-force, orientation, between sessions)', 'SHDR', 'SHDR /telemetry/', 'Device handling data between sessions — device condition. CLAUDE.md §5.1 boundary resolution.'),
    ('EMF shielding attenuation ratio (per fleet SHDR monitoring)', 'SHDR', 'SHDR /shielding/', 'Device shielding performance — device condition.'),
    ('USB-C insertion counter, PD negotiation log', 'SHDR', 'SHDR /telemetry/', 'Device connector wear — device condition.'),
    ('Supercapacitor cycle count, hub NTC aging estimate', 'SHDR', 'SHDR /supercap/', 'Device component aging — device condition.'),
    ('Firmware version history, OTA log', 'SHDR', 'SHDR /ota/', 'Device software history — device condition.'),
    ('Consumable session counts (sleeves, tips, pads, mesh)', 'SHDR', 'SHDR /consumables/', 'Device consumable status — device condition.'),
    ('Accessory authentication pass/fail (optical code + resistive sleeve)', 'SHDR', 'SHDR /telemetry/', 'Device accessory status — device condition.'),
    ('Calibration coefficient history (ADS1299, NTC, PBM PD)', 'SHDR', 'SHDR /calibration/', 'Device calibration state — device condition.'),
    ('Zone module FAULT events (contact fault, wrong zone, absent)', 'SHDR', 'SHDR /faults/', 'Device assembly status — device condition.'),
    ('Research study audit log (study ID, descriptor hash, byte count, session count)', 'SHDR', 'SHDR /research/', 'Device research participation metadata — device condition. No UHDR content.'),
]
for row_data in classification:
    row = t9.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val


# ── §13 UHDR Backup Protocol ──────────────────────────────────────────────────
add_heading(doc, '13.  UHDR BACKUP PROTOCOL')
add_normal(doc,
    'EMMC-BAK-01 (MANDATORY): Automated nightly incremental UHDR backup shall execute '
    'when: (a) device is connected to USB-C power (any PD level); AND (b) no session '
    'is in progress; AND (c) at least one UHDR session record has been written since '
    'the last successful backup. Backup shall not run on battery power — power '
    'interruption during backup could corrupt the backup target.')
add_normal(doc,
    'EMMC-BAK-02 (MANDATORY): Backup targets (user selects in app, at least one '
    'must be configured for backup to proceed):')
add_bullet(doc, 'USB-C local storage: a connected USB storage device (USB-A adapter via USB-C hub, or USB-C drive) receives an encrypted incremental backup archive. Backup is encrypted with the same UHDR AES-256-XTS key (K1, K2) — the backup archive is readable only with the user\'s biometric/PIN. NeurOne cannot decrypt a backup archive.')
add_bullet(doc, 'E2E encrypted cloud: UHDR backup archive transmitted to user-selected cloud provider (iCloud, Google Drive, or NeurOne-partnered encrypted backup service). Encrypted on-device before transmission. Cloud provider receives only ciphertext — they and NeurOne cannot decrypt. User-held key is the same biometric-derived Argon2id key.')
add_normal(doc,
    'EMMC-BAK-03 (MANDATORY): Backup format — incremental:')
add_bullet(doc, 'Backup manifest stored in UHDR /uhdr/backup_state/: last backed-up session counter, target identifier (hash of backup target path), last backup eMMC session count.')
add_bullet(doc, 'Incremental backup includes only session files with counter > last backed-up counter.')
add_bullet(doc, 'Backup archive format: tar.gz of UHDR session record files, then AES-256-XTS encrypted with UHDR key (same K1, K2). Archive includes a manifest header with device serial hash (no user identity), session count range, and HMAC-SHA256 of archive ciphertext.')
add_bullet(doc, 'Backup verification: after archive written to target, firmware reads back and verifies HMAC-SHA256 of archive ciphertext. Verification failure → retry once, then log SHDR fault.')
add_normal(doc,
    'EMMC-BAK-04 (MANDATORY): Backup failure is non-blocking — a failed backup does '
    'not prevent sessions. SHDR logs backup failure event. App displays notification '
    '"Backup failed — [reason]" with one-tap retry. After 7 consecutive failed backup '
    'attempts, app notification escalates to persistent banner (non-dismissible until '
    'backup succeeds or user explicitly defers for 30 days with confirmation).')
add_normal(doc,
    'EMMC-BAK-05 (MANDATORY): On USB-C DFU re-flash: UHDR partition is NOT erased '
    'by DFU mode (DFU only writes to Firmware Bank A). UHDR data survives firmware '
    'updates. The only operations that erase UHDR are: (a) explicit factory reset '
    '(user-initiated, double-confirmation); (b) UHDR partition failure requiring '
    're-initialisation (logs SHDR fault, prompts user). This makes firmware failure '
    'a hardware-swap event rather than a data-loss event (CLAUDE.md §7.2).')


# ── §14 Write Endurance Monitoring ────────────────────────────────────────────
add_heading(doc, '14.  WRITE ENDURANCE MONITORING (SHDR-BASED)')
add_normal(doc,
    'EMMC-WE-01 (MANDATORY): eMMC wear monitoring shall log the following to '
    'SHDR /shdr/telemetry/ at every boot and at 100-session checkpoints:')
t10 = doc.add_table(rows=1, cols=3)
t10.style = 'Table Grid'
for cell, hdr in zip(t10.rows[0].cells, ['Metric', 'Source', 'SHDR field']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
wear_metrics = [
    ('EXT_CSD[268] DEVICE_LIFE_TIME_EST_TYP_A', 'eMMC EXT_CSD register (SLC blocks)', 'emmc_life_type_a (uint8, 0x01=0–10%, 0x0A=90–100%)'),
    ('EXT_CSD[269] DEVICE_LIFE_TIME_EST_TYP_B', 'eMMC EXT_CSD register (MLC blocks, if present)', 'emmc_life_type_b (uint8)'),
    ('EXT_CSD[267] PRE_EOL_INFO', 'eMMC EXT_CSD (0x01=normal, 0x02=warning, 0x03=urgent)', 'emmc_pre_eol (uint8)'),
    ('LittleFS UHDR block utilisation', 'lfs_fs_stat().block_count − lfs_fs_stat().block_in_use', 'uhdr_blocks_free (uint32)'),
    ('LittleFS SHDR block utilisation', 'lfs_fs_stat() on SHDR mount', 'shdr_blocks_free (uint32)'),
    ('UHDR largest session record (bytes)', 'Max file size in /uhdr/sessions/', 'uhdr_max_session_bytes (uint64)'),
    ('UHDR write rate (bytes/session, rolling 10-session average)', 'Computed by firmware', 'uhdr_write_rate_avg (uint32)'),
    ('Estimated UHDR partition full date (session count)', 'uhdr_blocks_free / uhdr_write_rate_avg', 'uhdr_est_full_session (uint64)'),
]
for row_data in wear_metrics:
    row = t10.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val

add_normal(doc,
    'EMMC-WE-02 (MANDATORY): Predictive maintenance thresholds (reminder engine, '
    'per CLAUDE.md §5.2):')
add_bullet(doc, 'PRE_EOL_INFO = 0x02 (warning, ≥80% life consumed) → SHDR FAULT "eMMC approaching end of life". App: performance-critical notification, snooze max 3×.')
add_bullet(doc, 'PRE_EOL_INFO = 0x03 (urgent, ≥90% life consumed) → SHDR FAULT "eMMC replacement required". App: safety-critical notification, cannot be dismissed — blocks session start until user acknowledges and contacts support.')
add_bullet(doc, 'uhdr_est_full_session < 500 sessions remaining → notification "UHDR storage approaching capacity — consider archiving old sessions". Snooze max 5×.')
add_bullet(doc, 'shdr_blocks_free < 10% → notification "SHDR storage low — contact NeurOne support". Snooze max 3×.')
add_normal(doc,
    'EMMC-WE-03 (MANDATORY): Write amplification mitigation. LittleFS block_cycles = 500 '
    'controls wear leveling aggressiveness. Firmware shall monitor actual write amplification '
    'factor (WAF) by tracking: (bytes written to eMMC by firmware) / (bytes of new data '
    'written by application). WAF > 3× shall trigger a SHDR FAULT and is logged for '
    'fleet analysis. Expected WAF on LittleFS with SLC cache eMMC: 1.1–1.5×.')


# ── §15 Research Anonymization Workspace ──────────────────────────────────────
add_heading(doc, '15.  RESEARCH ANONYMIZATION WORKSPACE (SCRATCH PARTITION USAGE)')
add_normal(doc,
    'EMMC-RES-01 (MANDATORY): The research anonymization workflow (CLAUDE.md §5.3) '
    'uses the Scratch partition as a temporary workspace. The full sequence:')
add_bullet(doc, '(1) NeurOne server delivers signed study descriptor (study ID, approved UHDR element list, anonymization parameters) via USB-C or BT/Wi-Fi while user is authenticated. App verifies cryptographic signature on descriptor.')
add_bullet(doc, '(2) Firmware reads consented UHDR elements from UHDR partition (requires UHDR key — user must be authenticated). Copies raw consented elements to Scratch partition (Scratch is unencrypted — data exists in plaintext in Scratch only during processing, for minimum required time).')
add_bullet(doc, '(3) On-device anonymization engine applies transforms from study descriptor to Scratch data: k-anonymity grouping (k≥10), date/time rounding (≥1-week interval mapped to eMMC session count), direct identifier removal, quasi-identifier suppression. All transforms run in Scratch → Scratch (in-place rewrite).')
add_bullet(doc, '(4) Anonymized extract transmitted from Scratch to NeurOne research infrastructure. Extract is signed with device research key (separate from firmware signing key — prevents linkage) and includes study ID + extract HMAC.')
add_bullet(doc, '(5) Scratch zeroed (CMD38 ERASE + SANITIZE) immediately after transmission, before any other Scratch operation.')
add_bullet(doc, '(6) SHDR /shdr/research/ log updated: study ID, descriptor hash, extract byte count, eMMC session count. No UHDR content in SHDR research log.')
add_normal(doc,
    'EMMC-RES-02 (MANDATORY): Raw UHDR data NEVER leaves the Scratch partition. The '
    'only data transmitted to NeurOne research infrastructure is the post-anonymization '
    'extract in step (4). If transmission fails, Scratch is zeroed and the extract is '
    'not retransmitted automatically — the user must re-consent to regenerate the extract.')
add_normal(doc,
    'EMMC-RES-03 (MANDATORY): Consent withdrawal effect. Withdrawing research consent '
    'prevents the device from accepting any future study descriptors from NeurOne. '
    'Firmware rejects study descriptors for withdrawn studies at step (1). No extract '
    'is generated. This is effective for all future data periods, including historical '
    'sessions predating withdrawal — no extract is generated from any session if consent '
    'is withdrawn before step (1) of that study\'s data collection cycle.')


# ── §16 Processor Ownership ───────────────────────────────────────────────────
add_heading(doc, '16.  PROCESSOR OWNERSHIP')
t11 = doc.add_table(rows=1, cols=3)
t11.style = 'Table Grid'
for cell, hdr in zip(t11.rows[0].cells,
                     ['Firmware component', 'Owner processor', 'IEC 62304 class']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
ownership = [
    ('eMMC driver (low-level block read/write/erase)', 'i.MX RT1062 main processor', 'Class B'),
    ('LittleFS block device callbacks (all three instances)', 'i.MX RT1062 main processor', 'Class B'),
    ('AES-256-XTS encryption/decryption layer', 'i.MX RT1062 main processor (DCP hardware for AES-128 primitive)', 'Class B'),
    ('Argon2id key derivation (UHDR)', 'i.MX RT1062 main processor, LPSDR4 working memory', 'Class B'),
    ('HKDF-SHA256 SHDR key derivation', 'i.MX RT1062 main processor (at manufacturing / Config mount)', 'Class B'),
    ('SRAM key cache management (zero-on-session-end)', 'i.MX RT1062 main processor', 'Class B (safety-relevant — key zeroing on watchdog must be verified)'),
    ('Session data classification (UHDR vs SHDR tag)', 'i.MX RT1062 main processor, sensor/session manager layer', 'Class B'),
    ('UHDR session record write (EDF+ format, HMAC)', 'i.MX RT1062 main processor', 'Class B'),
    ('SHDR telemetry write (all categories)', 'i.MX RT1062 main processor', 'Class B'),
    ('OTA image verification (Ed25519 + SHA-256)', 'i.MX RT1062 main processor; bootloader (boot partitions)', 'Class B (bootloader) / Class B (firmware OTA)'),
    ('Safety MCU firmware update via SPI', 'i.MX RT1062 main processor (initiates) → Safety MCU (programs own flash)', 'Class B (main) / Class C (Safety MCU self-program)'),
    ('Boot bank flag (SNVS_LPGPR0)', 'i.MX RT1062 bootloader', 'Class B'),
    ('Scratch erase-on-boot (CMD38 + SANITIZE)', 'i.MX RT1062 main processor (early boot, before LittleFS mount)', 'Class B'),
    ('eMMC WP verification at boot', 'i.MX RT1062 main processor', 'Class B'),
    ('Research anonymization engine', 'i.MX RT1062 main processor', 'Class B'),
    ('UHDR backup scheduler and archive writer', 'i.MX RT1062 main processor', 'Class B'),
    ('Safety MCU: direct eMMC access', 'NONE — Safety MCU has NO eMMC interface', 'N/A — architectural boundary'),
]
for row_data in ownership:
    row = t11.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val


# ── §17 Cross-References and Gate Status ──────────────────────────────────────
add_heading(doc, '17.  CROSS-REFERENCES AND GATE STATUS')
t12 = doc.add_table(rows=1, cols=3)
t12.style = 'Table Grid'
for cell, hdr in zip(t12.rows[0].cells, ['Reference', 'Item', 'Status']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
xrefs = [
    ('CLAUDE.md §13.4 pending decision', 'eMMC partition architecture + separate UHDR/SHDR encryption in firmware specification before any storage code is written', 'CLOSED — this document (NP-FW-EMMC-001 Rev A)'),
    ('CLAUDE.md §13.4 pending decision', 'Dual-bank OTA bootloader (must be from first firmware line)', 'CLOSED — this document §8.1–§8.4'),
    ('CLAUDE.md §4.1', 'LittleFS filesystem; firmware partition write-protected; separate UHDR/SHDR partitions from first firmware line', 'ADDRESSED — §5 (LittleFS), §8.2 (write protection), §4.3 (separation)'),
    ('CLAUDE.md §5.1', 'UHDR AES-256 biometric-derived key; NeurOne does not hold decryption key; SHDR separate encryption', 'ADDRESSED — §6 (UHDR key), §7 (SHDR key)'),
    ('CLAUDE.md §5.3', 'Research anonymization on-device, per-study, fresh per request', 'ADDRESSED — §15'),
    ('NP-FW-REQ-001 Rev A §3.6', 'FW-APP-02: ZONE_ID FAULT events logged to SHDR with zone ID, fault type, ADC values, session count', 'CONSISTENT — SHDR /faults/ directory defined in §7.4; session count (not timestamp) as time reference'),
    ('CLAUDE.md §13.4 pending decision', 'eMMC partition architecture in firmware specification', 'CLOSED — ALSO closes the dual-bank OTA bootloader pending decision simultaneously'),
]
for row_data in xrefs:
    row = t12.add_row()
    for cell, val in zip(row.cells, row_data):
        cell.text = val


# ── §18 Revision History ──────────────────────────────────────────────────────
add_heading(doc, '18.  REVISION HISTORY')
t13 = doc.add_table(rows=1, cols=4)
t13.style = 'Table Grid'
for cell, hdr in zip(t13.rows[0].cells, ['Rev', 'Date', 'Author', 'Description']):
    cell.text = hdr
    for run in cell.paragraphs[0].runs:
        run.bold = True
rev_row = t13.add_row()
rev_row.cells[0].text = 'A'
rev_row.cells[1].text = '2026-05-11'
rev_row.cells[2].text = 'NeurOne Firmware'
rev_row.cells[3].text = (
    'Initial release. Closes CLAUDE.md §13.4 pending decisions: eMMC partition '
    'architecture + UHDR/SHDR encryption + dual-bank OTA bootloader. '
    'Specifies: 9-partition GPT layout (§4); LittleFS config (§5); UHDR AES-256-XTS + '
    'Argon2id key derivation + key lifecycle (§6); SHDR AES-256-XTS + HKDF device key '
    '+ SNVS OTPMK (§7); dual-bank OTA sequence + USB-C DFU recovery + Safety MCU '
    'firmware update (§8); Config partition (§9); Scratch with zero-on-boot (§10); '
    'boot sequence (§11); session data classification 27-element table (§12); '
    'UHDR backup protocol (§13); write endurance monitoring + predictive maintenance '
    'thresholds (§14); research anonymization workspace (§15); processor ownership (§16).'
)

doc.save('/home/user/NeurOne/docs/neurone_fw_emmc_001.docx')
print('Done: docs/neurone_fw_emmc_001.docx written')
