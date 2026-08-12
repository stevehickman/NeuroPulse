"""
Patch:
1. Update RISK-20 in risk register to reference NP-PROC-SUP-001.
2. Mitigate RISK-18: firmware debounce spec added to FPC spec + coord checklist.
3. Update risk register RISK-18 → MITIGATED.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

# ══════════════════════════════════════════════════════════════════════════════
# 1. Risk register — update RISK-20 spec ref + RISK-18 → MITIGATED
# ══════════════════════════════════════════════════════════════════════════════
RISK_PATH = "docs/superseded/np_risk_001.docx"
risk = Document(RISK_PATH)

for tbl in risk.tables:
    for row in tbl.rows:
        cells = row.cells
        if not cells:
            continue
        rid = cells[0].text.strip()

        if rid == "RISK-20" and len(cells) >= 9:
            # Update spec ref (col 7) and status (col 8)
            set_text(cells[7],
                "NP-PROC-SUP-001 §4 SUP-M-07, §6 SUP-B-01, §9\n"
                "NP-TOOL-ZM-001 §5 OI-1",
                size=8)
            set_text(cells[8],
                "OPEN — awaiting CFRP supplier written confirmation (Ra data + letter). "
                "NP-PROC-SUP-001 SUP-M-07 and SUP-B-01 are BLOCKING qualification items. "
                "Cannot release tooling until confirmed.",
                size=8)
            set_bg(cells[8], "FFF2CC")
            print("  RISK-20 spec ref + status updated")

        if rid == "RISK-18" and len(cells) >= 9:
            set_text(cells[1], "HIGH\n(MITIGATED)", bold=True,
                     color=(0x00, 0x70, 0x00), size=8)
            set_bg(cells[1], "E2EFDA")
            set_text(cells[5],
                "Three-layer mitigation: "
                "(1) Firmware: 3× ADC reads at 100 ms intervals before asserting FAULT — "
                "single transient contact resistance spike does not block session; "
                "only sustained open circuit triggers FAULT + audio alert. "
                "(2) FAI: FAI-K04 (pin 18 contact resistance ≤ 30 mΩ at 100 g insertion force) "
                "catches high-resistance contacts before field deployment. "
                "(3) App UX: FAULT condition triggers specific diagnostic message "
                "('Zone module contact issue — remove and reinsert') rather than generic error.",
                size=8)
            set_text(cells[7],
                "NP-HW-FPC-001 §5.2 pin 18; §11.4 ZONE_ID\n"
                "NP-FAI-ZM-001 FAI-K04\n"
                "NP-COORD-001 G2-07",
                size=8)
            set_text(cells[8],
                "MITIGATED — firmware debounce in spec; FAI-K04 validates; "
                "app UX diagnostic message specified",
                bold=True, color=(0x00, 0x70, 0x00), size=8)
            set_bg(cells[8], "92D050")
            print("  RISK-18 → MITIGATED")

risk.save(RISK_PATH)
print(f"  Risk register saved")

# ══════════════════════════════════════════════════════════════════════════════
# 2. FPC spec — add firmware debounce note to §5.2 ZONE_ID description
# ══════════════════════════════════════════════════════════════════════════════
FPC_PATH = "docs/superseded/np_hw_fpc_001.docx"
fpc = Document(FPC_PATH)

# Update pin 18 row in pinout table to add debounce note
for tbl in fpc.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 4 and cells[0].text.strip() == "18" and "ZONE_ID" in cells[1].text:
            # Append debounce note to description cell (col 3)
            existing = cells[3].text.strip()
            if "debounce" not in existing.lower():
                for para in cells[3].paragraphs:
                    for run in para.runs:
                        run.text = ""
                para = cells[3].paragraphs[0]
                run = para.add_run(
                    existing +
                    " Firmware debounce: Hub MCU performs 3 ADC reads at 100 ms intervals "
                    "before asserting FAULT — single transient open circuit does not block "
                    "session. Sustained open (all 3 reads fail) → FAULT + 'Zone module "
                    "contact issue' audio alert via bone conduction. (RISK-18 mitigation.)")
                run.font.size = Pt(8)
                print("  Pin 18 debounce note added to FPC pinout table")
            break

# Find the ZONE_ID description paragraph and add debounce spec after it
zone_id_para = None
for p in fpc.paragraphs:
    if "ZONE_ID" in p.text and "resistor" in p.text.lower() and "Layer 2" in p.text:
        zone_id_para = p
        break

if zone_id_para:
    # Insert debounce firmware paragraph after zone_id_para
    DEBOUNCE_TEXT = (
        "RISK-18 firmware debounce specification (ZONE_ID): The Hub MCU shall perform "
        "3 consecutive ADC reads of pin 18 at 100 ms intervals after zone module insertion "
        "before making a FAULT decision. Accept criterion: at least 2 of 3 reads must return "
        "a value within the valid ZONE_ID band (±5% of target) for the module to be accepted. "
        "A single transient reading outside band (e.g. from contact bounce at insertion) "
        "shall not trigger FAULT. If all 3 reads fail: assert FAULT, hold ENABLE low, "
        "issue bone conduction alert 'Zone module contact issue — remove and reinsert'. "
        "If ZONE_ID reads valid but does not match expected zone: assert FAULT, issue "
        "'Wrong module' alert (existing RISK-15 interlock).")
    new_p = OxmlElement("w:p")
    zone_id_para._p.addnext(new_p)
    for p_obj in fpc.paragraphs:
        if p_obj._p is new_p:
            run = p_obj.add_run(DEBOUNCE_TEXT)
            run.font.size = Pt(10)
            break
    print("  ZONE_ID firmware debounce spec paragraph added")

fpc.save(FPC_PATH)
print(f"  FPC spec saved")

# ══════════════════════════════════════════════════════════════════════════════
# 3. Coord checklist — add G1-07 (ZONE_ID firmware debounce spec)
# ══════════════════════════════════════════════════════════════════════════════
COORD_PATH = "docs/np_coord_001.docx"
coord = Document(COORD_PATH)

# Find G1-06 and insert G1-07 after it
for tbl in coord.tables:
    for row in tbl.rows:
        cells = row.cells
        if len(cells) >= 1 and cells[0].text.strip() == "G1-06":
            new_tr = OxmlElement("w:tr")
            row._tr.addnext(new_tr)
            g107_data = [
                "G1-07",
                "G1",
                "ZONE_ID firmware debounce specification: EE to add to firmware requirements "
                "document — 3× ADC reads at 100 ms intervals before FAULT assertion on pin 18. "
                "Accept: ≥2/3 reads in valid band. FAULT: all 3 reads fail. "
                "Audio alert text: 'Zone module contact issue — remove and reinsert'. "
                "Ref: RISK-18 mitigation. NP-HW-FPC-001 §11.4.",
                "EE",
                "CLOSED — spec in NP-HW-FPC-001 §11.4; firmware req to be written"
            ]
            widths_in = [0.65, 0.55, 3.5, 0.9, 1.2]
            for c_idx, txt in enumerate(g107_data):
                tc = OxmlElement("w:tc")
                tc_pr = OxmlElement("w:tcPr")
                w_el = OxmlElement("w:tcW")
                w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
                w_el.set(qn("w:type"), "dxa")
                tc_pr.append(w_el); tc.append(tc_pr)
                p_el = OxmlElement("w:p")
                r_el = OxmlElement("w:r")
                rpr_el = OxmlElement("w:rPr")
                sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
                rpr_el.append(sz_el); r_el.append(rpr_el)
                t_el = OxmlElement("w:t")
                t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
                r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
                new_tr.append(tc)
            for r2 in tbl.rows:
                if r2._tr is new_tr:
                    set_bg(r2.cells[4], "E2EFDA")  # green — spec written
                    if r2.cells[0].paragraphs[0].runs:
                        r2.cells[0].paragraphs[0].runs[0].bold = True
                    break
            print("  G1-07 ZONE_ID debounce item added")
            break

coord.save(COORD_PATH)
print(f"  Coord checklist saved")
print("Done.")
