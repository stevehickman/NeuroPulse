"""
Generate NP-PROC-SUP-001 Rev A: NeuroPulse Tooling Supplier Selection Checklist.
Covers injection moulding (zone module + shell), FPC fabrication, and PDMS bonding.
RISK-20 (CFRP Ra ≤ 1.6 µm) is a BLOCKING qualification item.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── helpers ───────────────────────────────────────────────────────────────────
def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_run(cell, text, bold=False, color=None, size=8):
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_heading(doc, text, color=(0x1F, 0x49, 0x7D), size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p

def add_body(doc, text, size=10, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def make_table(doc, headers, widths_in, hdr_bg="1F497D"):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    hrow = tbl.rows[0]
    for i, h in enumerate(headers):
        set_bg(hrow.cells[i], hdr_bg)
        r = hrow.cells[i].paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(8)
        hrow.cells[i].width = Inches(widths_in[i])
    return tbl

def add_row(tbl, data, widths_in, idx=0, col0_bold=True):
    row = tbl.add_row()
    bg = "FFFFFF" if idx % 2 == 0 else "F5F5F5"
    for i, txt in enumerate(data):
        cell = row.cells[i]
        set_bg(cell, bg)
        r = cell.paragraphs[0].add_run(txt)
        r.font.size = Pt(8)
        if i == 0 and col0_bold:
            r.bold = True
        cell.width = Inches(widths_in[i])
    return row

def add_block_row(tbl, data, widths_in, idx=0, bg_override=None, bold_col0=True):
    row = add_row(tbl, data, widths_in, idx, bold_col0)
    if bg_override:
        for cell in row.cells:
            set_bg(cell, bg_override)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    return row

# ═════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── Title block ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("NP-PROC-SUP-001")
r.bold = True; r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p = doc.add_paragraph()
r = p.add_run("NeuroPulse Tooling & Process Supplier Selection Checklist")
r.bold = True; r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

meta = [
    ("Document number:", "NP-PROC-SUP-001"),
    ("Revision:", "A (Draft)"),
    ("Status:", "ACTIVE — must be completed before supplier engagement"),
    ("Scope:", "Injection moulding (zone module body + headset shell), FPC fabrication, PDMS bonding"),
    ("Related specs:", "NP-HW-FPC-001, NP-DRV-SHELL-001, NP-TOOL-ZM-001, NP-PROC-FPC-001"),
    ("Date:", datetime.date.today().strftime("%Y-%m-%d")),
]
meta_tbl = doc.add_table(rows=len(meta), cols=2)
meta_tbl.style = "Table Grid"
for i, (label, val) in enumerate(meta):
    set_bg(meta_tbl.rows[i].cells[0], "1F497D")
    r = meta_tbl.rows[i].cells[0].paragraphs[0].add_run(label)
    r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(9)
    r = meta_tbl.rows[i].cells[1].paragraphs[0].add_run(val)
    r.font.size = Pt(9)
    meta_tbl.rows[i].cells[0].width = Inches(2.0)
    meta_tbl.rows[i].cells[1].width = Inches(4.8)

doc.add_paragraph()

# ── §1 Purpose ────────────────────────────────────────────────────────────────
add_heading(doc, "1.  PURPOSE AND SCOPE")
add_body(doc,
    "This checklist qualifies suppliers for three manufacturing processes critical to the "
    "NeuroPulse zone module and headset shell: injection moulding with co-moulding of "
    "silicone components, CFRP shell fabrication, and PDMS bonding. It must be completed "
    "for each candidate supplier before an NDA or purchase order is issued.")
add_body(doc,
    "BLOCKING items (marked ⛔) are disqualifying if the supplier cannot confirm compliance "
    "in writing. Non-blocking items are scored; a supplier scoring fewer than 80% of "
    "non-blocking items is placed on probation with a corrective action plan required "
    "within 30 days of engagement.")
add_body(doc,
    "RISK-20 traceability: This checklist is the primary mitigation document for RISK-20 "
    "(CFRP shell slot rim Ra ≤ 1.6 µm surface finish). The RISK-20 item (SUP-M-07) is "
    "BLOCKING — tooling must not be released without a written supplier confirmation "
    "that this surface finish is achievable with moulded CFRP at the slot rim geometry.")
doc.add_paragraph()

# ── §2 Supplier categories ────────────────────────────────────────────────────
add_heading(doc, "2.  SUPPLIER CATEGORIES")
add_body(doc,
    "This document applies to three supplier types. A single supplier may cover more than "
    "one category if they have demonstrated capability in all relevant sections.")

cat_hdrs = ["Category", "Scope", "Key Sections", "Min. Lead Time Required"]
cat_widths = [0.9, 3.0, 1.5, 1.3]
cat_tbl = make_table(doc, cat_hdrs, cat_widths)
CATS = [
    ("CAT-A\nMoulding", "Injection moulding of zone module body (PC/ABS) + co-moulding of Shore 40–50A silicone gasket. "
     "5 unique mould cavities (one per zone position). Volume: T1 ramp to ~10,000 units/year.",
     "§3, §4, §5", "12 weeks DFM → first article"),
    ("CAT-B\nCFRP Shell", "CFRP headset shell fabrication. Includes mu-metal liner bonding and FPC channel moulding. "
     "Surface finish requirement at zone slot rims: Ra ≤ 1.6 µm (RISK-20). "
     "Volume: same as zone module.",
     "§3, §4, §6", "16 weeks DFM → first article"),
    ("CAT-C\nPDMS Bonding", "PDMS optical window bonding to FPC coverlay (SiO₂ interlayer, O₂ plasma activation, "
     "120°C cure). Batch peel testing per lot. Biocompatibility certificate required.",
     "§3, §7", "8 weeks process validation"),
]
for idx, row_data in enumerate(CATS):
    add_row(cat_tbl, row_data, cat_widths, idx)

doc.add_paragraph()

# ── §3 General Qualification Requirements ─────────────────────────────────────
add_heading(doc, "3.  GENERAL QUALIFICATION REQUIREMENTS (all categories)")
add_body(doc,
    "Items in this section apply to all supplier categories. "
    "Complete this section first; a BLOCKING failure here ends evaluation.")
doc.add_paragraph()

gen_hdrs = ["Item", "Requirement", "Evidence Required", "Type", "Confirmed?"]
gen_widths = [0.7, 3.2, 1.7, 0.75, 0.85]
gen_tbl = make_table(doc, gen_hdrs, gen_widths)
GEN_ROWS = [
    ("SUP-G-01", "ISO 9001:2015 certification (or ISO 13485:2016 for medical-adjacent work). "
     "Certificate must be current (within 3-year cycle, last audit <18 months ago).",
     "Current ISO certificate", "⛔ BLOCKING", "☐"),
    ("SUP-G-02", "ESD-safe production environment for any operation involving bare PCBs, FPCs, or "
     "electronic components. ESD flooring, wrist straps, and ionisers documented.",
     "ESD audit report or site visit", "Required", "☐"),
    ("SUP-G-03", "Documented non-conforming material control process: quarantine, disposition, and "
     "8D corrective action capability.",
     "NCR procedure + sample 8D", "Required", "☐"),
    ("SUP-G-04", "Confidentiality: willing to sign NeuroPulse NDA before receiving any drawings. "
     "Drawings classified CONFIDENTIAL.",
     "NDA executed before DFM package", "⛔ BLOCKING", "☐"),
    ("SUP-G-05", "Lead time: able to deliver first-article samples within the lead times in §2 "
     "(CAT-A: 12 weeks, CAT-B: 16 weeks, CAT-C: 8 weeks) from DFM approval.",
     "Written schedule commitment", "Required", "☐"),
    ("SUP-G-06", "Pricing transparency: willing to provide itemised cost breakdown (tooling, "
     "per-part, secondary operations, inspection) within 10 business days of RFQ.",
     "RFQ response with itemised BOM", "Required", "☐"),
    ("SUP-G-07", "Supply chain resilience: no single-source raw material dependency without "
     "documented alternative. Lead times for all key materials stated.",
     "Material sourcing declaration", "Advisory", "☐"),
]
for idx, row_data in enumerate(GEN_ROWS):
    if "BLOCKING" in row_data[3]:
        add_block_row(gen_tbl, row_data, gen_widths, idx, bg_override="FFE7E7")
    else:
        add_row(gen_tbl, row_data, gen_widths, idx)

doc.add_paragraph()

# ── §4 Dimensional and Metrology Capability ───────────────────────────────────
add_heading(doc, "4.  DIMENSIONAL AND METROLOGY CAPABILITY")
add_body(doc,
    "NeuroPulse zone module and shell features include tolerances as tight as ±0.05 mm "
    "(braille dot height) and surface finish requirements as tight as Ra ≤ 1.6 µm "
    "(gasket seat). Supplier must demonstrate measurement capability before FAI approval.")
doc.add_paragraph()

dim_hdrs = ["Item", "Requirement", "Evidence Required", "Type", "Confirmed?"]
dim_widths = [0.7, 3.2, 1.7, 0.75, 0.85]
dim_tbl = make_table(doc, dim_hdrs, dim_widths)
DIM_ROWS = [
    ("SUP-M-01", "CMM or optical measurement system capable of ±0.02 mm measurement uncertainty "
     "for features ≤ 5 mm. Required for gasket groove depth (±0.1 mm) and PD2 aperture "
     "centre offset (≤0.2 mm).",
     "CMM calibration certificate + capability study", "⛔ BLOCKING", "☐"),
    ("SUP-M-02", "Contact profilometer or optical profilometer for surface roughness measurement "
     "(Ra). Measurement uncertainty ≤ 0.05 µm Ra. Required for slot rim Ra ≤ 1.6 µm (RISK-20).",
     "Profilometer calibration certificate", "⛔ BLOCKING", "☐"),
    ("SUP-M-03", "Gage R&R study on at least one representative feature before production approval "
     "(minimum 10 parts, 2 operators, 2 replicates). Gage R&R ≤ 30% of tolerance.",
     "Gage R&R report", "Required", "☐"),
    ("SUP-M-04", "First Article Inspection (FAI) report capability: able to produce balloon-annotated "
     "drawing with measured values for every dimensioned feature on the NeuroPulse zone module "
     "drawing (estimated 35–50 dimensions).",
     "Sample FAI report from prior project", "Required", "☐"),
    ("SUP-M-05", "In-process statistical process control (SPC) for critical dimensions: "
     "Cpk ≥ 1.33 target for all BLOCKING dimensions (gasket height, PD2 aperture, "
     "mechanical key geometry).",
     "SPC procedure + sample control chart", "Advisory", "☐"),
    ("SUP-M-06", "Colour measurement capability (spectrophotometer, ΔE ≤ 1.0) for zone module "
     "colour coding verification (5 colours, bulk material). "
     "Required before each production lot approval.",
     "Spectrophotometer calibration cert", "Required", "☐"),
    # ── RISK-20 BLOCKING ITEM ─────────────────────────────────────────────────
    ("SUP-M-07\n⛔ RISK-20",
     "[BLOCKING — RISK-20] CFRP shell slot rim surface finish: supplier must provide written "
     "confirmation, with supporting measurement data, that moulded CFRP can achieve "
     "Ra ≤ 1.6 µm at the zone slot rim contact surface (gasket mating face) without "
     "secondary machining operations.\n\n"
     "Background: moulded CFRP typically achieves Ra 1.6–3.2 µm depending on fibre "
     "orientation, resin system, and tool surface finish. The Ra ≤ 1.6 µm requirement "
     "comes from the RISK-16 Option A self-sealing gasket: insufficient surface finish "
     "at the gasket seat results in leakage paths and IPX4 non-compliance. "
     "If the supplier cannot achieve Ra ≤ 1.6 µm by moulding alone, they must propose "
     "a secondary operation (e.g. resin flooding, tool polishing, or post-mould machining) "
     "with confirmed cost and lead time impact before tooling is approved.\n\n"
     "Confirmation format: written statement on supplier letterhead specifying: "
     "(1) measured Ra achieved at equivalent CFRP geometry, (2) tool surface finish "
     "used (Ra of tool steel), (3) resin system, (4) whether secondary operation is "
     "required and at what cost. Measurements must be from a coupon or prior part, "
     "not theoretical. Supplier signature required.\n\n"
     "If Ra > 1.6 µm is confirmed unavoidable: escalate to ME + EE immediately; "
     "gasket geometry or shell slot design must be revised before mould release.",
     "Written confirmation on letterhead + Ra measurement data from representative CFRP coupon",
     "⛔ BLOCKING\nRISK-20", "☐"),
    ("SUP-M-08", "If secondary operation required for Ra ≤ 1.6 µm (per SUP-M-07): provide written "
     "process description, per-unit cost, lead time impact, and process control plan "
     "to ensure consistency across production lots.",
     "Process description + cost/lead time quote", "Conditional on SUP-M-07", "☐"),
]
for idx, row_data in enumerate(DIM_ROWS):
    if "BLOCKING" in row_data[3] or "RISK-20" in row_data[0]:
        add_block_row(dim_tbl, row_data, dim_widths, idx, bg_override="FFE7E7")
    else:
        add_row(dim_tbl, row_data, dim_widths, idx)

doc.add_paragraph()

# ── §5 Injection Moulding + Co-moulding (CAT-A) ──────────────────────────────
add_heading(doc, "5.  INJECTION MOULDING + CO-MOULDING CAPABILITY  (CAT-A)")
add_body(doc,
    "Required for zone module body supplier. The zone module body requires simultaneous "
    "co-moulding of Shore 40–50A silicone (gasket, feature F-05/F-06 in NP-TOOL-ZM-001) "
    "with a PC/ABS thermoplastic shell. Five distinct mould tools are required "
    "(one per zone position, unique mechanical key per zone).")
doc.add_paragraph()

mould_hdrs = ["Item", "Requirement", "Evidence Required", "Type", "Confirmed?"]
mould_widths = [0.7, 3.2, 1.7, 0.75, 0.85]
mould_tbl = make_table(doc, mould_hdrs, mould_widths)
MOULD_ROWS = [
    ("SUP-A-01", "Two-shot or insert moulding capability for silicone co-moulding with thermoplastic. "
     "Silicone material: Shore 40–50A UV-stable medical-grade silicone (ISO 10993-5 certified). "
     "Must demonstrate adhesion without delamination after 1,000 insertion cycles.",
     "Process capability study or equivalent prior project samples",
     "⛔ BLOCKING", "☐"),
    ("SUP-A-02", "Silicone primer application capability: Dow Corning 1200 OS Primer or equivalent, "
     "applied to thermoplastic gasket retention groove before co-moulding. "
     "Primer must be qualified for the selected thermoplastic (PC/ABS). "
     "Confirm primer compatibility in writing.",
     "Primer qualification data / supplier letter", "⛔ BLOCKING", "☐"),
    ("SUP-A-03", "Mould tool for small features: capable of producing braille dot diameter 1.5 mm "
     "± 0.1 mm, height 0.48 mm ± 0.05 mm, and inter-dot spacing 2.3 mm. "
     "EDM or high-precision milling required for braille cavity.",
     "Sample braille moulded part + measurements", "Required", "☐"),
    ("SUP-A-04", "Five unique mould tools (one per zone position, unique mechanical key geometry). "
     "Willing to manufacture all 5 with shared runner/gate design where possible "
     "to reduce per-part cost at volume.",
     "Tooling quote per zone variant", "Required", "☐"),
    ("SUP-A-05", "Bulk colour moulding in 5 colours (Blue/Red/Green/Yellow/Purple for ZM-01–05). "
     "Colour in material — not surface paint. Colour-fast under IPA wipe (1,000 cycles) "
     "and UV exposure (500 hours QUV-A).",
     "Material colour specification + QUV data or commitment", "Required", "☐"),
    ("SUP-A-06", "Moulded-in pull tab geometry (F-08): 8 × 5 mm raised tongue, 2 mm height, "
     "knurl texture Ra 1.6 µm on grip surface. Confirm draft angle feasibility.",
     "DFM review comment on pull tab", "Required", "☐"),
    ("SUP-A-07", "Gasket compression set validation capability: 70h at 70°C per ASTM D395. "
     "Accept criterion: compression set ≤ 20%. Must be performed on co-moulded gasket "
     "(not free silicone strip) for first-article qualification.",
     "Test capability confirmation or outsource plan", "Advisory", "☐"),
    ("SUP-A-08", "DFM review turnaround: provide formal DFM comments within 10 business days of "
     "receiving drawings. DFM must address all 8 features in NP-TOOL-ZM-001 §2.",
     "Commitment letter", "Required", "☐"),
]
for idx, row_data in enumerate(MOULD_ROWS):
    if "BLOCKING" in row_data[3]:
        add_block_row(mould_tbl, row_data, mould_widths, idx, bg_override="FFE7E7")
    else:
        add_row(mould_tbl, row_data, mould_widths, idx)

doc.add_paragraph()

# ── §6 CFRP Shell Fabrication (CAT-B) ────────────────────────────────────────
add_heading(doc, "6.  CFRP SHELL FABRICATION  (CAT-B)")
add_body(doc,
    "Required for headset shell supplier. The CFRP shell integrates zone module slot "
    "geometry, FPC routing channels, mu-metal liner bonding, and the gasket mating "
    "surface at each slot rim. Surface finish at the slot rim is the primary RISK-20 "
    "qualification criterion (see SUP-M-07 in §4 — BLOCKING).")
doc.add_paragraph()

cfrp_hdrs = ["Item", "Requirement", "Evidence Required", "Type", "Confirmed?"]
cfrp_widths = [0.7, 3.2, 1.7, 0.75, 0.85]
cfrp_tbl = make_table(doc, cfrp_hdrs, cfrp_widths)
CFRP_ROWS = [
    ("SUP-B-01",
     "[BLOCKING — RISK-20 primary item] Written confirmation of achievable Ra at CFRP "
     "slot rim: see SUP-M-07 in §4 for full requirement. Restated here for emphasis. "
     "No CFRP shell supplier is approved without SUP-M-07 written confirmation.",
     "See SUP-M-07 — letter + Ra data", "⛔ BLOCKING\nRISK-20", "☐"),
    ("SUP-B-02", "CFRP tooling experience: minimum 3 prior production programmes with moulded CFRP "
     "enclosures, internal channel features, and dimensional tolerances ≤ ±0.2 mm. "
     "Provide representative samples or part numbers (under NDA if required).",
     "Programme references + samples", "Required", "☐"),
    ("SUP-B-03", "Mu-metal liner bonding: confirmed process for bonding 0.2 mm mu-metal liner to "
     "inner CFRP surface with PETG laminate encapsulation and silicone RTV sealant at "
     "all cutout edges. No adhesive gap > 0.1 mm. Liner continuity verified by conductivity "
     "test after bonding.",
     "Process description + conductivity test method", "Required", "☐"),
    ("SUP-B-04", "FPC routing channel geometry: internal channels 12 mm wide × 3 mm deep, "
     "corner radii ≥ 5 mm, surface finish Ra ≤ 3.2 µm (less critical than slot rim). "
     "Three retention clip snap features per channel.",
     "DFM review on channel geometry", "Required", "☐"),
    ("SUP-B-05", "EEG cable routing channel (NP-DRV-SHELL-001 §2.4): separate 8 mm × 5 mm "
     "channel on outer CFRP inner surface, opposite side from FPC channels. "
     "Confirm no interference between the two channel systems in shell wall thickness.",
     "CAD section analysis confirming wall thickness ≥ 2 mm", "Required", "☐"),
    ("SUP-B-06", "EMF shielding fabric bonding: palladium-coated polyester inner liner, "
     "adhesive-bonded to CFRP inner surface after mu-metal. No wrinkles or gaps > 2 mm. "
     "Attenuation verification method (4-point probe resistance after bonding).",
     "Process description", "Required", "☐"),
    ("SUP-B-07", "Zone slot rim geometry: 5 zone slot openings with asymmetric mechanical key "
     "receiver profiles (one per zone), matching zone module keys (F-01 in NP-TOOL-ZM-001). "
     "Confirm machined vs. moulded receiver profile and achievable positional tolerance.",
     "DFM comment on receiver profile method + tolerance", "Required", "☐"),
    ("SUP-B-08", "Anchor post moulding: tethered interface cover anchor posts (5 zone slot posts + "
     "3 accessory port positions) moulded into shell at zero incremental tooling cost "
     "if specified before first cut. Confirm this is achievable.",
     "Confirmation in DFM review", "Advisory", "☐"),
]
for idx, row_data in enumerate(CFRP_ROWS):
    if "BLOCKING" in row_data[3]:
        add_block_row(cfrp_tbl, row_data, cfrp_widths, idx, bg_override="FFE7E7")
    else:
        add_row(cfrp_tbl, row_data, cfrp_widths, idx)

doc.add_paragraph()

# ── §7 PDMS Bonding (CAT-C) ──────────────────────────────────────────────────
add_heading(doc, "7.  PDMS OPTICAL WINDOW BONDING  (CAT-C)")
add_body(doc,
    "Required for PDMS window bonding supplier (may be the same as CAT-A moulder or a "
    "specialist optical bonding house). Full process sequence is defined in "
    "NP-HW-FPC-001 §9.2. Deviation from the specified sequence is not permitted.")
doc.add_paragraph()

pdms_hdrs = ["Item", "Requirement", "Evidence Required", "Type", "Confirmed?"]
pdms_widths = [0.7, 3.2, 1.7, 0.75, 0.85]
pdms_tbl = make_table(doc, pdms_hdrs, pdms_widths)
PDMS_ROWS = [
    ("SUP-C-01", "SiO₂ sputtering capability: RF magnetron sputtering of 75 nm SiO₂ on PI coverlay "
     "surface. Process temperature ≤ 80°C (FPC thermal limit). Thickness uniformity ±10%.",
     "Process spec + thickness measurement capability", "⛔ BLOCKING", "☐"),
    ("SUP-C-02", "O₂ plasma activation: 100 W, 30 sccm O₂, 60 s. Both PDMS and SiO₂ surfaces. "
     "Bond initiation within 10 minutes of plasma treatment (< 10 min delay spec). "
     "In-process timer control required.",
     "Process control plan with timer spec", "⛔ BLOCKING", "☐"),
    ("SUP-C-03", "Pressure application: 50 g/cm² uniform pressure during 120°C / 2 hr cure. "
     "Fixture designed to avoid pressure concentration at FPC component features. "
     "Temperature uniformity ±5°C across part.",
     "Fixture design + oven calibration cert", "Required", "☐"),
    ("SUP-C-04", "Batch peel test: minimum 5 coupons per production batch. Peel force > 150 N/m "
     "(per IPC-TM-650 2.4.9). Batch released only after peel test pass. "
     "Peel test data retained with batch record.",
     "Sample batch peel test record", "⛔ BLOCKING", "☐"),
    ("SUP-C-05", "PDMS biocompatibility: ISO 10993-5 cytotoxicity certificate for PDMS compound "
     "used. Certificate must cover the specific formulation (not generic PDMS category).",
     "ISO 10993-5 certificate for specific compound", "⛔ BLOCKING", "☐"),
    ("SUP-C-06", "Plasma-activated anti-fouling surface treatment: O₂ plasma to produce hydrophilic "
     "PDMS surface (contact angle < 20° immediately post-treatment). Must be the final "
     "process step before packaging. Confirm shelf life of treatment (target ≥ 6 months "
     "in sealed packaging).",
     "Contact angle measurement data + shelf life data", "Required", "☐"),
    ("SUP-C-07", "Optical transmission measurement: transmittance ≥ 85% at 660 nm and 810 nm "
     "through bonded PDMS window assembly. Measurement per batch using spectrophotometer.",
     "Spectrophotometer capability + sample data", "Required", "☐"),
]
for idx, row_data in enumerate(PDMS_ROWS):
    if "BLOCKING" in row_data[3]:
        add_block_row(pdms_tbl, row_data, pdms_widths, idx, bg_override="FFE7E7")
    else:
        add_row(pdms_tbl, row_data, pdms_widths, idx)

doc.add_paragraph()

# ── §8 Scoring Summary ────────────────────────────────────────────────────────
add_heading(doc, "8.  SCORING SUMMARY AND QUALIFICATION DECISION")
add_body(doc,
    "Complete this section after evaluating all applicable sections above. "
    "A supplier fails qualification if any BLOCKING item is not confirmed. "
    "For non-blocking items, calculate the percentage confirmed.")
doc.add_paragraph()

score_hdrs = ["Criterion", "Result", "Notes"]
score_widths = [3.0, 1.5, 2.2]
score_tbl = make_table(doc, score_hdrs, score_widths)
SCORE_ROWS = [
    ("All BLOCKING items confirmed (zero exceptions allowed)", "☐ YES  ☐ NO", ""),
    ("SUP-M-07 RISK-20 written confirmation received (letter + Ra data)", "☐ YES  ☐ NO", "BLOCKING — do not proceed without"),
    ("Non-blocking items confirmed (%)", "    / total", "Target ≥ 80%"),
    ("ISO 9001 / 13485 certificate verified", "☐ YES  ☐ NO", ""),
    ("NDA executed", "☐ YES  ☐ NO", "Required before drawings released"),
    ("DFM review turnaround commitment received", "☐ YES  ☐ NO", ""),
    ("QUALIFICATION DECISION", "☐ APPROVED  ☐ PROBATION  ☐ REJECTED", ""),
    ("Decision date", "", ""),
    ("Approved by (ME)", "", ""),
    ("Approved by (EE)", "", ""),
    ("Approved by (Procurement)", "", ""),
]
for idx, row_data in enumerate(SCORE_ROWS):
    row = add_row(score_tbl, row_data, score_widths, idx, col0_bold=False)
    if "QUALIFICATION DECISION" in row_data[0]:
        for cell in row.cells:
            set_bg(cell, "1F497D")
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()

# ── §9 RISK-20 Mitigation Status ─────────────────────────────────────────────
add_heading(doc, "9.  RISK-20 MITIGATION TRACKING")
add_body(doc,
    "RISK-20 (CFRP shell slot rim Ra ≤ 1.6 µm may require secondary operations) is tracked "
    "here as the primary mitigation record. Status must be updated when supplier "
    "confirmation is received.")
doc.add_paragraph()

r20_hdrs = ["Field", "Value"]
r20_widths = [2.2, 4.6]
r20_tbl = make_table(doc, r20_hdrs, r20_widths, hdr_bg="C00000")
R20_ROWS = [
    ("Risk ID", "RISK-20"),
    ("Severity", "HIGH — BLOCKING for tooling release"),
    ("Root cause", "Moulded CFRP typically achieves Ra 1.6–3.2 µm. Gasket seal requires Ra ≤ 1.6 µm "
     "at slot rim. Gap between achievable and required finish is unconfirmed."),
    ("Mitigation action", "Obtain written confirmation from CFRP shell tooling supplier (SUP-M-07 + SUP-B-01) "
     "that Ra ≤ 1.6 µm is achievable by moulding, or define and cost a secondary operation."),
    ("Mitigation document", "This document (NP-PROC-SUP-001), item SUP-M-07 and SUP-B-01"),
    ("Also referenced in", "NP-TOOL-ZM-001 §5 OI-1; NP-COORD-001 G2 gate"),
    ("Confirmation status", "OPEN — awaiting supplier engagement"),
    ("Supplier contacted", "(complete when engagement begins)"),
    ("Confirmation received", "☐ YES — date: ________   ☐ NO — secondary operation required"),
    ("If secondary op required", "Proposed operation: ________________  Cost impact: $______  Lead time: ______"),
    ("ME sign-off", "(sign when confirmed)"),
    ("EE sign-off", "(sign when confirmed)"),
    ("Risk closed?", "☐ YES — date: ________   ☐ NO — still open"),
]
for idx, row_data in enumerate(R20_ROWS):
    row = add_row(r20_tbl, row_data, r20_widths, idx, col0_bold=True)
    row.cells[0].width = Inches(2.2)
    row.cells[1].width = Inches(4.6)
    if "OPEN" in row_data[1] or "awaiting" in row_data[1]:
        set_bg(row.cells[1], "FFF2CC")

doc.add_paragraph()

# ── §10 Revision history ──────────────────────────────────────────────────────
add_heading(doc, "10.  REVISION HISTORY")
rev_hdrs = ["Rev", "Date", "Changes", "Author"]
rev_widths = [0.4, 1.0, 5.0, 0.9]
rev_tbl = make_table(doc, rev_hdrs, rev_widths)
add_row(rev_tbl,
    ("A", datetime.date.today().strftime("%Y-%m-%d"),
     "Initial release. Covers CAT-A (injection moulding + co-moulding), "
     "CAT-B (CFRP shell), CAT-C (PDMS bonding). "
     "SUP-M-07 and SUP-B-01 added as BLOCKING items for RISK-20 mitigation. "
     "§9 tracks RISK-20 mitigation status.",
     "Claude/ME"),
    rev_widths, 0)

OUT_PATH = "docs/neuropulse_supplier_selection_checklist.docx"
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
