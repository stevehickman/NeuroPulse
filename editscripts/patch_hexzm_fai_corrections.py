"""
Correct the First Article Inspection checklist (docs/neurone_fai_zone_module.docx,
NP-FAI-ZM-001) for the hex-tile/socket architecture (NP-HEX-ZM-001).

FAI-A01..A08 (§4a): CMM verification of EEG electrode position was written
against 5 fixed "Zone 1..5" module identities, each carrying one named 10-20
site (Fp1/F3, Fp2/F4, C3, C4, P3/P4) — undercounting the real T1 montage
(8-9 sites: Fp1/2, F3/4, C3/4, P3/4, +Oz) and asserting a socket-to-10-20
registration that NP-HEX-ZM-001 §7 REG-1 still lists as OPEN, not final.
Rewritten to verify CMM position of T1-B (EEG-carrying) modules at their
socket positions generically, explicitly gated on REG-1 rather than
asserting specific final socket numbers that don't exist yet.

FAI-A09..A14: the 5-layer zone-keying verification items (added by
patch_fai_risk15.py) test a mechanism that no longer exists — see the
RISK-15 correction in patch_hexzm_risk_corrections.py. Replaced with FAI
items for what actually needs verifying now: the universal orientation key
and the UID-based placement-check software gate.

FAI-R04: bend-radius physical check for "all 5 zones" — same open gap as
RISK-11/RISK-17 (patch_hexzm_risk_corrections.py): downgraded to NEEDS
REVIEW rather than left asserting a check against non-existent zone slots.

FAI-SY01/SY02: zone-illumination and zone-selective-deactivation system
tests referenced "Zone 1..5" by number. Rewritten against the real named
zones in protocols/predefined/00-zones.npps (the actual source of truth for
zone names), not invented zone numbers.

§9 cross-reference table: RISK-15 row description/status updated to match.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FAI_PATH = "docs/neurone_fai_zone_module.docx"

MARKER = "hex-tile/socket architecture (NP-HEX-ZM-001)"


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def find_table_with_row(doc, row_id):
    for tbl in doc.tables:
        for row in tbl.rows:
            if row.cells[0].text.strip() == row_id:
                return tbl, row
    return None, None


def already_patched(doc):
    return any(MARKER in p.text for tbl in doc.tables for row in tbl.rows for c in row.cells for p in c.paragraphs)


def set_row(row, item, test, method, accept):
    set_cell_text(row.cells[0], item, bold=True)
    set_cell_text(row.cells[1], test)
    set_cell_text(row.cells[2], method)
    set_cell_text(row.cells[3], accept)


def main():
    doc = Document(FAI_PATH)
    if already_patched(doc):
        print(f"  Already patched — skipping: {FAI_PATH}")
        return

    # ── FAI-A01..A05: EEG-site CMM checks — generalize + gate on REG-1 ──
    a01_tbl, _ = find_table_with_row(doc, "FAI-A01")
    for rid, site in [("FAI-A01", "Fp1/Fp2 (single shared socket — Fp row holds one tile, REG-1 open item)"),
                       ("FAI-A02", "F3/F4"),
                       ("FAI-A03", "C3/C4"),
                       ("FAI-A04", "P3/P4")]:
        _, row = find_table_with_row(doc, rid)
        if row:
            set_row(row, rid,
                f"[BLOCKING, PENDING REG-1] T1-B (EEG-carrying) module CMM measurement at the "
                f"socket assigned to 10-20 site {site}",
                "CMM / optical profilometer",
                "LED/electrode centroid within ±0.3 mm of the socket's nominal 10-20 target "
                "position — socket assignment per docs/np_hex_zm_001.md §3.2/§4a, PROVISIONAL "
                "pending gate REG-1")
            print(f"  {rid} corrected")
    # FAI-A05 was "Zone 5 (parietal, P3/P4)" — now redundant with A04 (both P3/P4 under old
    # scheme); repurpose as Oz (the montage site the old 5-zone scheme never covered).
    _, row = find_table_with_row(doc, "FAI-A05")
    if row:
        set_row(row, "FAI-A05",
            "[BLOCKING, PENDING REG-1] T1-B module CMM measurement at the socket assigned to Oz "
            "(required for the photoparoxysmal visual-stim safety interlock, docs/np_hex_zm_001.md §4a)",
            "CMM / optical profilometer",
            "Electrode centroid within ±0.3 mm of the Oz socket's nominal target position — "
            "PROVISIONAL pending gate REG-1")
        print("  FAI-A05 corrected (repointed to Oz — was a duplicate P3/P4 check)")

    _, row = find_table_with_row(doc, "FAI-A06")
    if row:
        set_row(row, "FAI-A06",
            "Inter-electrode spacing: confirm T1-B socket positions match 10-20 system "
            "inter-electrode distances",
            "CMM / derived from A01-A05 measurements",
            "All inter-electrode distances within ±0.5 mm of 10-20 system nominal — PROVISIONAL "
            "pending gate REG-1")
        print("  FAI-A06 corrected")

    for rid in ("FAI-A07", "FAI-A08"):
        _, row = find_table_with_row(doc, rid)
        if row:
            old_test = row.cells[1].text
            new_test = old_test.replace("Zone module", "Hex-tile module").replace("zone module", "hex-tile module").replace("each zone", "each sampled socket").replace("per zone", "per sampled socket")
            set_cell_text(row.cells[1], new_test)
            print(f"  {rid} generalized to hex-tile module")

    # ── FAI-A09..A14: retired keying system → orientation key + software gate ──
    tbl, ref_row = find_table_with_row(doc, "FAI-A09")
    if tbl and ref_row:
        # Replace A09/A10/A11/A13/A14 content in place; A12 becomes a placement-check test.
        replacements = {
            "FAI-A09": (
                "Universal orientation key verification: attempt to insert a hex-tile module "
                "rotated 60°/120°/180°/240°/300° from correct orientation into a socket",
                "Physical manual test — 5 wrong-orientation attempts per sampled socket",
                "Zero incorrect-orientation insertions possible at any of the 5 rotations. Key "
                "geometry alone (no vision, no reading, no firmware) must stop the attempt before "
                "electrical contact — verifies the blind-safe claim in docs/np_hex_zm_001.md §4a.",
            ),
            "FAI-A10": (
                "Correct-orientation insertion: confirm a module in the correct orientation seats "
                "fully and makes electrical contact on the first attempt",
                "Physical manual test, sampled sockets",
                "Full seating and contact confirmed on first attempt at every sampled socket. No "
                "false resistance at correct orientation.",
            ),
            "FAI-A11": (
                "UID auto-inventory verification: insert a known module type at a sampled socket; "
                "confirm np_module_map reports the correct type via UID-based auto-inventory "
                "(no ZONE_ID resistor — that mechanism is retired)",
                "Hub MCU diagnostic mode / app inventory view",
                "Reported module type at the socket matches the physical module inserted, within "
                "the poll interval specified in docs/np_hex_zm_001.md §4",
            ),
            "FAI-A13": (
                "[BLOCKING] Placement-check software gate: configure a protocol whose module map "
                "requires an EEG-capable module at a given socket; leave that socket unpopulated "
                "or populate it with a non-EEG module; confirm the app blocks the protocol and "
                "identifies the unmet socket",
                "Live app + firmware test — np_module_map_check_placement()",
                "Protocol is disabled with the specific unmet socket identified to the user. "
                "Protocol becomes runnable immediately once a matching module is inserted at that "
                "socket — no other action required.",
            ),
            "FAI-A14": (
                "Mixed-type insertion: populate a representative cluster with a mix of T1-A/B/C "
                "module types in arbitrary sockets; confirm all types insert, seat, and are "
                "correctly identified regardless of socket position",
                "Physical + firmware test",
                "Every type-agnostic combination inserts and is correctly identified. No socket "
                "position is type-restricted, per docs/np_hex_zm_001.md §4a \"any type inserts "
                "into any socket.\"",
            ),
        }
        for rid, (test, method, accept) in replacements.items():
            _, row = find_table_with_row(doc, rid)
            if row:
                set_row(row, rid, test, method, accept)
                print(f"  {rid} replaced (retired keying test → hex-tile equivalent)")

        # FAI-A12 previously the [BLOCKING] wrong-zone ENABLE test — retarget to the
        # placement-check gate's negative case (distinct from A13's positive/unmet case).
        _, row = find_table_with_row(doc, "FAI-A12")
        if row:
            set_row(row, "FAI-A12",
                "[BLOCKING] Type-mismatch rejection: attempt to run a protocol whose module map "
                "excludes the module type actually inserted at a required socket (e.g. T1-A where "
                "T1-B is required); confirm the app blocks the protocol before session start",
                "Live app + firmware test — np_module_map_check_placement()",
                "Session blocked before any stimulation/PBM enable. Unmet socket + required type "
                "reported to the user via the app.")
            print("  FAI-A12 replaced (wrong-zone ENABLE test → type-mismatch placement-check test)")

    # ── FAI-R04: bend radius, same open gap as RISK-11/17 ───────────────
    _, row = find_table_with_row(doc, "FAI-R04")
    if row:
        set_cell_text(row.cells[1],
            "NEEDS REVIEW — was scoped to the retired 5-zone-slot FPC routing architecture "
            "(RISK-11/RISK-17). The hex-tile/socket architecture (NP-HEX-ZM-001) per-socket "
            "interconnect this test would run against is not yet designed (see MECH-1/MECH-2/"
            "SMART-1 in docs/np_hex_zm_001.md §7). Cannot be executed as written.")
        set_cell_text(row.cells[3], "NEEDS RE-SCOPE once hex-tile interconnect design exists",
            bold=True, color=(0xBF, 0x60, 0x00))
        set_cell_bg(row.cells[3], "FFE699")
        print("  FAI-R04 flagged NEEDS REVIEW")

    # ── FAI-SY01/SY02: use real named zones from 00-zones.npps ──────────
    _, row = find_table_with_row(doc, "FAI-SY01")
    if row:
        set_row(row, "FAI-SY01",
            "All populated sockets illuminate at 25% duty cycle when the \"All\" zone "
            "(protocols/predefined/00-zones.npps) is activated — app confirms socket-level "
            "active status, photodiode returns nonzero dose reading per populated socket",
            "App UI, firmware telemetry",
            "Every populated socket reports active; J/cm² > 0 for each")
        print("  FAI-SY01 corrected to real 'All' zone")
    _, row = find_table_with_row(doc, "FAI-SY02")
    if row:
        set_row(row, "FAI-SY02",
            "Zone-selective deactivation: activate the \"Frontal Left\" zone only "
            "(protocols/predefined/00-zones.npps); confirm its member sockets are active and "
            "all sockets outside it remain inactive",
            "App UI",
            "Correct zone-membership isolation confirmed against the 00-zones.npps socket list")
        print("  FAI-SY02 corrected to real named zone")

    # ── §9 cross-reference table ─────────────────────────────────────────
    _, row = find_table_with_row(doc, "RISK-15")
    if row and len(row.cells) == 4:
        set_cell_text(row.cells[1], "Zone keying — eliminated by hex-tile/socket architecture (NP-HEX-ZM-001); universal orientation key + software placement-check", size=8)
        set_cell_text(row.cells[2], "FAI-A09, FAI-A10, FAI-A11, FAI-A12, FAI-A13, FAI-A14 (all corrected)", size=8)
        set_cell_text(row.cells[3], "MITIGATED (this FAI, corrected)", bold=True, color=(0x00, 0x70, 0x00), size=8)
        set_cell_bg(row.cells[3], "E2EFDA")
        print("  §9 RISK-15 cross-reference row corrected")

    doc.save(FAI_PATH)
    print(f"  FAI checklist saved: {FAI_PATH}")


if __name__ == "__main__":
    main()
