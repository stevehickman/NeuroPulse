"""
Generate NeurOne clinical trials strategy document and additional modalities document.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_col_width(cell, inches):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, size=9, bold=False, italic=False, space_before=0, space_after=4, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, size=9, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(size)
        run = p.add_run(text)
        run.font.size = Pt(size)
    else:
        run = p.add_run(text)
        run.font.size = Pt(size)
    return p

def add_table(doc, headers, rows, col_widths=None, header_bg="1F3864", alt_bg="EEF2F7"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_row = table.rows[0]
    for ci, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        set_bg(cell, header_bg)
        if col_widths:
            set_col_width(cell, col_widths[ci])
        para = cell.paragraphs[0]
        run = para.add_run(hdr)
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.add_row()
        bg = alt_bg if ri % 2 == 1 else "FFFFFF"
        for ci, (cell, txt) in enumerate(zip(row.cells, row_data)):
            set_bg(cell, bg)
            if col_widths:
                set_col_width(cell, col_widths[ci])
            para = cell.paragraphs[0]
            run = para.add_run(str(txt))
            run.font.size = Pt(8)
    return table

def add_divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: Clinical Trials Strategy
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# ── Title block ───────────────────────────────────────────────────────────────
t = doc.add_table(rows=1, cols=1)
set_bg(t.rows[0].cells[0], "1F3864")
p = t.rows[0].cells[0].paragraphs[0]
r = p.add_run("NeurOne — Clinical Evidence & Trials Strategy")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
p2 = t.rows[0].cells[0].add_paragraph()
r2 = p2.add_run("NP-CLIN-001 Rev A  |  Confidential — Pre-Decisional  |  May 2026")
r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xBB,0xCC,0xFF)
p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(6)
doc.add_paragraph()

# ── §1 Commercial Rationale ───────────────────────────────────────────────────
add_heading(doc, "1. Commercial Rationale — The All-in-One Proposition", level=1, color=(0x1F,0x38,0x64))

add_para(doc,
    "NeurOne does not require multi-modal synergy claims to succeed commercially. "
    "Each of the 8 T1 modalities (11 in T2) carries independent clinical evidence. "
    "The primary value proposition is straightforward:", size=9)

add_bullet(doc,
    "One device replaces up to 8 separate devices (PBM headset, neurofeedback system, "
    "tDCS/tACS stimulator, VNS device, audio entrainment player, visual stimulation "
    "goggles, intranasal probe, HRV monitor). Cumulative retail cost of equivalent "
    "individual devices: $8,000–$25,000+. NeurOne Home Standard: $849.",
    bold_prefix="Lower cost: ")
add_bullet(doc,
    "One wearable platform, one charging cable, one app, one subscription. "
    "No device management overhead for clinicians or patients.",
    bold_prefix="Less clutter: ")
add_bullet(doc,
    "Users and clinicians choose which modalities to use. Someone familiar only with "
    "tDCS uses only tDCS. A neurofeedback practitioner adds PBM on top without "
    "replacing their workflow. No multi-modal protocol required.",
    bold_prefix="Modality choice: ")
add_bullet(doc,
    "If simultaneous multi-modal use produces synergistic effects — and early evidence "
    "suggests it does for some combinations — NeurOne is the only consumer platform "
    "positioned to demonstrate and deliver those effects at scale.",
    bold_prefix="Multi-modal upside: ")

doc.add_paragraph()
add_heading(doc, "1.1 Why Vielight's 35+ RCTs Partially Support NeurOne", level=2)

add_para(doc,
    "Vielight has published or supported more than 35 clinical trials of transcranial "
    "and intranasal PBM. These trials are directly relevant to NeurOne — but with "
    "important distinctions that determine how they can be used in marketing and "
    "regulatory contexts.", size=9)

add_para(doc, "WHAT TRANSFERS:", size=9, bold=True)
add_bullet(doc,
    "The Vielight trials establish that transcranial + intranasal PBM at 810nm is safe, "
    "feasible for home use, and produces measurable neurological effects across cognition, "
    "PTSD/TBI, and sleep outcomes. NeurOne's bibliography cites these trials to "
    "establish that PBM is a legitimate, evidence-backed modality.")
add_bullet(doc,
    "For FDA general wellness positioning and FTC claims substantiation, the Vielight "
    "body of evidence is relevant prior art. NeurOne can cite it as background "
    "scientific context supporting the modality category.")
add_bullet(doc,
    "The Vielight Neuro EEG study (2025) — proving that itPBM reorganises neural "
    "oscillations and functional network connectivity — directly supports NeurOne's "
    "combined PBM + EEG neurofeedback architecture.")

add_para(doc, "WHAT DOES NOT TRANSFER:", size=9, bold=True)
add_bullet(doc,
    "Device equivalence: Vielight RCTs used Vielight devices at specific irradiance, "
    "duty cycle, and LED count. NeurOne (600 LEDs, dual-wavelength, 400 mW/cm² "
    "pulsed, 5 independently addressable zones) has no published comparator. Any "
    "marketing claim implying NeurOne produces outcomes 'at least as good as Vielight' "
    "requires head-to-head data or regulatory counsel opinion.")
add_bullet(doc,
    "Multi-modal interaction: No published trial tests PBM combined with BES, tDCS, "
    "VNS, audio entrainment, and visual stimulation simultaneously. Vielight trials "
    "are single-modality (PBM only). If NeurOne makes outcome claims implying "
    "benefit from the combined protocol, those require NeurOne-specific data.")
add_bullet(doc,
    "400 mW/cm² irradiance claim: RISK-03 (regulatory opinion pending) — this specific "
    "claim cannot appear in any public material until regulatory counsel clears it, "
    "regardless of what Vielight claims for their own devices.")

add_para(doc,
    "PRACTICAL RULE: Cite Vielight trials freely to establish that 'PBM is clinically "
    "studied and evidence-supported.' Do not cite them to claim that NeurOne produces "
    "equivalent or superior outcomes — that requires NeurOne-specific data or "
    "regulatory opinion.", size=9, italic=True, color=(0x1F,0x38,0x64))

doc.add_paragraph()
add_divider(doc)

# ── §2 Individual Modality Evidence ───────────────────────────────────────────
add_heading(doc, "2. Individual Modality Clinical Evidence", level=1, color=(0x1F,0x38,0x64))

modality_data = [
    ("T1",
     "Transcranial PBM\n660nm + 810nm",
     "Moderate-Strong",
     "MCI/Dementia, TBI/PTSD, Depression, Healthy cognition",
     "Rashidi-Ranjbar 2025 (n=15, MCI, multimodal neuroimaging, AJGPsych); "
     "Brazilian BDNF RCT 2024 (n=93, MoCA +3.20 vs +1.97, BDNF↑, p=0.03); "
     "Naeser 2023 (n=4, CTE/PTSD, 2 SD improvement); "
     "Chao/UCSF 2019 (n=8 dementia, DMN connectivity improved); "
     "Yao et al. Science Advances 2022 (1064nm, working memory capacity)",
     "Rashidi-Ranjbar (Toronto); Naeser (VA Boston); Hamblin (MGH); "
     "Jeffery (UCL, retinal)"),
    ("T1",
     "Intranasal PBM\n660nm + 810nm",
     "Moderate\n(always combined with transcranial in trials)",
     "MCI/Dementia, Healthy cognition",
     "Japanese RCT Protocol 2024 (Frontiers Neurology, 12-week sham-controlled, "
     "intranasal + transcranial, MCI/mild AD); "
     "Vielight EEG study 2025 (proves PBM reorganises neural oscillations — "
     "supports PBM+neurofeedback combination); "
     "Chun et al. 2025 (Korean RCT, MCI/mild AD)",
     "Lim (Vielight); Rashidi-Ranjbar; Hamblin"),
    ("T1",
     "EEG Neurofeedback",
     "Strong (ADHD)\nModerate (PTSD, anxiety)\nEmerging (depression)",
     "ADHD, PTSD, Anxiety, Depression, Cognitive enhancement",
     "Wu et al. Brain & Behavior 2024 (13 RCTs, n=1,370 children, TBR/SMR/SCP/HEG); "
     "Meta-analysis (21 RCTs, n=1,261, ADHD); "
     "Frontiers Psychiatry 2024 (PTSD moderate-to-large effect); "
     "MDPI Signals 2024 (EEG-adaptive closed-loop binaural beats, n=24, "
     "96% reached target vs 0% sham)",
     "Arns (Brainclinics NL); Strehl (Tübingen); Stinear (Auckland)"),
    ("T1",
     "tACS / BES\n(Brainwave Entrainment Stimulation)",
     "Moderate\n(contested by pub. bias analyses)",
     "Working memory, Depression, MCI/AD (40Hz)",
     "STM meta-analysis 2023 (102 studies, n=2,893, modest-to-moderate effect); "
     "Frohlich tACS-MDD RCT 2022 (n=255, remote, rapid improvement wk1, p=0.022); "
     "Frontiers Psychiatry meta-analysis 2023 (4 RCTs, n=224 MDD); "
     "Gamma-tACS review Transl Neurodegeneration 2024 (40Hz modifies AD progression)",
     "Frohlich (UNC Chapel Hill, depression); Reinhart (BU, WM); "
     "Herrmann (Oldenburg)"),
    ("T1",
     "tDCS\n(Cortical Priming Stimulation)",
     "Strong\n(depression, large RCTs)",
     "Depression (esp. MDD), MCI/AD cognition, Motor rehabilitation",
     "Brunoni et al. Nature Medicine 2024 (n=174 home tDCS, 44.9% remission vs 21.8% sham); "
     "Jog/UCLA 2025 (n=71, personalized HD-tDCS, significant mood improvement + "
     "gray matter changes, PMC12426800); "
     "JCM meta-analysis 2025 (14 RCTs, tDCS for AD cognition); "
     "DepressionDC Lancet 2023 (n=150, null result — add-on SSRIs)",
     "Jog (UCLA); Brunoni (São Paulo); Bikson (CCNY, safety/dosimetry)"),
    ("T1",
     "Auricular VNS + HRV\n(taVNS)",
     "Strong\n(insomnia, d=1.2)\nModerate\n(epilepsy, post-stroke)",
     "Insomnia, Epilepsy, Post-stroke depression, Working memory",
     "JAMA Network Open 2024 (n=72, insomnia RCT, PSQI −8.2 vs −3.9, d=1.2, "
     "sustained 20 weeks); "
     "Yang et al. (n=150, drug-resistant epilepsy, seizure↓); "
     "Pan et al. CNS N&T 2024 (n=28, working memory, EEG network stabilisation); "
     "Scoping review 2024 (109 studies, 21 clinical populations)",
     "Multiple groups; rapidly growing field"),
    ("T1",
     "Neural Audio Entrainment\n(binaural beats + bone conduction)",
     "Moderate\n(anxiety, sleep)\nWeak\n(cognition — entrainment inconsistent)",
     "Anxiety, Stress, Sleep, Cognitive enhancement (closed-loop)",
     "PLOS ONE systematic review 2023 (22 studies, 36% confirmed entrainment, "
     "medium effect cognition/anxiety/pain); "
     "MDPI Signals 2024 (first EEG-adaptive binaural beats RCT, n=24, 96% target); "
     "8-RCT meta-analysis (acoustic stimulation improves insomnia severity); "
     "Stress RCT (salivary cortisol + alpha-amylase, significant in 2 well-designed studies)",
     "Emerging; no dominant research group"),
    ("T1",
     "Visual Stimulation\n(40Hz gamma / EMDR / photic driving)",
     "Very Strong (40Hz AD)\nStrong (EMDR/PTSD)\nLimited (general photic driving)",
     "Alzheimer's disease, PTSD, Cognitive entrainment",
     "Cognito/OVERTURE Phase 2 2024 (n=76, mild-to-moderate AD, 76% ↓cognitive decline, "
     "69% ↓brain atrophy, 77% ↓functional decline); "
     "Chan et al. Phase 2A (n=15 mild AD, ↓ventricular dilation, ↑DMN, ↑memory); "
     "HOPE Phase 3 (n=670, enrollment complete July 2025, results mid-2026); "
     ">30 EMDR RCTs; WHO-recommended for PTSD",
     "Tsai (MIT Picower, GENUS; SAB candidate); Cognito Therapeutics (Phase 3)"),
    ("T2",
     "TMS\n(Transcranial Magnetic Stimulation)",
     "Very Strong\n(FDA-cleared, CPT codes)",
     "Treatment-resistant depression, MCI, OCD (FDA-cleared)",
     "George/Neuronetics pivotal RCT 2008 (n=301, FDA clearance); "
     "Meta-analysis JAMA Psychiatry 2017 (Cohen d=0.40); "
     "BRIGhTMIND Nature Medicine 2024 (n=255 TRD, connectivity-guided iTBS); "
     "George/MUSC 2023 (first EEG-synchronized rTMS RCT, alpha phase-locked stimulation); "
     "Pooled data: 46.81% response, 28.25% remission, Cohen d=1.4",
     "George (MUSC); Williams (Stanford, accelerated iTBS); "
     "Blumberger (CAMH, theta burst)"),
    ("T2",
     "1170nm Deep PBM\n(Laser diodes)",
     "Preclinical only\n(no human RCT at 1170nm)",
     "Subcortical / deep structure (theoretical)",
     "Theoretical depth advantage (35–40mm): no published human RCT at 1170nm "
     "specifically. 1064nm cognition data (Yao/Science Advances 2022) is closest. "
     "First-mover opportunity for NeurOne-sponsored human pilot.",
     "Hamblin (tPBM review author); Yao/UT Dallas (1064nm cognition)"),
]

add_table(doc,
    ["Tier", "Modality", "Evidence Strength", "Primary Indications", "Key Trials (n, result)", "Key Researchers"],
    modality_data,
    col_widths=[0.4, 0.95, 0.75, 1.0, 2.85, 1.25],
    header_bg="1F3864")

doc.add_paragraph()
add_divider(doc)

# ── §3 Multi-Modal Combination Evidence ───────────────────────────────────────
add_heading(doc, "3. Multi-Modal Combination Evidence", level=1, color=(0x1F,0x38,0x64))

add_para(doc,
    "The central research gap: no published RCT has tested PBM + BES + VNS + audio + "
    "visual simultaneously in adults. However, several two-modality combinations have "
    "published evidence, and the trajectory of the field is clearly toward combination "
    "protocols. NeurOne is the only consumer platform positioned to run these studies "
    "at scale.", size=9)
doc.add_paragraph()

# Best-evidenced combinations
add_heading(doc, "3.1 tDCS + taVNS — Simultaneous Synergy Demonstrated", level=2)
add_para(doc,
    "This is the strongest published multi-modal evidence directly applicable to NeurOne.",
    size=9, bold=True, color=(0x1F,0x38,0x64))
add_bullet(doc,
    "Aghajani et al., Brain Stimulation 2021 (PMID 33621676) — First simultaneous "
    "tDCS + taVNS fMRI study. Synergistic activation in bilateral thalamus, pallidum, "
    "parahippocampal gyrus, dorsal raphe, substantia nigra, periaqueductal gray. "
    "Combined effect > sum of individual effects.",
    bold_prefix="Synergy confirmed: ")
add_bullet(doc,
    "Frontiers in Neuroscience 2022 (PMC9344917) — First RCT, simultaneous tDCS + taVNS, "
    "working memory in healthy adults. Three-arm: tDCS alone / taVNS alone / combined. "
    "Combined > either alone; more consistent improvements in 3-back spatial accuracy. "
    "Effect depended on task difficulty, not stimulation site — suggests an attentional "
    "or arousal mechanism rather than purely local cortical effect.",
    bold_prefix="RCT: ")
add_bullet(doc,
    "Post-stroke cognitive impairment RCT (PMC11015246, 2024) — 3-arm protocol (tDCS alone "
    "/ taVNS alone / combined) recruiting at First Affiliated Hospital Yangtze University "
    "(n=66 target). Primary: MoCA. Results expected 2025.",
    bold_prefix="Active RCT: ")
add_bullet(doc,
    "NeurOne delivers tDCS (Cortical Priming Stimulation) and taVNS (auricular VNS) "
    "simultaneously from a single platform. This is the only published synergistic "
    "combination NeurOne can claim today, with RCT support.",
    bold_prefix="NeurOne relevance: ")
doc.add_paragraph()

add_heading(doc, "3.2 40Hz Audiovisual (GENUS) — Best-Evidenced Multi-Modal Protocol", level=2)
add_bullet(doc,
    "Cognito Therapeutics Phase 2A (n=15 mild AD, Chan et al. 2022/2025): combined 40Hz "
    "visual flicker + 40Hz audio click trains daily. Outcomes: ↓ventricular dilation, "
    "↓hippocampal atrophy, ↑DMN connectivity, ↑face-name delayed recall.",
    bold_prefix="Phase 2A: ")
add_bullet(doc,
    "OVERTURE Phase 2 (n=76, Frontiers Neurology 2024): 76% reduction in cognitive "
    "decline, 77% slowing of functional decline, 69% reduction in whole-brain atrophy. "
    "OLE (18 months): ~9 months time saved in cognitive function.",
    bold_prefix="Phase 2 (OVERTURE): ")
add_bullet(doc,
    "HOPE Phase 3 (n=670, enrollment complete July 2025, primary results expected mid-2026). "
    "Largest non-pharmacological AD device pivotal trial ever.",
    bold_prefix="Phase 3 (HOPE): ")
add_bullet(doc,
    "NeurOne delivers 40Hz audio (binaural beats + bone conduction) AND 40Hz visual "
    "(108 LEDs/lens, 6 zones/eye) simultaneously. This directly replicates the Cognito "
    "GENUS audiovisual protocol. NeurOne adds EEG-adaptive closed-loop verification "
    "that Cognito does not include — a meaningful architectural advantage.",
    bold_prefix="NeurOne relevance: ")
add_bullet(doc,
    "Tsai lab (MIT, 2023): 40Hz tactile vibration (wristband) also reduces AD pathology "
    "in mice. Multi-sensory (visual + audio + tactile simultaneously) is the next "
    "preclinical direction. NeurOne does not currently support 40Hz tactile vibration "
    "(see modalities expansion document).",
    bold_prefix="Emerging: ")
doc.add_paragraph()

add_heading(doc, "3.3 PBM + qEEG Neurofeedback + HRV Biofeedback — Published RCT (2025)", level=2)
add_bullet(doc,
    "ResearchGate January 2025: 'Efficiency of a Multimodal Neurotechnological Intervention "
    "(Cerebral Photobiomodulation, qEEG Neurofeedback, and Heart-Brain Biofeedback) in "
    "Enhancing Attention and Academic Performance in Middle School Students.' Nationally "
    "conducted RCT. Three modalities simultaneously — directly relevant to NeurOne's "
    "PBM + EEG + VNS/HRV stack. Population is children; adult replication needed.",
    bold_prefix="First published three-modality RCT: ")
add_bullet(doc,
    "Vielight 2025 publication (high-level journal): itPBM reorganises neural oscillations "
    "and functional network connectivity. Confirms that PBM directly changes the EEG "
    "signals that neurofeedback trains — making PBM + EEG neurofeedback mechanistically "
    "coherent, not merely additive.",
    bold_prefix="Mechanistic foundation: ")
add_bullet(doc,
    "PubMed 31647776, 2019: 'Treatment of Neurodegeneration: Integrating Photobiomodulation "
    "and Neurofeedback in Alzheimer's Dementia and Parkinson's' — review establishing "
    "the theoretical framework. PBM builds better neural infrastructure; neurofeedback "
    "trains connectivity within it.",
    bold_prefix="Review framework: ")
doc.add_paragraph()

add_heading(doc, "3.4 tDCS + EEG Neurofeedback — Emerging Evidence", level=2)
add_bullet(doc,
    "Frontiers in Neuroscience 2023: 'tDCS + Motor Imagery Neurofeedback — RCT.' "
    "10 min tDCS before NFB session. NFB+tDCS group showed improved motor imagery "
    "vs NFB alone. Priming cortical excitability enhances subsequent neurofeedback "
    "learning window.",
    bold_prefix="Published RCT: ")
add_bullet(doc,
    "ScienceDirect meta-analysis 2022: simultaneous tDCS + working memory training "
    "produces significant transfer effects vs training alone.",
    bold_prefix="Meta-analysis: ")
add_bullet(doc,
    "JMIR Research Protocols 2022: double-blind RCT (tDCS + NFB for declarative memory) "
    "enrolled July 2022–2023. Results expected.",
    bold_prefix="Active RCT: ")
add_bullet(doc,
    "Mechanism: tDCS primes neurons to be more responsive during the neurofeedback training "
    "window; timing (tDCS before or concurrent with NFB) is the key variable. "
    "NeurOne can sequence or layer these modalities under firmware control.",
    bold_prefix="Mechanism: ")
doc.add_paragraph()

add_heading(doc, "3.5 PBM + tDCS — Research Gap; High-Priority Opportunity", level=2)
add_bullet(doc,
    "JCM 2025 systematic review (17 RCTs, 2020–2024): tDCS and PBM each independently "
    "improve cognitive abilities in AD. Authors explicitly call for standardisation to "
    "enable future combination trials. No published RCT of PBM + tDCS together exists.",
    bold_prefix="Current status: ")
add_bullet(doc,
    "PBM increases mitochondrial ATP and BDNF. tDCS modulates cortical excitability via "
    "LTP/LTD-like neuroplasticity. These are complementary, non-overlapping mechanisms. "
    "Better-energized neurons may respond more strongly to tDCS-induced plasticity.",
    bold_prefix="Mechanistic rationale: ")
add_bullet(doc,
    "NeurOne is uniquely positioned to run the first PBM + tDCS combination RCT for "
    "MCI/AD cognition. Rashidi-Ranjbar (Toronto) and Jog (UCLA) are natural co-PIs.",
    bold_prefix="Research opportunity: ")
doc.add_paragraph()

# Summary table
add_heading(doc, "3.6 Multi-Modal Combination Evidence Summary", level=2)

combo_rows = [
    ("tDCS + taVNS (simultaneous)",
     "Synergy demonstrated\n(fMRI + WM RCT)",
     "Brain Stimulation 2021; Frontiers Neuroscience 2022; Post-stroke RCT enrolling 2024",
     "High — cite today"),
    ("40Hz visual + audio (GENUS)",
     "Strong\n(Phase 2 RCT + Phase 3 enrolling)",
     "Cognito/Chan 2022; OVERTURE 2024 (n=76); HOPE Phase 3 n=670",
     "High — NeurOne replicates this protocol"),
    ("PBM + qEEG NFB + HRV biofeedback",
     "Moderate\n(One 2025 RCT; children)",
     "ResearchGate 2025 (nationally conducted RCT)",
     "Moderate — adult replication needed"),
    ("tDCS + motor imagery neurofeedback",
     "Emerging\n(Published RCT 2023)",
     "Frontiers Neuroscience 2023; JMIR protocol 2022",
     "Moderate — motor rehabilitation context"),
    ("tDCS + cognitive training (WM)",
     "Moderate\n(Meta-analysis)",
     "ScienceDirect meta-analysis 2022",
     "Moderate"),
    ("PBM + tDCS",
     "Gap — no RCT\n(review only)",
     "JCM 2025 review; mechanistic rationale strong",
     "High research priority — NeurOne first-mover"),
    ("PBM + tACS",
     "Gap — no data",
     "Hypothesis only",
     "Research priority"),
    ("PBM + taVNS",
     "Gap — no data",
     "Hypothesis only",
     "Research priority"),
    ("tACS + EEG neurofeedback",
     "Framework only\n(Bikson closed-loop)",
     "bioRxiv 2023; tACS EEG personalisation paper 2025",
     "Research priority — EEG personalises tACS"),
    ("taVNS + EEG neurofeedback",
     "Implicit overlap\n(EEG signal change)",
     "Pan et al. 2024 (taVNS changes EEG during WM)",
     "Research priority"),
]

add_table(doc,
    ["Combination", "Evidence Level", "Best Citation(s)", "NeurOne Relevance"],
    combo_rows,
    col_widths=[1.5, 1.1, 2.6, 1.45],
    header_bg="1F3864")

doc.add_paragraph()
add_divider(doc)

# ── §4 Researchers ────────────────────────────────────────────────────────────
add_heading(doc, "4. Key Researchers by Modality", level=1, color=(0x1F,0x38,0x64))

researcher_rows = [
    ("Neda Rashidi-Ranjbar",
     "St. Michael's Hospital / Unity Health Toronto",
     "neda.rashidi-ranjbar@unityhealth.to",
     "tPBM (intranasal + transcranial), MCI/AD, multimodal neuroimaging",
     "PI of 2025 pilot RCT (n=15, AJGP). CIHR pathway. Most current PBM RCT PI in Canada. "
     "Early-career — motivated collaborator. Natural first contact.",
     "TRIAL PI", "1"),
    ("Mayank Jog",
     "UCLA Brain Mapping Center",
     "mjog@mednet.ucla.edu",
     "HD-tDCS, personalized stimulation, depression, structural neuroplasticity",
     "K99/R00 NIH MH128572 active. PI of n=71 personalized HD-tDCS RCT (PMC12426800, 2025). "
     "Gray matter plasticity data (Sci Reports 2023). JAMA Network Open 2025 RCT PI.",
     "TRIAL PI", "2"),
    ("Mark George",
     "MUSC Brain Stimulation Lab",
     "georgem@musc.edu\n843-876-5142",
     "TMS (all forms), EEG-synchronized TMS, accelerated rTMS, MCI",
     "55 active studies, 15 TMS machines. Pioneer: helped achieve FDA TMS clearance 2008. "
     "First-in-human EEG-synchronized rTMS RCT (Brain Stimulation 2023) — directly aligned "
     "with NeurOne EEG-adaptive architecture. MUSC MCI accelerated TMS trial 2023.",
     "TRIAL PI (T2)", "3"),
    ("Margaret Naeser",
     "VA Boston / BU Neurology",
     "mnaeser@bu.edu",
     "tPBM (TBI, PTSD, CTE, stroke, aphasia), veteran cohort",
     "47 years VA-funded. 2023 publication: n=4 ex-football players (CTE/PTSD), "
     "significant improvement PTSD/depression/pain/sleep. Virtual care TBI/PBM trial active. "
     "Existing veteran patient cohort = ready recruitment pool.",
     "TRIAL PI", "4"),
    ("Li-Huei Tsai",
     "MIT Picower Institute",
     "617-324-0305",
     "40Hz GENUS (visual + audio + tactile), Alzheimer's, multi-sensory gamma",
     "Founder of GENUS field (Nature 2016). Director MIT Aging Brain Initiative. "
     "SAB role ONLY — Cognito Therapeutics conflict prevents PI role. "
     "Expanding 40Hz multi-sensory to tactile vibration. "
     "Key scientific credentialing for NeurOne visual + audio modalities.",
     "SAB (highest priority)", "5"),
    ("Glen Jeffery",
     "UCL Institute of Ophthalmology",
     "g.jeffery@ucl.ac.uk",
     "Retinal PBM (670nm), age-related retinal decline, AMD",
     "World's leading retinal PBM researcher. Shinhmar 2021: single 3-min 670nm "
     "= 20% cone-mediated color vision improvement in older adults. "
     "AMD pilot 2020 (n=42). Natural SAB candidate for NeurOne Mode F "
     "(invisible NIR retinal PBM) scientific credentialing.",
     "SAB", "6"),
    ("Marom Bikson",
     "CCNY Neural Engineering Group",
     "neuralengr.org/bikson",
     "tDCS (safety, dosimetry, HD-tDCS), closed-loop neuromodulation architecture",
     "2023 bioRxiv: scalable closed-loop neuromodulation framework (DL + EEG), "
     "modality-agnostic (tACS/tDCS/tFUS/TMS). tDCS safety data cited by FDA. "
     "Technical advisory role for NeurOne's closed-loop architecture.",
     "TECHNICAL ADVISOR", "7"),
    ("Robert Reinhart",
     "Boston University",
     "reinhart@bu.edu",
     "tACS (theta/gamma), working memory, depression, aging",
     "Corresponding author: STM 2023 tACS meta-analysis (102 studies, n=2,893). "
     "Phase-specific tACS for working memory and depression.",
     "TRIAL PI", "8"),
    ("Flavio Frohlich",
     "UNC Chapel Hill",
     "frohlichlab.org",
     "tACS (alpha oscillations), depression, remote clinical trials",
     "Remote tACS RCT (n=255 MDD, JCP 2022). Phase-specific alpha tACS. "
     "Infrastructure for large remote tACS trials.",
     "TRIAL PI", "9"),
]

add_table(doc,
    ["Researcher", "Institution", "Contact", "Focus Area(s)",
     "Why NeurOne", "Role", "Priority"],
    researcher_rows,
    col_widths=[0.95, 1.1, 1.0, 1.05, 2.1, 0.85, 0.55],
    header_bg="1F3864")

doc.add_paragraph()
add_divider(doc)

# ── §5 Multi-Modal Research Proposals ─────────────────────────────────────────
add_heading(doc, "5. Researcher Proposals and Speculation on Multi-Modal Synergy", level=1, color=(0x1F,0x38,0x64))

add_para(doc,
    "The following represents published proposals, theoretical frameworks, and expert "
    "speculation about combining modalities — not yet tested in clinical trials.", size=9)
doc.add_paragraph()

add_heading(doc, "5.1 Closed-Loop Architecture as the Unifying Framework", level=2)
add_para(doc,
    "Bikson et al. (bioRxiv 2023) propose a deep-learning EEG-adaptive closed-loop "
    "framework agnostic to stimulation modality. The system reads EEG state in real time "
    "and adjusts stimulation parameters (frequency, phase, amplitude) across any modality. "
    "The key insight: the stimulation modality (tACS vs tDCS vs TMS) is less important "
    "than the real-time EEG guidance that ensures stimulation is delivered at the "
    "neurologically optimal moment.", size=9)
add_para(doc,
    "NeurOne's autonomous closed-loop EEG architecture is the consumer implementation "
    "of exactly this framework. The value is not that any single modality is better — "
    "it is that all modalities can be guided by the same real-time neural state.", size=9, italic=True)
doc.add_paragraph()

add_heading(doc, "5.2 PBM as a Neuroplasticity Primer for Other Modalities", level=2)
add_para(doc,
    "Proposed mechanism (PubMed 31647776, 2019; JCM 2025 review): PBM increases "
    "mitochondrial ATP production, BDNF, and synaptic density. This creates 'better "
    "neural infrastructure' that makes other plasticity-inducing modalities (tDCS, "
    "tACS, neurofeedback) more effective. Analogous to exercise priming BDNF before "
    "a learning session.", size=9)
add_para(doc,
    "If true: PBM delivered first in a session sequence, followed by tDCS or neurofeedback, "
    "would produce greater neuroplastic change than either modality alone. This is testable "
    "with NeurOne's sequencing firmware capabilities. No human RCT exists yet — "
    "NeurOne could commission the first.", size=9, italic=True)
doc.add_paragraph()

add_heading(doc, "5.3 taVNS as a Neuromodulatory Primer (Noradrenergic Enhancement)", level=2)
add_para(doc,
    "Proposed mechanism: taVNS activates the locus coeruleus (noradrenaline) and basal "
    "forebrain (acetylcholine), both of which modulate cortical plasticity and memory "
    "consolidation. Running taVNS concurrently with tDCS, PBM, or neurofeedback could "
    "enhance plasticity via neuromodulator gating — the same mechanism underlying the "
    "tDCS + taVNS synergy demonstrated in the 2021 Brain Stimulation fMRI study.", size=9)
add_para(doc,
    "The confirmed tDCS + taVNS synergy (thalamus, parahippocampal gyrus) provides "
    "proof-of-concept that taVNS amplifies effects of concurrent stimulation modalities. "
    "Extending this to PBM + taVNS or neurofeedback + taVNS is a logical next step "
    "that only NeurOne can test at consumer scale.", size=9, italic=True)
doc.add_paragraph()

add_heading(doc, "5.4 Multi-Sensory 40Hz — From Two Senses to Three", level=2)
add_para(doc,
    "Tsai lab (MIT, 2023): 40Hz whole-body vibrotactile stimulation reduces amyloid and "
    "tau in mouse models. The Tsai group is exploring whether combining visual, auditory, "
    "and tactile 40Hz stimulation simultaneously produces stronger effects than any "
    "two-sense combination. HOPE Phase 3 (n=670) tests visual + auditory; tactile is "
    "the next generation protocol.", size=9)
add_para(doc,
    "NeurOne currently delivers 40Hz visual and auditory simultaneously (replicating "
    "GENUS). If the HOPE Phase 3 results are positive, and if Tsai's tactile data "
    "translates to humans, adding a 40Hz vibrotactile component to a future NeurOne "
    "hardware revision would position NeurOne at the forefront of GENUS-class "
    "therapy. See additional modalities document.", size=9, italic=True)
doc.add_paragraph()

add_heading(doc, "5.5 EEG-Synchronized TMS (George, MUSC 2023) — Closed-Loop Precision", level=2)
add_para(doc,
    "George et al. (Brain Stimulation 2023): first-in-human RCT of EEG phase-locked rTMS. "
    "When rTMS pulses are synchronised to the patient's prefrontal alpha phase: "
    "(a) strong phase entrainment develops in DLPFC, (b) anterior cingulate cortex is "
    "more reliably engaged, (c) inter-individual variability in response is reduced. "
    "The EEG phase at the moment of stimulation — not the average EEG — determines outcome.", size=9)
add_para(doc,
    "NeurOne's EEG-adaptive architecture can implement this for all non-TMS modalities "
    "in T1 (tDCS, tACS, PBM sequencing) and for TMS in T2. This is the evidence base "
    "for NeurOne's core architectural claim: EEG-guided delivery is better than "
    "fixed-schedule delivery.", size=9, italic=True)
doc.add_paragraph()
add_divider(doc)

# ── §6 NeurOne Trial Priority ───────────────────────────────────────────────
add_heading(doc, "6. Recommended Clinical Trial Priority", level=1, color=(0x1F,0x38,0x64))

trial_rows = [
    ("1",
     "SBIR Phase I\nPBM for MCI/Cognition",
     "Rashidi-Ranjbar (Toronto)\nor Naeser (VA Boston)",
     "NeurOne tPBM vs sham. n=20–30, 12 weeks. "
     "Replicate Vielight design with NeurOne hardware. "
     "Generates first NeurOne-specific PBM data; enables comparative claims.",
     "$150K–$300K | 12–18 months | SBIR Phase I funding vehicle",
     "NOW"),
    ("2",
     "tDCS for Depression\n(Home-Based Replication)",
     "Jog (UCLA) or Brunoni group",
     "NeurOne tDCS (CPS) vs sham, MDD, home-based. "
     "n=40–60. Replicate Brunoni Nature Medicine 2024 design. "
     "Establishes NeurOne tDCS evidence; leverages existing protocol.",
     "$200K–$400K | 12–18 months",
     "YEAR 1"),
    ("3",
     "tDCS + taVNS Synergy\nfor Cognitive Impairment",
     "Any MCI research centre;\nUniversity collaboration",
     "Three-arm: NeurOne tDCS alone / taVNS alone / combined. "
     "n=45–60. Replicate Frontiers Neuroscience 2022 design in MCI population. "
     "First NeurOne-specific multi-modal synergy data.",
     "$250K–$400K | 18 months",
     "YEAR 1–2"),
    ("4",
     "PBM + tDCS Combination\nfor MCI (First-in-Class)",
     "Rashidi-Ranjbar + Jog\n(multi-site)",
     "Four-arm: PBM alone / tDCS alone / combined / sham. "
     "n=60–80. First human RCT of PBM + tDCS simultaneously. "
     "Genuine research first; high publication impact. "
     "JCM 2025 review explicitly calls for this trial.",
     "$400K–$600K | 18–24 months | R21 mechanism",
     "YEAR 2"),
    ("5",
     "EEG-Adaptive 40Hz AV\n(GENUS Replication + EEG)",
     "Tsai lab (SAB, not PI)\nor Rashidi-Ranjbar",
     "NeurOne 40Hz visual + audio (GENUS protocol) + EEG-adaptive "
     "closed-loop vs Cognito fixed-schedule. Mild AD, n=30–50. "
     "Tests whether EEG guidance improves on fixed GENUS delivery.",
     "$350K–$500K | 18 months",
     "YEAR 2"),
    ("6",
     "T2 — EEG-Synchronized\nTMS for TRD",
     "George (MUSC)\nTarget: SBIR Phase II",
     "EEG phase-locked rTMS via NeurOne T2 platform vs standard rTMS. "
     "Extends George's 2023 first-in-human RCT with larger n and "
     "NeurOne-specific hardware. Generates T2 TMS data.",
     "$800K–$2M | 24–36 months | SBIR Phase II",
     "YEAR 3"),
]

add_table(doc,
    ["Priority", "Trial", "Lead Researcher", "Design", "Budget / Timeline", "Start"],
    trial_rows,
    col_widths=[0.4, 1.1, 0.95, 2.35, 1.2, 0.65],
    header_bg="1F3864")

doc.add_paragraph()
add_divider(doc)

# ── §7 Key Bibliography ────────────────────────────────────────────────────────
add_heading(doc, "7. Key Bibliography", level=1, color=(0x1F,0x38,0x64))
add_para(doc, "Cite-ready format. See neurone_bibliography.docx for full 33-entry bibliography.", size=8, italic=True)
doc.add_paragraph()

bibs = [
    ("tPBM — MCI RCT",
     "Rashidi-Ranjbar N et al. 'tPBM for MCI.' Am J Geriatric Psychiatry 2025; "
     "preprint medRxiv 2025.08.19.25333989. (n=15, multimodal neuroimaging)"),
    ("tPBM — BDNF RCT",
     "Brazilian group. 'tPBM increases cognition and BDNF in adults over 50.' "
     "Photobiomodulation Photomedicine Laser Surg 2024. (n=93, MoCA Δ3.20 vs 1.97, p=0.03)"),
    ("tDCS — Depression (large)",
     "Brunoni AR et al. 'Home tDCS for MDD.' Nature Medicine 2024. "
     "(n=174, 44.9% remission vs 21.8% sham)"),
    ("tDCS — Depression (null)",
     "Brunoni AR et al. 'DepressionDC tDCS + SSRIs.' The Lancet 2023. (n=150, null result)"),
    ("tDCS — Personalized HD",
     "Jog MA et al. 'Personalized HD-tDCS for depression.' PMC12426800, 2025. "
     "(n=71, significant, gray matter changes)"),
    ("40Hz AV — Phase 2A",
     "Chan D et al. 'GENUS 40Hz audiovisual, mild AD.' Alzheimer's & Dementia 2025. "
     "(n=15, ↓atrophy, ↑DMN)"),
    ("40Hz AV — Phase 2 (OVERTURE)",
     "Cognito Therapeutics. 'OVERTURE.' Frontiers Neurology 2024. "
     "(n=76, 76% ↓cognitive decline, 69% ↓brain atrophy)"),
    ("taVNS — Insomnia",
     "JAMA Network Open 2024. 'taVNS for chronic insomnia.' (n=72, d=1.2, PSQI −8.2 vs −3.9)"),
    ("taVNS — Working memory",
     "Pan et al. 'taVNS + working memory in epilepsy.' CNS Neurosci Ther 2024. "
     "(n=28, 20 weeks, EEG network stabilisation)"),
    ("tDCS+taVNS — fMRI synergy",
     "Aghajani H et al. Brain Stimulation 2021. PMID 33621676. "
     "(fMRI synergy: thalamus, pallidum, parahippocampal gyrus)"),
    ("tDCS+taVNS — WM RCT",
     "Frontiers Neuroscience 2022 (PMC9344917). 'Simultaneous tDCS + taVNS on working "
     "memory.' RCT, combined > either alone."),
    ("PBM+NFB — Review",
     "PubMed 31647776, 2019. 'Integrating PBM and Neurofeedback in AD and Parkinson's.'"),
    ("PBM+NFB+HRV — RCT",
     "ResearchGate January 2025. 'Multimodal PBM + qEEG NF + HRV biofeedback for "
     "attention.' Nationally conducted RCT (middle schoolers)."),
    ("EEG NF — ADHD meta-analysis",
     "Wu et al. Brain & Behavior 2024. 13 RCTs, n=1,370. TBR/SMR/SCP/HEG."),
    ("tACS — Meta-analysis",
     "Science Translational Medicine 2023. 102 studies, n=2,893, modest-to-moderate "
     "cognitive improvement."),
    ("TMS — EEG-synchronized",
     "George MK et al. Brain Stimulation 2023. (PMC10872322). "
     "First EEG phase-locked rTMS RCT. Alpha-synchronized = better ACC engagement."),
    ("TMS — Connectivity-guided iTBS",
     "BRIGhTMIND. Sackeim HA et al. Nature Medicine 2024. (n=255 TRD)"),
    ("EMDR",
     "Systematic review PMC11111257, 2024. >30 RCTs. WHO-recommended for PTSD."),
    ("Closed-loop framework",
     "Bikson M et al. bioRxiv 2023. 'Scalable closed-loop neuromodulation with deep "
     "learning.' Modality-agnostic EEG-adaptive architecture."),
    ("tPBM + tDCS review",
     "JCM 2025. 'Efficacy of tDCS and PBM in AD.' 17 RCTs 2020–2024. Calls for "
     "combination trial."),
]

for category, citation in bibs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"[{category}] ")
    r1.bold = True; r1.font.size = Pt(8)
    r2 = p.add_run(citation)
    r2.font.size = Pt(8)

doc.save("docs/neurone_clinical_trials_strategy.docx")
print("NP-CLIN-001 saved: docs/neurone_clinical_trials_strategy.docx")


# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: Additional Modalities to Consider
# ══════════════════════════════════════════════════════════════════════════════

doc2 = Document()
for section in doc2.sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Title
t2 = doc2.add_table(rows=1, cols=1)
set_bg(t2.rows[0].cells[0], "1F3864")
p = t2.rows[0].cells[0].paragraphs[0]
r = p.add_run("NeurOne — Additional Modalities to Consider")
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
p2 = t2.rows[0].cells[0].add_paragraph()
r2 = p2.add_run("NP-MOD-EXT-001 Rev A  |  Confidential — Pre-Decisional  |  May 2026")
r2.font.size = Pt(9); r2.font.color.rgb = RGBColor(0xBB,0xCC,0xFF)
p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(6)
doc2.add_paragraph()

add_para(doc2,
    "This document lists modalities identified in multi-modal clinical trials and "
    "leading research that NeurOne does not currently support. For each, we note "
    "the evidence, the research context in which it was found, the hardware requirements, "
    "and a priority assessment for potential future inclusion.", size=9)
add_para(doc2,
    "Current NeurOne T1 stack: transcranial PBM (660+810nm), intranasal PBM, "
    "EEG neurofeedback (8-ch), BES/tACS, tDCS, auricular VNS + HRV, neural audio "
    "entrainment, visual stimulation (40Hz gamma + EMDR). T2 adds TMS, 1170nm deep PBM, "
    "21-ch qEEG, clinical tACS.", size=9, italic=True)
doc2.add_paragraph()
add_divider(doc2)

# Each modality
modalities_ext = [
    {
        "number": "1",
        "name": "40Hz Vibrotactile / Somatosensory Stimulation",
        "source": "Found in: Tsai lab (MIT) multi-sensory GENUS research; "
                  "MIT News 2023 (mouse models); Tsai PLOS Biology 2025 review",
        "evidence":
            "PRECLINICAL STRONG: 40Hz whole-body vibrotactile stimulation reduces amyloid "
            "plaques and tau tangles in mouse models (MIT News 2023). Tsai lab has shown that "
            "tactile 40Hz is the third sensory modality in the GENUS multi-sensory protocol "
            "(visual + audio + tactile). In mouse models, multi-sensory gamma is more effective "
            "than any single-sense delivery. No published human RCT at 40Hz tactile yet.\n\n"
            "The Cognito HOPE Phase 3 trial (n=670, results mid-2026) tests visual + audio only. "
            "Tsai is exploring tactile as the next generation. If HOPE is positive, tactile "
            "could be the next major addition to GENUS-class therapy.",
        "hardware":
            "A wristband or wearable pad delivering 40Hz vibration. Low engineering complexity "
            "relative to most NeurOne modalities. Mass-market fitness wearables already "
            "contain vibration motors; precision 40Hz control is trivial. BOM estimate: "
            "+$5–10/headset for a wrist module or temple-pad vibrotactile actuator.",
        "integration":
            "Could be a snap-on zone module accessory (mastoid or temporal position) or "
            "an accessory wristband paired via BT. Firmware: add 40Hz drive waveform to "
            "audio channel (same as bone conduction but haptic). "
            "No new MCU required.",
        "priority": "HIGH — if HOPE Phase 3 (mid-2026) is positive, 40Hz tactile becomes "
                    "the logical Rev B NeurOne addition. Plan hardware provision now. "
                    "Tsai SAB engagement is the critical path.",
        "action": "Await HOPE results (mid-2026). If positive: engage Tsai for tactile "
                  "protocol. Design vibrotactile module as hardware option Rev B or C.",
        "bg": "FFF2CC",
    },
    {
        "number": "2",
        "name": "1064nm Transcranial PBM",
        "source": "Found in: Yao et al., Science Advances 2022 (UT Dallas / Harvard); "
                  "tPBM cognition literature",
        "evidence":
            "MODERATE: Yao et al. Science Advances 2022 demonstrated that 1064nm tPBM applied "
            "to right prefrontal cortex improved visual working memory capacity in healthy "
            "adults, with EEG confirmation (increased contralateral delay activity, CDA). "
            "1064nm achieves deeper cortical penetration than 810nm due to reduced "
            "oxyhaemoglobin absorption. The UT Dallas / Harvard group has published multiple "
            "1064nm cognition studies.\n\n"
            "NeurOne currently uses 660nm + 810nm (T1) and adds 1170nm (T2). "
            "1064nm sits between these: better penetration than 810nm, less aggressive "
            "than 1170nm laser. It is distinct from both current wavelengths.",
        "hardware":
            "1064nm LED arrays are commercially available but less efficient than 810nm. "
            "Alternatively, 1064nm diode lasers (similar to 1170nm T2 hardware). "
            "A third LED wavelength on existing FPC strips would require new FPC layout "
            "and zone module mould revision. Significant BOM and tooling cost.",
        "integration":
            "Moderate complexity: third wavelength requires new LED procurement, FPC "
            "redesign, firmware channel, and potentially regulatory re-classification "
            "(laser vs LED).",
        "priority": "MEDIUM — 1064nm is evidence-backed but the penetration advantage "
                    "over 810nm in humans is not yet definitively demonstrated. "
                    "T2's 1170nm laser already covers the 'deep PBM' indication. "
                    "Monitor UT Dallas publications before committing.",
        "action": "Monitor Yao/UT Dallas 1064nm literature (2026–2027). "
                  "If large RCT confirms 1064nm advantage over 810nm in humans: "
                  "evaluate as T1.5 or T2 wavelength addition.",
        "bg": "EEF2F7",
    },
    {
        "number": "3",
        "name": "Transcranial Focused Ultrasound (tFUS)",
        "source": "Found in: Bikson et al. bioRxiv 2023 closed-loop framework "
                  "(listed as a modality alongside tDCS/tACS/TMS)",
        "evidence":
            "EMERGING: tFUS uses low-intensity focused ultrasound to modulate neural "
            "activity non-invasively. Achieves millimetre spatial precision (vs cm for "
            "tDCS/tACS). Multiple pilot studies show transient neuromodulatory effects. "
            "No large RCTs in psychiatric indications yet. FDA IDE trials beginning "
            "in 2024–2025. Key advantage: subcortical targeting without TMS's power "
            "requirements or safety concerns.\n\n"
            "Bikson's closed-loop framework explicitly includes tFUS as a modality that "
            "could be guided by EEG in real-time — the same framework NeurOne's "
            "architecture implements.",
        "hardware":
            "HIGH COMPLEXITY: tFUS requires a focused transducer (piezoelectric element "
            "at specific frequency, typically 250kHz–2MHz), acoustic coupling medium, "
            "and MRI-based targeting for precision. Consumer-grade tFUS is not yet "
            "feasible at price points compatible with NeurOne. Emerging startup "
            "companies (Cognito, Attune, NaviFUS) are building clinical-grade systems. "
            "BOM estimate: +$200–500 minimum; acoustic coupling = consumable.",
        "integration":
            "Not feasible for T1 or T2 in current form. Would require a new hardware "
            "platform revision and potentially separate regulatory pathway (Class II/III). "
            "Consumer feasibility: 5–8 years.",
        "priority": "LOW (WATCH) — Monitor tFUS regulatory and commercialisation "
                    "trajectory. If consumer-grade tFUS becomes feasible by 2029–2030, "
                    "evaluate for NeurOne T3 or successor platform.",
        "action": "Add to technology watch list. No hardware action now.",
        "bg": "EEF2F7",
    },
    {
        "number": "4",
        "name": "HRV Biofeedback as a Standalone Therapeutic Protocol",
        "source": "Found in: ResearchGate 2025 RCT (PBM + qEEG NF + Heart-Brain HRV "
                  "Biofeedback simultaneously); growing biofeedback literature",
        "evidence":
            "MODERATE: NeurOne currently monitors HRV via the VNS auricular clip "
            "(PPG + auricular taVNS in same accessory). However, HRV biofeedback as a "
            "distinct therapeutic protocol — where the user sees real-time HRV and "
            "performs resonance frequency breathing to maximise HRV amplitude — has "
            "its own clinical evidence base:\n\n"
            "- Meta-analysis (Applied Psychophysiology & Biofeedback, 2021): HRV biofeedback "
            "reduces anxiety (d=0.83) and depression (d=0.65) across 24 studies.\n"
            "- PTSD: HRV biofeedback + EMDR combination RCTs exist.\n"
            "- Performance: HRV biofeedback used by military, athletes, and first "
            "responders for stress inoculation.\n\n"
            "The 2025 multi-modal RCT combined PBM + qEEG NF + HRV biofeedback as "
            "three coordinated modalities. NeurOne has all the hardware to do this "
            "today — HRV biofeedback is a SOFTWARE addition, not a hardware addition.",
        "hardware":
            "ZERO additional hardware required. The VNS clip already provides PPG "
            "for HRV monitoring. Resonance frequency biofeedback is a display + "
            "breathing guide in the app. The clip already delivers taVNS, which "
            "itself enhances HRV (demonstrated in multiple studies).",
        "integration":
            "Software only. Add resonance frequency breathing guide (4–6 breaths/min "
            "target), real-time coherence score, and session summary to the app. "
            "Coordinate with taVNS delivery (taVNS at peak HRV = optimal timing). "
            "This would allow NeurOne to replicate the 2025 multi-modal RCT protocol "
            "in-app at no additional BOM cost.",
        "priority": "HIGH — no hardware cost; extends clinical evidence claim "
                    "(NeurOne can say: 'We replicate the combined protocol from the "
                    "2025 RCT — PBM + qEEG neurofeedback + HRV biofeedback'). "
                    "Should be in firmware roadmap for Year 1.",
        "action": "Add HRV biofeedback mode to firmware and app roadmap. "
                  "Coordinate taVNS delivery timing with HRV coherence peak. "
                  "No hardware change. No tooling change.",
        "bg": "E2EFDA",
    },
    {
        "number": "5",
        "name": "Personalized MRI-Guided HD-tDCS Neuronavigation",
        "source": "Found in: Jog et al. UCLA, PMC12426800, 2025; "
                  "BRIGhTMIND Nature Medicine 2024 (connectivity-guided iTBS)",
        "evidence":
            "STRONG (clinical context): Jog's HD-tDCS personalised RCT (n=71) used "
            "individual MRI to target left DLPFC precisely based on each patient's "
            "functional connectivity. This produced significant effects that standard "
            "F3-positioned tDCS (fixed position) did not reliably achieve in other trials. "
            "BRIGhTMIND (Nature Medicine 2024, n=255 TRD) showed connectivity-guided "
            "iTBS was superior to standard F3 targeting.\n\n"
            "The implication: NeurOne's current tDCS delivers fixed electrode positions. "
            "MRI-personalised targeting would require either (a) integrating with the "
            "user's existing MRI data, or (b) using EEG-derived source localisation "
            "(sLORETA, already in T2 stack) as a surrogate for MRI targeting.",
        "hardware":
            "Option A (MRI integration): No new hardware; app accepts MRI DICOM data "
            "and computes optimal electrode placement (software only, but requires MRI). "
            "Not consumer-feasible for most T1 users.\n\n"
            "Option B (sLORETA-guided, T2): T2 already includes sLORETA source imaging. "
            "Personalised electrode placement guidance based on sLORETA source maps is "
            "achievable within T2 platform (software only).\n\n"
            "Option C (cloud AI): App analyses session EEG and recommends electrode "
            "placement adjustments based on fleet response data. No additional hardware.",
        "integration":
            "T1: No action needed. T2: Add sLORETA-guided electrode placement recommendation "
            "to T2 app. Requires firmware access to sLORETA maps; coordinate with "
            "HD-tDCS multi-electrode system.",
        "priority": "MEDIUM (T2) — sLORETA-guided tDCS targeting is a meaningful T2 "
                    "clinical differentiator. Not relevant to T1. Add to T2 firmware "
                    "and app specification.",
        "action": "Include sLORETA-derived tDCS placement guidance in T2 app specification. "
                  "Engage Jog (UCLA) for T2 clinical collaboration on personalised "
                  "HD-tDCS protocol.",
        "bg": "EEF2F7",
    },
    {
        "number": "6",
        "name": "Cervical / Neck-Applied taVNS (Non-Invasive Cervical VNS)",
        "source": "Found in: Scoping review 2024 (109 studies); "
                  "electroCore K163334 510(k) predicate (T2 regulatory)",
        "evidence":
            "MODERATE: Transcutaneous cervical VNS (tcVNS) applies stimulation to the "
            "skin overlying the carotid sheath, stimulating the cervical vagus nerve "
            "directly rather than the auricular branch. electroCore's gammaCore "
            "(K163334, K173323) is FDA-cleared for cluster headache and migraine. "
            "Higher activation of vagus trunk vs auricular branch; used in clinical "
            "indications where auricular stimulation is insufficient.\n\n"
            "NeurOne T2 CLAUDE.md notes: '+ cervical option' under VNS. This "
            "suggests cervical VNS is already considered but not formally specified.",
        "hardware":
            "Separate neck-worn accessory with gel electrodes. Handheld gammaCore-style "
            "device or integrated neck band. More complex than auricular clip: "
            "skin-electrode interface, cervical positioning, safety (carotid proximity). "
            "FDA: would require separate 510(k) (gammaCore K163334 as predicate). "
            "BOM: +$30–60 for the accessory; regulatory pathway adds 12–18 months.",
        "integration":
            "T2 accessory only. New cable/connector from hub. Safety MCU oversight "
            "required (near carotid = cardiac risk if current path is incorrect). "
            "Would significantly strengthen T2 VNS clinical claim vs auricular-only.",
        "priority": "MEDIUM (T2) — auricular VNS is the correct T1 approach "
                    "(safe, simple, strong evidence). Cervical VNS is a meaningful T2 "
                    "clinical enhancement for patients where auricular stimulation is "
                    "insufficient. Add to T2 specification.",
        "action": "Add cervical VNS accessory specification to T2 hardware roadmap. "
                  "Reference gammaCore K163334 as regulatory predicate.",
        "bg": "EEF2F7",
    },
]

for m in modalities_ext:
    # Header row
    t_m = doc2.add_table(rows=1, cols=1)
    set_bg(t_m.rows[0].cells[0], "1F3864")
    p = t_m.rows[0].cells[0].paragraphs[0]
    r = p.add_run(f"  {m['number']}. {m['name']}")
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    doc2.add_paragraph()

    add_para(doc2, f"Source context: {m['source']}", size=8, italic=True, space_after=2)

    fields = [
        ("Evidence", m["evidence"]),
        ("Hardware requirements", m["hardware"]),
        ("Integration path", m["integration"]),
        ("Priority", m["priority"]),
        ("Recommended action", m["action"]),
    ]
    for label, text in fields:
        p = doc2.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"{label}: ")
        r1.bold = True; r1.font.size = Pt(9)
        r2 = p.add_run(text)
        r2.font.size = Pt(9)

    doc2.add_paragraph()
    add_divider(doc2)

doc2.save("docs/neurone_additional_modalities.docx")
print("NP-MOD-EXT-001 saved: docs/neurone_additional_modalities.docx")
print("\nDone.")
