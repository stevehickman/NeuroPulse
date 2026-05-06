"""Patch RISK-12 in risk register and update FPC spec §11.4."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# ── Risk register ──
RISK_PATH = "docs/neuropulse_fpc_zone_module_risks_revA.docx"
risk_doc = Document(RISK_PATH)

tbl = risk_doc.tables[0]
for row in tbl.rows:
    if row.cells[0].text.strip() == "RISK-12":
        set_cell_text(row.cells[1], "MEDIUM\n(MITIGATED)", bold=True, color=(0x00, 0x70, 0x00))
        set_cell_bg(row.cells[1], "E2EFDA")
        set_cell_text(row.cells[5],
            "MITIGATED: First Article Inspection checklist NP-FAI-ZM-001 Rev A created. "
            "FAI-A01 to FAI-A08 (Section 4a) mandate CMM measurement of all 5 zone module "
            "positions to ±0.3 mm tolerance (±0.2 mm margin against ±0.5 mm system requirement). "
            "FAI-A08 verifies re-insertion repeatability (±0.1 mm). "
            "FAI also requires datum features (reference holes/edges) on FPC stiffener (FAI-A07). "
            "Zone module alignment is re-verified after 50 connector insertion cycles (FAI-K14). "
            "FPC spec §11.4 references NP-FAI-ZM-001 as mandatory pre-production gate.")
        set_cell_text(row.cells[8],
            "MITIGATED — NP-FAI-ZM-001 Rev A created; CMM alignment check (FAI-A01–A08) "
            "required before production release",
            bold=True, color=(0x00, 0x70, 0x00))
        set_cell_bg(row.cells[8], "92D050")
        print("RISK-12 updated to MITIGATED")
        break

# Also update OI-10 in open issues table
open_issues_tbl = risk_doc.tables[1]
for row in open_issues_tbl.rows:
    if row.cells[0].text.strip() == "OI-10":
        set_cell_text(row.cells[1], "REQUIRED (addressed in NP-FAI-ZM-001)")
        set_cell_bg(row.cells[1], "E2EFDA")
        set_cell_text(row.cells[2],
            "Mechanical key geometry for all 5 zone slots — "
            "unique asymmetric key per zone still required in shell tooling spec. "
            "NP-FAI-ZM-001 §4a verifies alignment via CMM; mechanical key prevents "
            "wrong-zone insertion (colour-blind accessibility and incorrect alignment both mitigated). "
            "Key geometry must be in shell tooling spec before first cut.")
        print("OI-10 updated")
        break

risk_doc.save(RISK_PATH)
print(f"Risk register saved")

# ── FPC spec §11.4 — add FAI reference ──
SPEC_PATH = "docs/neuropulse_fpc_zone_module_spec_revA.docx"
spec = Document(SPEC_PATH)

for i, para in enumerate(spec.paragraphs):
    if "11.4" in para.text and "Alignment" in para.text:
        s114_idx = i
        break

# Find the §12 paragraph (next major section) to insert before it
s12_idx = None
for i, para in enumerate(spec.paragraphs):
    if para.text.strip().startswith("12."):
        s12_idx = i
        break

print(f"§11.4 at {s114_idx}, §12 at {s12_idx}")

if s12_idx:
    target = spec.paragraphs[s12_idx]

    # Insert a note referencing the FAI doc after the existing §11.4 content
    CONTENT = [
        ("[DRAWING REQUIREMENT] First Article Inspection: All 5 zone module positions must be "
         "verified by CMM measurement before production release. CMM verification protocol is "
         "defined in NP-FAI-ZM-001 Rev A (Zone Module FPC Assembly — FAI Checklist) §4a, "
         "items FAI-A01 to FAI-A08. The FAI checklist is the mandatory pre-production gate "
         "for zone module alignment.",
         True, (0xC0, 0x00, 0x00)),
        ("Datum features: the zone module stiffener shall include reference holes or edges "
         "as datum features for CMM fixturing. Datum feature positions must be defined on "
         "the FPC fabrication drawing before tooling release.",
         False, None),
        ("Re-insertion repeatability requirement: ±0.1 mm (module must self-locate to the "
         "same position on each insertion). This is verified by FAI-A08 (3 repeat insertions "
         "per zone, CMM re-measured each time).",
         False, None),
    ]

    for (text, bold, color) in reversed(CONTENT):
        new_p = OxmlElement("w:p")
        target._p.addprevious(new_p)
        for p_obj in spec.paragraphs:
            if p_obj._p is new_p:
                run = p_obj.add_run(text)
                run.bold = bold
                run.font.size = Pt(10)
                if color:
                    run.font.color.rgb = RGBColor(*color)
                break

    print("§11.4 FAI reference inserted")

spec.save(SPEC_PATH)
print(f"FPC spec saved")
print("Done.")
