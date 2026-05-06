"""Scan RISK-11 detail and FPC spec §11.2 content."""
from docx import Document

print("=== Risk Register — RISK-11 detail ===")
risks = Document("docs/neuropulse_fpc_zone_module_risks_revA.docx")
tbl = risks.tables[0]
for row in tbl.rows:
    if row.cells[0].text.strip() == "RISK-11":
        headers = ["ID","Severity","Title","Root Cause","Impact","Required Mitigation","Owner","Spec §","Status"]
        for h, cell in zip(headers, row.cells):
            print(f"  {h}: {cell.text.strip()!r}")
        break

print()
# Detailed section
for i, para in enumerate(risks.paragraphs):
    if "RISK-11" in para.text:
        print(f"  [{i}] {para.text[:120]!r}")
        # Print next 10 paragraphs
        for j in range(i+1, min(i+10, len(risks.paragraphs))):
            p = risks.paragraphs[j]
            if p.text.strip():
                print(f"  [{j}] {p.text[:120]!r}")
            if j > i+1 and ("RISK-12" in p.text or "OPEN ISSUES" in p.text):
                break

print()
print("=== FPC Spec — §11 Mechanical ===")
spec = Document("docs/neuropulse_fpc_zone_module_spec_revA.docx")
in_s11 = False
for i, para in enumerate(spec.paragraphs):
    t = para.text.strip()
    if "11." in t and ("Mechanical" in t or "Routing" in t or "Bend" in t or "11.1" in t or "11.2" in t):
        in_s11 = True
    if in_s11:
        if t:
            print(f"  [{i}] {t[:120]!r}")
        if "12." in t and i > 10:
            break
