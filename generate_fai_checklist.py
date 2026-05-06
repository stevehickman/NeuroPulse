"""
Generate NP-FAI-ZM-001 Rev A:
  Zone Module FPC Assembly — First Article Inspection (FAI) Checklist
Covers RISK-12 (alignment) and consolidates all inspection items from previous
risk mitigations: RISK-01, 02, 04, 05, 06, 07, 08, 10, 11, 12.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DATE = "2026-05-06"
DOC_NUM = "NP-FAI-ZM-001"
REV = "Rev A"

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def body(doc, text, space_after=4):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    if p.runs:
        p.runs[0].font.size = Pt(10)
    return p

def warn_box(doc, label, text, color_bg="FFF2CC", color_label=(0xBF, 0x60, 0x00)):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, color_bg)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(*color_label)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph()

def fai_table(doc, section_title, risk_refs, rows, col_widths):
    """
    FAI checklist table.
    rows: list of (item_id, description, method, accept_criterion, result_field, sign_field)
    """
    # Section title row + optional risk ref
    if section_title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(section_title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        if risk_refs:
            r2 = p.add_run(f"  [{risk_refs}]")
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    headers = ["Item", "Inspection / Test", "Method", "Accept Criterion", "Result", "Sign"]
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"

    # Header
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hdr.cells[i], "1F497D")
        run = hdr.cells[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8)

    for r_idx, row_data in enumerate(rows):
        item_id, desc, method, criterion, result, sign = row_data
        row = t.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
        data = [item_id, desc, method, criterion, result, sign]
        for c_idx, txt in enumerate(data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(txt))
            run.font.size = Pt(8)
            if c_idx == 0:
                run.bold = True

    for i, w in enumerate(col_widths):
        for row in t.rows:
            row.cells[i].width = Inches(w)

    doc.add_paragraph()
    return t

# ─────────────────────────────────────────────────────────────────────────────
# Document
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()
section = doc.sections[0]
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.85)

# ── Cover ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("NeuroPulse")
r.bold = True; r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Zone Module FPC Assembly\nFirst Article Inspection (FAI) Checklist")
r2.bold = True; r2.font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(
    f"{DOC_NUM}  {REV}  |  {DATE}  |  Pre-Production Draft\n"
    "References: NP-HW-FPC-001 Rev C · NP-PROC-FPC-001 Rev A · NP-DRV-SHELL-001 Rev A\n"
    "Risk mitigations consolidated: RISK-01, 02, 04, 05, 06, 07, 08, 10, 11, 12"
)
r3.font.size = Pt(9)
r3.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
doc.add_paragraph()

# ── Header fields table ──
t_hdr = doc.add_table(rows=4, cols=4)
t_hdr.style = "Table Grid"
hdr_fields = [
    ("Build lot / serial:", "", "FAI date:", ""),
    ("Headset serial number:", "", "Inspector name:", ""),
    ("FPC lot number:", "", "Reviewed by:", ""),
    ("Shell revision:", "", "Final result:", "  PASS  /  FAIL"),
]
for r_idx, (a, b, c, d) in enumerate(hdr_fields):
    row = t_hdr.rows[r_idx]
    for c_idx, txt in enumerate([a, b, c, d]):
        cell = row.cells[c_idx]
        set_cell_bg(cell, "F2F2F2" if c_idx % 2 == 0 else "FFFFFF")
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        if c_idx % 2 == 0:
            run.bold = True
for row in t_hdr.rows:
    row.cells[0].width = Inches(1.6)
    row.cells[1].width = Inches(1.6)
    row.cells[2].width = Inches(1.6)
    row.cells[3].width = Inches(1.6)
doc.add_paragraph()

warn_box(doc, "SCOPE:",
    "This checklist covers the first article of each zone module FPC assembly. "
    "ALL items must pass before production release. Any FAIL triggers a disposition "
    "review — no waiver without Engineering and Quality sign-off. "
    "Items marked [BLOCKING] must pass before any other FAI items proceed.")

# ── 1. Pre-Inspection Document Check ──
heading(doc, "SECTION 1 — PRE-INSPECTION DOCUMENT CHECK")
body(doc, "Verify all required documents and BOM revisions are current before physical inspection.")

fai_table(doc, None, None, [
    ("FAI-D01", "FPC specification NP-HW-FPC-001 Rev C on file", "Document check", "Rev C or later confirmed", "", ""),
    ("FAI-D02", "Procurement requirements NP-PROC-FPC-001 Rev A on file", "Document check", "Rev A or later confirmed", "", ""),
    ("FAI-D03", "Shell tooling design review NP-DRV-SHELL-001 Rev A fully signed off (all 8 signatories)", "Document check", "All sign-off boxes in §8 completed", "", ""),
    ("FAI-D04", "BOM confirms BCR421W (not BCR421U) — verify against approved BOM revision", "BOM review", "BCR421W in all LED driver positions", "", ""),
    ("FAI-D05", "BOM confirms Hirose FH34S-20S-0.5SH(50) — no Molex SlimStack or 0.35 mm pitch connectors", "BOM review", "FH34S or JAE FF03 only in connector positions", "", ""),
    ("FAI-D06", "FPC fab drawing specifies RA copper (IPC-4204/11 Type I) — check drawing fab notes", "Drawing review", "'ROLLED ANNEALED' or 'RA per IPC-4204/11' visible as fab note 1", "", ""),
    ("FAI-D07", "RA copper material certificate accompanies FPC lot", "Certificate check", "Cert references lot number and IPC-4204/11 Type I", "", ""),
    ("FAI-D08", "LED Vf bin distribution certificate accompanies LED lot (660 nm and 808–830 nm)", "Certificate check", "Cert present, lot mean ± range within ±0.10 V spec", "", ""),
], [0.65, 2.2, 1.3, 1.5, 0.7, 0.45])

# ── 2. Incoming Component Inspection ──
heading(doc, "SECTION 2 — INCOMING COMPONENT INSPECTION")
body(doc, "Verify purchased components before assembly. Failures here prevent assembly from proceeding.")

fai_table(doc, "2a. LED Emitters (660 nm and 808–830 nm)", "RISK-08 Vf binning · RISK-10 RA copper ref", [
    ("FAI-L01", "660 nm LED: measure Vf on 10-unit sample at 150 mA / 10 ms pulse / 25°C", "Benchtop measurement", "All 10 within ±0.10 V of lot mean on certificate", "", ""),
    ("FAI-L02", "808 nm LED: measure Vf on 10-unit sample at 150 mA / 10 ms pulse / 25°C", "Benchtop measurement", "All 10 within ±0.10 V of lot mean on certificate", "", ""),
    ("FAI-L03", "660 nm peak wavelength: spectrometer check on 5-unit sample", "Spectrometer", "λ_peak 655–670 nm, within ±5 nm of lot certificate", "", ""),
    ("FAI-L04", "808 nm peak wavelength: spectrometer check on 5-unit sample", "Spectrometer", "λ_peak 805–835 nm, within ±10 nm of lot certificate", "", ""),
    ("FAI-L05", "Reel label part number and date code check", "Visual", "Part number matches BOM; floor life not expired (MSL 3)", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

fai_table(doc, "2b. BCR421W LED Driver IC", "RISK-07 variant confusion", [
    ("FAI-C01", "[BLOCKING] Part marking: confirm 'W' suffix on every sampled IC", "Visual / magnifier (10×)", "Marking reads BCR421W. Any BCR421U → reject entire lot immediately", "", ""),
    ("FAI-C02", "Package type: SOT-323 (SC-70) 3-pin", "Visual / caliper", "SOT-323 body size 2.0 × 1.25 mm confirmed", "", ""),
    ("FAI-C03", "Authorised distributor CoC present", "Certificate check", "Distributor is Mouser, DigiKey, or Avnet; CoC on file", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

fai_table(doc, "2c. Hirose FH34S Connector", "RISK-01 connector family", [
    ("FAI-K01", "[BLOCKING] Part number on reel label: confirm FH34S-20S-0.5SH(50)", "Visual", "Exact part number confirmed. Any Molex or 0.35 mm connector → reject", "", ""),
    ("FAI-K02", "Position count: 20 contacts", "Visual / pin count", "Count = 20", "", ""),
    ("FAI-K03", "Lever actuation: back-flip lever opens and closes smoothly", "Manual test, 3 samples", "No stiffness, no binding, lever fully releases at 180°", "", ""),
    ("FAI-K04", "Contact resistance: insert FPC tail, measure all 20 contacts", "4-wire measurement", "All contacts < 50 mΩ after 5 insertion cycles", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

fai_table(doc, "2d. FPC Laminate", "RISK-10 RA copper · RISK-02 power trace", [
    ("FAI-F01", "[BLOCKING] RA copper material certificate: confirm 'Rolled Annealed' per IPC-4204/11 Type I", "Certificate check", "Certificate explicitly states RA. 'ED' certificate → reject lot", "", ""),
    ("FAI-F02", "Copper weight: 1 oz (35 µm ± 10%) confirmed on certificate", "Certificate check", "35 µm ± 3.5 µm", "", ""),
    ("FAI-F03", "Connector tail hard gold thickness: XRF measurement on 2 samples", "XRF (2 pads/unit)", "≥ 0.5 µm gold. Any pad < 0.4 µm → quarantine lot", "", ""),
    ("FAI-F04", "IPC class 2 certificate accompanies lot", "Certificate check", "Certificate references IPC-A-600 Class 2 or better", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

# ── 3. Assembled Zone Module Inspection ──
heading(doc, "SECTION 3 — ASSEMBLED ZONE MODULE INSPECTION")
body(doc, "Inspect first-article assembled zone module PCB before installation into shell.")

fai_table(doc, "3a. Power Trace and Thermal Verification", "RISK-02 current capacity", [
    ("FAI-T01", "Power trace width on VLED rail: measure from top layer photograph (calibrated image scale)", "Optical / caliper", "≥ 2.0 mm on all VLED/GND power bus segments", "", ""),
    ("FAI-T02", "Thermal imaging: run all 10 strings in one zone at full rated current (180 mA/string) for 60 s; capture IR image", "FLIR or equivalent IR camera", "No hotspot > 62°C junction temperature; uniform temperature gradient across strings", "", ""),
    ("FAI-T03", "Trace continuity on all 20 power pins: measure resistance pin-to-string LED", "4-wire measurement", "Per-trace resistance < 50 mΩ for any signal trace, < 10 mΩ for power rail", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

fai_table(doc, "3b. LED Driver Current Verification", "RISK-05 current hogging · RISK-07 BCR421W", [
    ("FAI-E01", "[BLOCKING] Per-string current: measure actual current in each of 10 strings for one wavelength at rated RSET", "Clamp meter / shunt, each string individually", "All 10 strings within ±5% of target (e.g. 180 mA ± 9 mA). Any string outside ±10% → investigate driver", "", ""),
    ("FAI-E02", "Across-wavelength match: compare 660 nm string mean vs 808 nm string mean current", "As above, both wavelengths", "Means within ±3% of each other at same RSET (same driver part)", "", ""),
    ("FAI-E03", "BCR421W thermal rise: measure IC case temperature at rated current for 60 s (no heat sink)", "NTC or contact thermocouple", "Case temperature ≤ 85°C. Above 80°C → evaluate thermal derating", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

fai_table(doc, "3c. Reference Photodiode Optical Isolation", "RISK-06 photodiode contamination", [
    ("FAI-P01", "[BLOCKING] Dark-field isolation test: cover all LEDs with opaque black tape; run photodiode readout in synchronous lock-in mode", "Bench test, firmware lock-in mode", "Photodiode ADC output < 1% of full-scale in absence of LED illumination", "", ""),
    ("FAI-P02", "Response test: uncover LEDs; run all 10 strings in zone at 25% duty cycle; verify photodiode responds", "Bench test", "Photodiode output increases ≥ 90× compared to dark reading; signal stable across 10 s", "", ""),
    ("FAI-P03", "LED aging simulation: attenuate LED output 30% with ND filter over PDMS window; verify photodiode detects reduction", "ND filter + bench test", "Reported J/cm² decreases by 28–32% (within ±2% of filter attenuation)", "", ""),
    ("FAI-P04", "PDMS window optical transmission: measure before and after assembly with integrating sphere or photodiode", "Integrating sphere / benchtop", "≥ 92% at 660 nm; ≥ 94% at 808–830 nm", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

fai_table(doc, "3d. PDMS Window Adhesion", "RISK-04 CTE delamination", [
    ("FAI-M01", "Visual inspection of PDMS–PI bond line: no delamination, no air bubbles > 0.5 mm, no discolouration", "Visual, 10× magnifier", "No defects visible. Any delamination ≥ 1 mm → reject", "", ""),
    ("FAI-M02", "Peel test (witness coupon from same batch): 90° peel test on 25 mm × 25 mm coupon", "Instron or similar, 90° peel", "Peel force ≥ 50 N/m. Below 20 N/m → reject batch (SiO₂ interlayer failure)", "", ""),
    ("FAI-M03", "Thermal cycle (accelerated): 5 cycles −20°C to +70°C on witness coupon; re-inspect bond line", "Thermal chamber", "No new delamination after 5 cycles. If delamination: halt production, escalate to process engineer", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

# ── 4. Shell + Alignment Verification ──
heading(doc, "SECTION 4 — SHELL INSTALLATION AND ALIGNMENT VERIFICATION")
body(doc, "RISK-12 primary mitigation. All 5 zone modules must be verified individually. "
     "The ±0.5 mm system alignment requirement is achieved via ±0.3 mm shell tooling tolerance (±0.2 mm margin).")

fai_table(doc, "4a. Zone Module Positional Accuracy", "RISK-12 alignment tolerance", [
    ("FAI-A01", "[BLOCKING] Zone 1 (frontal left — Fp1/F3): CMM measurement of LED array centroid position vs. shell datum", "CMM / optical profilometer", "LED array centroid within ±0.3 mm of nominal Fp1/F3 target position in X and Y", "", ""),
    ("FAI-A02", "[BLOCKING] Zone 2 (frontal right — Fp2/F4): CMM measurement", "CMM", "Within ±0.3 mm of nominal Fp2/F4 target", "", ""),
    ("FAI-A03", "[BLOCKING] Zone 3 (central left — C3): CMM measurement", "CMM", "Within ±0.3 mm of nominal C3 target", "", ""),
    ("FAI-A04", "[BLOCKING] Zone 4 (central right — C4): CMM measurement", "CMM", "Within ±0.3 mm of nominal C4 target", "", ""),
    ("FAI-A05", "[BLOCKING] Zone 5 (parietal — P3/P4): CMM measurement", "CMM", "Within ±0.3 mm of nominal P3/P4 target", "", ""),
    ("FAI-A06", "Zone-to-zone relative position: confirm inter-zone spacing matches 10-20 system inter-electrode distances", "CMM / derived from A01–A05 measurements", "All inter-zone distances within ±0.5 mm of 10-20 system nominal", "", ""),
    ("FAI-A07", "Zone module stiffener datum features (reference holes or edges): verify present and within ±0.1 mm of nominal", "CMM", "Datum features present; position ±0.1 mm", "", ""),
    ("FAI-A08", "Zone module re-insertion repeatability: remove and re-insert each zone module 3×; re-measure CMM position", "CMM, 3 repeats per zone", "Position repeatability ±0.1 mm (module self-locates to same position each insertion)", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

fai_table(doc, "4b. FPC Routing Path Verification (in-shell)", "RISK-11 bend radius", [
    ("FAI-R01", "Zone 1 FPC: visual inspection of routing path in assembled headset — no visible kink, tight bend, or pinch point", "Visual + endoscope (if needed)", "No kink visible; FPC follows channel smoothly", "", ""),
    ("FAI-R02", "Zone 2 FPC: same as FAI-R01 (highest risk — compound frontal curvature)", "Visual + endoscope", "No kink; smooth channel follow", "", ""),
    ("FAI-R03", "Zones 3–5 FPC: visual inspection", "Visual", "No kink on any zone", "", ""),
    ("FAI-R04", "FPC exit bend (Segment A) — physical radius check with radius gauge: all 5 zones", "Radius gauge, bend check tool", "≥ 25 mm at all Segment A points. Any zone < 20 mm → halt and escalate", "", ""),
    ("FAI-R05", "Zone 1 insertion/removal cycle test: insert/remove 10×; inspect FPC for crease, whitening, or deformation", "Manual test + visual", "No crease, whitening, or deformation after 10 cycles", "", ""),
    ("FAI-R06", "Cross-zone interference: remove Zone 1 module; confirm Zone 2 FPC is not disturbed", "Visual", "Zone 2 FPC remains in channel during Zone 1 removal", "", ""),
    ("FAI-R07", "ZIF lever access: confirm user can actuate ZIF lever (full 180° flip) without contacting FPC with fingertip", "Manual test, 3 operators", "All 3 operators actuate lever without FPC contact", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

# ── 5. Connector and Insertion Cycle Test ──
heading(doc, "SECTION 5 — CONNECTOR LIFECYCLE VALIDATION")
body(doc, "Zone modules are user field-replaceable. The Hirose FH34S connector is rated ≥1,000 cycles. "
     "FAI validates first-article connector performance at incoming (10-cycle) and production-release (50-cycle) levels.")

fai_table(doc, "5a. Connector Insertion Cycle Test", "RISK-01 connector family · RISK-11 routing path", [
    ("FAI-K10", "Baseline contact resistance (before cycling): measure all 20 contacts", "4-wire, all 20 pins", "All pins < 50 mΩ", "", ""),
    ("FAI-K11", "Insert/remove zone module 50× (simulated field replacement pace — one cycle per 15 s)", "Manual cycling, count logged", "50 cycles completed without mechanical failure of lever, body, or contacts", "", ""),
    ("FAI-K12", "Post-cycle contact resistance: measure all 20 contacts", "4-wire, all 20 pins", "All pins < 50 mΩ. Increase > 20 mΩ from baseline → flag for investigation", "", ""),
    ("FAI-K13", "Post-cycle visual: inspect connector body and FPC tail for wear marks, gold wear-through, or fretting", "Visual, 10× magnifier", "No exposed base metal on contact pads; no cracking of FPC coverlay at tail", "", ""),
    ("FAI-K14", "Post-cycle FPC position accuracy (re-run FAI-A01 to FAI-A05 on Zone 1 after 50 cycles)", "CMM", "Zone 1 position still within ±0.3 mm after 50 insertion cycles", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

# ── 6. Functional System Test ──
heading(doc, "SECTION 6 — FUNCTIONAL SYSTEM TEST (ASSEMBLED HEADSET)")
body(doc, "Run after all physical inspection items pass. Requires firmware-loaded headset connected to app.")

fai_table(doc, "6a. End-to-End Zone Function", "RISK-05 · RISK-06 · RISK-12", [
    ("FAI-SY01", "All 5 zones illuminate at 25% duty cycle — app confirms zone active, photodiode returns nonzero dose reading", "App UI, firmware telemetry", "All 5 zones report active; J/cm² > 0 for each zone", "", ""),
    ("FAI-SY02", "Zone-selective deactivation: deactivate Zone 1 only; confirm Zones 2–5 remain active and Zone 1 extinguishes", "App UI", "Correct zone isolation confirmed", "", ""),
    ("FAI-SY03", "PDMS fouling simulation: apply light smear of petroleum jelly to Zone 1 PDMS window; confirm photodiode detects reduced output and app flags maintenance reminder", "App UI + petroleum jelly test", "App shows reduced dose for Zone 1; maintenance reminder displayed within 2 sessions", "", ""),
    ("FAI-SY04", "EEG signal quality on all 8 electrodes: confirm impedance < 10 kΩ with hydrogel tips installed", "App impedance display", "All 8 electrodes < 10 kΩ (green)", "", ""),
    ("FAI-SY05", "Closed-loop session: run 5-minute EEG-adaptive PBM session; confirm adaptation events logged in UHDR", "App session log", "At least 1 adaptation event logged; UHDR session record created and encrypted", "", ""),
], [0.65, 2.3, 1.2, 1.55, 0.65, 0.45])

# ── 7. Non-Conformance Summary ──
heading(doc, "SECTION 7 — NON-CONFORMANCE SUMMARY")
body(doc,
    "List all items that did not pass during this FAI. All non-conformances must be dispositioned "
    "before production release — either by corrective action (re-inspect after fix) or by formal "
    "Engineering deviation approval. No BLOCKING item may be waived.")

t_nc = doc.add_table(rows=8, cols=5)
t_nc.style = "Table Grid"
hdr_nc = t_nc.rows[0]
for i, h in enumerate(["FAI Item", "Non-Conformance Description", "Disposition", "Corrective Action Ref", "Closed by / Date"]):
    set_cell_bg(hdr_nc.cells[i], "C00000")
    run = hdr_nc.cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(8)
for row in t_nc.rows[1:]:
    for cell in row.cells:
        set_cell_bg(cell, "FFFFFF")
        run = cell.paragraphs[0].add_run("")
        run.font.size = Pt(9)
for row in t_nc.rows:
    row.cells[0].width = Inches(0.7)
    row.cells[1].width = Inches(2.2)
    row.cells[2].width = Inches(1.0)
    row.cells[3].width = Inches(1.5)
    row.cells[4].width = Inches(1.4)
doc.add_paragraph()

# ── 8. Overall FAI Result and Sign-Off ──
heading(doc, "SECTION 8 — FAI RESULT AND SIGN-OFF")
body(doc,
    "All checklist items in Sections 1–6 must show PASS, or non-conformances in Section 7 must "
    "be fully closed, before production release is authorised.")

t_signoff = doc.add_table(rows=6, cols=4)
t_signoff.style = "Table Grid"
signoff_rows = [
    ("Overall FAI result:", "  PASS  /  FAIL", "Build lot approved for production:", "  YES  /  NO"),
    ("Manufacturing Engineer:", "", "Date:", ""),
    ("Hardware EE Lead:", "", "Date:", ""),
    ("Quality / Reliability:", "", "Date:", ""),
    ("Programme Manager:", "", "Date:", ""),
    ("Notes:", "", "", ""),
]
for r_idx, (a, b, c, d) in enumerate(signoff_rows):
    row = t_signoff.rows[r_idx]
    bg_l = "F2F2F2"
    for c_idx, txt in enumerate([a, b, c, d]):
        cell = row.cells[c_idx]
        set_cell_bg(cell, bg_l if c_idx % 2 == 0 else "FFFFFF")
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(9)
        if c_idx % 2 == 0:
            run.bold = True
for row in t_signoff.rows:
    for i, w in enumerate([1.5, 1.8, 1.5, 1.9]):
        row.cells[i].width = Inches(w)
doc.add_paragraph()

# ── 9. Risk Mitigation Cross-Reference ──
heading(doc, "SECTION 9 — RISK MITIGATION CROSS-REFERENCE")
body(doc, "All FAI items trace back to at least one risk register entry or specification requirement.")

t_xref = doc.add_table(rows=1, cols=4)
t_xref.style = "Table Grid"
for i, h in enumerate(["Risk ID", "Risk Title (abbreviated)", "FAI Items", "Status"]):
    set_cell_bg(t_xref.rows[0].cells[i], "1F497D")
    run = t_xref.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(8)

xref_data = [
    ("RISK-01", "Connector family — Hirose FH34S", "FAI-D05, FAI-K01–FAI-K04, FAI-K10–FAI-K14", "MITIGATED"),
    ("RISK-02", "FPC power trace current capacity", "FAI-T01, FAI-T02, FAI-T03", "MITIGATED"),
    ("RISK-04", "PDMS–PI CTE mismatch delamination", "FAI-M01, FAI-M02, FAI-M03, FAI-P04", "OPEN (process qualification pending)"),
    ("RISK-05", "LED current hogging (per-string CCC)", "FAI-E01, FAI-E02, FAI-E03, FAI-SY01", "MITIGATED"),
    ("RISK-06", "Reference photodiode optical isolation", "FAI-P01, FAI-P02, FAI-P03, FAI-SY03", "MITIGATED"),
    ("RISK-07", "BCR421U vs BCR421W variant", "FAI-C01, FAI-C02, FAI-D04", "MITIGATED"),
    ("RISK-08", "LED Vf binning not on PO", "FAI-L01–FAI-L05, FAI-D08", "MITIGATED"),
    ("RISK-10", "ED copper default on FPC fab", "FAI-D06, FAI-D07, FAI-F01, FAI-F02", "MITIGATED"),
    ("RISK-11", "Shell bend radius not specified", "FAI-R01–FAI-R07, FAI-K14", "MITIGATED"),
    ("RISK-12", "Zone alignment tolerance not validated", "FAI-A01–FAI-A08, FAI-SY04", "MITIGATED (this FAI)"),
]

for r_idx, (rid, title, items, status) in enumerate(xref_data):
    row = t_xref.add_row()
    bg = "FFFFFF" if r_idx % 2 == 0 else "F5F5F5"
    status_color = "E2EFDA" if "MITIGATED" in status and "pending" not in status.lower() else "FFF2CC"
    for c_idx, txt in enumerate([rid, title, items, status]):
        cell = row.cells[c_idx]
        set_cell_bg(cell, status_color if c_idx == 3 else bg)
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(8)
        if c_idx == 0:
            run.bold = True

for row in t_xref.rows:
    for i, w in enumerate([0.65, 2.0, 2.35, 1.8]):
        row.cells[i].width = Inches(w)
doc.add_paragraph()

# ── 10. Revision History ──
heading(doc, "SECTION 10 — REVISION HISTORY")
t_rev = doc.add_table(rows=2, cols=4)
t_rev.style = "Table Grid"
for i, h in enumerate(["Rev", "Date", "Author", "Description"]):
    set_cell_bg(t_rev.rows[0].cells[i], "1F497D")
    run = t_rev.rows[0].cells[i].paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(8)
for c_idx, txt in enumerate(["A", DATE, "NeuroPulse Engineering",
    "Initial release. RISK-12 primary mitigation (alignment FAI). "
    "Consolidates FAI items from RISK-01, 02, 04, 05, 06, 07, 08, 10, 11, 12. "
    "6 sections, 50+ inspection items."]):
    set_cell_bg(t_rev.rows[1].cells[c_idx], "F5F5F5")
    run = t_rev.rows[1].cells[c_idx].paragraphs[0].add_run(txt)
    run.font.size = Pt(8)
for row in t_rev.rows:
    row.cells[0].width = Inches(0.4)
    row.cells[1].width = Inches(0.9)
    row.cells[2].width = Inches(1.5)
    row.cells[3].width = Inches(4.0)

OUT = "docs/neuropulse_fai_zone_module.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
