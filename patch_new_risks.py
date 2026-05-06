"""
Add newly introduced risks (RISK-18 through RISK-24) to the risk register.
These are risks created by the mitigations for RISK-15, RISK-16, and RISK-17.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RISK_PATH = "docs/neuropulse_fpc_zone_module_risks_revA.docx"
doc = Document(RISK_PATH)

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

SEV_COLORS = {
    "HIGH":   "FF0000",
    "MEDIUM": "FF9900",
    "LOW":    "FFFF00",
}

NEW_RISKS = [
    {
        "id": "RISK-18",
        "severity": "HIGH",
        "title": "ZONE_ID ENABLE interlock creates session-blocking availability failure mode",
        "root_cause": (
            "Pin 18 (ZONE_ID) is now in the ENABLE assertion path: Hub MCU reads ADC(pin 18) "
            "before every ENABLE assertion. If pin 18 has an open-circuit fault — from a cracked "
            "solder joint, worn ZIF contact, or FPC trace break — the ADC reads out-of-range and "
            "Hub MCU asserts FAULT, holding ENABLE low. The user cannot run any session on the "
            "affected zone until the fault is diagnosed and repaired. A single bad contact "
            "on a low-priority signal pin now permanently blocks therapeutic use."
        ),
        "impact": (
            "User cannot complete sessions. For a daily-use wellness device used by people with "
            "neurological conditions, this is a significant availability failure. "
            "Contact resistance on ZIF connectors increases over the 1,000-cycle lifecycle; "
            "a marginal pin 18 connection may pass FAI but develop intermittent failures in the field."
        ),
        "mitigation": (
            "Mitigation (three-part): "
            "(1) Hub MCU firmware: implement debounced retry — read ZONE_ID ADC 3× at 100 ms intervals; "
            "only assert FAULT if all 3 readings are out-of-range. Single-sample transients do not block. "
            "(2) FAI-K04 (connector contact resistance test) must explicitly include pin 18 — "
            "contact resistance threshold ≤ 30 mΩ (tighter than the 50 mΩ general spec) to catch "
            "marginal contacts before production release. "
            "(3) App UX: on ZONE_ID FAULT, display diagnostic message distinguishing 'wrong module' "
            "from 'connection fault' — prompt user to re-seat module before assuming wrong part."
        ),
        "owner": "Firmware / HW EE",
        "spec_ref": "§5.2, pin 18",
        "status": "OPEN — firmware debounce and FAI-K04 pin 18 spec required",
        "introduced_by": "RISK-15 mitigation (Layer 2 — ZONE_ID resistor)",
    },
    {
        "id": "RISK-19",
        "severity": "HIGH",
        "title": "Co-moulded silicone gasket may delaminate from zone module body",
        "root_cause": (
            "The RISK-16 Option A gasket is co-moulded silicone bonded to the zone module "
            "thermoplastic body (likely ABS, PC, or PC-ABS). Silicone-to-thermoplastic co-moulding "
            "adhesion depends on surface preparation and compatible chemistries — without a primer "
            "or adhesion promoter, bond strength at silicone-to-ABS interfaces is typically "
            "3–8 N/mm (peel), which degrades with thermal cycling and repeated compression. "
            "After repeated zone module insertions and thermal cycles (body heat during sessions, "
            "cold storage), the gasket may delaminate from the module body edge."
        ),
        "impact": (
            "Delaminated gasket: (1) IPX4 seal is lost — water and sweat ingress to Hub PCB "
            "and ZIF connector during use; (2) free silicone strip becomes debris inside the "
            "headset shell, potentially jamming the ZIF lever or shorting exposed contacts; "
            "(3) user may not notice delamination — no warning, no SHDR alert. "
            "Failure mode is silent and potentially causes permanent hardware damage."
        ),
        "mitigation": (
            "Four mitigations required before tooling release: "
            "(1) Specify silicone primer/adhesion promoter on zone module body at gasket bond surface "
            "(e.g., Dow PR-1200 or equivalent) — include in zone module assembly process spec. "
            "(2) Add mechanical retention: groove or undercut in zone module body perimeter traps "
            "the gasket even if chemical bond fails — 'key' geometry prevents peel-off. "
            "(3) FAI gasket peel test (90°, 10 N/mm minimum) on production-representative co-mould "
            "sample — required before first production run. "
            "(4) SHDR monitoring: track per-session ZONE_ID ADC variance — gasket delamination "
            "that partially lifts the module changes pin 18 contact resistance; "
            "trend in SHDR flags mechanical degradation before full failure."
        ),
        "owner": "Mechanical Eng / Manufacturing Eng",
        "spec_ref": "§12",
        "status": "OPEN — mechanical retention geometry and peel test qualification required",
        "introduced_by": "RISK-16 mitigation (Option A — co-moulded gasket)",
    },
    {
        "id": "RISK-20",
        "severity": "HIGH",
        "title": "CFRP shell slot rim Ra ≤ 1.6 µm surface finish may require secondary operations",
        "root_cause": (
            "The RISK-16 Option A gasket spec requires shell slot rim Ra ≤ 1.6 µm for reliable "
            "sealing. Moulded CFRP typically achieves Ra 1.6–3.2 µm depending on mould tool finish "
            "and release agent. Ra ≤ 1.6 µm on CFRP requires either: a precision-polished mould tool "
            "(achievable but adds $5K–15K to shell tooling NRE and extends lead time), a gel-coat or "
            "paint layer on the slot interior (adds 0.05–0.15 mm to slot dimensions, affecting fit), "
            "or a moulded thermoplastic liner in the slot (adds a separate part and assembly step). "
            "This requirement was added without confirming feasibility with the shell tooling manufacturer."
        ),
        "impact": (
            "If slot rim Ra > 1.6 µm in production tooling: gasket sealing is inconsistent — "
            "microscale surface peaks penetrate the gasket, creating leak paths. "
            "IPX4 claim fails. 'Workout-safe' marketing claim cannot be made. "
            "Discovering this post-tooling requires either tooling rework (expensive), "
            "gel-coat rework (affects fit tolerances), or gasket redesign (requires re-qualification)."
        ),
        "mitigation": (
            "Required before tooling specification release: "
            "(1) Confirm achievable Ra on CFRP slot interior with shell tooling manufacturer — "
            "get written commitment to Ra ≤ 1.6 µm or identify secondary operation needed. "
            "(2) If secondary operation required: specify in shell tooling drawing and add to "
            "production process flow. Measure Ra on first-article shell with contact profilometer. "
            "(3) Alternative: if Ra ≤ 1.6 µm not reliably achievable, increase gasket durometer "
            "to Shore 50–60A and increase gasket compression to 25–30% — higher compression "
            "tolerates rougher surfaces. Requires gasket re-qualification (new FAI-IPX-02 test)."
        ),
        "owner": "Mechanical Eng / Manufacturing Eng",
        "spec_ref": "§12, NP-DRV-SHELL-001 REQ-ST-07",
        "status": "OPEN — CFRP surface finish feasibility confirmation required before tooling release",
        "introduced_by": "RISK-16 mitigation (Option A — gasket sealing spec)",
    },
    {
        "id": "RISK-21",
        "severity": "HIGH",
        "title": "EEG electrode cable routing path inside shell is unspecified — DRC-18 cannot be verified",
        "root_cause": (
            "RISK-17 mitigation (NP-DRV-SHELL-001 §2.3) requires ≥ 15 mm separation between "
            "zone module FPC tails and EEG electrode cables (DRC-18). However, the EEG electrode "
            "cable routing path inside the CFRP shell has never been specified — no routing document, "
            "no cable path drawing, no cable type or diameter spec exists. "
            "DRC-18 compares FPC paths against EEG cable paths in CAD overlay, but one of "
            "the two paths does not yet exist. RISK-17 mitigation is incomplete."
        ),
        "impact": (
            "Shell tooling is cut without EEG cable routing provisions. EEG cables, routed ad-hoc "
            "during assembly, may pass within 2–5 mm of FPC power tails. At 1.8 A RMS at ≥ 1 kHz, "
            "this produces > 100 µV of EEG artifact — masking delta and theta band activity entirely "
            "and producing artefactual gamma signals. Closed-loop algorithm operates on artifact. "
            "The core closed-loop differentiating claim fails. Additionally, EEG cable chafing "
            "against FPC tails over 1,000 sessions degrades both cable and FPC insulation."
        ),
        "mitigation": (
            "Add EEG cable routing specification to NP-DRV-SHELL-001 §2 as §2.4 before "
            "tooling specification release. Required content: EEG cable type (shielded twisted pair, "
            "AWG, outer diameter), routing path from each of 8 electrode pods to Hub PCB ADC, "
            "cable management provisions in shell tooling (dedicated channel or conduit separate "
            "from FPC channels), minimum separation from FPC channels at all crossing/parallel points. "
            "DRC-18 cannot be signed off until this spec exists and is verified in CAD. "
            "Add as NP-COORD-001 G2-11 (pre-tooling gate item)."
        ),
        "owner": "Mechanical Eng / HW EE",
        "spec_ref": "NP-DRV-SHELL-001 §2.3, DRC-18",
        "status": "OPEN — EEG cable routing spec required; blocks DRC-18 verification and tooling release",
        "introduced_by": "RISK-17 mitigation (exposed pre-existing gap)",
    },
    {
        "id": "RISK-22",
        "severity": "MEDIUM",
        "title": "Gasket compression force (2–4 N) and mechanical key alignment impede zone module replacement by users with reduced grip strength or tremor",
        "root_cause": (
            "Two RISK-15/16 mitigations add resistance to zone module insertion and extraction: "
            "(a) The co-moulded silicone compression gasket (RISK-16) adds 2–4 N extraction force — "
            "the gasket must be overcome to remove the module. "
            "(b) The asymmetric mechanical key (RISK-15 Layer 1) requires precise alignment before "
            "insertion — users must locate the key orientation before the module will enter. "
            "NeuroPulse's target populations include Parkinson's disease (tremor, rigidity), "
            "stroke rehabilitation (hemiparesis), TBI, and elderly users — all of whom may have "
            "significantly reduced grip strength (< 10 N) or hand tremor (5–15 Hz, 2–10 mm amplitude). "
            "A 2–4 N extraction force with a small module and a keyway alignment requirement "
            "may make field replacement impossible without caregiver assistance for these users."
        ),
        "impact": (
            "User cannot independently replace zone modules in the field. This is both a usability "
            "failure (reduces autonomy) and a clinical risk (user cannot address a fouled PDMS "
            "window or degraded module without help). "
            "If NeuroPulse markets autonomous home use as a differentiator, "
            "inability to self-service undermines that claim for a significant user segment."
        ),
        "mitigation": (
            "Three design recommendations for human factors review: "
            "(1) Add a lever or tab to zone module body for extraction grip — "
            "a 5–8 mm pull tab that can be grasped without fine motor control. "
            "(2) Reduce gasket compression to 15% (0.3 mm deflection, Shore 40A) — "
            "extraction force drops to ~1–2 N; re-qualify IPX4. "
            "(3) Mechanical key: ensure key chamfer (0.5 mm per §11.4) is generous; "
            "key tab should self-guide under even ±2 mm initial misalignment. "
            "Formative human factors testing (FDA 2016 HFE Guidance) required with "
            "representative users including tremor and reduced-grip participants. "
            "Document in Human Factors Engineering file (required for T2 510(k) anyway)."
        ),
        "owner": "UX / Mechanical Eng / Regulatory",
        "spec_ref": "§11.4, §12",
        "status": "OPEN — HFE formative testing required; pull tab geometry decision needed",
        "introduced_by": "RISK-15 (mechanical key) + RISK-16 (gasket extraction force)",
    },
    {
        "id": "RISK-23",
        "severity": "HIGH",
        "title": "Zone module mould requires 4 simultaneous co-moulded features — omission risk before first-cut",
        "root_cause": (
            "The zone module injection mould must now incorporate four distinct co-moulded or "
            "moulded-in features, each specified in a different document and driven by a different "
            "risk mitigation: "
            "(1) Braille abbreviation + raised numeral on module face (RISK-15 Layer 3, ISO 17049:2013); "
            "(2) Co-moulded silicone lip seal gasket on module perimeter (RISK-16 Option A); "
            "(3) Optical baffle housing/aperture for reference photodiode (RISK-06, §8.2); "
            "(4) Second photodiode aperture on scalp-facing surface (RISK-14 Option B — "
            "not yet implemented in spec). "
            "Each feature is documented separately. There is no single consolidated zone module "
            "tooling specification that lists all four. A tooling manufacturer receiving only "
            "part of the documentation will omit features not in their scope."
        ),
        "impact": (
            "Any feature omitted from zone module tooling cannot be retrofitted post-cut — "
            "a full mould retool is required. "
            "Retool cost: $20,000–60,000 per mould, 8–16 week lead time, programme delay. "
            "Worst case: all 4 features omitted from initial tooling specification = "
            "3–4 sequential retool cycles before production-ready zone module. "
            "This is a programme management risk, not just a technical one."
        ),
        "mitigation": (
            "Create a consolidated Zone Module Tooling Specification document (NP-TOOL-ZM-001) "
            "before tooling package is released to manufacturer. This document must list, "
            "in one place, every moulded-in feature of the zone module body and stiffener: "
            "all 4 co-moulded/moulded features above, plus mechanical key geometry (RISK-15 L1), "
            "datum features for CMM (RISK-12), module corner chamfers (§11.4), "
            "pull tab geometry (RISK-22 recommendation), and any other geometry driven by "
            "downstream documents. Gate: tooling package may not be released until "
            "NP-TOOL-ZM-001 is signed off by Mechanical Eng, HW EE, and Manufacturing Eng."
        ),
        "owner": "Mechanical Eng / Programme Manager",
        "spec_ref": "§8.2, §11.4, §12, NP-HW-FPC-001",
        "status": "OPEN — consolidated zone module tooling spec NP-TOOL-ZM-001 required before tooling release",
        "introduced_by": "Cumulative effect of RISK-06, RISK-12, RISK-14, RISK-15, RISK-16 mitigations",
    },
    {
        "id": "RISK-24",
        "severity": "MEDIUM",
        "title": "15 mm FPC-to-EEG cable separation requirement may be geometrically unachievable inside headset shell without design compromise",
        "root_cause": (
            "RISK-17 mitigation specifies ≥ 15 mm separation between zone module FPC power tails "
            "(1.8 A RMS) and EEG electrode signal cables. The NeuroPulse headset CFRP shell "
            "interior must simultaneously house: 5 FPC routing channels (each 3 mm wide × 1.5 mm deep), "
            "EEG electrode cables for 8 channels plus references, structural ribs, Boa cable routing, "
            "audio cable routing to over-ear speakers, and the Hub PCB compartment. "
            "The headset covers a 52–62 cm head circumference — interior shell width at crown is "
            "approximately 30–50 mm. Achieving 15 mm FPC-to-EEG separation while routing both "
            "within the same shell half may consume the entire available interior cross-section."
        ),
        "impact": (
            "If 15 mm separation is not geometrically achievable: "
            "(a) Shell geometry must be enlarged (increases headset size, reduces wearability); "
            "(b) Grounded barrier must be inserted between FPC and EEG channels "
            "(adds parts, weight, assembly complexity); "
            "(c) Separation requirement must be relaxed (increases EEG artifact risk). "
            "Any of these compromises must be resolved before shell tooling is released. "
            "Discovering the constraint after tooling is cut is a full retool scenario."
        ),
        "mitigation": (
            "Resolution requires shell CAD model (which does not yet exist — see RISK-21). "
            "Until CAD is available: "
            "(1) Architectural decision: route EEG cables and FPC tails on opposite halves of "
            "the shell interior (FPC tails on headset crown side, EEG cables on scalp side). "
            "This maximises separation without specifying exact dimensions. "
            "(2) If barrier is needed: 0.1 mm aluminium foil strip bonded to shell interior "
            "and connected to chassis GND — adds < 1 g weight, negligible cost. "
            "(3) DRC-18 in NP-DRV-SHELL-001 is the verification gate — cannot be signed off "
            "without EEG cable routing spec (RISK-21) and shell CAD. "
            "Add as explicit dependency in NP-COORD-001 G2-11."
        ),
        "owner": "Mechanical Eng / HW EE",
        "spec_ref": "NP-DRV-SHELL-001 §2.3, DRC-18",
        "status": "OPEN — CAD verification required; EEG cable routing spec (RISK-21) prerequisite",
        "introduced_by": "RISK-17 mitigation (15 mm separation requirement without CAD verification)",
    },
]

# ── Append new rows to the summary table ──
summary_tbl = doc.tables[0]
# Build header mapping from row 0
headers = [c.text.strip() for c in summary_tbl.rows[0].cells]
print(f"Table columns ({len(headers)}): {headers}")

for risk in NEW_RISKS:
    new_row = summary_tbl.add_row()
    cells = new_row.cells

    # Ensure 9 cells (some tables may merge — check)
    if len(cells) < 9:
        print(f"  WARNING: {risk['id']} row has only {len(cells)} cells")
        continue

    sev_bg = SEV_COLORS.get(risk["severity"], "FFFFFF")
    status_bg = "FFE699"  # OPEN = amber

    field_map = [
        (0, risk["id"],        False, "FFFFFF"),
        (1, risk["severity"],  True,  sev_bg),
        (2, risk["title"],     False, "FFFFFF"),
        (3, risk["root_cause"][:400], False, "FFFFFF"),
        (4, risk["impact"][:300],     False, "FFFFFF"),
        (5, risk["mitigation"][:500], False, "FFFFFF"),
        (6, risk["owner"],     False, "FFFFFF"),
        (7, risk["spec_ref"],  False, "FFFFFF"),
        (8, risk["status"],    False, status_bg),
    ]

    for (c_idx, txt, bold, bg) in field_map:
        cell = cells[c_idx]
        set_bg(cell, bg)
        run = cell.paragraphs[0].add_run(txt)
        run.bold = bold
        run.font.size = Pt(9)
        if c_idx == 1 and risk["severity"] == "HIGH":
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    print(f"  Added {risk['id']}: {risk['title'][:60]}")

# ── Update the header stats paragraph (near doc top) ──
for para in doc.paragraphs:
    if "risks identified" in para.text and "CRITICAL" in para.text:
        old = para.text
        for run in para.runs:
            run.text = ""
        total = 17 + len(NEW_RISKS)
        new_text = (
            f"{total} risks identified  |  3 CRITICAL  |  9 HIGH  |  9 MEDIUM  |  2 LOW\n"
            "All CRITICAL and HIGH items must be resolved before FPC layout release.\n"
            f"RISK-18 to RISK-24: newly introduced by mitigations — see detailed sections."
        )
        para.runs[0].text = new_text
        print(f"  Stats header updated")
        break

doc.save(RISK_PATH)
print(f"\nRisk register saved: {RISK_PATH}")
