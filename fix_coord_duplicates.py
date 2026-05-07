"""
Fix duplicate G1-08 entries in coord checklist introduced by patch_risk04.py.

Problems:
1. Table 7 row 2: spurious G1-08 inserted after G1-07 (OI-03 LED emitter).
   The real G1-08 in Table 7 is row 3 (OI-02 Hirose); restore its ID.
2. Table 15 rows 7-8: malformed G1-07/G1-08 dashboard entries (text in
   wrong columns). Remove them.
3. Table 15 row 14 (G1-12) already tracks RISK-04 SiO2 process house — update
   to reference G1-08 and NP-FAI-ZM-001 §3e.
4. Table 15 row 24 (G3-01) already tracks PDMS 200-cycle qual — update to
   reference FAI-TC02 and mark it as now captured at G1-08 (pre-layout gate).
5. G2-10 and G2-11 are missing from dashboard — add them.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_text(cell, text, bold=False, color=None, size=8):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

COORD_PATH = "docs/neuropulse_eng_coordination_checklist.docx"
coord = Document(COORD_PATH)

# ── 1. Table 7: remove spurious G1-08 (row 2) and fix row 3 ID back to G1-08 ─
tbl7 = coord.tables[7]
rows_to_remove = []
for ri, row in enumerate(tbl7.rows):
    rid = row.cells[0].text.strip()
    # The spurious row has G1-08 ID but has text crammed into wrong columns
    # (only 5 cells populated instead of 7, or description in Gate column)
    # Identify it: the real G1-08 in Table 7 has "OI-02" in coord item column
    if rid == "G1-08":
        item_text = row.cells[2].text.strip() if len(row.cells) > 2 else ""
        if "PDMS process qualification" in item_text and len(row.cells[1].text.strip()) <= 4:
            # This is the spurious one: Gate col = "G1", Owner col = "Mfg Eng + HW EE"
            rows_to_remove.append(row._tr)
            print(f"  Table 7 row {ri}: marking spurious G1-08 for removal")

for tr in rows_to_remove:
    tr.getparent().remove(tr)

# Now the old row 3 (OI-02 Hirose) is back at row 2 but may have ID G1-08 already
# Verify it looks correct
for ri, row in enumerate(tbl7.rows):
    if "OI-02" in row.cells[2].text or "Hirose" in row.cells[2].text:
        set_text(row.cells[0], "G1-08", bold=True, size=8)
        print(f"  Table 7: restored G1-08 = OI-02 Hirose row")
        break

# ── 2. Table 15 dashboard: remove malformed G1-07/G1-08 rows ─────────────────
tbl15 = coord.tables[15]
# Malformed rows: G1-07 row that has ZONE_ID debounce text in the Owner column
#                 G1-08 row that has PDMS text in the Owner column
# Both are identifiable because their Owner cell has long description text
# and they appear BEFORE the correct G1-07 (OI-03 LED emitter) dashboard row
bad_trs = []
for ri, row in enumerate(tbl15.rows):
    rid = row.cells[0].text.strip()
    if rid in ("G1-07", "G1-08"):
        owner_text = row.cells[2].text.strip()
        # Malformed rows: Owner cell has the long description (> 20 chars of description)
        if len(owner_text) > 30 and ("ZONE_ID" in owner_text or "PDMS" in owner_text
                                      or "debounce" in owner_text or "qualification" in owner_text):
            bad_trs.append((ri, row._tr))
            print(f"  Table 15 row {ri}: marking malformed {rid} for removal (owner='{owner_text[:50]}')")

for ri, tr in bad_trs:
    tr.getparent().remove(tr)

# ── 3. Update Table 15 G1-12 (SiO2 process house) to reference G1-08 ─────────
tbl15 = coord.tables[15]  # reload reference after removals
for row in tbl15.rows:
    if row.cells[0].text.strip() == "G1-12":
        set_text(row.cells[3],
            "PDMS SiO₂ process house confirmed; "
            "thermal cycling qual (FAI-TC02) BLOCKING (→ G1-08)",
            size=8)
        set_text(row.cells[4], "OPEN — see G1-08", size=8)
        set_bg(row.cells[4], "FFF2CC")
        print("  Table 15 G1-12 updated to cross-reference G1-08")
        break

# ── 4. Update Table 15 G3-01 (PDMS 200-cycle) to reference FAI-TC02 ──────────
for row in tbl15.rows:
    if row.cells[0].text.strip() == "G3-01":
        set_text(row.cells[3],
            "PDMS thermal cycling qualification 200 cycles — "
            "FAI-TC02 (BLOCKING at G1-08)",
            size=8)
        set_text(row.cells[4],
            "OPEN — blocked on G1-08 (CAT-C supplier selection)",
            size=8)
        set_bg(row.cells[4], "FFF2CC")
        print("  Table 15 G3-01 updated to reference FAI-TC02 / G1-08")
        break

# ── 5. Add G2-10 and G2-11 to dashboard (after G2-09) ────────────────────────
g209_tr = None
for row in tbl15.rows:
    if row.cells[0].text.strip() == "G2-09":
        g209_tr = row._tr
        break

if g209_tr:
    NEW_DASHBOARD_ROWS = [
        ("G2-10", "G2", "FW + HW EE",
         "Audio zone ID firmware — bone conduction zone name on insertion",
         "OPEN"),
        ("G2-11", "G2", "EE + ME",
         "EEG cable routing spec §2.4 in NP-DRV-SHELL-001; DRC-18 verified",
         "OPEN — §2.4 in progress (RISK-21)"),
    ]
    widths_in = [0.65, 0.55, 1.3, 2.8, 1.4]
    for row_data in reversed(NEW_DASHBOARD_ROWS):
        new_tr = OxmlElement("w:tr")
        for c_idx, txt in enumerate(row_data):
            tc = OxmlElement("w:tc")
            tc_pr = OxmlElement("w:tcPr")
            w_el = OxmlElement("w:tcW")
            w_el.set(qn("w:w"), str(int(widths_in[c_idx] * 1440)))
            w_el.set(qn("w:type"), "dxa")
            tc_pr.append(w_el); tc.append(tc_pr)
            p_el = OxmlElement("w:p")
            r_el = OxmlElement("w:r")
            rpr_el = OxmlElement("w:rPr")
            sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), "16")
            rpr_el.append(sz_el); r_el.append(rpr_el)
            t_el = OxmlElement("w:t")
            t_el.set(qn("xml:space"), "preserve"); t_el.text = txt
            r_el.append(t_el); p_el.append(r_el); tc.append(p_el)
            new_tr.append(tc)
        g209_tr.addnext(new_tr)
        # apply status bg
        for r2 in tbl15.rows:
            if r2._tr is new_tr:
                set_bg(r2.cells[4], "FFF2CC")
                if r2.cells[0].paragraphs[0].runs:
                    r2.cells[0].paragraphs[0].runs[0].bold = True
                break
        print(f"  Added {row_data[0]} to dashboard")

# ── Verify final state ────────────────────────────────────────────────────────
print("\nFinal G1-08 count across all tables:")
total = 0
for ti, tbl in enumerate(coord.tables):
    cnt = sum(1 for r in tbl.rows if r.cells[0].text.strip() == "G1-08")
    if cnt:
        print(f"  Table {ti}: {cnt} × G1-08")
        total += cnt
print(f"  Total: {total} (expected: 2 — Table 6 main item + Table 7 OI-02)")

coord.save(COORD_PATH)
print(f"\nNP-COORD-001 saved")
