"""
Generate NP-TOOL-SHELL-001 Rev A
NeurOne Headset Shell Tooling Specification
Covers: anchor posts (5 zone + 3 port), EEG cable channel, temporal wing anchor boss
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)

def cell_bold(cell, text, size=10, align=None, colour=None):
    para = cell.paragraphs[0]
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = RGBColor(*colour)

def cell_text(cell, text, size=10, bold=False, align=None):
    para = cell.paragraphs[0]
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_para(doc, text, size=10, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_notice(doc, text, colour='FFF3CD'):
    """Yellow/amber notice box via single-cell table."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, colour)
    para = cell.paragraphs[0]
    run  = para.add_run(text)
    run.font.size = Pt(9)
    run.bold = True
    doc.add_paragraph()  # spacer

def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        set_cell_bg(hdr.cells[i], '1F3864')
        cell_bold(hdr.cells[i], h, size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  colour=(255, 255, 255))
    # data rows
    for ri, row in enumerate(rows):
        fill = 'F2F2F2' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            set_cell_bg(c, fill)
            para = c.paragraphs[0]
            run  = para.add_run(str(val))
            run.font.size = Pt(9)
    # column widths
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph()  # spacer
    return t

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
title_tbl = doc.add_table(rows=1, cols=2)
title_tbl.style = 'Table Grid'
left  = title_tbl.rows[0].cells[0]
right = title_tbl.rows[0].cells[1]
set_cell_bg(left,  '1F3864')
set_cell_bg(right, '1F3864')
left.width  = Cm(11)
right.width = Cm(6)

p_brand = left.paragraphs[0]
r1 = p_brand.add_run('NeurOne')
r1.bold = True; r1.font.size = Pt(16); r1.font.color.rgb = RGBColor(255,255,255)
p_brand.add_run('\n')
r2 = p_brand.add_run('Headset Shell Tooling Specification')
r2.font.size = Pt(11); r2.font.color.rgb = RGBColor(200,210,230)

p_meta = right.paragraphs[0]
p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for line in [
    'NP-TOOL-SHELL-001',
    'Revision: A (Draft — Pre-Tooling)',
    'Status: OPEN — must be signed off',
    'before tooling first-cut',
    'Date: 2026-05-10',
]:
    run = p_meta.add_run(line + '\n')
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(255,255,255)
    if line.startswith('NP-TOOL'):
        run.bold = True; run.font.size = Pt(11)

doc.add_paragraph()

# companion notice
add_notice(doc,
    'STATUS:  This document is a REQUIRED TOOLING SPECIFICATION that must be '
    'completed and signed off before shell tooling first-cut. Features not '
    'specified before first cut cannot be retrofitted — tooling modification '
    'costs $15,000–40,000 and 6–8 weeks delay per feature. All four features '
    'below are MANDATORY.',
    colour='FDECEA')

# ══════════════════════════════════════════════════════════════════════════════
#  §1  PURPOSE AND SCOPE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1.  PURPOSE AND SCOPE')
add_para(doc,
    'This document specifies mandatory moulded features in the NeurOne headset '
    'CFRP shell that must be present in the tooling specification before first steel '
    'cut. It is the tooling-facing companion to NP-DRV-SHELL-001 (FPC routing design '
    'review) and NP-TOOL-ZM-001 (zone module tooling specification). Together these '
    'three documents constitute the complete shell + zone module tooling package.')
add_para(doc,
    'Scope: Headset outer CFRP shell and temporal stability wing. Not in scope: zone '
    'module body mould (NP-TOOL-ZM-001), control hub tooling, goggle arm tooling '
    '(lens rim guard anchor hook is on goggle arm — separate document).')

# reference table
add_heading(doc, 'Related Documents', level=2)
make_table(doc,
    ['Document Number', 'Title', 'Relationship'],
    [
        ['NP-DRV-SHELL-001 Rev A', 'Zone Module FPC Routing Path / Shell Tooling Design Review',
         'Parent design review — FPC routing channels, EEG separation (DRC-18). F-03 EEG channel in §2.4.'],
        ['NP-TOOL-ZM-001 Rev A', 'Zone Module Tooling Specification',
         'Zone module body features (F-01 through F-08). Zone colour codes for F-01 anchor post colour-match.'],
        ['NP-COORD-001 Rev A', 'Engineering Coordination Checklist',
         'G2 gate items. G2-01 (multi-FPC + EEG separation), G2-02 (keying), G2-11 (EEG routing spec) all cross-reference this document.'],
        ['NP-FAI-ZM-001 Rev A', 'First-Article Inspection Checklist — Zone Module',
         'FAI items FAI-CV-01–CV-04 (anchor post / tether inspection) and FAI-TW-01 (temporal wing boss) cross-reference this spec.'],
        ['CLAUDE.md §7.1', 'Interface protection covers — anchor post specification', 'Design intent and cover descriptions.'],
        ['CLAUDE.md §3b', '40Hz Mastoid LRA Pad — temporal wing anchor boss', 'Provisional accessory requiring boss provision at first cut.'],
        ['CLAUDE.md §8.3', 'Interface protection covers — cover types and counts', 'Counts: 5 zone + 3 port + 2 lens rim guards (lens rim guards on goggle arm).'],
    ],
    col_widths=[3.8, 6.5, 6.4])

# ══════════════════════════════════════════════════════════════════════════════
#  §2  MANDATORY SHELL FEATURES — CONSOLIDATED CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2.  MANDATORY SHELL FEATURES — CONSOLIDATED CHECKLIST')
add_para(doc,
    'The following table lists all moulded features required in the headset shell. '
    'Each feature is zero-cost when specified before first cut and unretrofittable '
    'after. The mould design review (NP-COORD-001 G2 gate) must confirm all items '
    'before steel is cut.')

make_table(doc,
    ['Feature ID', 'Feature Description', 'Risk / Ref', 'Confirmed?'],
    [
        ['F-01', 'Zone slot plug anchor posts (×5, colour-coded). See §2.1.',
         'CLAUDE.md §7.1 / §8.3', '☐'],
        ['F-02', 'Accessory port cover anchor posts (×3). See §2.2.',
         'CLAUDE.md §7.1 / §8.3', '☐'],
        ['F-03', 'EEG cable routing channel (8×5mm, outer CFRP surface). See §2.3.',
         'RISK-21 / NP-DRV-SHELL-001 §2.4 / DRC-18', '☐'],
        ['F-04', 'Temporal wing anchor boss (×2 bilateral, PROVISIONAL). See §2.4.',
         'CLAUDE.md §3b', '☐'],
    ],
    col_widths=[1.8, 8.5, 4.0, 2.4])

# ──────────────────────────────────────────────────────────────────────────────
# §2.1  F-01 Zone slot plug anchor posts
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, '2.1  F-01 — Zone Slot Plug Anchor Posts (×5)', level=2)
add_para(doc,
    'Five anchor posts, one adjacent to each zone module slot (ZM-01 through ZM-05). '
    'Each post serves as the retention anchor for the tethered Shore 30A medical '
    'silicone zone slot plug. The tether prevents the plug being permanently lost '
    'when a zone module is installed.')

add_para(doc, '[DRAWING REQUIREMENT — SHELL TOOLING]  Zone slot plug anchor posts:', bold=True)

bullet_items_f01 = [
    'Count: 5 (one per zone slot — ZM-01 Frontal Left, ZM-02 Frontal Right, '
     'ZM-03 Central Left, ZM-04 Central Right, ZM-05 Parietal/Occipital).',
    'Location: On the outer CFRP shell surface adjacent to each zone slot opening, '
     'on the side toward the headset crown (away from the module insertion direction). '
     'Post must not obstruct zone module insertion or removal path. Minimum clearance '
     'between post edge and module slot edge: ≥ 3 mm.',
    'Geometry: Cylindrical boss Ø 4.0 mm × 3.5 mm height. A 2.0 mm wide × 2.5 mm '
     'deep through-slot runs transversely across the boss crown, allowing a Shore 30A '
     'silicone tether loop to pass through and lock. Boss base radius: ≥ 1.0 mm '
     '(prevents stress concentration on CFRP surface at boss root).',
    'Colour coding: Boss moulded or over-moulded to match zone module colour. '
     'ZM-01 = Blue, ZM-02 = Red, ZM-03 = Green, ZM-04 = Yellow, ZM-05 = Purple. '
     'Colour must be in bulk material — not surface paint. Matches NP-TOOL-ZM-001 §3.1 '
     'colour assignments. Serves as visual + tactile zone identification at the shell '
     'surface, reinforcing the 5-layer zone keying system (RISK-15).',
    'Tether: Shore 30A medical-grade silicone loop, nominal 50 mm circumference, '
     '2.0 mm cross-section diameter. Loop threads through post slot before plug is '
     'assembled. Rated ≥ 500 tug cycles at 5 N without fracture or detachment.',
    'Sealing: Anchor post base must not create a water ingress path. Boss base to be '
     'flush-bonded or over-moulded with silicone sealant. IPX4 compliance must be '
     'maintained at all 5 post locations (per FAI item FAI-CV-02).',
    'Draft angle: ≥ 1.5° on all boss faces to ensure clean demoulding. Confirm with '
     'tooling manufacturer DFM review.',
]
for item in bullet_items_f01:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(item)
    run.font.size = Pt(9.5)

doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────────
# §2.2  F-02 Accessory port anchor posts
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, '2.2  F-02 — Accessory Port Cover Anchor Posts (×3)', level=2)
add_para(doc,
    'Three anchor posts, one adjacent to each accessory port: USB-C port, left hub '
    'accessory port, right hub accessory port. Each post serves as the retention '
    'anchor for the tethered Shore 40A TPE + encapsulated steel disc magnetic '
    'accessory port cover. Port covers attach magnetically (N42 steel disc, ~400 g '
    'pull per CLAUDE.md §8.3); the anchor post retains the cover when detached so it '
    'cannot be lost.')

add_para(doc, '[DRAWING REQUIREMENT — SHELL TOOLING]  Accessory port cover anchor posts:', bold=True)

bullet_items_f02 = [
    'Count: 3 (USB-C port, left hub accessory port, right hub accessory port).',
    'Location: On the control hub shell, adjacent to each port opening, positioned so '
     'the tethered cover can swing fully clear of the port opening for cable connection '
     'without the tether restricting the plug insertion angle. Minimum tether clearance '
     'from port centreline: 15 mm (allows full-size USB-C plug insertion with cover '
     'hanging freely to the side).',
    'Geometry: Same cylindrical boss spec as F-01 — Ø 4.0 mm × 3.5 mm height, '
     '2.0 mm × 2.5 mm through-slot for tether loop passage. Boss base radius ≥ 1.0 mm.',
    'Colour: Neutral shell colour (no zone colour coding — port covers are not '
     'zone-specific). Match hub shell body colour.',
    'Tether: Shore 40A TPE loop, nominal 50 mm circumference, 2.0 mm cross-section. '
     'Tether material matches cover body material for chemical compatibility and '
     'consistent aging. Rated ≥ 500 tug cycles at 5 N.',
    'Sealing: Post base must not compromise hub shell IPX4 rating. Confirm with '
     'hub shell sealing design.',
    'Cover retention is magnetic (not friction): port covers include encapsulated '
     'N42 steel disc (~400 g pull force) and Shore 20A silicone fins for dust '
     'exclusion per CLAUDE.md §8.3. Anchor post is a tether anchor only — it does '
     'not provide magnetic closing force.',
    'Draft angle: ≥ 1.5° on all boss faces.',
]
for item in bullet_items_f02:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(item)
    run.font.size = Pt(9.5)

doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────────
# §2.3  F-03 EEG cable routing channel
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, '2.3  F-03 — EEG Cable Routing Channel (8×5mm)', level=2)
add_para(doc,
    'A dedicated moulded channel on the outer CFRP inner surface (dorsal-outer side, '
    'opposite from the zone module FPC bundle) routes EEG electrode cables from the '
    'Hub PCB to the 8 electrode pod locations (Fp1/2, F3/4, C3/4, P3/4). This feature '
    'is fully specified in NP-DRV-SHELL-001 §2.4 (EEG Electrode Cable Routing — '
    'RISK-21) and is reproduced here as the tooling-facing requirement for the shell '
    'manufacturer. This is a PRE-TOOLING BLOCKING REQUIREMENT — shell tooling must '
    'not be released until NP-DRV-SHELL-001 §2.4 is approved by EE and ME (per '
    'NP-COORD-001 G2-11).')

add_para(doc, '[DRAWING REQUIREMENT — SHELL TOOLING]  EEG cable routing channel:', bold=True)

bullet_items_f03 = [
    'Cross-section: 8 mm wide × 5 mm deep, moulded channel on CFRP shell inner '
     'surface (outer/dorsal CFRP inner face).',
    'Routing surface: OUTER CFRP inner surface ONLY — the zone module FPC bundle '
     'routes on the INNER (scalp-facing) CFRP inner surface. These two routing '
     'surfaces are on OPPOSITE sides of the shell wall. Shell wall nominal 2.5 mm '
     'CFRP + mu-metal + EMI foam provides the separation.',
    'Wall separation: ≥ 2 mm wall thickness between EEG channel outer face and zone '
     'module FPC channel outer face at any common cross-section.',
    'FPC-to-EEG separation (DRC-18): Minimum 15 mm between any zone module FPC '
     'copper conductor and any EEG cable conductor, measured point-to-point in CAD. '
     'The shell wall routing geometry achieves approximately 18–22 mm separation. '
     'If geometry forces less than 15 mm at any point: grounded 35 µm aluminium '
     'foil barrier (adhesive-backed, tied to chassis GND at pin 20) must be '
     'specified in that region — barrier placement is zero-cost if moulded into '
     'shell at this stage (NP-DRV-SHELL-001 DRC-18b).',
    'Cable retention clips: 3 snap-in cable retention clips (PC/ABS, snap-moulded '
     'integral to channel wall) at approximately 80 mm intervals along the main '
     'channel run. Clip opening force ≤ 3 N (tool-free cable insertion). Clip '
     'internal radius ≥ 3 mm (prevents cable kinking).',
    'Channel wall corners: All internal corners radiused ≥ 1 mm.',
    'Branch points: Channel branches at electrode pod feed-through locations. Branch '
     'angle ≥ 90° (no acute bends). Branch cross-section: 5 mm wide × 4 mm deep '
     '(reduces from main channel at feed-through).',
    'Cable feed-throughs: At each electrode pod, a 5 mm diameter feed-through hole '
     'passes from the EEG channel to the electrode pod seating recess. Feed-through '
     'edge: chamfer 0.5 mm × 45° on both faces (prevents cable chafing).',
    'Status (NP-COORD-001 G2-11): EEG routing strategy confirmed (dedicated 8×5 mm '
     'channel, opposite CFRP surface from FPC bundle). CAD overlay measurement '
     '(DRC-18a) to proceed at G2 gate when shell CAD model is available (OI-09). '
     'DRC-22 and DRC-23 in NP-DRV-SHELL-001 currently OPEN.',
]
for item in bullet_items_f03:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(item)
    run.font.size = Pt(9.5)

doc.add_paragraph()

# ──────────────────────────────────────────────────────────────────────────────
# §2.4  F-04 Temporal wing anchor boss
# ──────────────────────────────────────────────────────────────────────────────
add_heading(doc, '2.4  F-04 — Temporal Wing Anchor Boss (×2 Bilateral, PROVISIONAL)', level=2)

add_notice(doc,
    'PROVISIONAL — F-04 release is contingent on HOPE Phase 3 results (Cognito '
    'Therapeutics, n=670, expected mid-2026). However, the anchor boss MUST be '
    'moulded into the temporal wing tooling at first cut regardless of study outcome. '
    'Adding the boss at first cut: zero incremental tooling cost. Retrofitting the '
    'boss post-cut: $15,000–40,000 tooling modification + 6–8 week delay. If HOPE '
    'Phase 3 is negative, the boss is a non-functional recessed cavity — no impact '
    'on headset form or function. If positive, mastoid pad ships within 6 months '
    'without tooling delay.',
    colour='FFF3CD')

add_para(doc,
    'Purpose: Provides the mechanical attachment point for the 40Hz mastoid '
    'vibrotactile LRA pad accessory (CLAUDE.md §3b). The pad clips to the temporal '
    'stability wing at the mastoid process location using a 30 mm diameter silicone '
    'over-moulded pad with a snap-fit clip. The clip engages the anchor boss. Without '
    'the boss, the temporal wing provides no pad attachment surface and the accessory '
    'cannot be offered as an add-on.')

add_para(doc, '[DRAWING REQUIREMENT — TEMPORAL WING TOOLING]  Temporal wing anchor boss:', bold=True)

bullet_items_f04 = [
    'Count: 2 — bilateral (left temporal wing, right temporal wing). One boss per wing.',
    'Location: At the mastoid process contact point on the temporal stability wing — '
     'the posterior-inferior area of the wing where it contacts the skull behind the '
     'ear. This is the anatomy-matched location for mastoid bone coupling, as specified '
     'in CLAUDE.md §3b: "Placement: Posterior temporal / mastoid process. Rationale: '
     'direct bone coupling to skull."',
    'Boss geometry: Recessed snap-fit clip receiver. External envelope: 34 mm wide '
     '× 34 mm tall × 5 mm proud of wing surface (provides engagement depth). '
     'Receiver opening: 30 mm diameter (matches 30 mm pad clip diameter per '
     'CLAUDE.md §3b). Snap detent: 0.5 mm undercut bead at 3 mm depth, full '
     '360° circumference. Snap release force: 3–6 N axial pull (holds pad securely '
     'during motion; user-removable without tools).',
    'Recessed profile when not in use: Receiver opening sits flush with (or 0.5 mm '
     'recessed below) temporal wing outer surface. No external protrusion when no '
     'pad is attached. No sharp exposed edges (all entry radii ≥ 1.5 mm).',
    'Connector pocket: Within the receiver, a 10 mm wide × 6 mm tall × 4 mm deep '
     'pocket at the 6 o\'clock position (inferior) accommodates the 3-pin pogo or '
     'JST connector from the pad to the hub accessory port cable. Pocket must '
     'provide ≥ 1.5 mm clearance around connector body.',
    'Load path: Boss must withstand 15 N axial pull (pad weight × safety factor '
     'of 50×) and 10 N lateral shear without visible deformation. Verify in FEA '
     'before first-cut (OI-05).',
    'Material compatibility: Boss is integral to temporal wing (same PC/ABS or '
     'ABS material as wing body). Snap detent in same material. Silicone '
     'pad contact face: add Shore 30A silicone over-mould insert at receiver '
     'opening face for vibration coupling comfort (0.5 mm overmould layer).',
    'IPX4 compliance: Receiver pocket must not create water ingress path into '
     'headset interior. Snap detent seals against pad silicone collar when '
     'pad is inserted. When no pad is inserted, a snap-fit dust cover (tethered, '
     'same anchor post tether spec as F-01) keeps the pocket closed.',
    'PROVISIONAL marking: Do not add any external product marking indicating '
     'mastoid pad compatibility until HOPE Phase 3 results are confirmed. '
     'Boss is non-labelled in Rev A tooling.',
    'Draft angle: ≥ 1° on all boss interior surfaces for clean demoulding.',
]
for item in bullet_items_f04:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(item)
    run.font.size = Pt(9.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  §3  CRITICAL DIMENSIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3.  CRITICAL DIMENSIONS')
add_para(doc,
    'The following dimensions are tolerance-critical and must appear as drawing '
    'callouts on all shell and temporal wing drawings submitted to the tooling '
    'manufacturer.')

make_table(doc,
    ['Dimension', 'Feature', 'Nominal', 'Tolerance', 'Critical for', 'Drawing callout?'],
    [
        ['Anchor post boss diameter', 'F-01, F-02', 'Ø 4.0 mm', '±0.2 mm',
         'Tether loop fit', 'YES'],
        ['Anchor post boss height', 'F-01, F-02', '3.5 mm', '+0.3 / −0.0 mm',
         'Tether loop retention', 'YES'],
        ['Anchor post through-slot width', 'F-01, F-02', '2.0 mm', '+0.1 / −0.0 mm',
         'Tether loop passage', 'YES'],
        ['Anchor post through-slot depth', 'F-01, F-02', '2.5 mm', '±0.2 mm',
         'Tether loop locking', 'YES'],
        ['Anchor post base radius', 'F-01, F-02', '≥ 1.0 mm', 'Min only',
         'CFRP stress concentration', 'YES'],
        ['Zone post to slot edge clearance', 'F-01', '≥ 3 mm', 'Min only',
         'Module insertion path', 'YES'],
        ['Port post to port CL clearance', 'F-02', '≥ 15 mm', 'Min only',
         'Cable plug access', 'YES'],
        ['EEG channel width', 'F-03', '8 mm', '+0.3 / −0.0 mm',
         'Cable bundle fit', 'YES'],
        ['EEG channel depth', 'F-03', '5 mm', '+0.3 / −0.0 mm',
         'Cable bundle fit', 'YES'],
        ['EEG channel wall corners', 'F-03', '≥ 1 mm radius', 'Min only',
         'Cable chafe prevention', 'YES'],
        ['EEG channel to FPC channel wall separation', 'F-03', '≥ 2 mm', 'Min only',
         'DRC-18 (15 mm FPC/EEG)', 'YES'],
        ['EEG channel retention clip spacing', 'F-03', '80 mm', '±10 mm',
         'Cable retention', 'YES'],
        ['EEG feed-through hole diameter', 'F-03', '5 mm', '+0.3 / −0.0 mm',
         'Cable feed-through', 'YES'],
        ['Temporal wing boss receiver diameter', 'F-04', 'Ø 30 mm', '±0.3 mm',
         'Pad clip fit', 'YES'],
        ['Temporal wing boss external envelope', 'F-04', '34×34×5 mm', '±0.5 mm',
         'Pad clip engagement', 'YES'],
        ['Temporal wing snap detent undercut', 'F-04', '0.5 mm', '+0.0 / −0.1 mm',
         'Pad retention (3–6 N)', 'YES'],
        ['Temporal wing snap detent depth', 'F-04', '3 mm from surface', '±0.3 mm',
         'Snap release force', 'YES'],
        ['Temporal wing connector pocket', 'F-04', '10×6×4 mm', '+0.5 / −0.0 mm',
         'Connector clearance', 'YES'],
        ['Temporal wing silicone overmould thickness', 'F-04', '0.5 mm', '±0.1 mm',
         'Coupling comfort', 'YES'],
    ],
    col_widths=[4.5, 1.6, 2.0, 2.2, 3.5, 2.9])

# ══════════════════════════════════════════════════════════════════════════════
#  §4  MATERIAL SPECIFICATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4.  MATERIAL SPECIFICATIONS')

make_table(doc,
    ['Material', 'Application', 'Spec', 'Standard'],
    [
        ['PC/ABS alloy, UL94-V0', 'Shell body, anchor post bosses (F-01, F-02), temporal wing boss (F-04)',
         'Colour-in-bulk for zone posts (F-01). Impact resistance: ≥ 45 kJ/m² (Charpy notched). HDT ≥ 90°C.',
         'UL 94, ISO 179'],
        ['Shore 30A medical silicone', 'Zone slot plug tethers (F-01) and temporal wing boss contact face overmould (F-04)',
         'ISO 10993-5 biocompatible. UV-stable. Compression set ≤ 15% after 70h at 70°C.',
         'ASTM D395, ISO 10993-5'],
        ['Shore 40A TPE', 'Port cover tethers (F-02)',
         'Chemical resistance to IPA 70%, saline, sebum simulant. No cracking after 1,000 wipe cycles.',
         'ISO 10993-5'],
        ['PC/ABS snap-in clips', 'EEG channel cable retention clips (F-03)',
         'Same PC/ABS as shell. Clip opening force ≤ 3 N. Moulded integral to channel wall — not separate parts.',
         'UL 94'],
        ['Silicone sealant', 'Anchor post base sealing (F-01, F-02), temporal wing IPX4 seal',
         'Dow Corning RTV 738 or equivalent. Shore A 25–35 after cure. Adhesion to CFRP and PC/ABS confirmed.',
         'IEC 60529 IPX4'],
    ],
    col_widths=[3.5, 4.5, 5.8, 2.9])

# ══════════════════════════════════════════════════════════════════════════════
#  §5  SHELL DESIGN REVIEW CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5.  SHELL DESIGN REVIEW CHECKLIST')
add_para(doc,
    'This checklist must be completed in CAD review before the shell tooling '
    'specification is released to the tooling manufacturer. Each item requires a '
    'named reviewer and a pass/fail with supporting evidence (CAD screenshot, '
    'measurement, or DFM confirmation letter). Items marked BLOCKING must pass '
    'before shell tooling is cut.')

make_table(doc,
    ['#', 'Check Item', 'Method', 'Pass Criterion', 'Feature', 'Result'],
    [
        ['SH-01', 'All 5 zone slot anchor posts (F-01) present in shell CAD; '
                  'positions adjacent to ZM-01 through ZM-05 slots',
         'CAD inspection', '5 bosses present, one per zone slot', 'F-01', '☐'],
        ['SH-02', 'Zone post clearance to slot edge ≥ 3 mm for all 5 posts',
         'CAD measurement', '≥ 3 mm at each post', 'F-01', '☐'],
        ['SH-03', 'Zone post colour-code confirmed in bulk material per zone '
                  '(Blue/Red/Green/Yellow/Purple); matches NP-TOOL-ZM-001 §3.1',
         'Material spec review', 'Colour in bulk, matches zone module colour', 'F-01', '☐'],
        ['SH-04', 'Zone post through-slot geometry allows Shore 30A tether loop '
                  'passage; loop retained on pull (prototype test)',
         'Physical prototype', 'Loop passes; 5 N tug × 10 cycles — no detachment', 'F-01', '☐'],
        ['SH-05', 'All 5 zone post locations maintain IPX4 at shell (no water '
                  'ingress path at boss base)',
         'IPX4 test on prototype', 'IPX4 pass at all 5 post locations', 'F-01', '☐'],
        ['SH-06', 'All 3 accessory port anchor posts (F-02) present in hub shell '
                  'CAD; one per port (USB-C, left accessory, right accessory)',
         'CAD inspection', '3 bosses present, one per port', 'F-02', '☐'],
        ['SH-07', 'Port post to port centreline clearance ≥ 15 mm for all 3 ports',
         'CAD measurement', '≥ 15 mm at each port', 'F-02', '☐'],
        ['SH-08', 'Port cover tether loop can swing port cover clear of port opening '
                  'for full USB-C plug insertion without tether obstruction',
         'CAD assembly simulation', 'Cover clears port; plug insertion unobstructed', 'F-02', '☐'],
        ['SH-09', 'EEG cable routing channel (F-03) present in shell CAD: '
                  '8×5 mm, outer CFRP surface, routing from Hub PCB to all 8 '
                  'electrode pod locations',
         'CAD inspection', 'Channel present; branches to all 8 pods', 'F-03', '☐'],
        ['SH-10', 'EEG channel on OPPOSITE shell surface from zone module FPC '
                  'bundle (DRC-18 routing separation)',
         'CAD cross-section', 'FPC bundle: scalp-facing interior. EEG channel: outer CFRP interior. Confirmed different surfaces.', 'F-03', '☐'],
        ['SH-11', 'EEG channel wall separation from FPC channels ≥ 2 mm at all '
                  'cross-sections (OI-10 in NP-DRV-SHELL-001)',
         'CAD measurement', '≥ 2 mm at tightest point', 'F-03', '☐'],
        ['SH-12', 'FPC-to-EEG copper separation ≥ 15 mm at all points (DRC-18a). '
                  'If < 15 mm: barrier design specified in same CAD view (DRC-18b)',
         'CAD overlay measurement', '≥ 15 mm throughout, or barrier specified', 'F-03', '☐'],
        ['SH-13', '3 integral snap-in cable retention clips present in EEG channel '
                  'at ≈80 mm intervals; clip opening force ≤ 3 N on prototype',
         'CAD inspection + prototype', 'Clips present; ≤ 3 N opening force', 'F-03', '☐'],
        ['SH-14', 'EEG feed-through holes (5 mm Ø) present at each electrode pod '
                  'location; chamfer 0.5 mm × 45° on both faces',
         'CAD inspection', 'Holes at all 8 pods; chamfer on both faces', 'F-03', '☐'],
        ['SH-15', 'Temporal wing boss (F-04) present in left and right temporal '
                  'wing CAD; at mastoid process anatomical location',
         'CAD inspection', '2 bosses present (L+R); mastoid position confirmed', 'F-04', '☐'],
        ['SH-16', 'Temporal wing boss receiver diameter = Ø 30 mm ±0.3 mm; '
                  'depth ≥ 5 mm; snap detent at 3 mm depth',
         'CAD measurement', 'Ø 30±0.3 mm; depth ≥ 5 mm; detent at 3 mm', 'F-04', '☐'],
        ['SH-17', 'Temporal wing boss flush or recessed when no pad is installed; '
                  'no external protrusion; no exposed sharp edges (entry radius ≥ 1.5 mm)',
         'CAD cross-section', 'Flush or recessed; radii ≥ 1.5 mm confirmed', 'F-04', '☐'],
        ['SH-18', 'Connector pocket (10×6×4 mm) present at inferior boss position; '
                  '≥ 1.5 mm clearance around reference connector body',
         'CAD measurement', 'Pocket present; ≥ 1.5 mm clearance', 'F-04', '☐'],
        ['SH-19', 'Temporal wing boss load path: FEA or analytical calculation '
                  'confirms ≥ 15 N axial pull and ≥ 10 N lateral shear without '
                  'visible deformation (OI-05)',
         'FEA or hand calc', '≥ 15 N axial; ≥ 10 N lateral without deformation', 'F-04', '☐'],
        ['SH-20', 'All critical dimensions in §3 appear as drawing callouts with '
                  'tolerances on shell and temporal wing drawings',
         'Drawing review', 'All 19 dimensions present with tolerances', 'All', '☐'],
        ['SH-21', 'Tooling manufacturer DFM review complete: all draft angles ≥ 1°–1.5°; '
                  'no undercuts outside of intentional snap features; gate locations '
                  'confirmed; written DFM sign-off on file',
         'DFM letter from supplier', 'Written DFM sign-off received', 'All', '☐'],
        ['SH-22', 'Cross-check with NP-DRV-SHELL-001: DRC-22 (EEG channel in CAD) '
                  'and DRC-23 (DRC-18 with EEG routing locked) both PASS',
         'NP-DRV-SHELL-001 checklist', 'DRC-22 PASS; DRC-23 PASS', 'F-03', '☐'],
        ['SH-23', 'Cross-check with NP-TOOL-ZM-001: zone module mechanical key '
                  'geometry (F-01 of NP-TOOL-ZM-001) physically prevents wrong-zone '
                  'insertion even with anchor post present — no interference',
         'CAD assembly test', 'Wrong-zone insertion blocked; post does not interfere', 'F-01', '☐'],
    ],
    col_widths=[1.2, 6.2, 2.8, 3.5, 1.3, 1.7])

# ══════════════════════════════════════════════════════════════════════════════
#  §6  FAI CROSS-REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6.  FAI CROSS-REFERENCE')
add_para(doc,
    'The following first-article inspection items validate features in this tooling '
    'specification. FAI must pass before production lot approval.')

make_table(doc,
    ['Feature', 'FAI Item', 'Pass Criterion'],
    [
        ['F-01 (zone anchor posts)', 'FAI-CV-01',
         'All 5 anchor post colours match zone module colours (visual comparison + spec)'],
        ['F-01 (zone anchor posts)', 'FAI-CV-02',
         'IPX4 splash test passes at all 5 post locations — no water ingress after 14 kPa, 10 min'],
        ['F-01 (zone anchor posts)', 'FAI-CV-03',
         'Tether loop 5 N tug × 500 cycles on Shore 30A loop — no fracture, no detachment from boss'],
        ['F-01 (zone anchor posts)', 'FAI-CV-04',
         'Zone module insertion/removal 20 cycles per zone — no interference from anchor post at any zone'],
        ['F-02 (port anchor posts)', 'FAI-CV-05',
         'Port cover tether allows full cable insertion at each of 3 ports — plug insertion unobstructed'],
        ['F-02 (port anchor posts)', 'FAI-CV-06',
         'IPX4 pass at all 3 port post locations'],
        ['F-03 (EEG channel)', 'FAI-EEG-01',
         'EEG cable bundle routes from Hub PCB to all 8 electrode pod locations without bending radius < 10 mm'],
        ['F-03 (EEG channel)', 'FAI-EEG-02',
         'DRC-18 oscilloscope test: EEG artifact < 5 µVpp at all frequencies with all zone LEDs at full PWM (per NP-DRV-SHELL-001 DRC-18c)'],
        ['F-03 (EEG channel)', 'FAI-EEG-03',
         'All 3 retention clips present; each clip retains cable at 5 N lateral pull without release'],
        ['F-04 (temporal wing boss)', 'FAI-TW-01',
         'PROVISIONAL — execute only if HOPE Phase 3 positive. Reference 30 mm pad clip inserted and removed 50× — snap detent releases at 3–6 N axial pull throughout'],
        ['F-04 (temporal wing boss)', 'FAI-TW-02',
         'PROVISIONAL — boss flush/recessed with no pad installed; no sharp edges (Ra ≤ 3.2 µm at entry)'],
    ],
    col_widths=[3.2, 2.4, 11.1])

# ══════════════════════════════════════════════════════════════════════════════
#  §7  OPEN ITEMS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7.  OPEN ITEMS')
add_notice(doc,
    'BLOCKING: None of the items below can be resolved in this document alone — they '
    'require CAD model updates, supplier engagement, or prototype test results. '
    'Items marked BLOCKING must be resolved before shell tooling first-cut.',
    colour='FDECEA')

make_table(doc,
    ['ID', 'Item', 'Owner', 'Required By', 'Status'],
    [
        ['OI-01', 'Shell CAD model updated with all 5 zone anchor posts (F-01), '
                  '3 port anchor posts (F-02), EEG cable channel (F-03), and '
                  'bilateral temporal wing bosses (F-04). SH-01 through SH-19 '
                  'cannot be verified until CAD model exists.',
         'Mechanical Eng', 'Shell tooling spec release', 'OPEN'],
        ['OI-02', 'SH-21 tooling manufacturer DFM review: confirm all anchor post '
                  'bosses (F-01, F-02) and temporal wing boss (F-04) draft angles '
                  '≥ 1° achievable; no unintended undercuts; gate locations acceptable.',
         'Mech Eng + Tooling', 'G2 Pre-Tooling gate', 'OPEN'],
        ['OI-03', 'EEG cable routing channel (F-03) CAD model update in NP-DRV-SHELL-001 '
                  '§2.4 (OI-09 in NP-DRV-SHELL-001, G2-11 in NP-COORD-001). DRC-22 and '
                  'DRC-23 cannot close until this CAD model exists. BLOCKING for '
                  'tooling release.',
         'ME + EE', 'G2 gate (BLOCKING)', 'OPEN — RISK-21'],
        ['OI-04', 'Zone post colour material confirmation: written confirmation from '
                  'PC/ABS moulding supplier that 5 distinct zone colours (Blue, Red, '
                  'Green, Yellow, Purple) are achievable in bulk material (not paint) '
                  'with saturated hue. Colour palette to pass accessibility check '
                  '(OI-05 in NP-TOOL-ZM-001 — same palette applies to posts).',
         'Mech Eng + Design', 'G1 Pre-Layout gate', 'OPEN'],
        ['OI-05', 'Temporal wing boss load analysis (SH-19): FEA or hand calculation '
                  'confirming ≥ 15 N axial pull and ≥ 10 N lateral shear without '
                  'deformation on PC/ABS boss. Required before boss geometry is frozen '
                  'in tooling spec.',
         'Mechanical Eng', 'G2 Pre-Tooling gate', 'OPEN'],
        ['OI-06', 'Tether loop supply: confirm Shore 30A silicone tether loop (F-01) '
                  'and Shore 40A TPE tether loop (F-02) suppliers and part numbers. '
                  'Confirm loop can be assembled through post through-slot before '
                  'cover assembly. Confirm ≥ 500-cycle tug test with production '
                  'material.',
         'Supply Chain + ME', 'G2 gate', 'OPEN'],
        ['OI-07', 'Temporal wing boss dust cover tether: same anchor post spec (boss '
                  'and tether) for the dust cover that seals the F-04 pocket when no '
                  'pad is installed. Confirm location on temporal wing for tether '
                  'anchor without interfering with wing snap-on mechanism.',
         'Mechanical Eng', 'G2 gate', 'OPEN'],
        ['OI-08', 'HOPE Phase 3 monitoring: track Cognito Therapeutics (n=670) results. '
                  'Expected mid-2026. If positive: initiate mastoid LRA pad engineering '
                  'and confirm F-04 boss geometry is compatible with final pad clip design. '
                  'If negative: mark F-04 as inactive in future revisions (boss remains '
                  'in tooling — no tooling change required).',
         'Product / Programme Mgr', 'Mid-2026', 'MONITORING — no action required yet'],
    ],
    col_widths=[1.2, 8.5, 2.5, 2.8, 1.7])

# ══════════════════════════════════════════════════════════════════════════════
#  §8  SIGN-OFF
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8.  SIGN-OFF')
add_para(doc,
    'This tooling specification is complete when all design review checklist items '
    '(§5, SH-01 through SH-23) pass and all blocking open items (§7) are resolved. '
    'Sign-off is required from Mechanical Engineering and Hardware Electrical '
    'Engineering leads before the shell tooling specification is released to the '
    'tooling manufacturer.')

make_table(doc,
    ['Role', 'Name', 'Date', 'Signature / Approval'],
    [
        ['Mechanical Engineering Lead', '', '', ''],
        ['Hardware EE Lead', '', '', ''],
        ['Manufacturing Engineering', '', '', ''],
        ['Product / Programme Manager', '', '', ''],
    ],
    col_widths=[4.5, 3.5, 2.5, 6.2])

# ══════════════════════════════════════════════════════════════════════════════
#  §9  RELATIONSHIP TO CLAUDE.MD PENDING DECISIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9.  CLAUDE.MD PENDING DECISIONS — CLOSURE STATUS')
add_para(doc,
    'The following items from CLAUDE.md §13.4 (Pending decisions — must resolve '
    'before tooling is cut) are addressed by this document. Items below are '
    'PARTIALLY CLOSED — specification is written but CAD verification and supplier '
    'confirmation (OI-01 through OI-08) remain OPEN.')

make_table(doc,
    ['CLAUDE.md §13.4 Item', 'Status after this document', 'Remaining action'],
    [
        ['Shell tooling: anchor posts for all interface covers '
         '(5 zone + 3 port positions, colour-coded for zones)',
         'SPECIFICATION WRITTEN — F-01 (§2.1) and F-02 (§2.2) define geometry, '
         'colour, dimensions, material, and FAI requirements',
         'CAD model update (OI-01); DFM review (OI-02); colour confirmation (OI-04)'],
        ['Shell tooling: EEG cable channel (§2.4)',
         'SPECIFICATION WRITTEN — F-03 (§2.3) specifies 8×5 mm channel, retention '
         'clips, feed-throughs, wall separations, DRC-18 cross-reference',
         'CAD model update (OI-01, OI-03); DRC-22 and DRC-23 close in NP-DRV-SHELL-001'],
        ['Goggle arm tooling: anchor hook for lens rim guard tether',
         'OUT OF SCOPE — lens rim guard anchors are on the goggle arm, not headset '
         'shell. Separate tooling document required for goggle arm.',
         'Create NP-TOOL-GOGGLE-001 for goggle arm anchor hook spec'],
        ['Hub tooling: probe dock + anchor posts + large-radius Boa cable channel '
         '+ tool-free fan (quarter-turn captive fastener)',
         'OUT OF SCOPE — hub tooling. Separate document required.',
         'Create NP-TOOL-HUB-001'],
    ],
    col_widths=[4.2, 6.5, 6.0])

# ══════════════════════════════════════════════════════════════════════════════
#  §10  REVISION HISTORY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, '10.  REVISION HISTORY')
make_table(doc,
    ['Rev', 'Date', 'Author', 'Description'],
    [
        ['A', '2026-05-10', 'NeurOne Engineering',
         'Initial release. Addresses Issue #16 (shell tooling: anchor posts 5 zone + '
         '3 port colour-coded + EEG cable channel + temporal wing anchor boss). '
         'Features F-01 through F-04 specified. 23-item design review checklist '
         '(SH-01 through SH-23). Companion to NP-DRV-SHELL-001 Rev A and '
         'NP-TOOL-ZM-001 Rev A. CLAUDE.md §13.4 shell tooling pending decisions '
         'addressed (CAD verification remains open — OI-01 through OI-07).'],
    ],
    col_widths=[1.2, 2.2, 3.8, 9.5])

# ── save ──────────────────────────────────────────────────────────────────────
out_path = '/home/user/NeurOne/docs/neurone_tool_shell_001.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
