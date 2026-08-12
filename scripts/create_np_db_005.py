"""
Generate NP-DB-005 Rev 6 — NeurOne Master Design Brief
Revision 6 of NP-DB-005, superseding Revision 5 (2026-05-16) of the same document.
Written in place to docs/np_db_005.docx per NP-CONV-001 Rev 6 §4.0 (the filename is the
document serial and nothing else) and §4.3 (no revision information in filenames).
Revision 5's content is retained by version control, not by a second file.

Every figure in this document traces to CLAUDE.md Rev 33, docs/status/*.md,
docs/np_dhf_001.md, or a named NP-* controlled document. Nothing is carried
forward from Rev 5 without re-verification against a source that still says so.

House format follows scripts/create_np_tool_shell_001.py.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

HDR_BG = 'D9E2F3'
RETIRE_BG = 'FBE4E4'
OK_BG = 'E2EFDA'
NOTICE_BG = 'FFF3CD'


# ── helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_colour):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)


def cell_text(cell, text, size=8.5, bold=False, align=None, colour=None):
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(1)
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if colour:
        run.font.color.rgb = RGBColor(*colour)


def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_para(text, size=9.5, bold=False, italic=False, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullets(items, size=9.5):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(it)
        run.font.size = Pt(size)


def add_notice(text, colour=NOTICE_BG, bold_lead=None):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, colour)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = para.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(9.5)
    r = para.add_run(text)
    r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_table(headers, rows, widths=None, header_bg=HDR_BG, row_bg=None, size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, header_bg)
        cell_text(c, h, size=size, bold=True)
    for r_i, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            if row_bg:
                set_cell_bg(cells[i], row_bg)
            cell_text(cells[i], v, size=size, bold=(i == 0))
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NeurOne')
r.bold = True
r.font.size = Pt(30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Master Design Brief — Revision 6')
r.bold = True
r.font.size = Pt(17)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Closed-Loop Multi-Modal Neuromodulation Platform')
r.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NP-DB-005 Rev 6  ·  2026-08-11  ·  Pre-tooling design phase')
r.font.size = Pt(9.5)
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('STATUS: DRAFT — PENDING APPROVAL')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xB0, 0x20, 0x20)

add_notice(
    'This revision has been authored but NOT reviewed or approved. 21 CFR 820.40(a) requires '
    'review and approval by a designated individual PRIOR to issue, and that has not happened. '
    'Until the approval block below is signed, Revision 5 of this document remains the last APPROVED '
    'revision of record and supersession does not take effect — but Revision 5 is already known '
    'to contain retired design (see §0.1), so neither document should be used as a design input '
    'for the retired areas without reading §0.1 first. Approving this revision is the action '
    'that closes that gap.',
    colour=RETIRE_BG,
    bold_lead='APPROVAL STATUS. ')

add_table(
    ['Role', 'Name', 'Signature', 'Date'],
    [
        ['Author', 'NeurOne Design Program', '(authored 2026-08-11)', '2026-08-11'],
        ['Reviewed by', '', '', ''],
        ['Approved by', 'Steve Hickman (CEO, interim Quality authority)', '', ''],
    ],
    widths=[3.4, 6.4, 4.4, 2.6])

add_notice(
    'On approval, this revision supersedes Revision 5 of NP-DB-005 (2026-05-16), which must '
    'not thereafter be used as a design input. Per NP-CONV-001 Rev 6 §4.0.4 and §4.3, Revision 6 '
    'is written in place at docs/np_db_005.docx and Revision 5 is retained by version control — '
    'a superseded REVISION of a live serial is not a second file, unlike a superseded DOCUMENT, '
    'which moves to docs/superseded/. Revision 5\'s headline zone-module engineering — the fixed five-slot, '
    '600-LED module architecture and everything addressed to it — was retired on 2026-07-28 '
    'and replaced, not revised. Section 0 below is the retirement ledger; read it before any '
    'other section.',
    colour=RETIRE_BG,
    bold_lead='SUPERSESSION NOTICE. ')

add_para(
    'Sources of record for this revision: CLAUDE.md Rev 33; docs/status/completed-decisions.md; '
    'docs/status/pending-decisions.md; docs/status/document-register.md; NP-DHF-001 Rev 28 '
    '(2026-08-04). Where a figure is gated by an open decision, the gate is named inline. '
    'Where two controlled documents disagree, the disagreement is reported rather than resolved.',
    size=9, italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# 0. WHAT CHANGED
# ══════════════════════════════════════════════════════════════════════════════
add_heading('0. What Changed Since Revision 5', level=1)

add_heading('0.1 Retired since Revision 5 — dead design, do not build to it', level=2)

add_para(
    'Each row below was stated as locked design in Revision 5. Each is now retired. The '
    'replacement column names what is authoritative today.', size=9.5)

add_table(
    ['Retired in Rev 5', 'Why retired', 'What replaced it'],
    [
        ['Fixed 5-slot zone-module architecture; 600 LEDs (300×660nm + 300×808–830nm); '
         '66×78 mm module footprint; "all-5-zone kit" pricing',
         'Position-unique modules meant a distinct SKU per helmet location, and the '
         'five-slot count was never derived from the tiling surface.',
         'One universal 40 mm hexagonal module SKU tiling an ~80-socket lattice '
         '(NP-HEX-ZM-001 Rev 2). Zones are protocol-defined socket sets authored in '
         'protocols/predefined/00-zones.npps, not hardware slots. Total LED count now '
         'scales with how many T1-A tiles a build populates.'],

        ['ZONE_ID 3.3 kΩ / resistor-ladder module detection (pin 18; ZM-01=10 kΩ … '
         'ZM-05=220 kΩ) and its 3×100 ms debounce',
         'A resistor ladder can only identify a fixed, small, position-keyed set. All '
         'sockets are now mechanically identical.',
         'UID-based auto-inventory: np_module_map (firmware/hub_control/np_module_map.{h,c}) '
         'polls for UID change and auto-inventories whatever is seated, plus '
         'np_module_map_check_placement() as the wrong-placement backstop.'],

        ['Per-slot LPI2C3 bus enable (five separate 400 kHz buses, one per slot)',
         'SMART-1 (2026-07-28) requires EVERY socket to be I2C/TIA-capable, not a subset. '
         'Five dedicated peripherals do not scale to ~80 sockets.',
         'One differential cluster bus (PCA9615-class) from the hub, with transactions '
         'tunnelled through each cluster controller\'s PCA9548A. No per-socket I2C '
         'peripheral exists.'],

        ['Hub PCB Rev B: five Vishay DG2788A per-slot TIA gain switches + NXP PCA9546A '
         'I2C mux + GAIN_SEL[0..4] GPIO on GPIO_B0_04..08',
         'Sized for five slots. Scaling the Rev B star to 80 sockets needs ~880 conductors '
         'through one blind-mate boss, 160 TIA channels and 160 ADC channels against ~32 '
         'available on the RT1062.',
         'Hub PCB Rev C (NP-HW-HUB-001 Rev 4, DRAFT): distributed cluster-controller tier. '
         'Rev 2 (Hub PCB Rev B) is archived verbatim as Appendix A; its DG2788A topology and TIA-saturation '
         'analysis carry forward onto the cluster carriers — only the count and the control '
         'source change. Zero GAIN_SEL GPIO remain.'],

        ['RISK-15 five-layer zone keying, layers 2–5 (ZONE_ID resistor, per-socket braille '
         'and raised numeral, shell tactile dot counts, bone-conduction announcement at '
         'insertion)',
         'The lattice tessellates: the shell surface between two sockets is a structural web '
         'a few millimetres wide, and an ISO 17049 braille cell needs ~6 × 10 mm — the '
         'artefact does not physically fit. An 80-member tactile-symbol set is not '
         'discriminable by touch. And bone conduction requires head contact, which is '
         'unavailable at the moment of insertion (the helmet must be off-head to reach a socket).',
         'NP-HFE-002 Rev 1 (2026-07-31): tactile marking follows TYPE onto the module (a '
         'closed 3–4 member set, 1–4 raised bars); POSITION moves to companion-app guided '
         'placement as the primary channel, over a tactile landmark grid on the shell. '
         'Bone conduction is retained for while-worn readback only.'],

        ['Multi-FPC bundle management (RISK-17): ≥2 mm inter-FPC separation, 3 anchor bosses '
         'per FPC, all five hub ZIF connectors on one PCB edge; ≥15 mm FPC-to-EEG separation '
         '(achieved 18–22 mm)',
         'Sized for five large discrete FPC tails. Hex tiles have no tail at all. Separately, '
         '≥15 mm EEG separation is unachievable in the hex architecture because electrodes now '
         'live inside T1-B tiles on the same L1 carrier as LED drive current.',
         'NP-DRV-SHELL-002 Rev 2 (2026-07-29) cluster-carrier interconnect. Modules present an '
         '18-contact back-face pad array — the module interconnect has zero dynamic-flex paths. '
         'The <5 µVpp injected-artifact requirement that separation served is retained verbatim '
         '(DRC-18c → SH2-DRC-16) and met by four other mechanisms.'],

        ['Safety MCU described as "~500 lines bare-metal"',
         'Understated by roughly 3×; corrected 2026-08-06.',
         '~1,600 physical lines across 9 modules as of 2026-08. '
         '`wc -l firmware/safety_mcu/src/*.c` is the source of truth, not any document.'],

        ['Per-zone safety enable bits NP_SAFETY_EN_PBM_ZONE_0..4',
         'Five was the retired slot count, and "zone" is a noun firmware must not model at '
         'all — a zone is a human-authored socket set whose membership can change with no '
         'hardware change. Because these symbols gate PBM and tES, a stale grouping is a '
         'wrong-site stimulation path.',
         'Single NP_SAFETY_EN_PBM_CRANIAL bit (bit 0); bits 1–4 reserved and excluded from '
         'NP_SAFETY_EN_ALL_MASK (0x3FFF → 0x3FE1), so bits 1–4 are stripped in '
         'np_spi_watchdog_tick and can never enable anything. Note what that guard is NOT: no hub '
         'can set a retired zone bit — no hub hardware exists, and the ZONE macros were deleted in '
         'the same change — so the exclusion is defence in depth against a future authoring error, '
         'not compatibility with a deployed hub (NP-HW-HUB-001 Rev 4 §7.2). The reservation itself '
         'binds on separate Class C grounds: enable-bit position ≡ current_ua[] slot ≡ '
         'charge-accumulator index, which the 40 µC/cm² limit rests on. PA4 double-assignment '
         'cleared.'],

        ['T2 clinical tACS as "16-channel arbitrary waveform"',
         'The 16-channel figure had no sourcing document, no BOM line and no named silicon; '
         'the placeholder mapping wrapped electrodes 16–20 back onto channels 11–15, aliasing '
         'geometric neighbours and making the M1_R ring undeliverable.',
         '21 channels, one per cap electrode. The T2 cap always carried 21 conductors; only '
         'the driver was 16-channel. Constraint of record if cost ever reintroduces sharing: '
         'never alias two electrodes within one ring radius (~60 mm).'],

        ['Single risk register of 25 risks in NP-RISK-001 (23 MITIGATED, 2 OPEN)',
         'NP-RISK-001 held one register for one artifact — the Zone Module FPC — and was then '
         'asked to hold the whole system. The artifact it was scoped to no longer exists.',
         'NP-RISK-001 is RETIRED to docs/superseded/ and re-baselined into three successors '
         'split by artifact, matching NP-ART-001 §2: NP-RISK-002 (risk file index, NP-RISK-001 '
         'disposition, and risks with no artifact home yet), NP-RISK-003 (hex-tile module — '
         'shell, window, gasket, FPC, emitters, metering), NP-RISK-004 (shell, socket, '
         'interconnect and hub). The RISK-NN sequence is append-only and CLOSED AT 26: no new '
         'hazard takes an RISK-NN number and no retired number is ever reused, because a reused '
         'number would silently re-point all thirty-eight inbound references. New hazards take '
         'artifact-prefixed IDs — RISK-HEX-nn, RISK-SHELL-nn, RISK-HUB-nn.'],

        ['PDMS–PI bond peel strength "≥150 N/m"',
         'Superseded by the measured range now carried in CLAUDE.md §3.',
         '174–860 N/m peel force from the 75 nm SiO₂ interlayer (RF magnetron sputter) plus '
         'O₂ plasma activation. The 200-cycle IEC 60068-2-14 qualification (FAI-TC02) remains '
         'BLOCKING before any production FPC.'],

        ['ACC presented as one of seven focal 4×1 HD-tDCS targets',
         'ACC (0, 28, 28) is 47.1 mm from Fz; its 2nd/3rd nearest electrodes are 49.1 mm and '
         'no 10-10 scalp position gets closer than ~37.9 mm (Fpz). The ≤35 mm focality '
         'criterion is unsatisfiable for ACC by ANY electrode placement.',
         'Targets carry NP_HD_TARGET_DEPTH_SURFACE / _DEEP. ACC stays selectable and keeps its '
         'anxiety indication, but a 4×1 there is indirect network modulation and must never be '
         'presented as focal stimulation of the structure. The other six targets are '
         'cortical-surface at 11–29 mm.'],
    ],
    widths=[4.6, 5.6, 6.6],
    header_bg=RETIRE_BG)

add_notice(
    'Firmware and documents carrying retired mechanisms are individually banner-marked at '
    'their source: NP-FW-PBM1064-001 Rev 3, NP-HW-FPC-001, NP-FW-ZA-001, NP-HW-HUB-001 '
    '(Rev 2 → Appendix A), NP-DRV-SHELL-001 Rev 2, NP-PROC-FPC-1064-001, and the '
    'neurone_fpc_zone_module_spec_revA.docx / docs/superseded/np_drv_shell_001.docx pair. '
    'The zone_announce DDS audio engine and debounce algorithm remain reusable but need '
    're-triggering from module-map events.',
    colour=NOTICE_BG)

add_heading('0.2 Added since Revision 5', level=2)

add_table(
    ['Area', 'What was added'],
    [
        ['Hex-tile module architecture',
         'NP-HEX-ZM-001 Rev 2 (design study, Option A committed): one universal 40 mm '
         'flat-to-flat hexagonal module, median curvature R_m = 87 mm, tiling an ~80-socket '
         'lattice recovered from a direct 3D scan of the reference helmet interior. Three T1 '
         'tile types on one mold. Principal directions CLUSTER-1, SYM-1 and CONTIG-1 recorded. '
         'Companion ISA at docs/np_hex_zm_isa.md; firmware np_module_map delivered.'],

        ['Hex-tile electrical / FPC spec',
         'NP-HW-HEXTILE-001 Rev 3 — the electrical and FPC counterpart for T1-A and T1-C; '
         '§8.2.1 carries the 18-cluster derivation and §8.2.2 the interconnect options for the '
         'count overflow. Cluster map diagram at docs/diagrams/np_hextile_cluster_map.svg.'],

        ['Hub PCB Rev C',
         'NP-HW-HUB-001 Rev 4 (DRAFT) — distributed cluster-controller tier replaces the '
         'five-slot star. The hub PCB no longer encodes the socket count at all, so REG-1 no '
         'longer blocks hub tooling.'],

        ['Shell socket interconnect',
         'NP-DRV-SHELL-002 Rev 2 (DRAFT) — new document number because the architecture is '
         'replaced, not revised. Five networks split by physics; 28-item DRC checklist; three '
         'EMI prohibitions locked.'],

        ['Helmet geometry from measurement',
         'NP-HELMET-GEOM-001 — socket coordinates placed on the measured scan surface and the '
         'helmet frame recovered from the actual scan, replacing the interim ellipsoid '
         'description. Companion ISA at docs/np_helmet_geom_isa.md.'],

        ['Thermal architecture',
         'NP-THERM-CFD-001 / -R1-001 / -C2-001 and NP-THERM-BEZEL-001: BN-boss conductive '
         'thermal export adopted as the base thermal design with a stagnant inter-bowl cavity; '
         'THERM-1a fan-health path selection (Path A rejected, Path B1 selected); '
         'NP-REQ-FANHEALTH-001 and NP-PLAN-FANHEALTH-001.'],

        ['PBM optical resolution floor',
         'NP-OPT-PSF-001 Rev 1 — Monte Carlo optical model establishing what boundary the '
         'hardware can actually produce at cortical depth. Now the required reference for zone '
         'sizing, lateralized protocols and any "targets region X" claim.'],

        ['Accessible position identification',
         'NP-HFE-002 Rev 1 — layered, app-primary design replacing RISK-15 layers 3/4/5; '
         'requirements HFE-R-01..15 enter NP-HFE-001\'s formative and summative plans.'],

        ['Human factors, design planning, traceability, post-market',
         'NP-HFE-001, NP-DP-001 (Design and Development Plan), NP-DT-001 (Input/Output '
         'Traceability Matrix, Rev 2 adds DI-SAFE-13), NP-PMS-001 (Post-Market Surveillance), '
         'NP-QMS-DC-001, NP-QMS-CAPA-001 — the QMS documents Rev 5 listed only as future targets.'],

        ['System FMEA',
         'NP-FMEA-001 and the geometry feeder NP-FMEA-GEOM-001, including FMEA-G07-01 '
         '(fan/heatsink loss) which became RISK-26, and FMEA-M08-04 (PA4 double-assignment).'],

        ['Privacy program',
         'NP-PRIV-001 Rev 1 (18 findings: 2 Critical, 6 High, 7 Medium, 3 Low) and '
         'NP-PRIV-REM-001 Rev 2 (30-step remediation calendar, extended by STEP-31/32 and '
         'STEP-34/35). Second- and third-pass code audits NP-PRIV-ANALYSIS-002 / -003. '
         'NP-PRIV-NOTICE-001 Rev 3, NP-PRIV-AUDIT-001, NP-PRIV-KEYSPLIT-001. '
         'Supporting: NP-SEC-BR-001 (breach response), NP-PROC-POA-001 (POA upload), '
         'NP-APP-TELEMETRY-001, NP-LEGAL-BAA-001, NP-IRB-001, NP-FW-ANON-001.'],

        ['Firmware supply chain under IEC 62304',
         'FreeRTOS-Kernel V11.3.0 (FreeRTOS-LTS 202604.00) vendored as Class B SOUP; '
         'ARM CMSIS-Core(M) 5.6.0 and ST CMSIS-Device STM32G0 v1.4.5 vendored as Class C SOUP; '
         'Monocypher 4.0.2 behind np_crypto. NP-SOUP-CMSIS-001 Rev 1 is the first IEC 62304 '
         '§7.1.2 anomaly evaluation performed for Class C software at NeurOne.'],

        ['Firmware cross-compile CI',
         'NP-SW-CI-001 — phased program taking the bootloader, main-firmware and Class C '
         'safety-MCU compile legs from advisory to gating, each phase demonstrated by a '
         'deliberately-broken probe and then reverted. Defects A–D closed.'],

        ['Hub control program',
         'firmware/hub_control/ — module registry, session runner, safety SPI, session-log '
         'eMMC backing (np_log_backend), BLE/USB-C protocol transport (np_transport), and '
         'np_mod_cvns as the sole hub↔cervical-VNS integration layer.'],

        ['Protocol language and reference',
         'NP-NPPS-REF-001 — NPPS protocol scripting reference; socket-indexed session '
         'descriptor v5; 00-zones.npps as the single home of zone definitions after ZONE-1 '
         'deleted the "lobe" concept from firmware, generator and web app.'],

        ['Naming and notation conventions',
         'NP-CONV-001 Rev 6 — the first consolidated statement of NeurOne naming rules, written '
         'after three name collisions reached the document set in one week (GUARD, /ALERT, '
         'SEAT_N), none caught by reading and all three caught by diffing tables mechanically. '
         'Binds active-low "#" suffixes (§1.1), the forbidden polarity suffixes and why each is '
         'already taken (§1.2), bracket index notation (§1.3), document IDs and filenames (§4), '
         'integer revisions (§4.1), and no revision information in filenames (§4.3). Enforced by '
         'scripts/check-doc-filenames.ts, which was falsified in both directions before it was '
         'trusted. THIS REVISION IS THE FIRST ISSUE OF THE DESIGN BRIEF UNDER THOSE RULES — see '
         '§12 for what changed about this document\'s own identity.'],

        ['Risk file re-baseline',
         'NP-RISK-002 / -003 / -004 replace the retired NP-RISK-001, split by artifact. See §0.1 '
         'and §8.'],

        ['Artifact register and FAI programme',
         'NP-ART-001 (artifact register — §4 found nine artifacts with no owning document, which '
         'is the gap-detection property the filename convention buys), NP-FAI-001, '
         'NP-FAI-HUB-001, NP-MOD-ID-001 (module identity), NP-REV-SHELL-001 (shell design '
         'review), NP-TOOL-HEXTILE-001 (hex-tile tooling, successor to the retired '
         'NP-TOOL-ZM-001 and NP-TOOL-ZM-SM-001).'],

        ['Clinical / regulatory additions',
         'NP-REG-CVNS-001 (510(k) Q-Sub), NP-REG-PBM1064-001 (RISK-03 scope expansion), '
         'NP-BIB-1064-001, NP-COND-LINK-001, NP-INT-FHIR-001, NP-API-001, NP-ANALYTICS-001, '
         'NP-INFRA-001, NP-ENV-001 / NP-ENV-OPRANGE-001, NP-PWR-BUDGET-001, '
         'NP-FEAS-FNIRS-001 (exploratory), NP-TOOL-HUB-001 Rev 1 (hub enclosure tooling).'],
    ],
    widths=[4.2, 12.6],
    header_bg=OK_BG)

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('1. Executive Summary', level=1)

add_para(
    'NeurOne is a two-tier closed-loop multi-modal neuromodulation platform. Tier 1 '
    '(NeurOne Home) is an FDA-exempt consumer wellness device combining 8 neuromodulation '
    'modalities in a single wired-first headset. Tier 2 (NeurOne Pro) is a clinical-grade '
    'device targeting FDA 510(k) clearance with 11 modalities including TMS, clinical tACS, '
    '21-channel qEEG, sLORETA-guided HD-tDCS, and cervical VNS. Both tiers share one chassis, '
    'one processor stack, one app, and USB-C connectivity.')

add_table(
    ['Tier', 'Name', 'Regulatory', 'Modalities', 'Price range', 'Timeline'],
    [
        ['T1', 'NeurOne Home', 'FDA-exempt wellness', '8', '$449–$1,199', '12–18 months'],
        ['T2', 'NeurOne Pro', 'FDA 510(k) target', '11', '$4,999–$13,999 + $1,800/yr',
         '18–36 months post T1'],
    ],
    widths=[1.4, 3.0, 3.4, 2.0, 3.6, 3.4])

add_para('Founding design principles (all locked):', bold=True)
add_bullets([
    'Shared platform — one production line, two markets (T1 consumer wellness + T2 clinical).',
    'Wired-first USB-C default — zero RF at scalp during tethered operation.',
    'Autonomous closed-loop EEG-adaptive stimulation without phone (primary competitive moat).',
    '5-layer EMF shielding + active Helmholtz cancellation — the only consumer brain wearable '
    'with measured, continuously fleet-verified shielding.',
    'Modular field-upgradeability — now via one universal hexagonal module SKU rather than '
    'position-unique zone modules.',
    'No mandatory subscription — all core functions offline-capable permanently.',
    'UHDR/SHDR data separation — user health data is never accessed by NeurOne under any '
    'circumstances.',
    'QMS established at company formation (2026-05-13) under 21 CFR Part 820 / ISO 13485:2016; '
    'DHF index at NP-DHF-001 Rev 28 (2026-08-04).',
])

add_para(
    'T1 — 8 modalities: transcranial PBM (660–670 + 808–830 nm, tiled across the hex-socket '
    'lattice) · intranasal PBM (bilateral Y-probe, 2λ) · 8-channel semi-dry EEG neurofeedback '
    '(24-bit ADS1299) · BES/tACS (consumer name: Brainwave Entrainment Stimulation) · tDCS '
    '(consumer name: Cortical Priming Stimulation) · auricular VNS + HRV + HRV biofeedback · '
    'neural audio entrainment (planar magnetic + bone conduction) · visual stimulation '
    '(108 micro-LEDs per lens, EMDR, retinal PBM Mode F).')

add_para(
    'T2 adds: 21-channel qEEG wet gel (10-20 + FC3/FC4 + Oz + A1/A2) · 1170 nm deep PBM '
    '(laser, 35–40 mm) · clinical tACS (≤4 mA, 21 channels, one per cap electrode) · '
    'sLORETA-guided HD-tDCS (4×1 ring) · TMS focal figure-8 coil (0.1–0.5 T, rTMS + TBS) · '
    'cervical VNS accessory (tcVNS) · HIPAA cloud, FHIR R4, multi-patient dashboard, '
    'scripting API, sLORETA source imaging.')

# ══════════════════════════════════════════════════════════════════════════════
# 2. DATA ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('2. Data Architecture — UHDR / SHDR Framework', level=1)

add_para(
    'The UHDR/SHDR framework is the foundational data architecture. All storage code, all '
    'session recording, and all research data sharing must comply with it. It is established '
    'in firmware from the first line of storage code (NP-FW-EMMC-001, extended by '
    'NP-FW-EMMC-002).')

add_table(
    ['', 'UHDR — User Health Data Record', 'SHDR — System Health Data Record'],
    [
        ['Definition',
         'Any record describing the user\'s biological state or therapeutic response.',
         'Any record describing device performance, wear state or calibration — with nothing '
         'that reveals user biology.'],
        ['Owner', 'The user, unconditionally.',
         'NeurOne. Linked to device ID and an opaque 256-bit TRNG warranty token only — never '
         'to user identity.'],
        ['NeurOne access',
         'NEVER — not for support, engineering, research, or regulatory submission.',
         'Yes, for warranty and fleet analytics.'],
        ['Storage',
         'On-device eMMC UHDR partition (6,903 MiB). AES-256-XTS. Two-layer key scheme: UKMD '
         '(32-byte TRNG master key) wrapped by WKMD (Argon2id from biometric/PIN, m=65536 KiB, '
         't=4, p=1 per RFC 9106). A biometric change re-encrypts only the 32-byte wrapper.',
         'On-device eMMC SHDR partition (512 MiB). AES-256-XTS, HKDF-SHA256 from the NeurOne '
         'manufacturing root key + eMMC CID hash, held under SNVS OTPMK.'],
        ['Consent subject',
         'The person wearing the device. Managed per user on-device by ConsentStore.',
         'The warranty owner — the entity that registered the device, which may be a clinic or '
         'institution and is NOT assumed to be the user. Managed by SHDRUploader.'],
    ],
    widths=[2.4, 7.2, 7.2])

add_para('Boundary case resolution rule: when in doubt → UHDR. Reclassification to SHDR '
         'requires positive demonstration that the record contains no user biology.', bold=True)

add_para('Boundary resolutions of record (the last four are new since Revision 5):', bold=True)
add_table(
    ['Signal', 'UHDR', 'SHDR'],
    [
        ['EEG impedance', 'Raw impedance', 'Derived trend slope'],
        ['Accelerometer', 'During active sessions',
         'Impact events between sessions — and per NP-FW-EMMC-002 §G only drop_detected: bool '
         'and maintenance_alert: bool. Raw series, g-force values, orientation vectors, '
         'timestamps and drop counts are all prohibited.'],
        ['Ambient light', 'Raw', 'Cumulative UV exposure index'],
        ['Auricular VNS impedance', 'Raw', 'Contact resistance trend'],
        ['Session count', 'Session timestamps', 'Unsigned integer count, no timestamps'],
        ['Cervical VNS per-electrode impedance (OI-CVNS-HUB-11)',
         'Raw tissue impedance in kΩ. Transferred device-internally (safety MCU → hub) for '
         'cross-validation only.',
         'The hub-vs-MCU divergence FLAG only (NP_CVNS_SHDR_EV_IMP_CROSSVAL, no kΩ values, '
         'timestamp suppressed). Raw kΩ is NEVER written to SHDR.'],
        ['Research anonymization pipeline failed_step',
         'UHDR / app-side only — it drives the user retry prompt. A per-device count of '
         'validate failures weakly signals the wearer is a re-identification outlier (a small '
         'anonymity set), which is health-adjacent under WA MHMD and GDPR Art. 9.',
         'If a device-health signal is needed at all, only a coarse anonymization_failed: bool '
         'without the stage.'],
        ['Safety-MCU fault latch',
         '—',
         'count (distinct-fault-transition tally) unconditionally. tick_ms in general — but '
         'SUPPRESSED to 0 when the latched fault carries NP_SAFETY_STATUS_CARDIAC, because a '
         'CVNS cardiac cutoff co-locates the event time with a UHDR-class health event. '
         'Enforced by np_fault_latch_report_tick_ms() / _report_count(); a future read command '
         'must marshal through these, not read s_latch directly.'],
    ],
    widths=[3.6, 6.6, 6.6])

add_para(
    'Research data anonymization (locked). All anonymization occurs on-device within the '
    'NeurOne app before any data leaves the device. NeurOne cannot access raw UHDR at any '
    'point. Per-study flow: (1) NeurOne sends a cryptographically signed study descriptor; '
    '(2) the app reads the encrypted UHDR partition in-app and applies on-device '
    'transformations (k ≥ 10, date rounding ≥ 1 week, direct identifier removal, quasi-identifier '
    'suppression); (3) only the pre-anonymized signed extract is transmitted; (4) no persistent '
    'per-user anonymized store; (5) no linkage table exists. Consent withdrawal permanently '
    'blocks the device from processing future study descriptors, for any data period including '
    'sessions predating withdrawal. Already-published extracts cannot be individually removed — '
    'the irreversibility notice is given at the L3 blanket consent screen AND at each '
    'per-project invitation. Planned upgrade: NP-FW-ANON-001 adds l-diversity (l ≥ 3) and '
    'differential privacy (ε ≤ 1.0), with external DP reviewer sign-off, before the first study '
    'descriptor is deployed.')

# ══════════════════════════════════════════════════════════════════════════════
# 3. CONSENT ENGINE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('3. Clinical Consent Engine', level=1)

add_heading('3.1 Two consent subjects (locked 2026-06-16)', level=2)

add_para(
    'NeurOne has two distinct consent subjects that must never be conflated. A clinic that '
    'registers a device warranty has NOT consented on behalf of any patient; a patient using a '
    'clinic-owned device has independent ConsentStore state. The two gates are '
    'code-structurally independent — SHDRUploader holds no reference to ConsentStore.')

add_table(
    ['Subject', 'Data', 'Consent granted at', 'Managed by'],
    [
        ['Warranty owner', 'SHDR fleet telemetry only', 'Device warranty registration',
         'SHDRUploader, gated by WarrantyAnalyticsGate (opaque token, no user identity)'],
        ['User', 'UHDR — EEG, HRV, PBM dose, adaptation events',
         'Per-user research consent flow (L1–L4)',
         'ConsentStore, per user, on-device; manages ResearchAnalyticsGate'],
    ],
    widths=[3.0, 4.2, 3.8, 5.8])

add_para(
    'Withdrawal scoping: withdrawing from a specific study or category stops that data flow '
    'only, and leaves app analytics untouched. Withdrawing blanket research consent (L3) stops '
    'ALL research data flows AND tears down research analytics — '
    'ConsentStore.withdrawBlanketResearchConsent() calls revokeResearchAnalytics() — because '
    'blanket withdrawal signals the user wants no data collection beyond basic device function. '
    'Revoking warranty consent has no effect on research participation, and vice versa.')

add_heading('3.2 Clinician subscription tiers', level=2)
add_table(
    ['Tier', 'Price', 'Use cases', 'UHDR elements', 'Target clinician'],
    [
        ['Monitor', '$49/mo/patient', 'Adherence monitoring, protocol compliance',
         'Session timestamps, duration, protocol parameters',
         'Primary care, wellness, coordinators'],
        ['Assess', '$149/mo/patient', 'All Monitor + EEG review, neurofeedback, efficacy',
         'Adds EEG waveforms, neurofeedback scores, dose logs',
         'Neurologists, psychiatrists'],
        ['Full Clinical', '$299/mo/patient', 'All Assess + HRV, closed-loop events, outcomes',
         'Adds HRV, PPG, adaptation events, outcome logs',
         'TMS clinics, neuromodulation programmes'],
        ['Research', '$599/mo/study', 'IRB-defined custom (NeurOne review required)',
         'IRB-approved minimum, k ≥ 10 anonymization, no IDs',
         'Academic trials, observational studies'],
    ],
    widths=[2.4, 2.6, 4.2, 4.4, 3.2])

add_para(
    'Key principle: clinicians select use cases, not data elements. The system determines the '
    'minimum necessary UHDR elements. Users receive a plain-language decision support document '
    'listing what the clinician CAN learn, CANNOT learn, and the privacy implications per '
    'element. Retroactive and prospective access are presented as separate consent decisions '
    'even when made simultaneously.')

add_heading('3.3 A priori research consent — four onboarding screens', level=2)
add_table(
    ['Layer', 'Question', 'If yes', 'If no'],
    [
        ['L1 — Contact', 'Can we reach you about future research opportunities?',
         'Contact method + frequency limit. POA holders upload an executed healthcare POA — '
         'human review, 3 business days, jurisdiction-flagged, annual re-verification '
         '(NP-PROC-POA-001).',
         'No contact. All features unchanged.'],
        ['L2 — Category',
         'Which research areas? (9: AD/dementia, depression, PTSD, TBI, sleep, attention, '
         'Parkinson\'s, healthy ageing, visual health)',
         'Per-project contact for selected categories only. Each project is a fresh decision.',
         'Not contacted for that category.'],
        ['L3 — Blanket', 'Pre-approve all NeurOne-reviewed research?',
         'Data included in all studies; still receives per-study engagement notifications '
         '(not consent requests). Irreversibility notice displayed at this screen.',
         'Per-category and per-project process applies.'],
        ['L4 — Results', 'Hear study results? Join the suggestion portal?',
         'Plain-language results notification per study including null results, plus paper '
         'link and access to the suggestion / voting / pledge portal.',
         'No results contact, no portal.'],
    ],
    widths=[2.6, 4.4, 5.4, 4.4])

add_para(
    'The research suggestion portal has three functions: a patient research agenda '
    '(plain-language study ideas, community voting); a pre-identified subject pool (a '
    '"would participate" intent flag creating a pre-screened, device-familiar cohort — '
    'recruitment is 40–60% of trial cost and is solved before the grant is written); and a '
    'crowdfunding catalyst (pledges are intent, not charges; escrow released to the '
    'institution research account only when a researcher confirms pilot feasibility).')

# ══════════════════════════════════════════════════════════════════════════════
# 4. CONFIGURATIONS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('4. Configurations + Pricing', level=1)

add_notice(
    'The BOM and COGS figures below were set before the hex-tile architecture. Two costed '
    'deltas are known and NOT yet netted into them: the cluster-controller tier at $114.12 '
    'per headset at n = 80 sockets plus roughly +$1.09 on the hub PCB itself (OI-HUB-C08 — the '
    'retired five-zone-module drive electronics already inside the $405 Home Standard BOM have '
    'not been subtracted), and the L1 electro-mechanical carrier assembly at an estimated '
    '+$125–216 per headset (OI-SHELL2-05, EE Lead confirmation required). Retail prices are '
    'unchanged and remain locked; the gross-margin column should be read as pre-hex and '
    'pending re-costing.',
    colour=NOTICE_BG,
    bold_lead='PRICING CAVEAT. ')

add_table(
    ['Config', 'BOM', 'COGS', 'Retail', 'GM%', 'Modalities included'],
    [
        ['Core — EEG only', '$168–169', '$258–260', '$449', '42%',
         '4-ch EEG · all connectivity · EMF shielding · processor stack · 8GB eMMC · 15W charger'],
        ['Home Lite', '$265–266', '$370–372', '$599', '38%',
         'Core + PBM tiles (660+810 nm) · 8-ch EEG · VNS+HRV clip · 30W GaN charger'],
        ['Home Standard ★', '$405', '$540', '$849', '36%',
         'All T1 modalities · hard clamshell case · braided aramid USB-C cable (spare in box) · '
         '45W NeurOne branded GaN charger · S1 opaque shade · interface covers (installed + '
         'spare set each type) · mesh cleaning brush · Boa replacement cable + hook tool · '
         'moisture-barrier electrode tip hydration caps · humidity indicator card · '
         'pre-impregnated cleaning cloth packets'],
        ['Home Premium', '$460', '$622', '$1,199', '48%',
         'All T1 + EC lens (+$89 value) · 2 yr warranty · priority support · 45W branded GaN'],
        ['Pro Entry', '$833', '$1,365', '$4,999', '73%',
         'All T1 + 21-ch qEEG · 1170 nm deep PBM · clinical tACS · HIPAA cloud · sLORETA · '
         '65W branded GaN'],
        ['Pro Full', '$1,506', '$2,628', '$13,999', '81%',
         'All T2 + TMS hub · multi-patient dashboard · scripting API · FHIR R4 · $1,800/yr '
         'service · 2× 65W GaN'],
    ],
    widths=[2.6, 1.8, 1.8, 1.8, 1.2, 7.6])

add_para(
    'Charger policy (locked): charger scaled to the peak draw of the configuration, '
    'auto-included at every upgrade by serial-number tracking, with an upfront 65W upgrade '
    'option at $19 at-cost offered at checkout as an intent signal (30W → PBM intent, '
    '45W → full T1 intent, 65W → T2 intent, triggering a human clinical sales call within 48 '
    'hours). EU note: chargers are branded recommendations, not proprietary requirements — any '
    'PD-compliant charger must work, and the app displays "power level: reduced" informatively, '
    'never blocking.')

add_para('Consumables and recurring revenue:', bold=True)
add_table(
    ['Item', 'Price', 'Interval', 'GM%'],
    [
        ['Intranasal sleeves (30-pack)', '$19/pack or $19/mo', 'Single use', '68–79%'],
        ['Electrode hydrogel tips (8-pack)', '$12–16 or $9.99/mo', '30–60 sessions', '60–72%'],
        ['VNS clip pads (2-pack)', '$8/pack', '20–40 sessions', '65%'],
        ['Audio cup foam (set)', '$24/set', '6–12 months', '58%'],
        ['Audio cup mesh frame (pair)', '$9.99/pair', 'Annual', '62%'],
        ['Interface protection covers (kit)', '$22.99 or $19.99/yr', 'Annual / as lost', '70%'],
        ['S3 prescription Rx insert', '$49–139', '12–24 months', 'Variable'],
        ['T2 service contract', '$1,800/yr', 'Annual', '~75%'],
    ],
    widths=[6.0, 4.0, 3.6, 2.4])

add_para(
    'Intranasal sleeves remain the only authenticated consumable, and the primary MRR driver. '
    'Authentication is optical code plus pogo-pin resistive sleeve — no NFC anywhere in the '
    'platform.', size=9)

# ══════════════════════════════════════════════════════════════════════════════
# 5. HARDWARE + FIRMWARE
# ══════════════════════════════════════════════════════════════════════════════
add_heading('5. Hardware + Firmware Specifications', level=1)

add_heading('5.1 Processor stack', level=2)
add_bullets([
    'Main: NXP i.MX RT1062 · Cortex-M7 · 600 MHz · FPU+DSP+SIMD · 1 MB on-chip SRAM + 32 MB '
    'LPSDR4 · USB-HS OTG · FreeRTOS-Kernel V11.3.0 (LTS 202604.00, vendored at '
    'firmware/vendor/freertos/) · ~1.1% CPU at full load.',
    'Safety MCU: STM32G071 (NOT G031 — G031 has only 8 KB SRAM, insufficient for EMF '
    'firmware) · Cortex-M0+ · 64 MHz · 36 KB SRAM · 128 KB flash · bare-metal · owns all '
    'stimulation GPIO enable lines · +$0.45 BOM.',
    'Cluster controllers (new, Hub PCB Rev C): one STM32G071 per cluster on the inner bowl. '
    'Same part as the safety MCU deliberately — it reuses the toolchain, BSP, errata and '
    'IEC 62304 qualification lineage rather than saving roughly $7 on a second family. '
    'Cluster-controller firmware (firmware/cluster_ctrl/) is a new IEC 62304 Class B unit.',
    'Storage: 8 GB industrial eMMC (SLC cache, 30,000+ P/E cycles) · LittleFS · 9-partition '
    'GPT layout · separate UHDR and SHDR partitions from the first firmware line.',
    'Connectivity: USB-C 3.2 Gen1 (default, zero RF, <1 ms) · BT 5.3 LE Audio · Wi-Fi 6 · '
    'antennas in the control hub, NOT the headset · single rear toggle.',
    'IEC 62304 classifications (NP-SW-001): Safety MCU = Class C · main processor = Class B · '
    'iOS/Android app = Class B · cluster controllers = Class B.',
])

add_heading('5.2 Hex-tile socket lattice (NP-HEX-ZM-001 Rev 2)', level=2)

add_para(
    'One universal hexagonal module SKU tiles the helmet interior: one mold, one part number, '
    'one inventory line for NeurOne and for customers. This replaces position-unique zone '
    'modules entirely.')

add_table(
    ['Parameter', 'Value', 'Basis'],
    [
        ['Tile width (flat-to-flat)', '40 mm',
         'Design point across a 34–46 mm sweep; balances coverage against curvature mismatch.'],
        ['Hex area', '13.86 cm²', 'A = (√3/2)·W².'],
        ['Vertex span / dome depth', '46 mm / 3.1 mm', 'NP-HEX-ZM-001 §3.1.'],
        ['Module curvature', 'R_m = 87 mm (κ = 0.0115 mm⁻¹)',
         'Curvature median — symmetrizes worst-case mismatch between temples (~65 mm) and '
         'crown (~130 mm). Worst-case astigmatic Δκ ≈ 0.0039 mm⁻¹, peak mismatch 1.04 mm.'],
        ['Active coverage', '77% (2.5 mm bezel)', 'NP-HEX-ZM-001 §3.1.'],
        ['Socket count', '~80 (shipped v1 lattice), PROVISIONAL',
         'Recovered from a direct 3D scan of the reference helmet interior (§3.4). Supersedes '
         'the earlier 30-socket skull-area model and 54-socket ellipsoid. Gated on REG-1 and '
         'ACT-1; re-cut by REGEN-1 on 2026-07-20 by explicit principal direction. Socket '
         'numbering is 1-based project-wide (NUMBER-1).'],
        ['Firmware addressing ceiling', 'NP_HEXMAP_MAX_SOCKETS = 128 (7-bit)',
         'The two-level bus tree arithmetic (8 branches × ≤2 clusters × ≤8 sockets) lands '
         'exactly on 128, so the physical tree and the firmware ceiling agree with no '
         'arbitrary sub-ceiling.'],
    ],
    widths=[3.6, 4.8, 8.4])

add_para('Module-type taxonomy — one size, one mechanical mold; a "type" differs only by '
         'element population on an identical footprint:', bold=True)
add_table(
    ['ID', 'Type', 'Elements', 'EEG', 'Covers'],
    [
        ['T1-A', 'Base PBM', '660–670 + 808–830 nm LEDs + PD1/PD2 + NTC', 'no',
         'PBM transcranial — bulk scalp coverage; the majority of sockets'],
        ['T1-B', 'EEG / electrode',
         'Dual-rated Ag/AgCl electrode + 660/808 PBM (reduced count for pod clearance) + PD + NTC',
         'yes',
         'EEG AND BES/tACS/tDCS on one electrode that both records and stimulates, plus PBM at '
         'electrode sites. Placed only at electrode positions — roughly 8–9 per build '
         '(the 8 T1 neurofeedback sites plus Oz when visual stimulation is used)'],
        ['T1-C', '1064 smart PBM',
         '660/808/1064 nm LEDs + on-module driver (Microchip ATtiny402 I2C slave + 3× Infineon '
         'IRLML6344 N-FETs) + InGaAs PD1/PD2 (Hamamatsu G12180-010A) + NTC', 'no',
         'Premium deep PBM, three-tier depth stack; substituted at protocol-configured depth '
         'sockets, not capped at a fixed count'],
    ],
    widths=[1.4, 3.0, 5.6, 1.2, 5.6])

add_para(
    'Decisions baked in: T1-B keeps base 660/808 PBM so coverage stays continuous at electrode '
    'sites rather than leaving PBM dead spots; T1-C carries no EEG, so deep 1064 nm PBM at an '
    'electrode site would need a fourth type, deferred as a grow-to-4 option built only if the '
    '1064 zones and EEG sites actually overlap. At 40 mm the Fp row holds a single tile, so '
    'Fp1/Fp2 share one socket — an open item against REG-1.', size=9)

add_para(
    'SMART-1 (2026-07-28): every socket is I2C/TIA-capable, not a subset. Full coverage was '
    'chosen over cheaper partial coverage because the research mission — SBIR trials, '
    'IRB-approved custom studies, the research suggestion portal — depends on placing any '
    'current or future smart module type at any socket for arbitrary montage design. Subset '
    'coverage would bake a placement constraint into shell and hub tooling permanently, for '
    'every smart module type this platform will ever ship.')

add_heading('5.3 Cluster architecture', level=2)

add_para(
    'The cluster is simultaneously the mechanical clamp unit, the electrical aggregation '
    'boundary, the I2C bus segment, and the safety-cut domain. Three principal directions '
    'define it:')
add_table(
    ['Direction', 'Rule', 'Reason'],
    [
        ['CLUSTER-1 (2026-07-30)',
         'The 7-hex flower is the cluster unit. Where the lattice allows a full flower, use '
         'one; where it does not, use a partial flower — not a larger patch and not a 3-hex '
         'triad.',
         'Mechanical. The flower minimises the cluster\'s longest dimension, and both '
         'clamp-plate stress and stored curvature scale off it. An 8th tile must attach on the '
         'perimeter, forcing span 122.2 → 161.8 mm, bending stress ×1.75, plate deflection '
         '×3.07, dome depth 25.1 → 55.0 mm.'],
        ['SYM-1 (2026-08-04)',
         'The cluster partition is mirror-symmetric about the sagittal midline. Every cluster '
         'is either self-symmetric (centred on a midline socket) or one of a left/right mirror '
         'pair.',
         'A cluster holding a midline socket must equal its own mirror image, so six midline '
         'clusters are forced, absorbing 30 sockets; the two residual lateral bands of 25 need '
         '6 each. Cost: 6 extra clusters and about $38, accepted because the cluster is both '
         'the safety-cut domain and the I2C segment — an asymmetric partition would make a '
         'per-cluster safety cut asymmetric across hemispheres.'],
        ['CONTIG-1 (2026-08-04)',
         'A cluster\'s petals must form a contiguous arc — no pendant petal (one whose only '
         'in-cluster contact is the centre, with foreign sockets on both flanks).',
         'The clamp plate cannot bridge a gap, because the gap socket belongs to a '
         'differently-actuated cluster. It would need a 40 mm cantilever arm necked to 23.09 mm '
         'following 2.33 mm of dome — under-seating that plunger and creating a withdrawal snag '
         'path. Resolved at zero cost in cluster count by reassignment.'],
    ],
    widths=[3.0, 6.4, 7.4])

add_para(
    'Resulting count: 18 clusters at the shipped 80-socket lattice — provably minimal under '
    'CLUSTER-1 + SYM-1 + CONTIG-1, mirror-symmetric, maximum size 6, zero pendants. Derivation '
    'at NP-HW-HEXTILE-001 §8.2.1; verified by HT-DRC-21; diagram at '
    'docs/diagrams/np_hextile_cluster_map.svg.', bold=True)

add_notice(
    'The 18-cluster count post-dates NP-DRV-SHELL-002 Rev 2 and NP-HW-HUB-001 Rev 3 (still '
    'uncorrected at Rev 4, which was a documentation-only change), both of '
    'which were written against ~12. Consequences not yet corrected in those documents '
    '(OI-HEXTILE-14): cluster-tier BOM $76.08 → $114.12; analog front ends scale with 18, not '
    '12; and 18 exceeds BOTH the 16 provisioned connector positions AND the 8 branches × ≤2 '
    'clusters I2C tree. Interconnect options and a recommendation are recorded at '
    'NP-HW-HEXTILE-001 §8.2.2. The five-network split, the cluster-carrier selection and the '
    'zero-dynamic-flex conclusion are unaffected — only the counts and the enable ownership.',
    colour=NOTICE_BG,
    bold_lead='COUNT CONFLICT, OPEN. ')

add_heading('5.4 Hub PCB Rev C — distributed cluster-controller tier', level=2)

add_para(
    'Key finding: the binding constraint is interconnect, not I2C mux count. Rev B is a star — '
    'every socket\'s analog and digital lines terminate at the hub — so scaling it to 80 sockets '
    'needs roughly 11 conductors × 80 ≈ 880 conductors through the single posterior blind-mate '
    'boss, 160 TIA channels and 160 ADC channels against about 32 available on the RT1062, plus '
    '80 GAIN_SEL GPIO. A deeper PCA9546A mux tree — the obvious answer — fixes address collision '
    'and leaves all of that untouched, so it was rejected, along with per-socket address '
    'translation (LTC4316, roughly $200) and a hub-side scanning MCU. Both are still stars.')

add_para('Each cluster board carries:', bold=True)
add_bullets([
    'One STM32G071 cluster controller.',
    'One NXP PCA9548A for I2C isolation — an exact 8-channel fit. Tier-1 controllers have '
    'unique strapped addresses, so no mux is needed there at all.',
    'One 16:1 low-leakage mux feeding PD1/PD2 currents into one shared switchable-gain TIA.',
    'The per-socket LED drive stage (still unwritten — the largest open item, OI-HUB-C01…C19).',
])

add_para('Four consequences worth recording:', bold=True)
add_bullets([
    'The hub PCB no longer encodes the socket count at all — no per-socket footprint, net or '
    'n_sockets-derived constant — so REG-1 no longer blocks hub tooling, and the lattice can '
    'move anywhere in 30–128 sockets by re-tooling inner-bowl boards only.',
    'Zero GAIN_SEL GPIO. Gain is a cluster-MCU pin set per sample from live UID inventory '
    'inside the scan loop, which also removes Rev B\'s stale-gain-after-hot-swap hazard, with '
    'no shift register or GPIO expander.',
    'PD1/PD2 ratio accuracy improves. Both photodiodes now traverse the same mux, Rf, op-amp and '
    'ADC, so gain error is common-mode and cancels exactly in the ratio the fouling-versus-aging '
    'discriminant rests on — retiring Rev B\'s ≤0.5% per-slot matching requirement.',
    'Calibration coefficients K_PD1 / K_PD2 / K_ratio_nom must be keyed to MODULE UID in the '
    'existing UID-keyed np_module_map NVRAM record, never re-indexed by socket — a swappable '
    'module re-indexed by socket would silently apply the wrong calibration while cal_source '
    'still read NP_CAL_FACTORY (OI-HUB-C06).',
])

add_para(
    'Hub ↔ inner bowl is one differential I2C bus (PCA9615-class) plus an ATTN line, with a '
    'firmware bus-quiet window in the ADS1299 inter-conversion gaps. Differential was chosen '
    'chiefly because a single-ended 400 kHz open-drain bus running past µV EEG leads would '
    'repeat the mistake NP-HEX-ZM-001 §5.3.1 rejects for the Helmholtz coil. BOM: $6.34 per '
    'cluster board, roughly +$1.09 net on the hub PCB itself. Status: DRAFT — the socket count '
    'and every part number are explicitly open; none of the hub-side fan-out is implemented yet.')

add_heading('5.5 Shell socket interconnect (NP-DRV-SHELL-002 Rev 2, DRAFT)', level=2)

add_para(
    'The inherited assumption that was retired: "the Hub PCB is where module signals are '
    'conditioned." Retiring it is what collapses the SMART-1 residual — analog front ends now '
    'scale with cluster count, the same order as the five they replace, rather than with socket '
    'count. Rejected alternatives, with reasons: per-socket tails (80 ZIF receptacles need about '
    '960 mm of hub connector edge and put 1,600 conductors on one board — rejected on '
    'arithmetic); a monolithic L1 backplane (makes the socket layer a single non-repairable '
    'part, fights the compound curvature, and maximises conductive mass on the one layer '
    'constrained to non-magnetic / low-eddy because the fluxgates mount there).')

add_para('Five networks split by physics:', bold=True)
add_table(
    ['Network', 'Carries', 'Constraint'],
    [
        ['N1 — power', 'LED and stimulation rails', 'Broadside VLED+/PGND tree; ≤25 mm² loop '
         'area per cluster feed. May not form a closed ring on L1 (bowl-scale loop antenna).'],
        ['N2 — control', 'I2C', 'Two-level tree: 8 branches × ≤2 clusters × ≤8 sockets = 128.'],
        ['N3 — local sense', 'PD1/PD2 at 14–72 µA; EEG at sub-µV',
         '≤50 mm, never leaves its own cluster carrier. High-impedance analog cannot travel; '
         'the cluster is the nearest conversion point that already exists.'],
        ['N4 — electrode', 'EEG / tES conductors',
         'Sized by channel count (8 or 21), not socket count.'],
        ['N5 — safety', 'Per-cluster hardware enable',
         'SAFE_EN_n gates each cluster\'s high-side load switch, so no module firmware state, '
         'stuck I2C transaction or main-processor hang can produce emission from a de-energized '
         'cluster. CONTESTED against the single-bit cranial enable — see below.'],
    ],
    widths=[3.0, 4.6, 9.2])

add_para('Three EMI prohibitions locked:', bold=True)
add_bullets([
    'The shell, shield stack, mu-metal and any DRL-driven structure may NOT carry LED or tES '
    'return current (REQ-EMI-08). The shell is a driven EEG shield, so returning '
    'therapeutic-band current through it injects that current into the EEG reference.',
    'N1 may NOT form a closed ring on L1 (REQ-EMI-09).',
    'Spread-spectrum / dithered PWM is PROHIBITED (REQ-EMI-03/04) — inverting standard EMC '
    'practice. Dither smears bus energy into broadband noise across the EEG band that neither '
    'the notch nor subtraction can remove, whereas a deterministic phase-locked artifact is a '
    'known, subtractable line.',
])

add_para(
    'The therapeutic band cannot be filtered off the bus, and this is arithmetic rather than '
    'opinion: holding a 40 Hz / 25% pulse at 0.25 A within 100 mV needs 15.6 mF per tile, while '
    'the PWM carrier needs about 62 µF and the edges about 10 µF. On-tile decoupling therefore '
    'kills the carrier and the edges and cannot touch the 0.5–100 Hz envelope — which IS the '
    'therapy, and sits in exactly the EEG band and the ELF band the Helmholtz loop cancels. It '
    'must be handled geometrically and computationally, never by filtering.')

add_para(
    'Modules lose their tail. A hex tile presents an 18-contact back-face pad array seated by '
    'the cluster clamp plunger, so the module interconnect has zero dynamic-flex paths — the '
    'retired architecture\'s dominant failure mode (fatigue cracking at a stiffener edge over '
    '1,000 swap cycles) is designed out rather than mitigated. Non-obvious coupling recorded: '
    'contact count drives clamp-plate load (roughly 38–63 N at 18 contacts per 7-tile cluster), '
    'so contact count is a RISK-22 accessibility variable.')

add_para(
    'Conductor count, stated honestly: 144 conductors at the hub PCB (EEG included) versus 110 '
    'for the retired design and 1,600 for per-socket tails. That is about 1.3× the retired count '
    'for 16× the sockets — not a reduction, and the document does not claim one. Cancellation '
    'calibration is now configuration-dependent, because any tile type may occupy any socket: '
    'recalibration fires on np_module_map rebuild, and feed-forward self-field subtraction must '
    'bind to measured PD/dose telemetry, not commanded current, since LED aging, PD-metered '
    'compensation and the 62 °C throttle all make actual current diverge from commanded.')

add_heading('5.6 Safety architecture', level=2)
add_bullets([
    'The safety MCU physically owns all stimulation enable GPIO — an app crash cannot cause '
    'unsafe stimulation.',
    'SPI heartbeat from the main processor every 200 ms; 1.5 s watchdog → all-stimulation '
    'cutoff <50 ms.',
    'Dual-processor isolation: IEC 62304 Class C (safety MCU, bare-metal — ~1,600 physical '
    'lines across 9 modules as of 2026-08; wc -l firmware/safety_mcu/src/*.c is the source of '
    'truth) and Class B (main processor) separately certified. The safety-MCU enable gates each '
    'cluster\'s LED rail UPSTREAM of the cluster MCU (HUB-REQ-C02), preserving dual-processor '
    'isolation across the new tier.',
    'Session protocols are cryptographically signed by the app; the headset rejects unsigned or '
    'corrupted protocols.',
    'Enable word: bit 0 is NP_SAFETY_EN_PBM_CRANIAL; bits 1–4 are reserved holes EXCLUDED from '
    'NP_SAFETY_EN_ALL_MASK (0x3FFF → 0x3FE1), so bits 1–4 are stripped in np_spi_watchdog_tick '
    'and can never enable anything. That exclusion is defence in depth against a FUTURE authoring '
    'error re-introducing those positions — NOT compatibility with a deployed or legacy hub. No '
    'hub can set a retired zone bit: no hub hardware exists, and the NP_SAFETY_EN_PBM_ZONE_0..4 '
    'macros were deleted in the same change, surviving only in "Formerly …" comments '
    '(NP-HW-HUB-001 Rev 4 §7.2, corrected 2026-08-12). The reservation of bits 1–4 is unaffected '
    'and remains correct on independent Class C grounds: the bit position is '
    'deliberately NOT compacted, because enable-bit position ≡ current_ua[] slot ≡ '
    'charge-accumulator index is a three-way Class C identity that the 40 µC/cm² limit rests on, '
    'and that binds regardless of whether any hub exists.',
])

add_notice(
    'Whether the safety layer can cut ONE cluster or only the whole cranial lattice is open. '
    'NP-DRV-SHELL-002 specifies per-cluster SAFE_EN_n (18 GPIO); NP-HW-HUB-001 Rev 4 §7.2 '
    'specifies a single cranial bit (1 GPIO, all-or-nothing). Per-cluster policy bits would need '
    'the enable word widened to uint32_t (18 cluster + 9 modality = 27 > 16). Tracked as '
    'OI-HUB-C07 / OI-HEXTILE-13; proposed resolution at NP-HW-HEXTILE-001 §8.4.1 — split the '
    'enable by IEC 62304 class, with Class B per-cluster gates for availability and one Class C '
    'bit in series as the hard interlock.',
    colour=NOTICE_BG,
    bold_lead='SAFETY GRANULARITY, OPEN. ')

add_para('Modality-specific interlocks:', bold=True)
add_table(
    ['Modality', 'Interlock', 'Implementation'],
    [
        ['EEG + visual', 'Photoparoxysmal detection → goggle halt',
         'Oz electrode, <200 ms, clinician-unlock required for 3–30 Hz'],
        ['BES / tDCS', '40 µC/cm² charge density limit',
         'Safety MCU hardware — the app cannot override'],
        ['Visual / retinal', 'IEC 62471 MPE ceiling',
         'IR proximity + Hall sensor + hardware current limit (3 independent layers)'],
        ['PBM scalp', 'IEC 60601 42 °C applied-part limit',
         'NTC per sense domain → hardware current throttle at 62 °C junction'],
        ['PBM scalp — fan loss', 'Scalp-facing NTC (RISK-26, Path B1)',
         'NTC co-located with PD2 (NP-REQ-FANHEALTH-001 §4a); halt or trickle PBM on fan loss'],
        ['TMS', 'Coil protection',
         'EMF cancellation gated off 5 ms before pulse, 50 ms hold'],
        ['Auricular VNS', 'Contact confirmation',
         'Safety MCU reads impedance; holds if contacts are not confirmed'],
        ['Cervical VNS (T2)', 'Cardiac rhythm interlock',
         'Safety MCU owns the enable GPIO and monitors the R-peak GPIO. HR change >15 BPM '
         'within a rolling 5 s window → GPIO LOW within <5.1 ms worst-case (spec <100 ms). '
         '30 s re-enable lockout + app confirm + repeat impedance. Additionally, per-electrode '
         'impedance is cross-validated hub-versus-MCU (OI-CVNS-HUB-11), fail-closed beyond '
         '1.0 kΩ divergence.'],
    ],
    widths=[3.2, 4.4, 9.2])

add_heading('5.7 EMF shielding (5-layer passive + active)', level=2)
add_bullets([
    'Layer 1: CFRP outer (30–50 dB RF).',
    'Layer 2: 0.2 mm mu-metal liner (15–25 dB ELF magnetic), PETG laminate encapsulation, '
    'silicone RTV sealant at all cutout edges.',
    'Layer 3: palladium-coated polyester inner liner (40–60 dB RF) — replaces silver, which '
    'tarnishes in 12–18 months. Palladium is tarnish-immune for device lifetime, making the '
    'shielding claim permanent and verifiable by fleet SHDR attenuation monitoring.',
    'Layer 4: carbon-loaded EMI absorber foam (cavity resonance suppression).',
    'Layer 5: USB-C and accessory port filters (30–50 dB).',
    'Active: 3-axis fluxgate magnetometers + Helmholtz coil pairs. Combined 35–45 dB ELF '
    'magnetic, 40–60 dB RF. Shell bonded to the EEG DRL output as an active EEG shield. '
    'Non-conductive CFRP window at the TMS coil site.',
])

add_heading('5.8 Thermal architecture', level=2)

add_para(
    'BN-boss conductive thermal export is the adopted base thermal design (NP-THERM-CFD-R1-001, '
    '2026-07-22). Module heat is exported from the boron-nitride thermal boss through a solid '
    'conductive via to an EXTERNAL fan-cooled heatsink, with the inter-bowl cavity left '
    'STAGNANT — the shielded, sealed interior is never ventilated. THERM-1a first-pass analysis '
    '(perfect-sink bound) shows this carries about 90% of module heat and holds the LED junction '
    'throttle-free at 54 °C against the 62 °C throttle at full T2-peak flux. Cavity ventilation '
    'was rejected on shield integrity (apertures through mu-metal and RF liners), ingress and '
    'hygiene, optics fouling, acoustic and EEG noise, power and MTBF, and safety-case inversion. '
    'Formalized as NP-DT-001 DI-SAFE-13.')

add_notice(
    'The conductive via is a HEALTHY-STATE junction-cooling strategy only. It gives zero benefit '
    'under heatsink or fan loss — the fault face reaches 60.3 °C with or without it — so it does '
    'NOT replace the fan-health interlock. THERM-1a rejected Path A: the 62 °C junction throttle '
    'leaves the scalp-facing face at about 60 °C and scalp skin at 52–55 °C in the worst fault '
    '(fan off, 43.3 °C ambient, low perfusion), which is 14–21 °C over the IEC 60601 42 °C '
    'applied-part limit, with the 1D and 2D models in agreement. Path B1 (scalp-facing NTC '
    'co-located with PD2) is committed. Provisional constants: SR-FAN-03 natural-convection-safe '
    'ceiling ≈ 4.5 mW/cm² at 43.3 °C ambient, tending to zero as ambient approaches 42 °C; '
    'SR-FAN-04 τ_face ≈ 35–45 min. The steady ceiling, not response speed, governs. All '
    'constants are provisional pending verification-grade 2D CFD and the THERM-1b bench; path '
    'selection is firm. Tracked as RISK-26 / FMEA-G07-01, OPEN pending SW01-M04 implementation.',
    colour=NOTICE_BG,
    bold_lead='RISK-26 — fan/heatsink-loss scalp overheat. ')

add_heading('5.9 eMMC partition architecture and storage encryption', level=2)
add_bullets([
    '9-partition GPT layout across an 8,192 MiB user area: Firmware Bank A (128 MiB, '
    'write-protected in production) · Firmware Bank B (128 MiB, OTA target) · Safety MCU FW '
    'staging (4 MiB) · Config/Calibration (16 MiB) · SHDR (512 MiB) · UHDR (6,903 MiB, about '
    '84% of the user area) · Scratch (500 MiB, zeroed on boot).',
    'UHDR: AES-256-XTS under a two-layer key scheme — UKMD (32-byte TRNG master key) held in '
    'the Config partition encrypted by WKMD (Argon2id from biometric/PIN). A biometric change '
    're-encrypts only the 32-byte wrapper; UHDR ciphertext is unchanged. NeurOne never holds '
    'the key. Implemented via PHC phc-winner-argon2 vendored as an SPM C target.',
    'SHDR: AES-256-XTS, HKDF-SHA256 from the manufacturing root key + eMMC CID hash under SNVS '
    'OTPMK. NeurOne can decrypt for warranty and fleet analytics.',
    'Scratch partition encryption for research anonymization: per-task AES-256-CTR key, '
    'SRAM-only, memset_explicit on completion, eMMC SANITIZE post-run (NP-FW-EMMC-002 §D).',
    'Dual-bank OTA: Ed25519 signature verification on all images before boot; SNVS_LPGPR0 boot '
    'bank flag; automatic revert after 3 failed boot attempts; USB-C DFU recovery always '
    'available. The manufacturing root public key is a placeholder and must be replaced at a '
    'secure build step before any production image is signed.',
    'Safety MCU firmware is NEVER updated by automated OTA — explicit user confirmation plus '
    'page-by-page SPI flash with readback is required.',
    'Device factory reset: eMMC SANITIZE on UHDR + zero SHDR + clear Config + new TRNG salt + '
    'new warranty token, power-loss resilient via a reset_in_progress flag in SNVS_LPGPR1.',
    'Session data classification table now spans 32 elements (extended from 27 by the '
    'fault-latch tick_ms and count entries).',
])

add_heading('5.10 T2 clinical firmware', level=2)

add_para('sLORETA-guided HD-tDCS (NP-FW-HD-001):', bold=True)
add_bullets([
    'Scalar sLORETA weight matrix W (2,447 voxels × 21 channels, 7 mm MNI grid, about 205 KB) '
    'precomputed offline and loaded into LPSDR4 at session start. On-device: 64-epoch '
    'covariance accumulation → per-voxel source power map → peak detection → MNI-to-10-20 '
    'electrode mapping (Jurcak 2007 atlas).',
    '4×1 ring montage: nearest-electrode anode plus 4 cathodes at −I/4 each, with a ≥2 quadrant '
    'angular spread constraint and driver-channel distinctness enforced during ring selection. '
    'Bilateral 4×1 enforces electrode distinctness across both rings. 30 s ramp up and down, '
    '100 ms tick.',
    'Driver: 21 channels, one per cap electrode — no sharing, no aliasing. Constraint of record '
    'if cost ever reintroduces sharing: never alias two electrodes within one ring radius '
    '(~60 mm). NP-HW-TCAP-001 still requires authoring to become the controlled hardware source '
    'of truth; the firmware table is the interim authority, and the BOM delta for 5 additional '
    'precision current sources is not yet costed.',
    'Safety: 40 µC/cm² charge density enforced by the safety MCU; software monitor aborts at '
    '95% of the limit; ≤2 mA per electrode. Concurrent EEG uses 200 µs blanking on current step '
    'edges against a ≥20 dB SNR spec.',
])

add_notice(
    'Localization is not reachability. sLORETA resolves deep sources, but a 4×1 ring is focal '
    'only for cortical-surface targets (about 1.5 cm FWHM at 10 mm depth). ACC sits 47.1 mm from '
    'its nearest scalp electrode (Fz), with 2nd/3rd nearest at 49.1 mm and no 10-10 position '
    'closer than about 37.9 mm (Fpz) — the ≤35 mm criterion is unsatisfiable for ACC by ANY '
    'electrode placement. Targets therefore carry NP_HD_TARGET_DEPTH_SURFACE or _DEEP. ACC stays '
    'selectable and keeps its anxiety indication, but a 4×1 there is indirect network modulation, '
    'and callers must never present it as focal stimulation of the structure. The FAI criterion '
    'is checked in BOTH directions — surface ≤35 mm, deep >35 mm — so reclassifying ACC to '
    'SURFACE to silence the test fails instead of passing vacuously. Deliberately NOT resolved by '
    'raising the 35 mm threshold, which would have made the FAI assert focal reach the physics '
    'does not support. OPEN: the app must surface the DEEP classification when a clinician '
    'selects ACC — the firmware exposes the class but nothing consumes it yet.',
    colour=NOTICE_BG,
    bold_lead='HD-tDCS DEPTH CLASSIFICATION. ')

add_para('Cervical VNS (NP-FW-CVNS-001 Rev 2):', bold=True)
add_bullets([
    'Safety MCU owns the CVNS enable GPIO independently of the main processor. R-peak from '
    'main-processor PPG detection drives an RPEAK_IN GPIO pulse; the MCU measures R-R via TIM2 '
    'and runs a TIM6 ISR at 200 Hz.',
    'Baseline cross-validation: main-processor PPG-derived HR and safety-MCU GPIO-timer HR must '
    'agree within 5 BPM before enable is granted. The cardiac baseline is mandatory and cannot '
    'be waived by the session descriptor.',
    'Rev 2 reconciled §5 with the Class C firmware: §5 had documented the safety MCU using the '
    'MAIN PROCESSOR\'s constants. The two sides are meant to be independent implementations that '
    'cross-validate, so attributing one side\'s constants to the other was an error independent '
    'of which values are correct. No number in either the document or the firmware was changed '
    'to make the other side agree, and no firmware behaviour changed. Four open items were '
    'raised rather than papered over: OI-CVNS-08 (pin map, settled only by the G1 layout '
    'package), OI-CVNS-09 (single versus dual enable — a clinical/regulatory question against '
    'the gammaCore predicate), OI-CVNS-10 (baseline window 8 on the MCU versus 5 on the main '
    'processor — explicitly do not harmonise by editing one number), OI-CVNS-11 (the MCU applies '
    'no R-R validity filter).',
    'RPEAK_IN on PB0 remains the one pin-map row that cannot work in either reading, PB0 still '
    'being NP_EN_BES_PIN — a cardiac input on a stimulation output.',
    'First host coverage of the Class C cardiac interlock now exists '
    '(np_cardiac_interlock_tests, 35 assertions), falsified against four deliberately-broken '
    'variants before acceptance.',
])

add_heading('5.11 Software supply chain and CI', level=2)
add_table(
    ['Component', 'Version', 'IEC 62304 class', 'Notes'],
    [
        ['FreeRTOS-Kernel', 'V11.3.0 (FreeRTOS-LTS 202604.00)', 'Class B SOUP',
         'Byte-exact subset vendored 2026-07-11. croutine.c and all non-CM7 ports excluded. '
         'Config and hooks are first-party in firmware/hub_control/.'],
        ['ARM CMSIS-Core(M)', '5.6.0 (from CMSIS_5 tag 5.9.0)', 'Class C SOUP',
         'Five headers, no translation unit, vendored 2026-08-09. Present because '
         'stm32g071xx.h includes core_cm0plus.h unconditionally. Tag 5.9.0 chosen as the '
         'terminal 5.x release so its anomaly list is closed.'],
        ['ST CMSIS-Device STM32G0', 'v1.4.5', 'Class C SOUP',
         'Three headers, no translation unit. The ST HAL and LL drivers are deliberately NOT '
         'vendored — SW-01 makes zero HAL_* and zero LL_* calls, and the HAL would add tens of '
         'thousands of lines of third-party logic to the Class C surface for no benefit.'],
        ['Monocypher', '4.0.2', 'Class C + Class B',
         'Backs np_ed25519_verify in firmware/crypto/. Small-subgroup guard: an all-zero public '
         'key is rejected before crypto_ed25519_check.'],
        ['PHC phc-winner-argon2', '(vendored SPM C target)', 'App (Class B)',
         'Backs UHDRKeyManager. Replaced a PBKDF2 placeholder entirely.'],
    ],
    widths=[3.4, 4.2, 2.8, 6.4])

add_para(
    'NP-SOUP-CMSIS-001 Rev 1 is the first IEC 62304 §7.1.2 anomaly evaluation performed for '
    'Class C software at NeurOne. Its principal finding is a scoping one: CMSIS_5 carried 188 '
    'open issues at review, essentially all in components not vendored (DSP, NN, RTOS, DAP, '
    'SVDConv, Pack) — that is the return on the narrow-vendoring decision. The evaluation is '
    'bound to the pinned tags; moving either tag voids it and requires a new revision. Stated '
    'limitations are recorded deliberately: a keyword-driven issue search is not exhaustive, and '
    'ARM publishes no numbered errata for CMSIS-Core.')

add_para(
    'NP-SW-CI-001 established firmware cross-compile CI as a phased program: the bootloader '
    'cross-build, the main-firmware leg and the Class C safety-MCU compile leg were each promoted '
    'from advisory to gating, and each promotion was demonstrated by a deliberately-broken probe '
    'that was then reverted in the same PR — rather than asserted from a green check. Defects '
    'A–D closed (missing device headers, freestanding memcpy/memset, linker-owned OCRAM staging, '
    'and the main-firmware gate).')

# ══════════════════════════════════════════════════════════════════════════════
# 6. DURABILITY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('6. Durability + Maintenance Engineering', level=1)

add_para('Pre-tooling critical items, current status:', bold=True)
add_table(
    ['Change', 'Status', 'BOM delta', 'Notes'],
    [
        ['ITO → AgNW conductive lens coating', 'PRE-TOOLING CRITICAL', '+$8–12/lens',
         'ITO has 0.5% strain-to-failure and cracks on any flex or point impact; AgNW tolerates '
         '5–10%. Supplier qualification is BLOCKING before lens tooling release '
         '(NP-TOOL-LENS-001 Rev 2 P-01).'],
        ['PDMS–PI SiO₂ interlayer bonding', 'PRE-TOOLING CRITICAL (BLOCKING)',
         'Process cost (CAT-C supplier)',
         '75 nm RF magnetron sputtered SiO₂ on PI coverlay + O₂ plasma activation, achieving '
         '174–860 N/m peel force. 200-cycle IEC 60068-2-14 thermal cycling qualification '
         '(FAI-TC02) is BLOCKING before any production FPC. CAT-C supplier not yet selected '
         '(RISK-04 OPEN); 8-week lead time from selection.'],
        ['CFRP shell slot rim Ra ≤ 1.6 µm', 'OPEN — BLOCKING', '$0 (supplier engagement)',
         'Written confirmation required from the CFRP tooling supplier — letter plus Ra '
         'measurement from a representative coupon. Supplier qualification item SUP-M-07. If '
         'Ra > 1.6 µm is unavoidable, escalate to ME + EE; gasket or shell geometry must be '
         'revised. RISK-20.'],
        ['Dual photodiode per module (RISK-14 Option B)', 'CARRIED FORWARD',
         '+$0.75–1.50/headset',
         'PD1 behind the PDMS window measures forward emission; PD2 on the scalp-facing surface '
         'measures backscattered tissue power. The PD1/PD2 ratio separates PDMS fouling '
         '(PD1 falls, PD2 stable) from LED aging (both fall). Ratio accuracy improves under Rev C '
         'because both photodiodes now share one signal chain. PD2 is now also the co-location '
         'point for the RISK-26 scalp-facing NTC.'],
        ['Self-sealing co-molded silicone gasket (RISK-16 Option A)', 'CARRIED FORWARD',
         '+$0.30–0.60/module',
         'Shore 40–50A medical silicone, D-section, 20% compression when seated. No user RTV. '
         'IPX4 compliant after 10 field swap cycles (FAI-IPX-02 BLOCKING).'],
        ['Module extraction accessibility (RISK-22)', 'REDESIGNED',
         'Included in cluster clamp',
         'The retired per-module sliding eject lever is superseded by the cluster clamp — one '
         'actuator per cluster instead of a lever per module. Contact count drives clamp-plate '
         'load (roughly 38–63 N at 18 contacts per 7-tile cluster), so contact count is now an '
         'accessibility variable, closed jointly with MECH-2 (OI-SHELL2-03). The Parkinson\'s '
         'Hoehn & Yahr II–III accessibility target is unchanged; the HFE formative study with 5 '
         'subjects remains outstanding.'],
        ['Hard clamshell case; intranasal probe hub dock', 'CARRIED FORWARD', '+$8–14',
         'Lens scratching from bag contents is certain within the first month without a hard '
         'case; the case doubles as the shipping container. The Y-probe dock is molded into the '
         'hub from day one (NP-TOOL-HUB-001 Rev 1 F-01) because a dropped probe fractures the '
         'Y-junction and this cannot be retrofitted.'],
        ['Hub enclosure tooling (NP-TOOL-HUB-001 Rev 1)', 'ISSUED 2026-07-27', 'Tooling',
         'Four molded features on a single hub housing tool: F-01 intranasal probe dock; F-02 '
         'tethered anchor posts bounded so a detached cover geometrically cannot reach the fan '
         'intake grille; F-03 large-radius Boa cable channel; F-04 tool-free fan/heatsink '
         'service door with two diagonally-opposed quarter-turn captive fasteners, giving the '
         'user a way to act on the SR-FAN-05 predictive-maintenance alert before the safety '
         'derate engages. Hub placement (occipital arch, external to the shielded shell) is '
         'INFERRED from three already-locked facts, not newly asserted; the hub-to-shell '
         'mechanical interface remains open (OI-HTOOL-01).'],
        ['Palladium-coated EMF shielding fabric', 'DECIDED', '+$6/headset',
         'Silver tarnishes in 12–18 months; palladium is tarnish-immune for device lifetime, '
         'and fleet SHDR verifies stable attenuation — a permanent, marketable claim.'],
        ['Interface protection covers, all tethered', 'PRE-TOOLING CRITICAL', '+$8–9 total',
         'Anchor posts mold into the shell at zero cost if specified before first cut. Loss '
         'prevention by design — covers cannot be permanently lost without deliberate cutting.'],
        ['Sliding rail lens mount', 'DECIDED', '+$1.20',
         'T-slot groove in the lens rim plus a tongue in the goggle arm; ≤5 N extraction, '
         '500-cycle rated. Eliminates the Tier A alignment jig; user self-install.'],
    ],
    widths=[4.0, 3.0, 2.6, 7.2])

add_para(
    'Retired from the durability program: the per-zone molded FPC channels (REQ-ST-01..07) are '
    'deleted from shell tooling, with complexity moving into L1 lamination; L1 becomes an '
    'electro-mechanical assembly requiring a new rigid-flex-into-molded-carrier supplier category '
    'in NP-PROC-SUP-001 (OI-SHELL2-04). Cluster-level rework replaces bowl scrap.')

add_para(
    'Predictive maintenance runs on SHDR fleet telemetry in three phases: population-average '
    'survival analysis (0–1,000 devices, Year 1), a fleet-trained LSTM on sensor trajectories '
    '(1,000–10,000 devices, Year 2), and Bayesian personalization with continuously revised RUL '
    'predictions (10,000+ devices, Year 3+). Models are version-stamped by hardware revision and '
    'deployed back to devices by OTA, so the moat grows automatically with fleet size. All '
    'reminders are measurement-triggered, never calendar-triggered; safety-critical reminders '
    'cannot be dismissed and block session start.')

# ══════════════════════════════════════════════════════════════════════════════
# 7. SERVICE NETWORK
# ══════════════════════════════════════════════════════════════════════════════
add_heading('7. Service Network', level=1)
add_table(
    ['Tier', 'Examples', 'Service tasks', 'Certification', 'Revenue/yr', 'Launch'],
    [
        ['A — Optical centers', 'LensCrafters, Pearle Vision, independent opticians',
         'S3 Rx clip manufacture and fitting (primary); lens replacement; calibration (secondary)',
         '4 hr initial, 1 hr annual online', '$8K–35K', 'Year 1'],
        ['B — Electronics repair', 'uBreakiFix/Asurion, iFixit partners',
         'Cluster carrier and module swap; DFU recovery; eMMC data recovery; impact inspection; '
         'fluxgate calibration', '6 hr initial, 2 hr annual practical', '$4K–18K',
         'Year 2, major metros first'],
        ['C — Retail triage', 'Best Buy Geek Squad, Apple Authorized Service',
         'Warranty intake and triage; routing to Tier A/B; consumable sales',
         '2 hr initial, 30 min annual online', '$1K–7K', 'Year 1–2'],
        ['NeurOne Depot', 'Mail-in backstop',
         'All tasks; T2 same-day loaner; precision fluxgate calibration; eMMC specialist recovery',
         'Full internal training', 'Highest margin per task', 'Day 1'],
    ],
    widths=[3.0, 3.2, 4.8, 2.8, 2.0, 1.8])

add_para(
    'Structural, accepted and managed: fluxgate calibration requires a Tier B visit every 3–5 '
    'years — geomagnetic field magnitude comparison detects severe drift, and fleet SHDR predicts '
    'the need with months of lead time, so the Tier B network is established before first visits '
    'are needed. iOS/Android OS update dependency is managed by developer beta participation, a '
    '7-day OS compatibility SLA, and autonomous Mode 3 as the structural fallback.', size=9)

# ══════════════════════════════════════════════════════════════════════════════
# 8. RISKS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('8. Risk Register + Open Blocking Items', level=1)

add_para(
    'Risk register: the RISK-01…RISK-26 sequence is closed and append-only, under ISO 14971 '
    'change control per NP-RM-001 Rev 1. The register itself was re-baselined on 2026-08-11: '
    'NP-RISK-001 is retired and split by artifact into NP-RISK-002 (index and disposition), '
    'NP-RISK-003 (hex-tile module) and NP-RISK-004 (shell, socket, interconnect, hub). Of the '
    '26, twenty are CARRIED into the successors; three remain OPEN or unmitigated at the new '
    'numbers — RISK-03 (regulatory opinion, never commissioned, externally blocked since '
    '2026-05-06 — OI-RISK2-04), RISK-16 (IPX4 after field replacement; the only risk whose '
    'exposure the architecture change INCREASED, no seam-length budget for ~30 seals) and '
    'RISK-20 (CFRP rim finish). RISK-26 (fan/heatsink airflow loss) is CARRIED to NP-RISK-004 '
    'at HIGH / ALARP. RISK-25 (cervical VNS cardiac reflex) is MITIGATED pending the FAI-CV02 '
    'hardware bench. New hazards raised against the hex architecture take artifact-prefixed '
    'IDs rather than RISK-NN — for example RISK-SHELL-02, the continuous dissipation of 18 '
    'cluster controllers behind the dominant outward thermal resistance, which is unbudgeted.',
    bold=True)

add_table(
    ['Severity', 'Issue', 'Status and mitigation'],
    [
        ['CRITICAL', 'Zero published clinical trials — a 35-trial gap versus Vielight',
         'Every clinical adoption decision, investor meeting and regulatory filing needs clinical '
         'data. SBIR Phase I to be initiated at company formation. First contacts: '
         'Rashidi-Ranjbar → Jog → Naeser. 2–3 years to published data.'],
        ['CRITICAL', 'RISK-03 — 400 mW/cm² regulatory opinion not obtained (BLOCKING)',
         'This number cannot appear in ANY public-facing material, investor deck or marketing '
         'copy until cleared. Outside regulatory counsel (PBM/digital health specialist), '
         '$8,000–15,000, 3–5 weeks. Scope expanded by NP-REG-PBM1064-001 to cover 1064 nm '
         'irradiance, the 600 mW/cm² aggregate ceiling (OI-PBM-05), T2 combined 1064+1170 nm, and '
         'the depth-tier penetration claim. A Q-13 addition covering Mode F cumulative retinal '
         'dose is still to be added. Note: a NIR wavelength change would likely require the '
         'opinion\'s scope to be re-checked rather than assumed to still cover it.'],
        ['CRITICAL', '"NeurOne" is an uncleared placeholder name',
         'Trademark search and clearance required for US, EU, Canada and Australia before ANY '
         'external conversation. $15,000–25,000.'],
        ['CRITICAL', 'RISK-20 — CFRP shell slot rim Ra ≤ 1.6 µm unconfirmed (BLOCKING)',
         'Tooling cannot be released until written supplier confirmation with Ra measurement data '
         'is received. Supplier qualification item SUP-M-07.'],
        ['CRITICAL', 'RISK-04 — PDMS CAT-C supplier not selected (BLOCKING)',
         'Production cannot start until the CAT-C supplier is selected and the 200-cycle '
         'IEC 60068-2-14 thermal cycling qualification (FAI-TC02) passes. 8-week lead time from '
         'selection.'],
        ['HIGH (OPEN)', 'RISK-26 — fan/heatsink-loss scalp overheat (FMEA-G07-01)',
         'Path B1 selected (scalp-facing NTC co-located with PD2). OPEN pending SW01-M04 '
         'implementation and the THERM-1b bench. All thermal constants provisional pending '
         'verification-grade 2D CFD; path selection is firm.'],
        ['HIGH (OPEN)', 'OI-LED-W1 — no shortlisted high-power NIR LED sits inside the '
         '808–830 nm spec',
         'Decision brief written 2026-07-28, NOT yet decided. See §9.2.'],
        ['HIGH (OPEN)', 'Hub-side fan-out is specified but not implemented (OI-HUB-C01…C19)',
         'None of the Rev C cluster-controller fan-out exists in hardware or firmware yet. The '
         'per-socket FPC pinout and LED drive stage are the largest unwritten items. Rev C is '
         'DRAFT; the socket count and every part number are explicitly open.'],
        ['HIGH (OPEN)', 'Cluster count 18 exceeds the provisioned interconnect (OI-HEXTILE-14)',
         '18 clusters exceed both the 16 provisioned connector positions and the 8 branches × ≤2 '
         'clusters I2C tree. Options at NP-HW-HEXTILE-001 §8.2.2; peer documents not yet revised.'],
        ['HIGH (OPEN)', 'Safety-cut granularity contested (OI-HUB-C07 / OI-HEXTILE-13)',
         'Per-cluster SAFE_EN_n (18 GPIO) versus a single cranial bit (1 GPIO, all-or-nothing). '
         'Proposed resolution at NP-HW-HEXTILE-001 §8.4.1.'],
        ['HIGH (OPEN)', 'OI-CHARGE-04 — tDCS electrode-area assumption diverges',
         'The 40 µC/cm² charge-density limit is computed against a 25 cm² assumption in one place '
         'and 35 cm² in another, between the app pre-flight check and the safety MCU enforcer. '
         'Decision owner: Regulatory / Safety SW.'],
        ['HIGH (OPEN)', 'OI-DOC-01 — NP-FW-HUB-001 is registered and cited as governing but was '
         'never written',
         'docs/np_fw_hub_001.md does not exist. Decision owner: Design Authority.'],
        ['HIGH (OPEN)', 'NP-PWR-BUDGET-001 — TMS has no electrical spec',
         'The design study reversed the derivation direction used by every other power and '
         'thermal document and found that the T2-peak figure never included TMS. The CLAUDE.md '
         '§4.5 power table is therefore incomplete for T2.'],
        ['HIGH (OPEN)', 'T2 hardware FAI items pending',
         'FAI-CV01, FAI-CV02, FAI-CV03, FAI-HD01, FAI-HD03, FAI-HD04 all require a T2 prototype '
         'bench. FAI-CV02 is required for full G3-08 closure; FAI-HD01/03/04 for G3-07.'],
        ['MODERATE', 'SAB not formed — no scientific credentialing',
         'Tsai outreach first (SAB role only — a Cognito Therapeutics conflict prevents a PI '
         'role). Jeffery and Naeser are natural SAB candidates. Budget $50,000–80,000/yr for a '
         '5-person SAB. OI-LED-W1 needs SAB sign-off, so SAB formation is now on the critical '
         'path for a component decision, not only for credentialing.'],
        ['MODERATE', 'QMS — Quality Lead (VP Quality) not yet hired',
         'Month 6 target; interim authority is the CEO. eQMS platform selection targeted Month 9 '
         '(must be 21 CFR Part 11 compliant with DHF and IEC 62304 modules).'],
        ['MODERATE', 'SW-01 safety-MCU unit-level FMEA still outstanding',
         'Required per NP-SW-001 §11 (IEC 62304 §7.1 for Class C). NP-FMEA-001 exists at system '
         'level and carries SW01 module deltas, but the dedicated unit-level FMEA for '
         'SW01-M01..M08 is not complete.'],
        ['MODERATE', 'tFUS / LIFU — modality watch (elevated MEDIUM 2026-05-13)',
         'Science basis confirmed (PIEZO1 mechanism, PNAS 2023; n > 400 human subjects with no '
         'serious AEs at MI < 1.9; MiniUlTra bioadhesive hydrogel coupling, Nature Communications '
         '2025). Cannot enter T1 — no wellness regulatory pathway for brain-targeted therapeutic '
         'ultrasound. A T2 add-on needs a 510(k) predicate that does not yet exist. '
         'Re-evaluation trigger: SPIRE DIADEM FDA clearance (submission 2026, decision '
         '2027–2028). Do NOT include in T1 product claims, BOM or marketing.'],
    ],
    widths=[2.4, 5.2, 9.2])

# ══════════════════════════════════════════════════════════════════════════════
# 9. OPEN ITEMS
# ══════════════════════════════════════════════════════════════════════════════
add_heading('9. Open Items', level=1)

add_heading('9.1 Architecture gates that block downstream work', level=2)
add_table(
    ['Gate', 'What it decides', 'What it blocks'],
    [
        ['REG-1', 'Whether the socket lattice registers to the 10-20 system within tolerance '
                  '(NP-HEX-ZM-001 §7).',
         'Authoring any Fp2/F4-precise zone; the clinical-03 evidence-grade claim; the '
         '"standard electrode site" tactile markers in NP-HFE-002 L1; final zone boundaries. It '
         'no longer blocks hub tooling — Rev C removed that dependency.'],
        ['ACT-1', 'Cluster clamp actuation, jointly with the lattice provisionality.',
         'Treating the 80-socket lattice as final.'],
        ['MECH-1', 'Blind-mate boss tooling.',
         'Time-boxes OI-SHELL2-02 (segregated returns in the boss) — cheap now, expensive once '
         'the boss is tooled.'],
        ['MECH-2', 'Clamp mechanism sizing.',
         'Contact-count-versus-accessibility closure (OI-SHELL2-03).'],
        ['G1 layout package', 'The STM32G071 pin map.',
         'OI-CVNS-08. RPEAK_IN on PB0 versus NP_EN_BES_PIN cannot be settled by reading either '
         'document.'],
    ],
    widths=[2.4, 6.0, 8.4])

add_heading('9.2 OI-LED-W1 — NIR wavelength decision (OPEN, raised 2026-07-28)', level=2)

add_para(
    'No shortlisted high-power NIR LED (≥350 mA class, thermal-pad ceramic SMD) sits inside the '
    '808–830 nm spec. The two electrically-workable candidates — Lumileds L1IZ-0850 (850 nm) and '
    'ams-OSRAM SFH 4718A (860 nm) — are 20 nm and 30 nm over the window. The one in-window '
    'candidate, ams-OSRAM SFH 4703AS (820 nm), has a Vf of about 3.0–3.6 V against the '
    '1.6–2.2 V the existing 660 nm string and driver architecture assumes. Pulse-current rating '
    'is NOT the blocker — that was closed 2026-07-28 from manufacturer datasheets, with the '
    '660 nm primary (ams-OSRAM GH CSSRM5.24) showing 7.8× margin at the 180 mA operating ceiling. '
    'The blocker is wavelength versus electrical compatibility.')

add_para('This is bigger than component sourcing, because 810 nm is already load-bearing:', bold=True)
add_bullets([
    'CLAUDE.md §3 states 808–830 nm in eight separate places.',
    'The competitive position document states the marketed wavelength as 810 nm, not a range, '
    'and one headline claim asserts that 810 nm IS a CCO absorption peak in direct comparison '
    'to Vielight. Shipping 850–860 nm would not widen a spec; it would falsify a stated '
    'comparative claim.',
    'The NIR clinical evidence in the bibliography is almost entirely Vielight Neuro Gamma '
    'studies run at 810 nm specifically. Moving to 850–860 nm weakens the "we replicate the '
    'clinically-tested protocol" argument that underwrites those citations.',
    'Option A (extend to 850/860 nm) does not even solve the electrical problem it is trying to '
    'avoid: L1IZ-0850 has the same elevated Vf (2.7–3.7 V, from DS190 Table 2), so it trades '
    'clinical and marketing fidelity without buying back string/driver compatibility.',
    'Option B (stay in-window with SFH 4703AS at 820 nm) preserves every claim but requires a '
    'separate higher-voltage driver rail, a string topology change and a hub revision — real NRE '
    'and schedule cost. Its own pulse-current rating at 120–180 mA / ≤25% duty was not '
    'independently verified, so that is residual work if Option B is chosen.',
    'Option C (keep searching) is low-probability on two research passes — true 800–830 nm '
    'AlGaAs LEDs at ≥350 mA in a thermal-pad ceramic package appear to be a genuinely '
    'under-served segment, with in-band vendors topping out at 60–100 mA while the high-power '
    'ceramic NIR market has concentrated at 850/940 nm for biometrics and ToF volume. Runs in '
    'parallel as a contingency via OI-LED-07.',
])

add_para(
    'Not decided here. Decision owners: SAB for the science call, Regulatory Counsel for the '
    'RISK-03 scope question, EE Lead for the driver-rework cost. The standing precedent from the '
    'clinical-03 decision applies — withhold the matching-evidence claim rather than quietly '
    'redefine parameters to fit available hardware.', bold=True)

add_heading('9.3 Dose fidelity and optical resolution', level=2)

add_para(
    'NP-OPT-PSF-001 Rev 1 establishes what boundary the hardware can actually produce. The '
    '1064 nm edge-spread function is about 26 mm (10–90%) at cortical depth. Two decisions follow, '
    'both already taken:')
add_bullets([
    'clinical-03 lateralization (2026-07-20): the protocol was retargeted from an 11-socket '
    '"Frontal Right" zone that put 16.3% of its cortical energy contralateral, to an 8-socket '
    '"Frontal Right (excl. midline)" zone at 3.6% — a 78% reduction, at a cost of one zone '
    'definition and zero firmware. The audit that found it is now a build check in '
    'scripts/sync-socket-map.ts: any protocol naming a bare single-hemisphere zone fails the run '
    'until it is narrowed or explicitly accepted with a reason.',
    'Partial-module inclusion is dominated and must not be revisited as a lateralization tool. '
    'Lighting only the ipsilateral half of a midline module still spills 22.1% contralateral '
    'versus 50% lit whole — it removes about 56% of the leak, while simply omitting the module '
    'takes it to zero. The honest form of the argument is that the optics does not prove masking '
    'useless, it proves masking cannot reach zero; since zone narrowing reaches zero for free, '
    'that is enough to close the question.',
])

add_notice(
    'The evidence "site" in the Gonzalez-Lima / Barrett lineage is a single ~13.6 cm² aperture, '
    'and one 40 mm hex module is 13.86 cm² — the same size to within 2%. clinical-03 as now '
    'shipped covers 8 modules, or 111 cm², which is 8.2× the evidence footprint. Narrowing the '
    'zone fixed lateralization; it did not fix this. Separately, the 1064 nm channel is '
    'firmware-capped at 25% duty (np_pbm1064_drive_set_duty(), and the npps schema itself only '
    'accepts duty_cycle 1–25), giving about 0.10 W/cm² average against the literature\'s '
    '0.25 W/cm² CW — 2.5× below — and capped at 36 J/cm² per session against the literature\'s '
    '60 J/cm² per site. Offered the choice between relabeling honestly and raising the firmware '
    'ceilings, the owner chose relabeling; raising safety ceilings would require its own safety '
    'and regulatory review. Decision (2026-07-27, option iii): WITHHOLD the evidence-grade / '
    '"Grade A" / gold-standard characterization from clinical-03 until REG-1 lands and the zone '
    'is re-authored to the true 1–2 module Fp2/F4 footprint. REG-1 landing alone does NOT lift '
    'the gate — the duty and dose shortfall is independent of zone geometry and must be '
    'separately re-evaluated, or explicitly accepted as a permanent disclosed limitation, before '
    'any gold-standard claim is reinstated.',
    colour=NOTICE_BG,
    bold_lead='clinical-03 evidence-grade claim — GATED. ')

add_heading('9.4 Privacy and regulatory program items', level=2)
add_bullets([
    'EU-US Data Privacy Framework self-certification — before any EU resident\'s personal data '
    'is transmitted to US infrastructure. Month 3 target.',
    'Clinician BAA template (NP-LEGAL-BAA-001), including the consent-revocation deletion '
    'cascade clause with a 30-day deletion obligation — before any clinician subscription is '
    'sold.',
    'Research anonymization upgrade (NP-FW-ANON-001): l-diversity l ≥ 3 plus differential '
    'privacy ε ≤ 1.0, with external DP reviewer sign-off — before the first study descriptor is '
    'deployed.',
    'HIPAA Expert Determination certifier engagement (STEP-31) and a signed '
    'NP-ANON-CERT-[study_id] per study (STEP-32) — a BLOCKING gate enforced in the deployment '
    'pipeline before any study descriptor is cryptographically signed.',
    'IRB protocol (NP-IRB-001, template authored 2026-07-27, DRAFT) — before the first study '
    'descriptor reaches any device. Month 9 target.',
    'App Store privacy nutrition label and Google Play Data Safety — before first TestFlight or '
    'Play internal test.',
    'EU data residency for the T2 clinical cloud — Month 9, when the vendor is selected.',
    'T2 scripting API independent security audit (NP-SEC-PENTEST-002) — before any clinical API '
    'key is issued.',
    'NP-SEC-BR-001 annual tabletop exercise — before T1 launch and annually thereafter.',
    'SBOM (NP-SBOM-001) — Year 2; required for the 510(k) cybersecurity submission.',
    'OI-PA-04: plain-language copy for the Adaptive Adjustments card requires Privacy Lead '
    'sign-off before it ships.',
    'OI-BLE-01: SHDRUploader upgrade to the hub-provisioned TRNG warranty token is implemented '
    'app-side on iOS and Android; the residual is external hub firmware (OI-WA-03).',
])

add_para(
    'Geo-gated protections were universalized on 2026-07-10 under the durable principle that any '
    'jurisdiction\'s user protection applies to all users everywhere and is never geo-gated. The '
    'BIPA biometric written release and the Washington MHMD no-sale of SHDR behavioural data now '
    'ship to all users; regulation citations are retained as the WHY, but the controls apply '
    'globally, and the outstanding legal opinions confirm scope rather than gate shipping. '
    'Deliberately left jurisdictional: transfer MECHANISMS (EU-US DPF self-certification, EU data '
    'residency), which are inherently tied to EU data — not user protections.')

# ══════════════════════════════════════════════════════════════════════════════
# 10. COMPETITIVE POSITION
# ══════════════════════════════════════════════════════════════════════════════
add_heading('10. Competitive Position', level=1)

add_notice(
    'Every claim below carries its gate. Claims marked GATED must not appear in any '
    'public-facing material, investor deck, marketing copy or FTC-substantiable statement until '
    'the named gate closes. Two former headline claims are further affected by OI-LED-W1: any '
    'claim asserting 810 nm as the marketed NIR wavelength, or as a CCO absorption peak in '
    'comparison to Vielight, is contingent on Option B being selected.',
    colour=NOTICE_BG,
    bold_lead='CLAIMS GATE. ')

add_table(
    ['Claim', 'Gate'],
    [
        ['Real-time J/cm² dose metering — the only device that shows the exact dose delivered '
         '(dual-PD, per-module, firmware-computed)',
         'UNGATED. Ratio accuracy improves under Hub Rev C because both photodiodes share one '
         'signal chain.'],
        ['The only consumer brain device with palladium-fabric EMF shielding verified by '
         'continuous fleet monitoring', 'UNGATED.'],
        ['Autonomous closed-loop operation from any USB-C power bank — no phone required (Mode 3)',
         'UNGATED.'],
        ['One universal module SKU — any module type in any socket, with software placement '
         'validation', 'UNGATED as an architecture claim. Socket count is PROVISIONAL '
         '(REG-1/ACT-1), so do not quote a specific socket or LED count.'],
        ['LED-count comparisons against Vielight and Neuronic',
         'GATED. The 600-LED figure these claims were built on is retired. Total LED count now '
         'scales with how many T1-A tiles a build populates, so any count claim must be '
         're-derived per configuration before use.'],
        ['400 mW/cm² peak pulsed irradiance; 600 mW/cm² aggregate three-channel ceiling',
         'GATED on RISK-03. Cannot appear in ANY public material until the opinion letter is '
         'received.'],
        ['Three independent penetration depths — 660 nm surface / 1064 nm cortical / 1170 nm deep',
         'GATED on RISK-03 (depth-tier penetration claim is explicitly in the counsel scope).'],
        ['Working memory protocol backed by Yao et al., Science Advances 2022 (right PFC, '
         '1064 nm, n = 20)',
         'GATED on RISK-03 AND on the clinical-03 decision — the evidence-grade / "Grade A" '
         'characterization is withheld until REG-1 lands AND the duty/dose shortfall is '
         'separately resolved or explicitly accepted as a disclosed limitation.'],
    ],
    widths=[8.4, 8.4])

add_para(
    'Critical gap, unchanged: no clinical trial evidence exists anywhere for the combined '
    'multi-modal protocol (PBM + BES + VNS + audio + visual simultaneously). This is '
    'simultaneously the highest-value research opportunity and the primary FTC substantiation '
    'risk for any system-level marketing claim.', bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# 11. REGULATORY
# ══════════════════════════════════════════════════════════════════════════════
add_heading('11. Regulatory Strategy', level=1)

add_para('T1 — FDA-exempt wellness pathway:', bold=True)
add_bullets([
    'General wellness device, the same category as Muse, sens.ai and Apollo Neuro.',
    'Consumer naming: "Brainwave Entrainment Stimulation" (not tACS) and "Cortical Priming '
    'Stimulation" (not tDCS) — deliberate, to avoid triggering FDA medical device classification.',
    'Required standards: IEC 60601-1, IEC 60601-2-10, IEC 62471, IEC 62133, FCC Part 15.',
    'Cybersecurity: SBOM (Year 2), vulnerability disclosure policy, documented OTA update '
    'approach.',
    'FTC claims substantiation: the 39-entry bibliography plus the 6-entry 1064 nm addendum maps '
    'each marketing claim to citations. RISK-03 blocks all public irradiance figures.',
])

add_para('T2 — FDA 510(k):', bold=True)
add_bullets([
    'Modular predicate strategy: TMS (NeuroStar K083538, BrainsWay K122288) + tACS (Soterix '
    'K142485, Neuroelectrics K173185) + taVNS (electroCore K163334, K173323).',
    'Cervical VNS 510(k): predicate electroCore gammaCore K163334 + K173323. Substantial '
    'equivalence argued on same intended use plus same-or-safer technology (≤2 mA versus the '
    'predicate\'s ≤24 mA). Type B Q-Sub meeting targeted Month 20.',
    'Timeline 18–36 months from T1 launch; budget $2–5M.',
    'QMS established 2026-05-13 under 21 CFR Part 820 / ISO 13485:2016; DHF index NP-DHF-001 '
    'Rev 28.',
    'IEC 62304 classifications locked: safety MCU Class C · main processor Class B · app Class B '
    '· cluster controllers Class B.',
    'Human factors per FDA 2016 HFE guidance: NP-HFE-001 plus NP-HFE-002 requirements '
    'HFE-R-01..15 enter the formative and summative plans. No human-factors data exists yet, and '
    'the primary empirical risk — tactile counting on a concave dome rather than a flat page — '
    'must be tested on helmet-representative curvature or the result will be falsely reassuring.',
    'Clinical data for the TMS modality is required; seeding T2 units into research institutions '
    'in Years 2–3 generates it.',
])

add_notice(
    'Development-mode posture: there is no hardware, no bench data and no UHDR record in '
    'existence yet. Document-control corrections to BASELINED specifications are handled as '
    'document control, not as CAPA defect filings, and no QMS defect record is opened for them. '
    'docs/status/completed-decisions.md is APPEND-ONLY — a superseded entry is evidence of what '
    'was believed when it was written, and is never edited or deleted. Specification documents '
    '(docs/np_*.md) have the opposite rule: they ARE edited in place and carry their history in '
    'a Rev header.',
    colour=NOTICE_BG,
    bold_lead='PROCESS NOTE. ')

# ══════════════════════════════════════════════════════════════════════════════
# 12. DOCUMENT REGISTER
# ══════════════════════════════════════════════════════════════════════════════
add_heading('12. Document Register — Design Brief Lineage', level=1)

add_para(
    'The authoritative registers are NP-DHF-001 Rev 28 (formal DHF index, 21 CFR §820.30(j)) and '
    'docs/status/document-register.md. Only the design-brief lineage is reproduced here; the '
    'register below is a pointer, not a duplicate index.', size=9, italic=True)

add_para(
    'Numbering note — RESOLVED by NP-CONV-001 Rev 6 (2026-08-11), which Revision 5 predates. '
    'Revisions 1–5 of this brief were issued as five distinct SERIALS (NP-DB-001 … NP-DB-005) '
    'when each had inherited the previous one\'s decisions and should have taken a new REVISION '
    'of one serial. NP-CONV-001 §4 states the test — new serial when the architecture is '
    'replaced, new revision when decisions are inherited — and §4.0.4 resolves the legacy '
    'explicitly: the five NP-DB serials stay exactly as written, because they are referenced '
    'from the DHF and fall under the 21 CFR §820.180 retention duty, and the correction applies '
    'going forward rather than retroactively (OI-CONV-06, closed on that basis). This document '
    'is therefore NP-DB-005 Revision 6 — a new revision of the live serial, not NP-DB-006 — and '
    'it is written in place at docs/np_db_005.docx. There is no NP-DB-006 and none should be '
    'created for a future revision of this brief.', size=9, italic=True)

add_table(
    ['Revision', 'File', 'Status'],
    [
        ['NP-DB-001 Rev 1', 'docs/superseded/np_db_001.docx', 'SUPERSEDED by NP-DB-005'],
        ['NP-DB-002 Rev 2', 'docs/superseded/np_db_002.docx', 'SUPERSEDED by NP-DB-005'],
        ['NP-DB-003 Rev 3', 'docs/superseded/np_db_003.docx', 'SUPERSEDED by NP-DB-005'],
        ['NP-DB-004 Rev 4', 'docs/superseded/np_db_004.docx', 'SUPERSEDED by NP-DB-005'],
        ['NP-DB-005 Rev 5', 'docs/np_db_005.docx (this path, prior revision)',
         'LAST APPROVED REVISION — retained by version control, not as a second file. '
         'Superseded by Rev 6 on Rev 6 approval. Its retired content is enumerated in §0.1; do '
         'not use those areas as a design input regardless of approval state.'],
        ['NP-DB-005 Rev 6 — THIS DOCUMENT', 'docs/np_db_005.docx',
         'DRAFT — PENDING APPROVAL. Authored 2026-08-11; not yet reviewed or approved.'],
    ],
    widths=[4.6, 7.0, 5.2])

add_para('Primary controlled documents referenced by this brief:', bold=True)
add_table(
    ['Document', 'Rev / status', 'Subject'],
    [
        ['NP-DHF-001', 'Rev 28, ACTIVE', 'Design History File index'],
        ['NP-CONV-001', 'Rev 6, ACTIVE', 'Naming and notation conventions; §4 document IDs, filenames and integer revisions'],
        ['NP-QMS-001 / -DC-001 / -CAPA-001', 'Rev 1, ACTIVE', 'QMS manual, design controls, CAPA'],
        ['NP-RM-001', 'Rev 1, ACTIVE', 'ISO 14971 risk management plan'],
        ['NP-RISK-002 / -003 / -004', 'Rev 1, ACTIVE', 'Risk file index and NP-RISK-001 disposition; hex-tile module risks; shell/socket/interconnect/hub risks. NP-RISK-001 is retired to docs/superseded/'],
        ['NP-SW-001', 'Rev 3, ACTIVE', 'IEC 62304 software development plan; SR-FAN-01…06 in §6.2'],
        ['NP-SW-CI-001', 'Rev 8, DRAFT', 'Firmware cross-compile CI program'],
        ['NP-SOUP-CMSIS-001', 'Rev 1, DRAFT', 'First Class C IEC 62304 §7.1.2 anomaly evaluation'],
        ['NP-HEX-ZM-001', 'Rev 2, DESIGN STUDY', 'Hex-tile module architecture; CLUSTER-1/SYM-1/CONTIG-1'],
        ['NP-HW-HEXTILE-001', 'Rev 3, DESIGN STUDY', 'Hex-tile electrical/FPC spec; 18-cluster derivation'],
        ['NP-HW-HUB-001', 'Rev 4, DRAFT', 'Cluster-controller tier (Hub PCB Rev B archived as Appendix A)'],
        ['NP-DRV-SHELL-002', 'Rev 2, DRAFT', 'Cluster-carrier shell socket interconnect. NP-DRV-SHELL-001 retired to docs/superseded/'],
        ['NP-HELMET-GEOM-001', '—', 'Scan-derived helmet geometry and socket coordinates'],
        ['NP-OPT-PSF-001', 'Rev 1, ACTIVE', 'PBM optical resolution floor; lateralization model'],
        ['NP-HFE-001 / NP-HFE-002', 'Rev 1', 'Human factors plan; accessible position identification'],
        ['NP-THERM-CFD-001 / -R1 / -C2; NP-THERM-BEZEL-001', '—',
         'Thermal architecture; BN-boss export; fan-health path'],
        ['NP-REQ-FANHEALTH-001 / NP-PLAN-FANHEALTH-001', '—', 'RISK-26 mitigation requirements'],
        ['NP-FMEA-001 / NP-FMEA-GEOM-001', '—', 'System FMEA and geometry feeder'],
        ['NP-FW-EMMC-001 / -002', 'Rev 1', 'Storage architecture; warranty token; factory reset; Mode F'],
        ['NP-FW-HD-001 / -CVNS-001 / -HRV-001 / -ZA-001 / -PBM1064-001 / -ANON-001 / -POE-001',
         'various', 'Firmware specifications'],
        ['NP-PRIV-001 / -REM-001 / -NOTICE-001 / -AUDIT-001 / -ANALYSIS-002 / -003 / -KEYSPLIT-001',
         'various', 'Privacy analysis, remediation, notice and audit'],
        ['NP-SEC-BR-001 / NP-PROC-POA-001 / NP-APP-TELEMETRY-001 / NP-LEGAL-BAA-001 / NP-IRB-001',
         'Rev 1–2', 'Breach response, POA, telemetry, BAA, IRB'],
        ['NP-REG-CVNS-001 / NP-REG-PBM1064-001', 'Rev 1', '510(k) Q-Sub; RISK-03 scope expansion'],
        ['NP-TOOL-SHELL-001 / -LENS-001 / -HUB-001 / -HEXTILE-001', 'Rev 1–2', 'Tooling specifications. NP-TOOL-ZM-001 and NP-TOOL-ZM-SM-001 retired to docs/superseded/; NP-TOOL-HEXTILE-001 is the hex-tile successor'],
        ['NP-PROC-SUP-001 / NP-PROC-FPC-001 / NP-PROC-FPC-1064-001', 'Rev 1–2', 'Supplier and FPC procurement'],
        ['NP-NPPS-REF-001', 'Rev 2', 'NPPS protocol scripting reference'],
        ['NP-ART-001 / NP-FAI-001 / NP-FAI-HUB-001 / NP-MOD-ID-001 / NP-REV-SHELL-001', 'Rev 1', 'Artifact register; FAI programme; hub FAI; module identity; shell design review — all authored after Revision 5'],
        ['NP-PWR-BUDGET-001 / NP-ENV-001 / NP-ENV-OPRANGE-001', 'Rev 1',
         'Power budget; environmental and operating range'],
        ['NP-DP-001 / NP-DT-001 / NP-PMS-001', 'Rev 1–2',
         'Design plan; input/output traceability (DI-SAFE-13); post-market surveillance'],
    ],
    widths=[6.4, 3.2, 7.2], size=8)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'NP-DB-005 Rev 6  ·  2026-08-11  ·  Pre-tooling design phase  ·  Supersedes Rev 5  ·  '
    'Decisions marked "locked" have been through full design review; items marked OPEN require '
    'resolution before first tooling cut (docs/status/pending-decisions.md)')
r.font.size = Pt(8)
r.italic = True

out = 'docs/np_db_005.docx'
doc.save(out)
print('WROTE', out)
